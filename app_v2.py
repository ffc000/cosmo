"""
CosmoTools — app.py v2
Plataforma de herramientas DI REPA / ARCA
"""

import os, sqlite3, io, uuid, threading, bcrypt, logging, subprocess, xml.etree.ElementTree as ET
from datetime import datetime, timedelta
from functools import wraps
from flask import Flask, render_template, request, jsonify, send_file, session, redirect, url_for
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "sintia-repa-2026")
app.permanent_session_lifetime = timedelta(hours=4)

# ── Config ─────────────────────────────────────────────────────────────────────
DB_PATH       = os.environ.get("DB_PATH",     "/data/pad.db")
API_KEY       = os.environ.get("ANTHROPIC_API_KEY", "")
APP_USER      = os.environ.get("APP_USER",    "cosmo")
APP_PASS      = os.environ.get("APP_PASS",    "")
APP_USER2     = os.environ.get("APP_USER2",   "")
APP_PASS2     = os.environ.get("APP_PASS2",   "")
OUTPUT_FOLDER = "/data/informes"
HIST_DB       = "/data/historial.db"
os.makedirs(OUTPUT_FOLDER, exist_ok=True)
os.makedirs("/data/minutas", exist_ok=True)
os.makedirs("/tmp/sintia_uploads", exist_ok=True)

# ── Logging ────────────────────────────────────────────────────────────────────
logging.basicConfig(filename="/data/accesos.log", level=logging.INFO,
    format="%(asctime)s %(message)s", datefmt="%Y-%m-%d %H:%M:%S")

# ── Rate limiter ───────────────────────────────────────────────────────────────
limiter = Limiter(key_func=get_remote_address, app=app,
    default_limits=[], storage_uri="memory://")

# ── Job queue ──────────────────────────────────────────────────────────────────
job_status = {}

# ── Historial DB ───────────────────────────────────────────────────────────────
def init_historial():
    con = sqlite3.connect(HIST_DB)
    con.execute("""CREATE TABLE IF NOT EXISTS historial (
        id TEXT PRIMARY KEY, fecha TEXT, usuario TEXT, pais TEXT,
        anio TEXT, mes_d TEXT, mes_h TEXT, uso_ia INTEGER,
        archivo_word TEXT, archivo_excel TEXT, revisado INTEGER DEFAULT 0,
        tipo TEXT DEFAULT 'sintia', descripcion TEXT DEFAULT ''
    )""")
    con.execute("""CREATE TABLE IF NOT EXISTS vua_cronologia (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        fecha TEXT, actividad TEXT, participantes TEXT,
        estado TEXT DEFAULT 'Pendiente', orden INTEGER DEFAULT 0,
        creado TEXT, modificado TEXT
    )""")
    con.execute("""CREATE TABLE IF NOT EXISTS vua_minutas (
        id TEXT PRIMARY KEY, fecha TEXT, asunto TEXT, lugar TEXT,
        participantes TEXT, temas TEXT, acuerdos TEXT, proximos TEXT,
        archivo TEXT, creado_por TEXT, creado TEXT
    )""")
    con.execute("""CREATE TABLE IF NOT EXISTS vua_ejes (
        id TEXT PRIMARY KEY, nombre TEXT, estado TEXT, orden INTEGER DEFAULT 0
    )""")
    cur = con.cursor()
    cur.execute("SELECT COUNT(*) FROM vua_cronologia")
    if cur.fetchone()[0] == 0:
        cronologia = [
            ("08/10/2025","Designación del referente de la DGA ante VUCEA para el proyecto VUA","DI REPA","Completado",1),
            ("14/10/2025","Primera reunión: VUCEA expone el nuevo alcance del proyecto VUA Carga","DI REPA, VUCEA","Completado",2),
            ("04/12/2025","Reunión de relevamiento del proceso de carga aérea","DI REPA, VUCEA","Completado",3),
            ("18/12/2025","Reunión de análisis de la RG 5797/2025 y su implicancia en VUA","DI REPA, VUCEA","Completado",4),
            ("11/02/2026","Demo sistema VUA: primera versión del formulario Plan de Vuelo","DI REPA, VUCEA, SENASA, ORSNA, DI ADEZ, Aerolíneas Argentinas, PSA","Completado",5),
            ("25/02/2026","Reunión presencial en Ezeiza: relevamiento operativo de circuitos de carga","DI REPA, DI ADEZ, VUCEA","Completado",6),
            ("17/03/2026","Reunión presencial con TCA: relevamiento del circuito de desconsolidación","DI REPA, DI ADEZ, VUCEA, TCA","Completado",7),
            ("25/03/2026","Reunión sobre Manifiesto Desconsolidado de Importación","DI REPA, DI ADEZ, VUCEA, TCA","Completado",8),
            ("11/05/2026","Reunión ARCA-SENASA: integración de organismos de control en PAD","DI REPA, SENASA","Completado",9),
            ("14/05/2026","Mesa de trabajo: IA, tablero de vuelos, MANE y marco normativo","DI REPA, DI SADU, DI ADEZ, VUCEA","Completado",10),
            ("21/05/2026","Reunión ampliada con aerolíneas: IA, circuito importación, desconsolidado y DAI","DGA, DI REPA, DI SADU, DI ADEZ, VUCEA, Aerolíneas Argentinas, JURCA, IATA","Completado",11),
            ("A definir","Reunión específica análisis del MANE","DI REPA, DI ADEZ, VUCEA","Pendiente",12),
            ("A definir","Análisis normativo rol de VUCEA como intermediario","DI REPA, áreas legales DGA y VUCEA","Pendiente",13),
            ("A definir","Corrección de formularios Guía Madre por parte de VUCEA","VUCEA","Pendiente",14),
        ]
        con.executemany("INSERT INTO vua_cronologia (fecha,actividad,participantes,estado,orden,creado,modificado) VALUES (?,?,?,?,?,datetime('now'),datetime('now'))", cronologia)
    cur.execute("SELECT COUNT(*) FROM vua_ejes")
    if cur.fetchone()[0] == 0:
        ejes = [
            ("4.1","Transmisión de información anticipada — XML, sujetos obligados y marco sancionatorio","En análisis — requiere definición normativa",1),
            ("4.2","Tablero de programación de vuelos","En análisis técnico interno",2),
            ("4.3","Manifiesto de Exportación (MANE)","Pendiente — sin normativa vigente para IA de exportación",3),
            ("4.4","Manifiestos desconsolidados de importación","Pendiente — sin normativa vigente",4),
            ("4.5","Estándar de transmisión XML — Guía Madre (XFWB)","Postura definida — observaciones comunicadas a VUCEA",5),
        ]
        con.executemany("INSERT INTO vua_ejes VALUES (?,?,?,?)", ejes)
    con.commit(); con.close()

init_historial()

# ── Auth ───────────────────────────────────────────────────────────────────────
def check_password(plain, hashed):
    try: return bcrypt.checkpw(plain.encode(), hashed.encode())
    except: return plain == hashed

def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get("logged_in"):
            return redirect(url_for("login"))
        last = session.get("last_active")
        if last and datetime.now().timestamp() - last > 14400:
            session.clear(); return redirect(url_for("login"))
        session["last_active"] = datetime.now().timestamp()
        session.permanent = True
        return f(*args, **kwargs)
    return decorated

def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if session.get("role") != "admin":
            return jsonify({"ok": False, "error": "Sin permiso"})
        return f(*args, **kwargs)
    return decorated

@app.route("/login", methods=["GET", "POST"])
@limiter.limit("20 per 15 minutes", error_message="Demasiados intentos. Esperá 15 minutos.")
def login():
    error = None
    if request.method == "POST":
        u = request.form.get("user","").strip()
        p = request.form.get("pass","")
        ip = request.headers.get("X-Real-IP", request.remote_addr)
        if u == APP_USER and check_password(p, APP_PASS):
            session.update({"logged_in":True,"role":"admin","username":u,"last_active":datetime.now().timestamp()})
            session.permanent = True
            logging.info(f"LOGIN OK | user={u} | ip={ip}")
            return redirect(url_for("index"))
        elif APP_USER2 and u == APP_USER2 and check_password(p, APP_PASS2):
            session.update({"logged_in":True,"role":"readonly","username":u,"last_active":datetime.now().timestamp()})
            session.permanent = True
            logging.info(f"LOGIN OK | user={u} | ip={ip}")
            return redirect(url_for("index"))
        else:
            logging.warning(f"LOGIN FAIL | user={u} | ip={ip}")
            error = "Usuario o contraseña incorrectos"
    return render_template("login.html", error=error)

@app.route("/logout")
def logout():
    logging.info(f"LOGOUT | user={session.get('username','?')} | ip={request.headers.get('X-Real-IP', request.remote_addr)}")
    session.clear()
    return redirect(url_for("login"))

# ── Index ──────────────────────────────────────────────────────────────────────
@app.route("/")
@login_required
def index():
    db_exists = os.path.exists(DB_PATH)
    db_size   = round(os.path.getsize(DB_PATH)/(1024**3),2) if db_exists else 0
    hoy = datetime.today()
    mes_ult = str(hoy.month-1).zfill(2) if hoy.month > 1 else "12"
    meses = {"01":"Enero","02":"Febrero","03":"Marzo","04":"Abril","05":"Mayo","06":"Junio",
             "07":"Julio","08":"Agosto","09":"Septiembre","10":"Octubre","11":"Noviembre","12":"Diciembre"}
    pendientes = []
    try:
        con = sqlite3.connect(HIST_DB); con.row_factory = sqlite3.Row
        limite = (datetime.now()-timedelta(days=10)).strftime("%Y-%m-%d")
        pendientes = [dict(r) for r in con.execute(
            "SELECT * FROM historial WHERE revisado=0 AND fecha < ? ORDER BY fecha ASC",(limite,)).fetchall()]
        con.close()
    except: pass
    return render_template("index.html",
        db_exists=db_exists, db_size=db_size, now=hoy, mes_ult=mes_ult, meses=meses,
        api_key=bool(API_KEY), role=session.get("role","admin"),
        username=session.get("username",""), pendientes=pendientes)

# ── DB Status ──────────────────────────────────────────────────────────────────
@app.route("/api/db-status")
@login_required
def db_status():
    if not os.path.exists(DB_PATH): return jsonify({"exists":False})
    try:
        con = sqlite3.connect(DB_PATH); cur = con.cursor()
        cur.execute("SELECT name FROM sqlite_master WHERE type='table'")
        tables = [r[0] for r in cur.fetchall()]
        info = {}
        for t in tables:
            try: cur.execute(f"SELECT COUNT(*) FROM {t}"); info[t] = cur.fetchone()[0]
            except: pass
        con.close()
        return jsonify({"exists":True,"tables":info,"size_gb":round(os.path.getsize(DB_PATH)/(1024**3),2)})
    except Exception as e:
        return jsonify({"exists":True,"error":str(e)})

# ── Upload BD ──────────────────────────────────────────────────────────────────
@app.route("/api/upload-db", methods=["POST"])
@login_required
@admin_required
def upload_db():
    if "file" not in request.files: return jsonify({"ok":False,"error":"No se recibió archivo"})
    f = request.files["file"]
    if not f.filename.endswith(".db"): return jsonify({"ok":False,"error":"El archivo debe ser .db"})
    os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)
    f.save(DB_PATH)
    logging.info(f"BD UPLOAD | user={session.get('username')} | size={os.path.getsize(DB_PATH)}")
    return jsonify({"ok":True,"size_gb":round(os.path.getsize(DB_PATH)/(1024**3),2)})

# ── Import SQL ─────────────────────────────────────────────────────────────────
@app.route("/api/import-sql", methods=["POST"])
@login_required
@admin_required
def import_sql():
    if "file" not in request.files: return jsonify({"ok":False,"error":"No se recibió archivo"})
    f = request.files["file"]
    if not f.filename.endswith(".sql"): return jsonify({"ok":False,"error":"El archivo debe ser .sql"})
    os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)
    tmp_sql = "/tmp/import_cosmo.sql"
    f.save(tmp_sql)
    try:
        if os.path.exists(DB_PATH): os.remove(DB_PATH)
        result = subprocess.run(["sqlite3", DB_PATH],
            stdin=open(tmp_sql,"r",encoding="utf-8",errors="replace"),
            capture_output=True, text=True, timeout=1800)
        os.remove(tmp_sql)
        if result.returncode != 0 and result.stderr:
            return jsonify({"ok":False,"error":result.stderr[:500]})
        size = round(os.path.getsize(DB_PATH)/(1024**3),2)
        logging.info(f"SQL IMPORT | user={session.get('username')} | size={size}GB")
        return jsonify({"ok":True,"size_gb":size})
    except subprocess.TimeoutExpired:
        return jsonify({"ok":False,"error":"Timeout — archivo muy grande."})
    except Exception as e:
        return jsonify({"ok":False,"error":str(e)})

# ── Update CSV (async) ─────────────────────────────────────────────────────────
@app.route("/api/update-csv", methods=["POST"])
@login_required
def update_csv():
    tabla = request.form.get("tabla","").strip()
    anio  = request.form.get("anio", str(datetime.today().year))
    if tabla == "DAT": tabla = f"DAT_{anio}"
    if "file" not in request.files:
        return jsonify({"ok":False,"error":"No se recibió archivo"})
    if not os.path.exists(DB_PATH):
        return jsonify({"ok":False,"error":"La BD no está cargada"})
    f = request.files["file"]
    tmp_path = f"/tmp/upload_{uuid.uuid4().hex[:8]}.txt"
    f.save(tmp_path)
    size_kb = round(os.path.getsize(tmp_path)/1024,1)
    job_id = str(uuid.uuid4())[:8]
    job_status[job_id] = {"status":"running","log":[f"Archivo recibido: {size_kb} KB"],"files":[]}
    logging.info(f"CSV UPLOAD | tabla={tabla} | size={size_kb}KB")
    t = threading.Thread(target=_run_csv_job, args=(job_id, tmp_path, tabla))
    t.start()
    return jsonify({"ok":True,"job_id":job_id})

def _run_csv_job(job_id, tmp_path, tabla):
    log = job_status[job_id]["log"]
    try:
        _procesar_csv(tmp_path, tabla, log)
        job_status[job_id]["status"] = "done"
    except Exception as e:
        log.append(f"✗ Error: {e}")
        job_status[job_id]["status"] = "error"
    finally:
        try: os.remove(tmp_path)
        except: pass

def _procesar_csv(tmp_path, tabla, log):
    raw = open(tmp_path, encoding="utf-8", errors="replace").read()
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    try:
        if tabla.startswith("DAT_"):
            log.append("Procesando archivo DAT...")
            lines = [l for l in raw.splitlines()
                     if l.strip() and ";" in l
                     and not l.strip().startswith("---")
                     and not (l.strip().upper().startswith("REGISTRO") and ";" not in l)]
            if len(lines) < 2:
                log.append("✗ No se encontraron datos válidos"); con.close(); return
            headers = [h.strip() for h in lines[0].split(";")]
            log.append(f"Columnas: {len(headers)} | Filas a insertar: {len(lines)-1:,}")
            rows = []
            for line in lines[1:]:
                vals = [v.strip() for v in line.split(";")]
                while len(vals) < len(headers): vals.append(None)
                rows.append(vals[:len(headers)])
            cur.execute(f"DROP TABLE IF EXISTS {tabla}")
            cols_def = ", ".join([f'"{h}" TEXT' for h in headers])
            cur.execute(f"CREATE TABLE {tabla} ({cols_def})")
            placeholders = ", ".join(["?" for _ in headers])
            batch_size = 2000
            for i in range(0, len(rows), batch_size):
                cur.executemany(f"INSERT INTO {tabla} VALUES ({placeholders})", rows[i:i+batch_size])
                if (i+batch_size) % 20000 == 0 or i+batch_size >= len(rows):
                    con.commit()
                    log.append(f"  {min(i+batch_size,len(rows)):,} / {len(rows):,} filas insertadas...")
            log.append("Calculando fechas ISO...")
            for col in ["FECHA_INGRESO_ISO","FECHA_TRANS_ISO"]:
                try: cur.execute(f"ALTER TABLE {tabla} ADD COLUMN {col} TEXT")
                except: pass
            cur.execute(f"""UPDATE {tabla} SET
                FECHA_INGRESO_ISO = CASE
                    WHEN FECHA_INGRESO IS NOT NULL AND length(FECHA_INGRESO)>=10 THEN
                        substr(FECHA_INGRESO,7,4)||'-'||substr(FECHA_INGRESO,4,2)||'-'||substr(FECHA_INGRESO,1,2)||
                        CASE WHEN instr(FECHA_INGRESO,' ')>0 THEN ' '||substr(FECHA_INGRESO,instr(FECHA_INGRESO,' ')+1,5) ELSE ' 00:00' END
                    ELSE NULL END,
                FECHA_TRANS_ISO = CASE
                    WHEN FECHA_TRANS IS NOT NULL AND FECHA_TRANS NOT IN ('-','') AND length(FECHA_TRANS)>=10 THEN
                        substr(FECHA_TRANS,7,4)||'-'||substr(FECHA_TRANS,4,2)||'-'||substr(FECHA_TRANS,1,2)||
                        CASE WHEN instr(FECHA_TRANS,' ')>0 THEN ' '||substr(FECHA_TRANS,instr(FECHA_TRANS,' ')+1,5) ELSE ' 00:00' END
                    ELSE NULL END""")
            con.commit()
            log.append("Creando índice...")
            try: cur.execute(f"CREATE UNIQUE INDEX IF NOT EXISTS idx_{tabla}_key ON {tabla}(OPERACION_PAD_EXT, MIC, TIPO_REGISTRO)")
            except: pass
            con.commit(); con.close()
            log.append(f"✓ {tabla}: {len(rows):,} registros, fechas ISO calculadas, índice creado")

        elif tabla == "RECHAZOS":
            log.append("Procesando archivo RECHAZOS...")
            lines = []
            for l in raw.splitlines():
                s = l.strip()
                if not s or "@" not in s: continue
                if len(s) >= 3 and s[:2].isdigit() and s[2] == "@": s = s[3:]
                parts = [p.strip() for p in s.split("@")][:6]
                if len(parts) >= 2: lines.append(parts)
            if len(lines) < 2:
                log.append("✗ No se encontraron datos válidos"); con.close(); return
            col_map = {"Pais Emisor":"PaisEmisor","Metodo":"Metodo","Nro. MIC/DTA":"NroMic","Fecha":"Fecha","Mensaje":"Mensaje"}
            raw_headers = [col_map.get(h.strip(),h.strip()) for h in lines[0]]
            cols_utiles = ["PaisEmisor","Metodo","NroMic","Fecha","Mensaje"]
            idx_cols = [i for i,h in enumerate(raw_headers) if h in cols_utiles]
            cols_finales = [raw_headers[i] for i in idx_cols]
            rows = [[parts[i].strip() if i < len(parts) else None for i in idx_cols] for parts in lines[1:]]
            log.append(f"Filas a insertar: {len(rows):,}")
            cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='RECHAZOS'")
            if not cur.fetchone():
                cur.execute("CREATE TABLE RECHAZOS (PaisEmisor TEXT, Metodo TEXT, NroMic TEXT, Fecha TEXT, Mensaje TEXT, Fecha_ISO TEXT, Mes TEXT, Anio TEXT)")
            else:
                for col in ["Fecha_ISO","Mes","Anio"]:
                    try: cur.execute(f"ALTER TABLE RECHAZOS ADD COLUMN {col} TEXT")
                    except: pass
            placeholders = ", ".join(["?" for _ in cols_finales])
            cols_str = ", ".join(cols_finales)
            inserted = 0
            batch_size = 5000
            for i in range(0, len(rows), batch_size):
                batch = rows[i:i+batch_size]
                try: cur.executemany(f"INSERT INTO RECHAZOS ({cols_str}) VALUES ({placeholders})", batch); inserted += len(batch)
                except:
                    for row in batch:
                        try: cur.execute(f"INSERT INTO RECHAZOS ({cols_str}) VALUES ({placeholders})", row); inserted += 1
                        except: pass
                con.commit()
                log.append(f"  {min(i+batch_size,len(rows)):,} / {len(rows):,} filas procesadas...")
            log.append("Calculando fechas ISO...")
            cur.execute("""UPDATE RECHAZOS SET
                Fecha_ISO = printf('%04d-%02d-%02d',
                    CAST(SUBSTR(CAST(Fecha AS TEXT),LENGTH(CAST(Fecha AS TEXT))-3,4) AS INT),
                    CAST(SUBSTR(CAST(Fecha AS TEXT),LENGTH(CAST(Fecha AS TEXT))-5,2) AS INT),
                    CAST(SUBSTR(CAST(Fecha AS TEXT),1,LENGTH(CAST(Fecha AS TEXT))-6) AS INT)),
                Anio = CAST(SUBSTR(CAST(Fecha AS TEXT),LENGTH(CAST(Fecha AS TEXT))-3,4) AS INTEGER),
                Mes = CASE CAST(SUBSTR(CAST(Fecha AS TEXT),LENGTH(CAST(Fecha AS TEXT))-5,2) AS INTEGER)
                    WHEN 1 THEN 'ENERO' WHEN 2 THEN 'FEBRERO' WHEN 3 THEN 'MARZO' WHEN 4 THEN 'ABRIL'
                    WHEN 5 THEN 'MAYO' WHEN 6 THEN 'JUNIO' WHEN 7 THEN 'JULIO' WHEN 8 THEN 'AGOSTO'
                    WHEN 9 THEN 'SEPTIEMBRE' WHEN 10 THEN 'OCTUBRE' WHEN 11 THEN 'NOVIEMBRE' WHEN 12 THEN 'DICIEMBRE'
                    END
                WHERE Fecha IS NOT NULL AND LENGTH(CAST(Fecha AS TEXT))>=6""")
            con.commit(); con.close()
            log.append(f"✓ RECHAZOS: {inserted:,} registros insertados, fechas calculadas")
        else:
            con.close()
            log.append(f"✗ Tabla no reconocida: {tabla}")
    except Exception as e:
        log.append(f"✗ Error: {e}")
        try: con.close()
        except: pass

# ── Generar informe (async) ────────────────────────────────────────────────────
def run_job(job_id, pais, anio, mes_d, mes_h, usar_ia, username):
    log = job_status[job_id]["log"]
    try:
        from generar import generar_informe
        archivos = generar_informe(
            ruta_db=DB_PATH, pais=pais, anio=anio, mes_d=mes_d, mes_h=mes_h,
            usar_ia=usar_ia, api_key=API_KEY, carpeta=OUTPUT_FOLDER,
            log_fn=lambda msg: log.append(msg))
        # Guardar en historial
        hist_id = str(uuid.uuid4())[:8]
        word  = next((a for a in archivos if a.endswith(".docx")),"")
        excel = next((a for a in archivos if a.endswith(".xlsx")),"")
        con = sqlite3.connect(HIST_DB)
        con.execute("INSERT INTO historial VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)",
            (hist_id, datetime.now().strftime("%Y-%m-%d %H:%M:%S"), username,
             pais, anio, mes_d, mes_h, int(usar_ia), word, excel, 0, 'sintia',
             f"{pais} {mes_d}-{mes_h}/{anio}"))
        con.commit(); con.close()
        logging.info(f"INFORME OK | user={username} | pais={pais} | {mes_d}-{mes_h}/{anio}")
        job_status[job_id]["status"] = "done"
        job_status[job_id]["files"]  = archivos
    except Exception as e:
        log.append(f"✗ Error: {e}")
        job_status[job_id]["status"] = "error"

@app.route("/api/generar", methods=["POST"])
@login_required
def api_generar():
    if not os.path.exists(DB_PATH):
        return jsonify({"ok":False,"error":"La BD no está cargada"})
    data = request.json or {}
    pais    = data.get("pais","").upper()
    anio    = data.get("anio", str(datetime.today().year))
    mes_d   = str(data.get("mes_d","01")).zfill(2)
    mes_h   = str(data.get("mes_h","12")).zfill(2)
    usar_ia = data.get("usar_ia",True) and bool(API_KEY)
    username = session.get("username","?")
    job_id  = str(uuid.uuid4())[:8]
    job_status[job_id] = {"status":"running","log":["Iniciando generación..."],"files":[]}
    t = threading.Thread(target=run_job, args=(job_id, pais, anio, mes_d, mes_h, usar_ia, username))
    t.start()
    return jsonify({"ok":True,"job_id":job_id})

@app.route("/api/job/<job_id>")
@login_required
def job_poll(job_id):
    info = job_status.get(job_id)
    if not info: return jsonify({"error":"Job no encontrado"})
    return jsonify(info)

@app.route("/api/download/<job_id>/<int:idx>")
@login_required
def download_file(job_id, idx):
    files = job_status.get(job_id,{}).get("files",[])
    if idx >= len(files): return "Archivo no encontrado",404
    path = files[idx]
    if not os.path.exists(path): return "Archivo no encontrado",404
    return send_file(path, as_attachment=True, download_name=os.path.basename(path))

# ── Historial ──────────────────────────────────────────────────────────────────
@app.route("/api/historial")
@login_required
def api_historial():
    try:
        con = sqlite3.connect(HIST_DB); con.row_factory = sqlite3.Row
        rows = [dict(r) for r in con.execute(
            "SELECT * FROM historial ORDER BY fecha DESC LIMIT 100").fetchall()]
        con.close()
        return jsonify({"ok":True,"rows":rows})
    except Exception as e:
        return jsonify({"ok":False,"error":str(e)})

@app.route("/api/historial/completo")
@login_required
def historial_completo():
    con = sqlite3.connect(HIST_DB); con.row_factory = sqlite3.Row
    sintia = [dict(r) for r in con.execute(
        "SELECT id, fecha, usuario, 'sintia' as tipo, pais||' '||mes_d||'-'||mes_h||'/'||anio as descripcion, archivo_word, archivo_excel, revisado FROM historial ORDER BY fecha DESC LIMIT 100").fetchall()]
    minutas = [dict(r) for r in con.execute(
        "SELECT id, creado as fecha, creado_por as usuario, 'minuta' as tipo, asunto as descripcion, archivo as archivo_word, '' as archivo_excel, 1 as revisado FROM vua_minutas ORDER BY creado DESC LIMIT 50").fetchall()]
    con.close()
    todos = sorted(sintia+minutas, key=lambda x: x["fecha"], reverse=True)
    return jsonify({"ok":True,"rows":todos})

@app.route("/api/historial/<hist_id>/revisar", methods=["POST"])
@login_required
def revisar_historial(hist_id):
    accion = (request.json or {}).get("accion","conservar")
    try:
        con = sqlite3.connect(HIST_DB)
        if accion == "eliminar":
            row = con.execute("SELECT archivo_word, archivo_excel FROM historial WHERE id=?",(hist_id,)).fetchone()
            if row:
                for f in [row[0],row[1]]:
                    if f and os.path.exists(f):
                        try: os.remove(f)
                        except: pass
            con.execute("DELETE FROM historial WHERE id=?",(hist_id,))
        else:
            con.execute("UPDATE historial SET revisado=1 WHERE id=?",(hist_id,))
        con.commit(); con.close()
        return jsonify({"ok":True})
    except Exception as e:
        return jsonify({"ok":False,"error":str(e)})

@app.route("/api/historial/<hist_id>/download/<tipo>")
@login_required
def download_historial(hist_id, tipo):
    try:
        con = sqlite3.connect(HIST_DB)
        row = con.execute("SELECT archivo_word, archivo_excel FROM historial WHERE id=?",(hist_id,)).fetchone()
        con.close()
        if not row: return "No encontrado",404
        path = row[0] if tipo == "word" else row[1]
        if not path or not os.path.exists(path): return "Archivo no encontrado",404
        return send_file(path, as_attachment=True, download_name=os.path.basename(path))
    except: return "Error",500

# ── VUA ────────────────────────────────────────────────────────────────────────
ROLES_PREDEFINIDOS = {
    "Diego Bugallo": "Jefe Dpto. Facilitación y Simplificación de Comercio (DI REPA)",
    "Martín Macías": "Jefe Div. Modernización de Procesos Aduaneros (DI REPA)",
    "Hernán Cascón": "Supervisor de Informática Aduanera (DI SADU)",
    "Maximiliano Luengo": "Consejero técnico (DI ADEZ)",
    "Pablo Gómez Valdez": "Consejero técnico (DI ADEZ)",
    "Fabiola Cochello": "Directora VUCEA",
    "Vanesa Franco": "Jefa de Procesos VUCEA",
    "Federico Cáceres": "Sec. Simplificación de Procesos Operativos (DI REPA)",
}

CAMPOS_XFWB = [
    {"tab":"Información general","campo":"Número de guía aérea","campo_xml":"masterDocumentNumber","norma":"IATA Cargo-XML XFWB schema","obligatorio":True,"observacion":"Identificador único de la guía aérea master."},
    {"tab":"Información general","campo":"CUIT del consignatario","campo_xml":"OCI/AR/IMP//CUIT...","norma":"Art. 3° RG 4517/2019","obligatorio":True,"observacion":"Formato exacto: OCI/AR/IMP//CUIT12345678901. Sin este campo el XFWB es inválido para Aduana."},
    {"tab":"Información general","campo":"Notificación (notifyParty)","campo_xml":"ConsignmentType/NotifyParty","norma":"IATA Cargo-XML XFWB schema","obligatorio":True,"observacion":"Nombre, número de cuenta y dirección."},
    {"tab":"Información general","campo":"Agente de aduana","campo_xml":"FreightForwarder rol CustomsBroker","norma":"IATA Cargo-XML, RG 3596/2014","obligatorio":True,"observacion":"Distinto del FreightForwarderParty."},
    {"tab":"Información general","campo":"Remitente — Nombre","campo_xml":"ShipperParty/Name","norma":"IATA Cargo-XML XFWB schema","obligatorio":True,"observacion":""},
    {"tab":"Información general","campo":"Destinatario — Nombre","campo_xml":"ConsigneeParty/Name","norma":"IATA Cargo-XML XFWB schema","obligatorio":True,"observacion":""},
    {"tab":"Carga","campo":"Descripción de la mercancía","campo_xml":"GoodsDescription","norma":"IATA Cargo-XML XFWB schema","obligatorio":True,"observacion":""},
    {"tab":"Carga","campo":"Peso bruto total","campo_xml":"TotalGrossWeight","norma":"IATA Cargo-XML XFWB schema","obligatorio":True,"observacion":"En kilogramos."},
    {"tab":"Vuelo","campo":"Número de vuelo","campo_xml":"FlightBooking/FlightIdentifier","norma":"IATA Cargo-XML XFWB schema","obligatorio":True,"observacion":""},
    {"tab":"Vuelo","campo":"Aeropuerto de origen","campo_xml":"DepartureLocation","norma":"IATA Cargo-XML XFWB schema","obligatorio":True,"observacion":"Código IATA de 3 letras."},
    {"tab":"Vuelo","campo":"Aeropuerto de destino","campo_xml":"ArrivalLocation","norma":"IATA Cargo-XML XFWB schema","obligatorio":True,"observacion":"Código IATA de 3 letras."},
]

REGLAS_BPMN = {
    "EXPO": [
        {"id":"EXPO-001","descripcion":"Generación y Registración del MANE deben ser un único nodo en carril ATA MT","patron":["Generación del MANE","Registración del MANE"],"tipo":"error","norma":"RG 5756/2025"},
        {"id":"EXPO-002","descripcion":"Confirmación de Partida debe estar en carril ATA MT","patron":["Confirmación del MANE"],"tipo":"error","norma":"RG 5756/2025 Art. 3"},
        {"id":"EXPO-004","descripcion":"El carril de bodega compartida debe llamarse ATA CBC, no ATA CVC","patron":["ATA CVC"],"tipo":"advertencia","norma":"RG 5756/2025 Anexo §1.2"},
        {"id":"EXPO-005","descripcion":"Debe existir nodo de Ratificación de Autoría con token NF4","patron_ausente":["Ratificación de autoría","Ratificacion de autoria"],"tipo":"error","norma":"RG 5756/2025 §2.2"},
    ],
    "IMPO": [
        {"id":"IMPO-001","descripcion":"Transmisión IA del ATA MT debe incluir XFWB además del XFFM","patron":["XFFM"],"patron_ausente":["XFWB"],"tipo":"error","norma":"RG 3596/2014"},
        {"id":"IMPO-003","descripcion":"El timer de presentación automática debe ser 15 minutos desde confirmación de arribo","patron_ausente":["15 min","15min"],"tipo":"advertencia","norma":"RG 4517/2019 Art. 7"},
        {"id":"IMPO-005","descripcion":"Ratificación de Autoría debe requerir token NF4","patron_ausente":["NF4","nivel 4","token"],"tipo":"error","norma":"RG 4517/2019"},
    ]
}

SYSTEM_NORMATIVA = """Sos un experto en normativa aduanera argentina especializado en carga aérea.
Normativa clave: RG 3596/2014 (IA vía aérea, XFFM+XFWB, plazo 4hs), RG 4517/2019 (MANI SIM, generación automática, token NF4, 15min post-arribo), RG 5756/2025 (MANE exportación, registro post puesta a bordo, 3hs de partida, ATA CBC).
Respondé citando artículos. Si hay ambigüedad normativa, señalala explícitamente."""

SYSTEM_CORREOS = """Sos un asistente de redacción de correos institucionales para DI REPA de ARCA.
Contexto: proyecto VUA, circuito de carga aérea, XML IATA, MANI SIM y MANE.
Estilo: formal con externos/superiores; informal con colegas (primer nombre, "Abrazo" al cerrar).
Federico Cáceres firma como "Fede" en correos informales.
Incluí siempre: ASUNTO: [texto] al inicio."""

@app.route("/vua")
@login_required
def vua_index():
    con = sqlite3.connect(HIST_DB); con.row_factory = sqlite3.Row
    cronologia = [dict(r) for r in con.execute("SELECT * FROM vua_cronologia ORDER BY orden ASC, id ASC").fetchall()]
    ejes       = [dict(r) for r in con.execute("SELECT * FROM vua_ejes ORDER BY orden ASC").fetchall()]
    minutas    = [dict(r) for r in con.execute("SELECT id,fecha,asunto,lugar,creado_por,creado FROM vua_minutas ORDER BY creado DESC LIMIT 20").fetchall()]
    con.close()
    return render_template("vua.html", roles=ROLES_PREDEFINIDOS, cronologia=cronologia,
        ejes=ejes, minutas=minutas, campos_xfwb=CAMPOS_XFWB,
        role=session.get("role","admin"), username=session.get("username",""))

@app.route("/api/vua/minuta", methods=["POST"])
@login_required
def vua_minuta():
    import json as _json
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    data = request.json or {}
    asunto       = data.get("asunto","")
    fecha        = data.get("fecha", datetime.today().strftime("%d/%m/%Y"))
    lugar        = data.get("lugar","")
    participantes = data.get("participantes",[])
    temas        = data.get("temas",[])
    acuerdos     = data.get("acuerdos",[])
    proximos     = data.get("proximos",[])

    def set_cell_color(cell, hex_color):
        tc=cell._tc; tcPr=tc.get_or_add_tcPr(); shd=OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hex_color); tcPr.append(shd)

    doc = Document()
    for section in doc.sections:
        section.top_margin=Cm(2.5); section.bottom_margin=Cm(2.5)
        section.left_margin=Cm(3); section.right_margin=Cm(2.5)
    titulo = doc.add_paragraph(); titulo.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("ACTA DE REUNIÓN"); run.bold=True; run.font.size=Pt(16)
    run.font.color.rgb=RGBColor(0x24,0x2D,0x4F)
    doc.add_paragraph()
    for label, valor in [("Asunto:",asunto),("Fecha:",fecha),("Lugar:",lugar)]:
        p=doc.add_paragraph(); r1=p.add_run(f"{label} "); r1.bold=True; r1.font.size=Pt(11)
        r2=p.add_run(valor); r2.font.size=Pt(11)
    doc.add_paragraph()
    doc.add_paragraph().add_run("Participantes").bold=True
    table=doc.add_table(rows=1,cols=2); table.style="Table Grid"
    hdr=table.rows[0]
    for i,txt in enumerate(["Nombre","Cargo/Organismo"]):
        hdr.cells[i].text=txt
        hdr.cells[i].paragraphs[0].runs[0].bold=True
        hdr.cells[i].paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        set_cell_color(hdr.cells[i],"242D4F")
    for p in participantes:
        row=table.add_row()
        row.cells[0].text=p.get("nombre","")
        row.cells[1].text=p.get("cargo",ROLES_PREDEFINIDOS.get(p.get("nombre",""),""))
    for titulo_sec, items in [("Temas tratados",temas),("Acuerdos",acuerdos),("Próximos pasos",proximos)]:
        doc.add_paragraph()
        doc.add_paragraph().add_run(titulo_sec).bold=True
        for item in items:
            p=doc.add_paragraph(style="List Bullet"); p.add_run(item).font.size=Pt(11)
    minuta_id = str(uuid.uuid4())[:8]
    fname = f"Acta_{fecha.replace('/','_')}_{asunto[:30].replace(' ','_')}_{minuta_id}.docx"
    ruta = os.path.join("/data/minutas", fname)
    doc.save(ruta)
    con = sqlite3.connect(HIST_DB)
    con.execute("INSERT INTO vua_minutas VALUES (?,?,?,?,?,?,?,?,?,?,datetime('now'))",
        (minuta_id,fecha,asunto,lugar,_json.dumps(participantes),_json.dumps(temas),
         _json.dumps(acuerdos),_json.dumps(proximos),ruta,session.get("username","?")))
    con.commit(); con.close()
    return send_file(ruta, as_attachment=True, download_name=fname,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.route("/api/vua/correo", methods=["POST"])
@login_required
def vua_correo():
    data = request.json or {}
    instruccion = data.get("instruccion","")
    if not instruccion: return jsonify({"ok":False,"error":"Instrucción vacía"})
    try:
        import anthropic, httpx
        client = anthropic.Anthropic(api_key=API_KEY, http_client=httpx.Client(follow_redirects=True))
        msg = client.messages.create(model="claude-haiku-4-5-20251001", max_tokens=1000,
            system=SYSTEM_CORREOS, messages=[{"role":"user","content":instruccion}])
        return jsonify({"ok":True,"texto":msg.content[0].text})
    except Exception as e:
        return jsonify({"ok":False,"error":str(e)})

@app.route("/api/vua/normativa", methods=["POST"])
@login_required
def vua_normativa():
    data = request.json or {}
    pregunta = data.get("pregunta","").strip()
    if not pregunta: return jsonify({"ok":False,"error":"Pregunta vacía"})
    try:
        import anthropic, httpx
        client = anthropic.Anthropic(api_key=API_KEY, http_client=httpx.Client(follow_redirects=True))
        msg = client.messages.create(model="claude-haiku-4-5-20251001", max_tokens=1200,
            system=SYSTEM_NORMATIVA, messages=[{"role":"user","content":pregunta}])
        return jsonify({"ok":True,"respuesta":msg.content[0].text})
    except Exception as e:
        return jsonify({"ok":False,"error":str(e)})

@app.route("/api/vua/bpmn", methods=["POST"])
@login_required
def vua_bpmn():
    if "archivo" not in request.files: return jsonify({"ok":False,"error":"No se recibió archivo"})
    archivo = request.files["archivo"]
    circuito = request.form.get("circuito","AUTO")
    try:
        xml_content = archivo.read().decode("utf-8")
        ET.fromstring(xml_content)
    except Exception as e:
        return jsonify({"ok":False,"error":f"XML inválido: {e}"})
    if circuito == "AUTO":
        upper = xml_content.upper()
        if "EXPORTACI" in upper and "MANE" in upper: circuito = "EXPO"
        elif "IMPORTACI" in upper and "MANI" in upper: circuito = "IMPO"
        else: return jsonify({"ok":False,"error":"No se pudo detectar el circuito."})
    reglas = REGLAS_BPMN.get(circuito,[])
    errores=[]; advertencias=[]
    for regla in reglas:
        hallado=False
        if "patron" in regla:
            for p in regla["patron"]:
                if p.lower() in xml_content.lower():
                    (errores if regla["tipo"]=="error" else advertencias).append(
                        {"id":regla["id"],"descripcion":regla["descripcion"],"norma":regla["norma"],"patron_encontrado":p})
                    hallado=True; break
        if "patron_ausente" in regla and not hallado:
            if not any(p.lower() in xml_content.lower() for p in regla["patron_ausente"]):
                (errores if regla["tipo"]=="error" else advertencias).append(
                    {"id":regla["id"],"descripcion":regla["descripcion"],"norma":regla["norma"],"patron_encontrado":"ausente"})
    return jsonify({"ok":True,"circuito":circuito,"errores":errores,"advertencias":advertencias,"total":len(errores)+len(advertencias)})

@app.route("/api/vua/xfwb")
@login_required
def vua_xfwb():
    import openpyxl
    from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    wb=openpyxl.Workbook(); ws=wb.active; ws.title="Checklist XFWB"
    HDR=PatternFill("solid",fgColor="242D4F"); HDR_F=Font(bold=True,color="FFFFFF",size=10)
    bs=Side(style="thin",color="CCCCCC"); BORDER=Border(left=bs,right=bs,top=bs,bottom=bs)
    headers=["Tab","Campo","Campo XML","Norma","Obligatorio","Observación"]
    ws.append(headers)
    for ci,h in enumerate(headers,1):
        cell=ws.cell(1,ci); cell.fill=HDR; cell.font=HDR_F
        cell.alignment=Alignment(horizontal="center"); cell.border=BORDER
    ALT=PatternFill("solid",fgColor="EEF2F7")
    for ri,campo in enumerate(CAMPOS_XFWB,2):
        row=[campo["tab"],campo["campo"],campo["campo_xml"],campo["norma"],
             "Sí" if campo["obligatorio"] else "No",campo["observacion"]]
        for ci,val in enumerate(row,1):
            cell=ws.cell(ri,ci,val); cell.border=BORDER
            cell.fill=ALT if ri%2==0 else PatternFill(); cell.font=Font(size=10)
    for ci in range(1,len(headers)+1):
        ws.column_dimensions[get_column_letter(ci)].width=25
    buf=io.BytesIO(); wb.save(buf); buf.seek(0)
    return send_file(buf,as_attachment=True,download_name="Checklist_XFWB.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

@app.route("/api/vua/informe")
@login_required
def vua_informe():
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    con=sqlite3.connect(HIST_DB); con.row_factory=sqlite3.Row
    cronologia=[dict(r) for r in con.execute("SELECT * FROM vua_cronologia ORDER BY orden ASC").fetchall()]
    ejes=[dict(r) for r in con.execute("SELECT * FROM vua_ejes ORDER BY orden ASC").fetchall()]
    con.close()
    doc=Document()
    for section in doc.sections:
        section.top_margin=Cm(2.5); section.bottom_margin=Cm(2.5)
        section.left_margin=Cm(3); section.right_margin=Cm(2.5)
    titulo=doc.add_paragraph(); titulo.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=titulo.add_run("PROYECTO VUA — ESTADO DE SITUACIÓN")
    run.bold=True; run.font.size=Pt(16); run.font.color.rgb=RGBColor(0x24,0x2D,0x4F)
    doc.add_paragraph().add_run(f"Generado: {datetime.today().strftime('%d/%m/%Y')}").font.size=Pt(10)
    doc.add_paragraph()
    doc.add_paragraph().add_run("Ejes de trabajo").bold=True
    table=doc.add_table(rows=1,cols=3); table.style="Table Grid"
    for i,h in enumerate(["ID","Eje","Estado"]):
        c=table.rows[0].cells[i]; c.text=h
        c.paragraphs[0].runs[0].bold=True
        c.paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        tc=c._tc; tcPr=tc.get_or_add_tcPr()
        shd=OxmlElement("w:shd"); shd.set(qn("w:val"),"clear")
        shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),"242D4F"); tcPr.append(shd)
    for eje in ejes:
        row=table.add_row()
        row.cells[0].text=eje["id"]; row.cells[1].text=eje["nombre"]; row.cells[2].text=eje["estado"]
    doc.add_paragraph()
    doc.add_paragraph().add_run("Cronología de actividades").bold=True
    table2=doc.add_table(rows=1,cols=4); table2.style="Table Grid"
    for i,h in enumerate(["Fecha","Actividad","Participantes","Estado"]):
        c=table2.rows[0].cells[i]; c.text=h
        c.paragraphs[0].runs[0].bold=True
        c.paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        tc=c._tc; tcPr=tc.get_or_add_tcPr()
        shd=OxmlElement("w:shd"); shd.set(qn("w:val"),"clear")
        shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),"242D4F"); tcPr.append(shd)
    for item in cronologia:
        row=table2.add_row()
        row.cells[0].text=item["fecha"]; row.cells[1].text=item["actividad"]
        row.cells[2].text=item["participantes"]; row.cells[3].text=item["estado"]
    buf=io.BytesIO(); doc.save(buf); buf.seek(0)
    return send_file(buf,as_attachment=True,download_name="Informe_VUA_Estado.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.route("/api/vua/cronologia", methods=["GET"])
@login_required
def vua_cronologia_get():
    con=sqlite3.connect(HIST_DB); con.row_factory=sqlite3.Row
    rows=[dict(r) for r in con.execute("SELECT * FROM vua_cronologia ORDER BY orden ASC, id ASC").fetchall()]
    con.close(); return jsonify({"ok":True,"rows":rows})

@app.route("/api/vua/cronologia", methods=["POST"])
@login_required
def vua_cronologia_add():
    data=request.json or {}
    con=sqlite3.connect(HIST_DB); cur=con.cursor()
    cur.execute("SELECT MAX(orden) FROM vua_cronologia")
    max_orden=cur.fetchone()[0] or 0
    cur.execute("INSERT INTO vua_cronologia (fecha,actividad,participantes,estado,orden,creado,modificado) VALUES (?,?,?,?,?,datetime('now'),datetime('now'))",
        (data.get("fecha","A definir"),data.get("actividad",""),data.get("participantes",""),data.get("estado","Pendiente"),max_orden+1))
    new_id=cur.lastrowid; con.commit(); con.close()
    return jsonify({"ok":True,"id":new_id})

@app.route("/api/vua/cronologia/<int:item_id>", methods=["PUT"])
@login_required
def vua_cronologia_update(item_id):
    data=request.json or {}
    con=sqlite3.connect(HIST_DB)
    con.execute("UPDATE vua_cronologia SET fecha=?,actividad=?,participantes=?,estado=?,modificado=datetime('now') WHERE id=?",
        (data.get("fecha"),data.get("actividad"),data.get("participantes"),data.get("estado"),item_id))
    con.commit(); con.close(); return jsonify({"ok":True})

@app.route("/api/vua/cronologia/<int:item_id>", methods=["DELETE"])
@login_required
@admin_required
def vua_cronologia_delete(item_id):
    con=sqlite3.connect(HIST_DB)
    con.execute("DELETE FROM vua_cronologia WHERE id=?",(item_id,))
    con.commit(); con.close(); return jsonify({"ok":True})

@app.route("/api/vua/ejes/<eje_id>", methods=["PUT"])
@login_required
def vua_ejes_update(eje_id):
    data=request.json or {}
    con=sqlite3.connect(HIST_DB)
    con.execute("UPDATE vua_ejes SET nombre=?,estado=? WHERE id=?",(data.get("nombre"),data.get("estado"),eje_id))
    con.commit(); con.close(); return jsonify({"ok":True})

@app.route("/api/vua/minutas", methods=["GET"])
@login_required
def vua_minutas_list():
    con=sqlite3.connect(HIST_DB); con.row_factory=sqlite3.Row
    rows=[dict(r) for r in con.execute("SELECT id,fecha,asunto,lugar,creado_por,creado FROM vua_minutas ORDER BY creado DESC").fetchall()]
    con.close(); return jsonify({"ok":True,"rows":rows})

@app.route("/api/vua/minutas/<minuta_id>/download")
@login_required
def vua_minuta_download(minuta_id):
    con=sqlite3.connect(HIST_DB); con.row_factory=sqlite3.Row
    row=con.execute("SELECT * FROM vua_minutas WHERE id=?",(minuta_id,)).fetchone()
    con.close()
    if not row: return "No encontrada",404
    row=dict(row)
    if row.get("archivo") and os.path.exists(row["archivo"]):
        return send_file(row["archivo"],as_attachment=True,download_name=os.path.basename(row["archivo"]))
    return "Archivo no encontrado",404

@app.route("/api/vua/minutas/<minuta_id>", methods=["DELETE"])
@login_required
def vua_minuta_delete(minuta_id):
    con=sqlite3.connect(HIST_DB)
    row=con.execute("SELECT archivo FROM vua_minutas WHERE id=?",(minuta_id,)).fetchone()
    if row and row[0] and os.path.exists(row[0]):
        try: os.remove(row[0])
        except: pass
    con.execute("DELETE FROM vua_minutas WHERE id=?",(minuta_id,))
    con.commit(); con.close(); return jsonify({"ok":True})

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=False)
