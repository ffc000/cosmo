"""
Tests de actas.py — generación de Word compartida entre VUA y SENASA.
Sin dependencia de Flask/BD: solo prueba generar_acta_word() directamente.
"""
import os
import sys
import tempfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from docx import Document  # noqa: E402
import actas  # noqa: E402


def test_genera_documento_con_titulo_y_datos_basicos():
    doc = actas.generar_acta_word(
        "ACTA DE REUNIÓN", "09/07/2026", "Reunión de prueba", "ARCA",
        participantes=[], secciones=[])
    textos = [p.text for p in doc.paragraphs]
    assert any("ACTA DE REUNIÓN" in t for t in textos)
    assert any("Reunión de prueba" in t for t in textos)
    assert any("09/07/2026" in t for t in textos)
    assert any("ARCA" in t for t in textos)


def test_tabla_participantes_con_nombre_cargo_organismo():
    doc = actas.generar_acta_word(
        "ACTA", "09/07/2026", "Asunto", "Lugar",
        participantes=[{"nombre": "Fer", "cargo": "Analista", "organismo": "DI REPA"}],
        secciones=[])
    assert len(doc.tables) == 1
    tabla = doc.tables[0]
    assert [c.text for c in tabla.rows[0].cells] == ["Nombre", "Cargo", "Organismo"]
    assert [c.text for c in tabla.rows[1].cells] == ["Fer", "Analista", "DI REPA"]


def test_participante_como_string_usa_roles_predefinidos():
    doc = actas.generar_acta_word(
        "ACTA", "09/07/2026", "Asunto", "Lugar",
        participantes=["Diego Bugallo"],
        secciones=[],
        roles_predefinidos={"Diego Bugallo": "Jefe Dpto. Facilitación"})
    tabla = doc.tables[0]
    assert tabla.rows[1].cells[0].text == "Diego Bugallo"
    assert tabla.rows[1].cells[1].text == "Jefe Dpto. Facilitación"


def test_sin_participantes_no_crea_tabla():
    doc = actas.generar_acta_word("ACTA", "09/07/2026", "Asunto", "Lugar",
                                   participantes=[], secciones=[])
    assert len(doc.tables) == 0


def test_secciones_vacias_se_omiten():
    """Regresión: una sección con items=[] no debe dejar un título 'Acuerdos'
    colgado sin contenido debajo."""
    doc = actas.generar_acta_word(
        "ACTA", "09/07/2026", "Asunto", "Lugar", participantes=[],
        secciones=[("Temas tratados", ["Tema 1"]), ("Acuerdos", []), ("Próximos pasos", ["Paso 1"])])
    textos = [p.text for p in doc.paragraphs]
    assert "Temas tratados" in textos
    assert "Próximos pasos" in textos
    assert "Acuerdos" not in textos


def test_secciones_variables_vua_vs_senasa():
    """VUA manda 3 secciones, SENASA manda 4 (con 'Conclusiones' de más) —
    la función no debe asumir una cantidad fija."""
    doc_vua = actas.generar_acta_word(
        "ACTA VUA", "09/07/2026", "Asunto", "Lugar", participantes=[],
        secciones=[("Temas tratados", ["a"]), ("Acuerdos", ["b"]), ("Próximos pasos", ["c"])])
    doc_senasa = actas.generar_acta_word(
        "ACTA SENASA", "09/07/2026", "Asunto", "Lugar", participantes=[],
        secciones=[("Temas tratados", ["a"]), ("Conclusiones", ["b"]),
                   ("Compromisos", ["c"]), ("Próximos pasos", ["d"])])
    textos_senasa = [p.text for p in doc_senasa.paragraphs]
    assert "Conclusiones" in textos_senasa
    assert "Compromisos" in textos_senasa


def test_documento_se_puede_guardar_y_reabrir():
    doc = actas.generar_acta_word(
        "ACTA", "09/07/2026", "Asunto", "Lugar",
        participantes=[{"nombre": "Test", "cargo": "", "organismo": ""}],
        secciones=[("Temas tratados", ["Tema único"])])
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc.save(path)
    reabierto = Document(path)
    assert any("Tema único" in p.text for p in reabierto.paragraphs)
    os.remove(path)
