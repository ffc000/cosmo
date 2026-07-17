"""
Tests para el informe con análisis de IA (Word) del panel Aduanas del País.
Mockea la llamada real a la API de Anthropic (no gasta tokens reales) pero
ejercita el flujo completo: job async, polling, generación del .docx real,
y reapertura del archivo para confirmar contenido.
"""
import io
import sqlite3
from unittest.mock import patch, MagicMock


def _armar_dat_para_informe(app_module):
    con = sqlite3.connect(app_module.DB_PATH)
    con.execute("DROP TABLE IF EXISTS DAT_2026")
    con.execute("""CREATE TABLE DAT_2026 (
        ADUANA TEXT, MIC TEXT, EST_MIC TEXT, CARGADO TEXT, TIPO_REGISTRO TEXT,
        FECHA_INGRESO_ISO TEXT, ULT_ESTADO TEXT, FECHA_ULT_INT TEXT, OPERACION_PAD_EXT TEXT
    )""")
    con.executemany("INSERT INTO DAT_2026 VALUES (?,?,?,?,?,?,?,?,?)", [
        ("601", "MIC1", "TRANS", "SI", "I", "2026-07-01 08:00", "SAL", "01-07-2026 10:00:00", "OP1"),  # 2h
        ("601", "MIC2", "TRANS", "SI", "I", "2026-07-02 08:00", "SAL", "07-07-2026 13:00:00", "OP2"),  # outlier
    ])
    con.commit(); con.close()

    con = sqlite3.connect(app_module.HIST_DB)
    con.execute("DELETE FROM ref_aduanas"); con.execute("DELETE FROM ref_dira")
    con.execute("INSERT INTO ref_dira VALUES ('MIS', 'DIRA Misiones', 1)")
    con.execute("INSERT INTO ref_aduanas VALUES ('601', 'Posadas', 'MIS')")
    con.commit(); con.close()


def _login(client, test_user):
    client.post("/login", data={"user": test_user["username"], "pass": test_user["password"]})


def _fake_anthropic_texto(texto):
    """Mock del cliente anthropic.Anthropic -- evita llamar a la API real."""
    def _fake_create(*a, **kw):
        resp = MagicMock()
        resp.content = [MagicMock(text=texto)]
        return resp
    mock_client_class = MagicMock()
    mock_client_class.return_value.messages.create = _fake_create
    return mock_client_class


def _esperar_job(client, job_id, intentos=30):
    import time
    for _ in range(intentos):
        r = client.get(f"/api/job/{job_id}")
        d = r.get_json()
        if d["status"] in ("done", "error"):
            return d
        time.sleep(0.1)
    return d


def test_ruta_informe_registrada(app_module):
    rutas = {r.rule for r in app_module.app.url_map.iter_rules()}
    assert "/api/sintia/aduanas_nacional/informe" in rutas


def test_informe_requiere_login(client):
    resp = client.post("/api/sintia/aduanas_nacional/informe", json={})
    assert resp.status_code in (302, 303)


def test_informe_sin_api_key_termina_en_error_prolijo(client, test_user, app_module, monkeypatch):
    _armar_dat_para_informe(app_module)
    monkeypatch.delenv("ANTHROPIC_API_KEY", raising=False)
    _login(client, test_user)

    r = client.post("/api/sintia/aduanas_nacional/informe", json={"anio": "2026", "dira": "MIS"})
    job_id = r.get_json()["job_id"]
    d = _esperar_job(client, job_id)
    assert d["status"] == "error"
    assert any("API key" in m for m in d["log"])


def test_informe_completo_genera_word_descargable(client, test_user, app_module, monkeypatch):
    """El flujo completo: job -> IA (mockeada) -> Word real -> descarga.
    Reabre el .docx generado y confirma que tiene las secciones esperadas,
    no solo que el archivo existe."""
    monkeypatch.setenv("ANTHROPIC_API_KEY", "fake-key-de-test")
    _armar_dat_para_informe(app_module)
    _login(client, test_user)

    texto_ia = ("Resumen ejecutivo: la aduana de Posadas registró un total de dos operaciones "
                "en el período. Se detectó una operación excluida del promedio por demora extendida.")

    with patch("anthropic.Anthropic", _fake_anthropic_texto(texto_ia)):
        r = client.post("/api/sintia/aduanas_nacional/informe",
                         json={"anio": "2026", "dira": "MIS", "umbral_dias": 2})
        assert r.get_json()["ok"] is True
        job_id = r.get_json()["job_id"]
        d = _esperar_job(client, job_id)

    assert d["status"] == "done", f"log: {d['log']}"
    assert any("Informe generado" in m for m in d["log"])
    assert len(d["files"]) == 1
    assert d["files"][0].endswith(".docx")

    resp_dl = client.get(f"/api/download/{job_id}/0")
    assert resp_dl.status_code == 200
    assert "wordprocessingml" in resp_dl.content_type

    from docx import Document
    doc = Document(io.BytesIO(resp_dl.data))
    textos = [p.text for p in doc.paragraphs]
    assert any("ADUANAS DEL PAÍS" in t for t in textos)
    assert any("DIRA Misiones" in t for t in textos)
    assert any("2 días" in t for t in textos)  # umbral usado
    assert any(texto_ia[:30] in t for t in textos)  # la narrativa de la IA está en el doc
    assert any("Metodología" in t for t in textos)
    assert len(doc.tables) == 4  # Indicadores, Evolución mensual, Detalle por aduana, y la sección de Posadas (tiene alerta)


def test_informe_incluye_indicadores_reales_en_la_tabla(client, test_user, app_module, monkeypatch):
    monkeypatch.setenv("ANTHROPIC_API_KEY", "fake-key-de-test")
    _armar_dat_para_informe(app_module)
    _login(client, test_user)

    with patch("anthropic.Anthropic", _fake_anthropic_texto("Análisis de prueba.")):
        r = client.post("/api/sintia/aduanas_nacional/informe",
                         json={"anio": "2026", "dira": "MIS", "umbral_dias": 2})
        d = _esperar_job(client, r.get_json()["job_id"])

    from docx import Document
    doc = Document(io.BytesIO(client.get(f"/api/download/{r.get_json()['job_id']}/0").data))
    tabla_indicadores = doc.tables[0]
    valores = {row.cells[0].text: row.cells[1].text for row in tabla_indicadores.rows[1:]}
    assert valores["Operaciones totales"] == "2"
    assert valores["Salieron (SAL)"] == "2"
    assert valores["En alerta — salió tarde"] == "1"  # la operación de 6 días


def test_verificar_narrativa_aduanas_detecta_numero_no_rastreable(app_module):
    """El chequeo liviano post-IA debe marcar en el log un número que no
    aparece en ningún indicador -- señal de que la IA pudo haber inventado
    algo. No debe bloquear ni modificar el texto, solo avisar."""
    import app as app_module_real
    indicadores = {"total_operaciones": 42, "total_sali": 40, "en_alerta_total": 2}
    avisos = []
    app_module_real._verificar_narrativa_aduanas(
        "Se registraron 42 operaciones, de las cuales 999 mostraron algo raro.",
        indicadores, avisos.append)
    assert any("999" in a for a in avisos)
    assert not any("42" in a for a in avisos)  # 42 sí es un valor real, no debe marcarse


def test_verificar_narrativa_aduanas_sin_numeros_sospechosos_no_avisa(app_module):
    import app as app_module_real
    indicadores = {"total_operaciones": 42, "total_sali": 40, "en_alerta_total": 2}
    avisos = []
    app_module_real._verificar_narrativa_aduanas(
        "Se registraron 42 operaciones, de las cuales 40 salieron correctamente.",
        indicadores, avisos.append)
    assert avisos == []


# ── Evolución mensual por aduana en alerta (Word + Excel) ───────────────────

def _armar_dat_con_y_sin_alerta(app_module):
    """601 tiene una operación en alerta (outlier de 6 días), 602 no tiene
    ninguna -- para confirmar que solo la aduana con alerta aparece en la
    sección nueva."""
    con = sqlite3.connect(app_module.DB_PATH)
    con.execute("DROP TABLE IF EXISTS DAT_2026")
    con.execute("""CREATE TABLE DAT_2026 (
        ADUANA TEXT, MIC TEXT, EST_MIC TEXT, CARGADO TEXT, TIPO_REGISTRO TEXT,
        FECHA_INGRESO_ISO TEXT, ULT_ESTADO TEXT, FECHA_ULT_INT TEXT, OPERACION_PAD_EXT TEXT
    )""")
    con.executemany("INSERT INTO DAT_2026 VALUES (?,?,?,?,?,?,?,?,?)", [
        ("601", "MIC1", "TRANS", "SI", "I", "2026-06-01 08:00", "SAL", "01-06-2026 10:00:00", "OP1"),  # 2h, normal
        ("601", "MIC2", "TRANS", "SI", "I", "2026-07-01 08:00", "SAL", "07-07-2026 13:00:00", "OP2"),  # 6 días, alerta
        ("602", "MIC3", "TRANS", "SI", "I", "2026-07-01 08:00", "SAL", "01-07-2026 09:00:00", "OP3"),  # 1h, sin alerta
    ])
    con.commit(); con.close()

    con = sqlite3.connect(app_module.HIST_DB)
    con.execute("DELETE FROM ref_aduanas"); con.execute("DELETE FROM ref_dira")
    con.execute("INSERT INTO ref_dira VALUES ('MIS', 'DIRA Misiones', 1)")
    con.execute("INSERT INTO ref_aduanas VALUES ('601', 'Posadas', 'MIS')")
    con.execute("INSERT INTO ref_aduanas VALUES ('602', 'Iguazu', 'MIS')")
    con.commit(); con.close()


def test_word_incluye_seccion_solo_para_aduanas_con_alerta(client, test_user, app_module, monkeypatch):
    monkeypatch.setenv("ANTHROPIC_API_KEY", "fake-key-de-test")
    _armar_dat_con_y_sin_alerta(app_module)
    _login(client, test_user)

    with patch("anthropic.Anthropic", _fake_anthropic_texto("Análisis de prueba.")):
        r = client.post("/api/sintia/aduanas_nacional/informe",
                         json={"anio": "2026", "dira": "MIS", "umbral_dias": 2})
        d = _esperar_job(client, r.get_json()["job_id"])
    assert d["status"] == "done", f"log: {d['log']}"

    from docx import Document
    doc = Document(io.BytesIO(client.get(f"/api/download/{r.get_json()['job_id']}/0").data))
    textos = [p.text for p in doc.paragraphs]
    assert any("Evolución mensual — aduanas en alerta" in t for t in textos)
    assert any(t.strip() == "Posadas" for t in textos)   # tiene alerta
    assert not any(t.strip() == "Iguazu" for t in textos)  # no tiene alerta, no debe aparecer

    # 4 tablas: Indicadores, Evolución mensual nacional, Detalle por aduana, y la de Posadas
    assert len(doc.tables) == 4
    # el gráfico de Posadas quedó embebido como imagen
    assert len(doc.inline_shapes) == 1


def test_word_sin_ninguna_alerta_no_agrega_la_seccion(client, test_user, app_module, monkeypatch):
    """Si ninguna aduana tiene alertas, la sección entera (y sus gráficos)
    no debe aparecer -- nada que mostrar, nada que generar."""
    monkeypatch.setenv("ANTHROPIC_API_KEY", "fake-key-de-test")
    con = sqlite3.connect(app_module.DB_PATH)
    con.execute("DROP TABLE IF EXISTS DAT_2026")
    con.execute("""CREATE TABLE DAT_2026 (
        ADUANA TEXT, MIC TEXT, EST_MIC TEXT, CARGADO TEXT, TIPO_REGISTRO TEXT,
        FECHA_INGRESO_ISO TEXT, ULT_ESTADO TEXT, FECHA_ULT_INT TEXT, OPERACION_PAD_EXT TEXT
    )""")
    con.execute("INSERT INTO DAT_2026 VALUES (?,?,?,?,?,?,?,?,?)",
        ("601", "MIC1", "TRANS", "SI", "I", "2026-07-01 08:00", "SAL", "01-07-2026 10:00:00", "OP1"))
    con.commit(); con.close()

    con = sqlite3.connect(app_module.HIST_DB)
    con.execute("DELETE FROM ref_aduanas"); con.execute("DELETE FROM ref_dira")
    con.execute("INSERT INTO ref_dira VALUES ('MIS', 'DIRA Misiones', 1)")
    con.execute("INSERT INTO ref_aduanas VALUES ('601', 'Posadas', 'MIS')")
    con.commit(); con.close()

    _login(client, test_user)
    with patch("anthropic.Anthropic", _fake_anthropic_texto("Análisis de prueba.")):
        r = client.post("/api/sintia/aduanas_nacional/informe",
                         json={"anio": "2026", "dira": "MIS", "umbral_dias": 10})
        d = _esperar_job(client, r.get_json()["job_id"])
    assert d["status"] == "done", f"log: {d['log']}"

    from docx import Document
    doc = Document(io.BytesIO(client.get(f"/api/download/{r.get_json()['job_id']}/0").data))
    textos = [p.text for p in doc.paragraphs]
    assert not any("aduanas en alerta" in t for t in textos)
    assert len(doc.inline_shapes) == 0


def test_excel_incluye_hoja_de_tablas_por_aduana_en_alerta_sin_graficos(client, test_user, app_module):
    _armar_dat_con_y_sin_alerta(app_module)
    _login(client, test_user)

    resp = client.get("/api/sintia/aduanas_nacional/export?anio=2026&dira=MIS&umbral_dias=2")
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(resp.data))
    assert "Evolución - aduanas en alerta" in wb.sheetnames

    ws = wb["Evolución - aduanas en alerta"]
    valores = [row for row in ws.iter_rows(values_only=True) if any(v is not None for v in row)]
    nombres = [v[0] for v in valores]
    assert "Posadas" in nombres
    assert "Iguazu" not in nombres  # sin alerta, no debe tener su propia tabla

    # sin gráficos en el Excel -- openpyxl expone las imágenes insertadas en ws._images
    assert len(ws._images) == 0


def test_excel_sin_ninguna_alerta_no_crea_la_hoja(client, test_user, app_module):
    con = sqlite3.connect(app_module.DB_PATH)
    con.execute("DROP TABLE IF EXISTS DAT_2026")
    con.execute("""CREATE TABLE DAT_2026 (
        ADUANA TEXT, MIC TEXT, EST_MIC TEXT, CARGADO TEXT, TIPO_REGISTRO TEXT,
        FECHA_INGRESO_ISO TEXT, ULT_ESTADO TEXT, FECHA_ULT_INT TEXT, OPERACION_PAD_EXT TEXT
    )""")
    con.execute("INSERT INTO DAT_2026 VALUES (?,?,?,?,?,?,?,?,?)",
        ("601", "MIC1", "TRANS", "SI", "I", "2026-07-01 08:00", "SAL", "01-07-2026 10:00:00", "OP1"))
    con.commit(); con.close()

    con = sqlite3.connect(app_module.HIST_DB)
    con.execute("DELETE FROM ref_aduanas"); con.execute("DELETE FROM ref_dira")
    con.execute("INSERT INTO ref_dira VALUES ('MIS', 'DIRA Misiones', 1)")
    con.execute("INSERT INTO ref_aduanas VALUES ('601', 'Posadas', 'MIS')")
    con.commit(); con.close()

    _login(client, test_user)
    resp = client.get("/api/sintia/aduanas_nacional/export?anio=2026&dira=MIS&umbral_dias=10")
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(resp.data))
    assert "Evolución - aduanas en alerta" not in wb.sheetnames


def test_regresion_informe_queda_registrado_en_historial(client, test_user, app_module, monkeypatch):
    """Regresión: el Word se guarda en OUTPUT_FOLDER, y _limpiar_archivos_
    huerfanos() borra a las 48hs todo lo que haya ahí sin fila en
    historial.archivo_word -- sin este registro, el informe generado
    desaparecería solo aunque el job siguiera diciendo 'done'."""
    monkeypatch.setenv("ANTHROPIC_API_KEY", "fake-key-de-test")
    _armar_dat_para_informe(app_module)
    _login(client, test_user)

    with patch("anthropic.Anthropic", _fake_anthropic_texto("Análisis de prueba.")):
        r = client.post("/api/sintia/aduanas_nacional/informe",
                         json={"anio": "2026", "dira": "MIS", "umbral_dias": 2})
        d = _esperar_job(client, r.get_json()["job_id"])
    assert d["status"] == "done", f"log: {d['log']}"

    ruta_generada = d["files"][0]
    with app_module.get_db(app_module.HIST_DB, row_factory=True) as con:
        fila = con.execute(
            "SELECT * FROM historial WHERE archivo_word=?", (ruta_generada,)).fetchone()
    assert fila is not None, "el informe no quedó registrado en historial -- lo va a borrar la limpieza de huérfanos"
    assert fila["tipo"] == "aduanas_pais"
    assert "DIRA Misiones" in fila["descripcion"]


# ── Caché de 24hs (mejora 3) ──────────────────────────────────────────────────

def test_segunda_llamada_con_mismos_parametros_usa_cache(client, test_user, app_module, monkeypatch):
    monkeypatch.setenv("ANTHROPIC_API_KEY", "fake-key-de-test")
    _armar_dat_para_informe(app_module)
    with sqlite3.connect(app_module.HIST_DB) as con:
        con.execute("DELETE FROM historial WHERE tipo='aduanas_pais'")
    _login(client, test_user)

    llamadas = []
    def _fake_create(*a, **kw):
        llamadas.append(kw["messages"][0]["content"])
        resp = MagicMock()
        resp.content = [MagicMock(text="Análisis de prueba.")]
        return resp
    mock_client_class = MagicMock()
    mock_client_class.return_value.messages.create = _fake_create

    with patch("anthropic.Anthropic", mock_client_class):
        r1 = client.post("/api/sintia/aduanas_nacional/informe",
                          json={"anio": "2026", "dira": "MIS", "umbral_dias": 10})
        assert r1.get_json()["cached"] is False
        d1 = _esperar_job(client, r1.get_json()["job_id"])
        assert d1["status"] == "done"

        r2 = client.post("/api/sintia/aduanas_nacional/informe",
                          json={"anio": "2026", "dira": "MIS", "umbral_dias": 10})
        assert r2.get_json()["cached"] is True
        d2 = client.get(f"/api/job/{r2.get_json()['job_id']}").get_json()
        assert d2["status"] == "done"  # inmediato, sin pasar por el thread

    assert len(llamadas) == 1, "no debería haber llamado a la IA una segunda vez"
    assert d1["files"][0] == d2["files"][0]  # mismo archivo reutilizado


def test_parametros_distintos_no_usan_cache(client, test_user, app_module, monkeypatch):
    """Cambiar el umbral es un informe distinto -- no debería servir el
    cacheado de otro umbral."""
    monkeypatch.setenv("ANTHROPIC_API_KEY", "fake-key-de-test")
    _armar_dat_para_informe(app_module)
    with sqlite3.connect(app_module.HIST_DB) as con:
        con.execute("DELETE FROM historial WHERE tipo='aduanas_pais'")
    _login(client, test_user)

    with patch("anthropic.Anthropic", _fake_anthropic_texto("Análisis de prueba.")):
        r1 = client.post("/api/sintia/aduanas_nacional/informe",
                          json={"anio": "2026", "dira": "MIS", "umbral_dias": 10})
        _esperar_job(client, r1.get_json()["job_id"])

        r2 = client.post("/api/sintia/aduanas_nacional/informe",
                          json={"anio": "2026", "dira": "MIS", "umbral_dias": 2})
        assert r2.get_json()["cached"] is False


def test_buscar_informe_cacheado_respeta_ttl(app_module):
    """Un informe generado hace más de 24hs no debe considerarse cacheado."""
    import app as app_module_real
    with app_module_real.get_db(app_module_real.HIST_DB) as con:
        con.execute("DELETE FROM historial WHERE tipo='aduanas_pais'")
    with app_module_real.get_db(app_module_real.HIST_DB) as con:
        con.execute(
            "INSERT INTO historial VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)",
            ("test-viejo", "2020-01-01 00:00:00", "tester", "dira=MIS;umbral=10",
             "2026", "", "", 1, "/tmp/no_existe_pero_no_importa.docx", "", 0,
             "aduanas_pais", "Aduanas del país — DIRA Misiones (2026)"))
    resultado = app_module_real._buscar_informe_aduanas_cacheado("2026", "MIS", 10)
    assert resultado is None  # tiene más de 24hs, no debe encontrarlo


# ── Prompt específico para Misiones (mejora 4) ───────────────────────────────

def test_prompt_menciona_proyecto_de_ley_solo_para_misiones(app_module, monkeypatch):
    monkeypatch.setenv("ANTHROPIC_API_KEY", "fake-key-de-test")
    import app as app_module_real

    prompts = []
    def _fake_create(*a, **kw):
        prompts.append(kw["messages"][0]["content"])
        resp = MagicMock()
        resp.content = [MagicMock(text="Análisis.")]
        return resp
    mock_client_class = MagicMock()
    mock_client_class.return_value.messages.create = _fake_create

    indicadores = {"total_operaciones": 2, "total_sali": 2, "sali_dentro_umbral": 2,
                    "demora_media_dias": 0.5, "demora_media_fmt": "12h 00m 00s",
                    "en_alerta_bandeja": 0, "en_alerta_demora_larga": 0, "en_alerta_total": 0}

    with patch("anthropic.Anthropic", mock_client_class):
        app_module_real._generar_narrativa_aduanas_ia(
            "2026", "DIRA Misiones", 10, indicadores, [], [], "fake-key")
        app_module_real._generar_narrativa_aduanas_ia(
            "2026", "DIRA Buenos Aires", 10, indicadores, [], [], "fake-key")

    assert "proyecto de ley" in prompts[0]
    assert "24 horas" in prompts[0]
    assert "proyecto de ley" not in prompts[1]


def test_prompt_misiones_calcula_comparacion_en_python_no_en_la_ia(app_module, monkeypatch):
    """La comparación (mayor/menor/igual a 24hs) tiene que venir ya resuelta
    en el prompt -- no se le pide a la IA que haga la cuenta."""
    monkeypatch.setenv("ANTHROPIC_API_KEY", "fake-key-de-test")
    import app as app_module_real

    prompts = []
    def _fake_create(*a, **kw):
        prompts.append(kw["messages"][0]["content"])
        resp = MagicMock()
        resp.content = [MagicMock(text="Análisis.")]
        return resp
    mock_client_class = MagicMock()
    mock_client_class.return_value.messages.create = _fake_create

    # demora media de 12hs -- claramente MENOR a 24hs
    indicadores = {"total_operaciones": 2, "total_sali": 2, "sali_dentro_umbral": 2,
                    "demora_media_dias": 0.5, "demora_media_fmt": "12h 00m 00s",
                    "en_alerta_bandeja": 0, "en_alerta_demora_larga": 0, "en_alerta_total": 0}

    with patch("anthropic.Anthropic", mock_client_class):
        app_module_real._generar_narrativa_aduanas_ia(
            "2026", "DIRA Misiones", 10, indicadores, [], [], "fake-key")

    assert "MENOR" in prompts[0]
