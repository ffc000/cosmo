"""
blueprints/pad_acuatico.py — Módulo Pad Acuático.

Creado a pedido (22/07/2026) reutilizando la misma estructura que ya
funciona en producción para VUA y SENASA: Cronología, Ejes de trabajo y
Minutas están tomados de blueprints/senasa.py (mismo esquema de datos,
mismos endpoints, solo cambia el prefijo de tabla/URL); Glosario está
tomado de blueprints/vua.py. Repositorio NO tiene endpoints acá porque ya
es genérico en app.py (/api/repositorio/<modulo>) -- alcanza con que
"pad_acuatico" esté en MODULOS_REPOSITORIO (core.py).

A propósito NO se incluyeron "Acuerdos/Compromisos" ni "Informe" ni
"Integrantes" (SENASA y VUA sí los tienen) porque no fueron pedidos para
este módulo -- si hacen falta más adelante, agregarlos siguiendo el mismo
patrón que senasa.py.
"""
import os
import json
import uuid
import logging
import threading

from actas import generar_acta_word
from datetime import datetime

from flask import Blueprint, request, jsonify, render_template, session, send_file

from core import (
    HIST_DB, OUTPUT_FOLDER, login_required, modulo_required,
    get_api_key, contexto_repositorio, notificar_telegram,
    job_status, job_create, job_get, _job_persist,
    _normalizar_fecha_a_ddmmaaaa, _validar_fecha_ddmmaaaa,
    validar_enum, ESTADOS_TAREA, get_db,
)

pad_acuatico_bp = Blueprint("pad_acuatico", __name__)

# MÓDULO PAD ACUÁTICO
# ══════════════════════════════════════════════════════════════════════════════

@pad_acuatico_bp.route("/pad_acuatico")
@login_required
@modulo_required("pad_acuatico")
def pad_acuatico_index():
    return render_template("pad_acuatico.html", username=session.get("username", ""),
        role=session.get("role", "admin"))

# ── Cronología ────────────────────────────────────────────────────────────────
@pad_acuatico_bp.route("/api/pad_acuatico/cronologia", methods=["GET"])
@login_required
@modulo_required("pad_acuatico")
def pad_acuatico_crono_list():
    with get_db(HIST_DB, row_factory=True) as con:
        rows = [dict(r) for r in con.execute(
            "SELECT *, CASE WHEN fecha GLOB '[0-9][0-9]/[0-9][0-9]/[0-9][0-9][0-9][0-9]' "
            "THEN substr(fecha,7,4)||substr(fecha,4,2)||substr(fecha,1,2) ELSE '00000000' END as _ord "
            "FROM pad_acuatico_cronologia ORDER BY (estado='Pendiente') DESC, _ord DESC, id ASC").fetchall()]
    for r in rows: r.pop("_ord", None)
    return jsonify({"ok": True, "rows": rows})

@pad_acuatico_bp.route("/api/pad_acuatico/cronologia", methods=["POST"])
@login_required
@modulo_required("pad_acuatico")
def pad_acuatico_crono_add():
    data = request.json or {}
    ok, err = validar_enum(data.get("estado"), ESTADOS_TAREA, "estado")
    if not ok: return jsonify({"ok": False, "error": err}), 400
    fecha = _normalizar_fecha_a_ddmmaaaa(data.get("fecha","")) if data.get("fecha") else "A definir"
    with get_db(HIST_DB) as con:
        cur = con.cursor()
        cur.execute("SELECT MAX(orden) FROM pad_acuatico_cronologia")
        max_o = cur.fetchone()[0] or 0
        cur.execute("INSERT INTO pad_acuatico_cronologia (fecha,actividad,participantes,estado,orden) VALUES (?,?,?,?,?)",
            (fecha, data.get("actividad",""),
             data.get("participantes",""), data.get("estado","Pendiente"), max_o+1))
        new_id = cur.lastrowid
    return jsonify({"ok": True, "id": new_id})

@pad_acuatico_bp.route("/api/pad_acuatico/cronologia/<int:iid>", methods=["PUT"])
@login_required
@modulo_required("pad_acuatico")
def pad_acuatico_crono_update(iid):
    data = request.json or {}
    if "fecha" in data and not _validar_fecha_ddmmaaaa(data["fecha"]):
        return jsonify({"ok": False, "error": f"Fecha inválida: '{data['fecha']}'. Formato requerido: dd/mm/aaaa."}), 400
    ok, err = validar_enum(data.get("estado"), ESTADOS_TAREA, "estado")
    if not ok: return jsonify({"ok": False, "error": err}), 400
    campos = {k: v for k, v in data.items() if k in ("fecha","actividad","participantes","estado")}
    if not campos:
        return jsonify({"ok": False, "error": "Nada para actualizar"}), 400
    set_clause = ", ".join(f"{k}=?" for k in campos)
    with get_db(HIST_DB) as con:
        con.execute(f"UPDATE pad_acuatico_cronologia SET {set_clause}, modificado=datetime('now') WHERE id=?",
            (*campos.values(), iid))
    return jsonify({"ok": True})

@pad_acuatico_bp.route("/api/pad_acuatico/cronologia/<int:iid>", methods=["DELETE"])
@login_required
@modulo_required("pad_acuatico")
def pad_acuatico_crono_delete(iid):
    with get_db(HIST_DB) as con:
        con.execute("DELETE FROM pad_acuatico_cronologia WHERE id=?", (iid,))
    return jsonify({"ok": True})

# ── Ejes de trabajo ───────────────────────────────────────────────────────────
@pad_acuatico_bp.route("/api/pad_acuatico/ejes", methods=["GET"])
@login_required
@modulo_required("pad_acuatico")
def pad_acuatico_ejes_list():
    with get_db(HIST_DB, row_factory=True) as con:
        rows = [dict(r) for r in con.execute(
            "SELECT * FROM pad_acuatico_ejes ORDER BY orden ASC").fetchall()]
    return jsonify({"ok": True, "rows": rows})

@pad_acuatico_bp.route("/api/pad_acuatico/ejes/<int:iid>", methods=["PUT"])
@login_required
@modulo_required("pad_acuatico")
def pad_acuatico_ejes_update(iid):
    data = request.json or {}
    with get_db(HIST_DB) as con:
        con.execute("UPDATE pad_acuatico_ejes SET nombre=?,descripcion=?,estado=? WHERE id=?",
            (data.get("nombre",""), data.get("descripcion",""), data.get("estado",""), iid))
    return jsonify({"ok": True})

@pad_acuatico_bp.route("/api/pad_acuatico/ejes", methods=["POST"])
@login_required
@modulo_required("pad_acuatico")
def pad_acuatico_ejes_create():
    data = request.json or {}
    with get_db(HIST_DB, row_factory=True) as con:
        max_orden = con.execute("SELECT MAX(orden) FROM pad_acuatico_ejes").fetchone()[0] or 0
        con.execute("INSERT INTO pad_acuatico_ejes (nombre, descripcion, estado, orden) VALUES (?,?,?,?)",
            (data.get("nombre",""), data.get("descripcion",""), data.get("estado","Pendiente"), max_orden + 1))
    return jsonify({"ok": True})

@pad_acuatico_bp.route("/api/pad_acuatico/ejes/<int:iid>", methods=["DELETE"])
@login_required
@modulo_required("pad_acuatico")
def pad_acuatico_ejes_delete(iid):
    with get_db(HIST_DB) as con:
        con.execute("DELETE FROM pad_acuatico_ejes WHERE id=?", (iid,))
    return jsonify({"ok": True})

# ── Minutas ───────────────────────────────────────────────────────────────────
@pad_acuatico_bp.route("/api/pad_acuatico/minutas", methods=["GET"])
@login_required
@modulo_required("pad_acuatico")
def pad_acuatico_minutas_list():
    with get_db(HIST_DB, row_factory=True) as con:
        rows = [dict(r) for r in con.execute(
            "SELECT id,fecha,asunto,lugar,creado_por,creado FROM pad_acuatico_minutas ORDER BY creado DESC").fetchall()]
    return jsonify({"ok": True, "rows": rows})

@pad_acuatico_bp.route("/api/pad_acuatico/minuta", methods=["POST"])
@login_required
@modulo_required("pad_acuatico")
def pad_acuatico_minuta_create():
    data = request.json or {}
    minuta_id = str(uuid.uuid4())[:8]
    fecha       = data.get("fecha","")
    asunto      = data.get("asunto","")
    lugar       = data.get("lugar","")
    participantes = data.get("participantes",[])
    temas         = data.get("temas",[])
    conclusiones  = data.get("conclusiones",[])
    compromisos   = data.get("compromisos",[])
    proximos      = data.get("proximos",[])

    doc = generar_acta_word(
        "ACTA DE REUNIÓN — PAD ACUÁTICO / ARCA", fecha, asunto, lugar, participantes,
        secciones=[("Temas tratados", temas), ("Conclusiones", conclusiones),
                   ("Compromisos", compromisos), ("Próximos pasos", proximos)],
    )

    os.makedirs("/data/minutas_pad_acuatico", exist_ok=True)
    fname = f"Acta_PadAcuatico_{fecha.replace('/','_')}_{minuta_id}.docx"
    ruta  = os.path.join("/data/minutas_pad_acuatico", fname)
    doc.save(ruta)

    with get_db(HIST_DB) as con:
        con.execute("INSERT INTO pad_acuatico_minutas VALUES (?,?,?,?,?,?,?,?,?,?,?,datetime('now'))",
            (minuta_id, fecha, asunto, lugar,
             json.dumps(participantes), json.dumps(temas), json.dumps(conclusiones),
             json.dumps(compromisos), json.dumps(proximos), ruta, session.get("username","?")))

    partic_str = ", ".join(
        p.get("nombre",p) if isinstance(p,dict) else str(p) for p in participantes)
    return jsonify({
        "ok": True, "minuta_id": minuta_id,
        "download_url": f"/api/pad_acuatico/minutas/{minuta_id}/download",
        "fname": fname,
        "cronologia_sugerida": {
            "fecha": fecha, "actividad": asunto,
            "participantes": partic_str, "estado": "Completado"
        }
    })

@pad_acuatico_bp.route("/api/pad_acuatico/minutas/<minuta_id>/download")
@login_required
@modulo_required("pad_acuatico")
def pad_acuatico_minuta_download(minuta_id):
    with get_db(HIST_DB, row_factory=True) as con:
        row = con.execute("SELECT archivo FROM pad_acuatico_minutas WHERE id=?", (minuta_id,)).fetchone()
    if not row or not os.path.exists(row["archivo"]):
        return jsonify({"ok": False, "error": "Archivo no encontrado"}), 404
    return send_file(row["archivo"], as_attachment=True,
        download_name=os.path.basename(row["archivo"]),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ── IA (estructurar notas de reunión) ─────────────────────────────────────────
@pad_acuatico_bp.route("/api/pad_acuatico/minuta_ia", methods=["POST"])
@login_required
@modulo_required("pad_acuatico")
def pad_acuatico_minuta_ia():
    api_key = get_api_key()
    if not api_key: return jsonify({"ok": False, "error": "API key no configurada"})
    data = request.json or {}
    notas = data.get("notas","").strip()
    if not notas: return jsonify({"ok": False, "error": "Sin notas"})
    try:
        import anthropic, httpx
        client = anthropic.Anthropic(api_key=api_key, http_client=httpx.Client(follow_redirects=True))
        msg = client.messages.create(
            model="claude-haiku-4-5-20251001", max_tokens=1500,
            system="Sos asistente de DI REPA (ARCA Argentina). Estructurás minutas de reuniones del proyecto Pad Acuático. Respondés solo con JSON válido." + contexto_repositorio("pad_acuatico"),
            messages=[{"role":"user","content":(
                f"Estructurá estas notas de reunión de Pad Acuático en JSON:\n{notas}\n\n"
                'Devolvé: {"asunto":"...","temas":["..."],"conclusiones":["..."],"compromisos":["ORG — compromiso..."],"proximos":["..."]}'
            )}])
        texto = msg.content[0].text.strip().replace("```json","").replace("```","").strip()
        return jsonify({"ok": True, "resultado": json.loads(texto)})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)})

# ── Glosario ──────────────────────────────────────────────────────────────────
@pad_acuatico_bp.route("/api/pad_acuatico/glosario", methods=["GET"])
@login_required
@modulo_required("pad_acuatico")
def pad_acuatico_glosario_list():
    with get_db(HIST_DB, row_factory=True) as con:
        rows = [dict(r) for r in con.execute("SELECT * FROM pad_acuatico_glosario ORDER BY orden, termino").fetchall()]
    return jsonify({"ok": True, "rows": rows})

@pad_acuatico_bp.route("/api/pad_acuatico/glosario", methods=["POST"])
@login_required
@modulo_required("pad_acuatico")
def pad_acuatico_glosario_create():
    data = request.json or {}
    with get_db(HIST_DB) as con:
        con.execute("INSERT INTO pad_acuatico_glosario (termino, definicion, categoria) VALUES (?,?,?)",
            (data.get("termino",""), data.get("definicion",""), data.get("categoria","general")))
    return jsonify({"ok": True})

@pad_acuatico_bp.route("/api/pad_acuatico/glosario/<int:gid>", methods=["PUT"])
@login_required
@modulo_required("pad_acuatico")
def pad_acuatico_glosario_update(gid):
    data = request.json or {}
    with get_db(HIST_DB) as con:
        fields = []; params = []
        for f in ["termino","definicion","categoria"]:
            if f in data: fields.append(f + "=?"); params.append(data[f])
        if fields:
            params.append(gid)
            con.execute("UPDATE pad_acuatico_glosario SET " + ", ".join(fields) + " WHERE id=?", params)
    return jsonify({"ok": True})

@pad_acuatico_bp.route("/api/pad_acuatico/glosario/<int:gid>", methods=["DELETE"])
@login_required
@modulo_required("pad_acuatico")
def pad_acuatico_glosario_delete(gid):
    with get_db(HIST_DB) as con:
        con.execute("DELETE FROM pad_acuatico_glosario WHERE id=?", (gid,))
    return jsonify({"ok": True})

# ── Informe (async) ───────────────────────────────────────────────────────────
@pad_acuatico_bp.route("/api/pad_acuatico/informe")
@login_required
@modulo_required("pad_acuatico")
def pad_acuatico_informe():
    """Genera el informe Pad Acuático en background — misma arquitectura
    de job que VUA/SENASA, y mismo formato visual que los informes SINTIA
    (portada compuesta, índice, encabezado/pie institucional, fuente
    unificada -- ver generar_documento.py) a pedido, 22/07/2026. A
    diferencia de SENASA (que tiene "compromisos pendientes"), acá no hay
    tabla de acuerdos, así que en su lugar el informe incluye el
    glosario."""
    with get_db(HIST_DB, row_factory=True) as con:
        datos = {
            "cronologia": [dict(r) for r in con.execute("SELECT * FROM pad_acuatico_cronologia ORDER BY orden").fetchall()],
            "ejes":       [dict(r) for r in con.execute("SELECT * FROM pad_acuatico_ejes ORDER BY orden").fetchall()],
            "minutas":    [dict(r) for r in con.execute("SELECT * FROM pad_acuatico_minutas ORDER BY creado DESC LIMIT 10").fetchall()],
            "glosario":   [dict(r) for r in con.execute("SELECT * FROM pad_acuatico_glosario ORDER BY orden, termino").fetchall()],
        }
    job_id = str(uuid.uuid4())[:8]
    job_create(job_id, "Generando informe Pad Acuático...", username=session.get("username", "?"))

    def _run(jid, datos):
        log = job_status[jid]["log"]
        try:
            from docx import Document as DocxDoc
            from docx.shared import Cm, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from generar_documento import (
                _unificar_fuente_documento, _agregar_encabezado, _agregar_pie_pagina,
                _ocultar_encabezado_portada, _generar_portada_compuesta,
                _agregar_imagen_portada_ajustada, _insertar_indice, _heading_indexado,
                agregar_tabla_word, kpi_box,
            )
            from generar_utils import pl

            ejes, cronologia, minutas, glosario = datos["ejes"], datos["cronologia"], datos["minutas"], datos["glosario"]
            hoy = datetime.today().strftime("%d/%m/%Y")

            doc = DocxDoc()
            _unificar_fuente_documento(doc)
            for section in doc.sections:
                section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
                section.left_margin = Cm(3); section.right_margin = Cm(2.5)
            _agregar_encabezado(doc, "Dirección de Reingeniería de Procesos Aduaneros")
            _agregar_pie_pagina(doc, "Informe de Avance — Pad Acuático")
            _ocultar_encabezado_portada(doc)

            # Portada
            imagen_portada = _generar_portada_compuesta(
                titulo="INFORME DE AVANCE — PAD ACUÁTICO",
                subtitulo=f"Actualizado al {hoy}",
                meta_lineas=[
                    "Dirección de Reingeniería de Procesos Aduaneros (DG ADUA)",
                    f"Última modificación: {hoy}",
                    "Elaborado por: Sección Simplificación de Procesos Operativos — DI REPA",
                ])
            if imagen_portada:
                _agregar_imagen_portada_ajustada(doc, imagen_portada)
            else:
                titulo = doc.add_heading("INFORME DE AVANCE — PAD ACUÁTICO", 0)
                titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for txt, sz in [("Dirección de Reingeniería de Procesos Aduaneros (DG ADUA)", 12),
                                 (f"Actualizado al {hoy}", 11)]:
                    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(txt); run.font.size = Pt(sz); run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
                doc.add_paragraph()
                dest = doc.add_paragraph(); dest.alignment = WD_ALIGN_PARAGRAPH.CENTER
                dest.add_run("Elaborado por: ").bold = True
                dest.add_run("Sección Simplificación de Procesos Operativos — DI REPA")
            doc.add_page_break()

            # Índice -- Ejes y Cronología son fijos (1. y 2.); Minutas y
            # Glosario son condicionales y su número depende de qué haya
            # antes, así que se arman con un contador en vez de hardcodear
            # "3."/"4." (a diferencia del patrón de generar_documento.py,
            # que sí hardcodea porque ahí la mayoría de las secciones son
            # opcionales y no vale la pena calcularlo bien).
            secciones = [(1, "Resumen Ejecutivo"), (1, "1.  Ejes de trabajo"), (1, "2.  Cronología de reuniones")]
            num = 2  # último número fijo usado (2. Cronología) -- el primer condicional es num+1 = 3.
            idx_doc = 4  # posición secuencial para _heading_indexado (bookmarks); ver mismo criterio en generar_documento.py
            idx_minutas = idx_glosario = None
            num_minutas = num_glosario = None
            if minutas:
                num += 1; num_minutas = num
                secciones.append((1, f"{num}.  Minutas generadas")); idx_minutas = idx_doc; idx_doc += 1
            if glosario:
                num += 1; num_glosario = num
                secciones.append((1, f"{num}.  Glosario")); idx_glosario = idx_doc; idx_doc += 1
            _insertar_indice(doc, secciones)

            # Resumen ejecutivo
            _heading_indexado(doc, "Resumen Ejecutivo", 1, 1)
            doc.add_paragraph(
                "El presente informe resume el estado de avance del proyecto Pad Acuático al "
                f"{hoy}: ejes de trabajo definidos, cronología de reuniones realizadas y "
                "pendientes, minutas generadas y glosario de términos del proyecto.")
            en_curso = sum(1 for e in ejes if "curso" in e.get("estado", "").lower() or "análisis" in e.get("estado", "").lower() or "diseño" in e.get("estado", "").lower())
            completados = sum(1 for e in ejes if "completado" in e.get("estado", "").lower())
            pendientes_crono = sum(1 for c in cronologia if c.get("estado") == "Pendiente")
            kpi_box(doc, [
                ("EJES DE TRABAJO", str(len(ejes)), f"{completados} completados, {en_curso} en curso"),
                ("REUNIONES", str(len(cronologia)), f"{pendientes_crono} pendientes"),
                ("MINUTAS GENERADAS", str(len(minutas)), ""),
                ("TÉRMINOS EN GLOSARIO", str(len(glosario)), ""),
            ])
            doc.add_page_break()

            # 1. Ejes de trabajo
            _heading_indexado(doc, "1.  Ejes de trabajo", 1, 2)
            if ejes:
                doc.add_paragraph(f"Se relevaron {len(ejes)} {pl(len(ejes), 'eje de trabajo', 'ejes de trabajo')} del proyecto.")
                for e in ejes:
                    p = doc.add_paragraph(); p.add_run(e["nombre"]).bold = True
                    if e.get("descripcion"): doc.add_paragraph(e["descripcion"])
                    ep = doc.add_paragraph(); ep.add_run("Estado: ").bold = True; ep.add_run(e["estado"])
            else:
                doc.add_paragraph("Todavía no hay ejes de trabajo cargados.")
            doc.add_page_break()

            # 2. Cronología de reuniones
            _heading_indexado(doc, "2.  Cronología de reuniones", 1, 3)
            if cronologia:
                doc.add_paragraph(f"Se registraron {len(cronologia)} {pl(len(cronologia), 'evento', 'eventos')} en la cronología del proyecto.")
                agregar_tabla_word(doc, ["FECHA", "ACTIVIDAD", "PARTICIPANTES", "ESTADO"],
                    [[c["fecha"], c["actividad"], c.get("participantes", ""), c["estado"]] for c in cronologia],
                    col_widths=[2.2, 6.5, 4.5, 2.3])
            else:
                doc.add_paragraph("Todavía no hay entradas en la cronología.")

            # 3. Minutas generadas
            if minutas:
                doc.add_page_break()
                _heading_indexado(doc, f"{num_minutas}.  Minutas generadas", 1, idx_minutas)
                doc.add_paragraph(f"Últimas {len(minutas)} {pl(len(minutas), 'minuta generada', 'minutas generadas')} para el proyecto.")
                agregar_tabla_word(doc, ["FECHA", "ASUNTO", "LUGAR", "GENERADA POR"],
                    [[m["fecha"], m["asunto"], m.get("lugar", ""), m.get("creado_por", "")] for m in minutas],
                    col_widths=[2.2, 6.5, 4, 2.8])

            # 4. Glosario
            if glosario:
                doc.add_page_break()
                _heading_indexado(doc, f"{num_glosario}.  Glosario", 1, idx_glosario)
                for g in glosario:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(g["termino"] + ": ").bold = True
                    p.add_run(g["definicion"])

            os.makedirs(OUTPUT_FOLDER, exist_ok=True)
            fname = f"Informe_PadAcuatico_{datetime.today().strftime('%Y%m%d_%H%M')}_{jid}.docx"
            dest  = os.path.join(OUTPUT_FOLDER, fname)
            doc.save(dest)
            job_status[jid]["files"] = [dest]
            log.append(f"✓ Informe generado: {fname}")
            job_status[jid]["status"] = "done"
            _job_persist(jid)
            notificar_telegram(f"✓ Informe Pad Acuático listo ({job_status[jid].get('username','?')})")
        except Exception as e:
            log.append(f"✗ {e}")
            job_status[jid]["status"] = "error"
            _job_persist(jid)
            notificar_telegram(f"⚠️ Informe Pad Acuático falló ({job_status[jid].get('username','?')}): {e}")

    threading.Thread(target=_run, args=(job_id, datos)).start()
    return jsonify({"ok": True, "job_id": job_id})

@pad_acuatico_bp.route("/api/pad_acuatico/informe/download/<job_id>")
@login_required
@modulo_required("pad_acuatico")
def pad_acuatico_informe_download(job_id):
    import re, glob
    job = job_get(job_id)
    if job and job.get("status") == "done" and job.get("files") and os.path.exists(job["files"][0]):
        return send_file(job["files"][0], as_attachment=True,
            download_name=os.path.basename(job["files"][0]),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    if not re.match(r'^[a-zA-Z0-9]{1,32}$', job_id or ""):
        return jsonify({"ok": False, "error": "job_id inválido"}), 400
    archivos = sorted(glob.glob(os.path.join(OUTPUT_FOLDER, f"*{job_id}*.docx")), key=os.path.getmtime, reverse=True)
    if archivos:
        return send_file(archivos[0], as_attachment=True, download_name=os.path.basename(archivos[0]),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    return jsonify({"ok": False, "error": "Informe no encontrado"}), 404
