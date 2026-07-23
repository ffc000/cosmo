"""
blueprints/senasa.py — Módulo SENASA: integración PAD-SENASA (SIG-Embalajes,
NIMF N°15), cronología, ejes de trabajo, minutas, acuerdos, informe de avance.

Cuarto blueprint extraído de app.py en la Fase 2 de profesionalización.
Sigue finanzas (el más grande y el que requiere más cuidado por los datos
personales/DDJJ).

Nota: /api/integrantes/* NO está acá — es compartido con VUA/SINTIA y se
queda en app.py.
"""
import os
import re
import json
import uuid
import logging
import threading

from actas import generar_acta_word
from datetime import datetime, date, timedelta

from flask import Blueprint, request, jsonify, render_template, session, send_file

from core import (
    HIST_DB, OUTPUT_FOLDER, login_required, modulo_required,
    get_api_key, contexto_repositorio, notificar_telegram,
    job_status, job_create, job_get, _job_persist,
    _normalizar_fecha_a_ddmmaaaa, _validar_fecha_ddmmaaaa,
    validar_enum, ESTADOS_TAREA, get_db,
)

senasa_bp = Blueprint("senasa", __name__)

# MÓDULO SENASA
# ══════════════════════════════════════════════════════════════════════════════

@senasa_bp.route("/senasa")
@login_required
@modulo_required("senasa")
def senasa_index():
    return render_template("senasa.html", username=session.get("username",""),
        role=session.get("role","admin"))

# ── Cronología SENASA ─────────────────────────────────────────────────────────
@senasa_bp.route("/api/senasa/cronologia", methods=["GET"])
@login_required
@modulo_required("senasa")
def senasa_crono_list():
    with get_db(HIST_DB, row_factory=True) as con:
        rows = [dict(r) for r in con.execute(
            "SELECT *, CASE WHEN fecha GLOB '[0-9][0-9]/[0-9][0-9]/[0-9][0-9][0-9][0-9]' "
            "THEN substr(fecha,7,4)||substr(fecha,4,2)||substr(fecha,1,2) ELSE '00000000' END as _ord "
            "FROM senasa_cronologia ORDER BY (estado='Pendiente') DESC, _ord DESC, id ASC").fetchall()]
    for r in rows: r.pop("_ord", None)
    return jsonify({"ok": True, "rows": rows})

@senasa_bp.route("/api/senasa/cronologia", methods=["POST"])
@login_required
@modulo_required("senasa")
def senasa_crono_add():
    data = request.json or {}
    ok, err = validar_enum(data.get("estado"), ESTADOS_TAREA, "estado")
    if not ok: return jsonify({"ok": False, "error": err}), 400
    fecha = _normalizar_fecha_a_ddmmaaaa(data.get("fecha","")) if data.get("fecha") else "A definir"
    with get_db(HIST_DB) as con:
        cur = con.cursor()
        cur.execute("SELECT MAX(orden) FROM senasa_cronologia")
        max_o = cur.fetchone()[0] or 0
        cur.execute("INSERT INTO senasa_cronologia (fecha,actividad,participantes,estado,orden) VALUES (?,?,?,?,?)",
            (fecha, data.get("actividad",""),
             data.get("participantes",""), data.get("estado","Pendiente"), max_o+1))
        new_id = cur.lastrowid
    return jsonify({"ok": True, "id": new_id})

@senasa_bp.route("/api/senasa/cronologia/<int:iid>", methods=["PUT"])
@login_required
@modulo_required("senasa")
def senasa_crono_update(iid):
    data = request.json or {}
    if "fecha" in data and not _validar_fecha_ddmmaaaa(data["fecha"]):
        return jsonify({"ok": False, "error": f"Fecha inválida: '{data['fecha']}'. Formato requerido: dd/mm/aaaa."}), 400
    ok, err = validar_enum(data.get("estado"), ESTADOS_TAREA, "estado")
    if not ok: return jsonify({"ok": False, "error": err}), 400
    campos = {k: v for k, v in data.items() if k in ("fecha","actividad","participantes","estado")}
    if not campos:
        return jsonify({"ok": False, "error": "Nada para actualizar"}), 400
    set_clause = ", ".join(f"{k}=?" for k in campos)
    with get_db(HIST_DB) as con:
        con.execute(f"UPDATE senasa_cronologia SET {set_clause}, modificado=datetime('now') WHERE id=?",
            (*campos.values(), iid))
    return jsonify({"ok": True})

@senasa_bp.route("/api/senasa/cronologia/<int:iid>", methods=["DELETE"])
@login_required
@modulo_required("senasa")
def senasa_crono_delete(iid):
    with get_db(HIST_DB) as con:
        con.execute("DELETE FROM senasa_cronologia WHERE id=?", (iid,))
    return jsonify({"ok": True})

# ── Ejes SENASA ───────────────────────────────────────────────────────────────
@senasa_bp.route("/api/senasa/ejes", methods=["GET"])
@login_required
@modulo_required("senasa")
def senasa_ejes_list():
    with get_db(HIST_DB, row_factory=True) as con:
        rows = [dict(r) for r in con.execute(
            "SELECT * FROM senasa_ejes ORDER BY orden ASC").fetchall()]
    return jsonify({"ok": True, "rows": rows})

@senasa_bp.route("/api/senasa/ejes/<int:iid>", methods=["PUT"])
@login_required
@modulo_required("senasa")
def senasa_ejes_update(iid):
    data = request.json or {}
    with get_db(HIST_DB) as con:
        con.execute("UPDATE senasa_ejes SET nombre=?,descripcion=?,estado=? WHERE id=?",
            (data.get("nombre",""), data.get("descripcion",""), data.get("estado",""), iid))
    return jsonify({"ok": True})

@senasa_bp.route("/api/senasa/ejes", methods=["POST"])
@login_required
@modulo_required("senasa")
def senasa_ejes_create():
    data = request.json or {}
    with get_db(HIST_DB, row_factory=True) as con:
        max_orden = con.execute("SELECT MAX(orden) FROM senasa_ejes").fetchone()[0] or 0
        con.execute("INSERT INTO senasa_ejes (nombre, descripcion, estado, orden) VALUES (?,?,?,?)",
            (data.get("nombre",""), data.get("descripcion",""), data.get("estado","Pendiente"), max_orden + 1))
    return jsonify({"ok": True})

@senasa_bp.route("/api/senasa/ejes/<int:iid>", methods=["DELETE"])
@login_required
@modulo_required("senasa")
def senasa_ejes_delete(iid):
    with get_db(HIST_DB) as con:
        con.execute("DELETE FROM senasa_ejes WHERE id=?", (iid,))
    return jsonify({"ok": True})

# ── Minutas SENASA ────────────────────────────────────────────────────────────
@senasa_bp.route("/api/senasa/minutas", methods=["GET"])
@login_required
@modulo_required("senasa")
def senasa_minutas_list():
    with get_db(HIST_DB, row_factory=True) as con:
        rows = [dict(r) for r in con.execute(
            "SELECT id,fecha,asunto,lugar,creado_por,creado FROM senasa_minutas ORDER BY creado DESC").fetchall()]
    return jsonify({"ok": True, "rows": rows})

@senasa_bp.route("/api/senasa/minuta", methods=["POST"])
@login_required
@modulo_required("senasa")
def senasa_minuta_create():
    data = request.json or {}
    minuta_id = str(uuid.uuid4())[:8]
    fecha       = data.get("fecha","")
    asunto      = data.get("asunto","")
    lugar       = data.get("lugar","")
    participantes = data.get("participantes",[])
    temas         = data.get("temas",[])
    conclusiones  = data.get("conclusiones",[])
    compromisos   = data.get("compromisos",[])
    proximos      = data.get("proximos",[])

    doc = generar_acta_word(
        "ACTA DE REUNIÓN — SENASA / ARCA", fecha, asunto, lugar, participantes,
        secciones=[("Temas tratados", temas), ("Conclusiones", conclusiones),
                   ("Compromisos", compromisos), ("Próximos pasos", proximos)],
    )

    os.makedirs("/data/minutas_senasa", exist_ok=True)
    fname = f"Acta_SENASA_{fecha.replace('/','_')}_{minuta_id}.docx"
    ruta  = os.path.join("/data/minutas_senasa", fname)
    doc.save(ruta)

    with get_db(HIST_DB) as con:
        con.execute("INSERT INTO senasa_minutas VALUES (?,?,?,?,?,?,?,?,?,?,?,datetime('now'))",
            (minuta_id, fecha, asunto, lugar,
             json.dumps(participantes), json.dumps(temas), json.dumps(conclusiones),
             json.dumps(compromisos), json.dumps(proximos), ruta, session.get("username","?")))

    partic_str = ", ".join(
        p.get("nombre",p) if isinstance(p,dict) else str(p) for p in participantes)
    return jsonify({
        "ok": True, "minuta_id": minuta_id,
        "download_url": f"/api/senasa/minutas/{minuta_id}/download",
        "fname": fname,
        "cronologia_sugerida": {
            "fecha": fecha, "actividad": asunto,
            "participantes": partic_str, "estado": "Completado"
        }
    })

@senasa_bp.route("/api/senasa/minutas/<minuta_id>/download")
@login_required
@modulo_required("senasa")
def senasa_minuta_download(minuta_id):
    with get_db(HIST_DB, row_factory=True) as con:
        row = con.execute("SELECT archivo FROM senasa_minutas WHERE id=?", (minuta_id,)).fetchone()
    if not row or not os.path.exists(row["archivo"]):
        return jsonify({"ok": False, "error": "Archivo no encontrado"}), 404
    return send_file(row["archivo"], as_attachment=True,
        download_name=os.path.basename(row["archivo"]),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ── IA SENASA ─────────────────────────────────────────────────────────────────
@senasa_bp.route("/api/senasa/minuta_ia", methods=["POST"])
@login_required
@modulo_required("senasa")
def senasa_minuta_ia():
    api_key = get_api_key()
    if not api_key: return jsonify({"ok": False, "error": "API key no configurada"})
    data = request.json or {}
    notas = data.get("notas","").strip()
    if not notas: return jsonify({"ok": False, "error": "Sin notas"})
    try:
        import anthropic, httpx
        client = anthropic.Anthropic(api_key=api_key, http_client=httpx.Client(follow_redirects=True))
        msg = client.messages.create(
            model="claude-haiku-4-5-20251001", max_tokens=1500,
            system="Sos asistente de DI REPA (ARCA Argentina). Estructurás minutas de reuniones con SENASA. Respondés solo con JSON válido." + contexto_repositorio("senasa"),
            messages=[{"role":"user","content":(
                f"Estructurá estas notas de reunión SENASA-ARCA en JSON:\n{notas}\n\n"
                'Devolvé: {"asunto":"...","temas":["..."],"conclusiones":["..."],"compromisos":["ORG — compromiso..."],"proximos":["..."]}'
            )}])
        texto = msg.content[0].text.strip().replace("```json","").replace("```","").strip()
        return jsonify({"ok": True, "resultado": json.loads(texto)})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)})

# ── Acuerdos SENASA ───────────────────────────────────────────────────────────
@senasa_bp.route("/api/senasa/acuerdos", methods=["GET"])
@login_required
@modulo_required("senasa")
def senasa_acuerdos_list():
    with get_db(HIST_DB, row_factory=True) as con:
        rows = [dict(r) for r in con.execute(
            "SELECT * FROM senasa_acuerdos ORDER BY estado ASC, orden ASC").fetchall()]
    return jsonify({"ok": True, "rows": rows})

def _validar_fecha_compromiso(valor):
    """Devuelve (ok, valor_normalizado_o_mensaje_de_error).
    Vacío es válido (compromiso sin fecha límite), igual que 'A definir'
    (mismo sentinel que usa vua_cronologia). Cualquier otro valor debe ser
    una fecha real -- se acepta dd/mm/aaaa y las variantes que ya soporta
    _normalizar_fecha_a_ddmmaaaa (aaaa-mm-dd, dd-mm-aaaa). Antes esto se
    guardaba tal cual sin validar, y _chequear_acuerdos_vencidos lo
    ignoraba en silencio si no matcheaba dd/mm/aaaa exacto."""
    valor = (valor or "").strip()
    if not valor or valor == "A definir":
        return True, valor
    normalizada = _normalizar_fecha_a_ddmmaaaa(valor)
    if normalizada == "A definir":
        return False, f"Fecha de compromiso inválida: '{valor}'. Formato esperado: dd/mm/aaaa."
    return True, normalizada


@senasa_bp.route("/api/senasa/acuerdos", methods=["POST"])
@login_required
@modulo_required("senasa")
def senasa_acuerdos_add():
    data = request.json or {}
    ok, err = validar_enum(data.get("estado"), ESTADOS_TAREA, "estado")
    if not ok: return jsonify({"ok": False, "error": err}), 400
    ok, fecha_o_err = _validar_fecha_compromiso(data.get("fecha_compromiso", ""))
    if not ok: return jsonify({"ok": False, "error": fecha_o_err}), 400
    with get_db(HIST_DB) as con:
        cur = con.cursor()
        cur.execute("SELECT MAX(orden) FROM senasa_acuerdos")
        max_o = cur.fetchone()[0] or 0
        cur.execute("INSERT INTO senasa_acuerdos (descripcion,responsable,fecha_compromiso,estado,orden) VALUES (?,?,?,?,?)",
            (data.get("descripcion",""), data.get("responsable",""),
             fecha_o_err, data.get("estado","Pendiente"), max_o+1))
        new_id = cur.lastrowid
    return jsonify({"ok": True, "id": new_id})

@senasa_bp.route("/api/senasa/acuerdos/<int:iid>", methods=["PUT"])
@login_required
@modulo_required("senasa")
def senasa_acuerdos_update(iid):
    data = request.json or {}
    ok, err = validar_enum(data.get("estado"), ESTADOS_TAREA, "estado")
    if not ok: return jsonify({"ok": False, "error": err}), 400
    ok, fecha_o_err = _validar_fecha_compromiso(data.get("fecha_compromiso", ""))
    if not ok: return jsonify({"ok": False, "error": fecha_o_err}), 400
    with get_db(HIST_DB) as con:
        # Si se reprograma la fecha de compromiso, permitir que se vuelva a
        # avisar si el nuevo vencimiento también queda en el pasado (evita
        # que quede silenciado para siempre por haberse avisado una vez).
        actual = con.execute("SELECT fecha_compromiso FROM senasa_acuerdos WHERE id=?", (iid,)).fetchone()
        reset_alerta = actual and fecha_o_err and fecha_o_err != actual[0]
        con.execute("UPDATE senasa_acuerdos SET descripcion=?,responsable=?,fecha_compromiso=?,estado=?"
                    + (",alerta_vencido_enviada=0" if reset_alerta else "") + " WHERE id=?",
            (data.get("descripcion",""), data.get("responsable",""),
             fecha_o_err, data.get("estado","Pendiente"), iid))
    return jsonify({"ok": True})

@senasa_bp.route("/api/senasa/acuerdos/<int:iid>", methods=["DELETE"])
@login_required
@modulo_required("senasa")
def senasa_acuerdos_delete(iid):
    with get_db(HIST_DB) as con:
        con.execute("DELETE FROM senasa_acuerdos WHERE id=?", (iid,))
    return jsonify({"ok": True})


# ── Alerta de compromisos vencidos (Fase 6: alertas proactivas) ──────────────
def _chequear_acuerdos_vencidos():
    """Avisa por Telegram los compromisos SENASA cuyo fecha_compromiso ya
    pasó y siguen sin estado 'Completado'. Cada acuerdo se avisa una sola
    vez (columna alerta_vencido_enviada, ver migración 003) — si se
    reprograma la fecha, senasa_acuerdos_update resetea el flag."""
    while True:
        try:
            hoy = date.today()
            with get_db(HIST_DB, row_factory=True) as con:
                rows = con.execute(
                    "SELECT id, descripcion, responsable, fecha_compromiso FROM senasa_acuerdos "
                    "WHERE estado != 'Completado' AND COALESCE(alerta_vencido_enviada,0)=0 "
                    "AND fecha_compromiso IS NOT NULL AND fecha_compromiso != ''").fetchall()
                vencidos = []
                for r in rows:
                    try:
                        fc = datetime.strptime(r["fecha_compromiso"], "%d/%m/%Y").date()
                    except ValueError:
                        continue  # fecha_compromiso es texto libre — si no matchea dd/mm/aaaa, se ignora
                    if fc < hoy:
                        vencidos.append(r)
                if vencidos:
                    lineas = "\n".join(
                        f"• {r['descripcion']} ({r['responsable'] or 's/d'}) — vencía {r['fecha_compromiso']}"
                        for r in vencidos[:15])
                    extra = f"\n… y {len(vencidos) - 15} más" if len(vencidos) > 15 else ""
                    notificar_telegram(f"⏰ SENASA — {len(vencidos)} compromiso(s) vencido(s):\n\n{lineas}{extra}")
                    con.executemany("UPDATE senasa_acuerdos SET alerta_vencido_enviada=1 WHERE id=?",
                                     [(r["id"],) for r in vencidos])
        except Exception:
            logging.exception("Error en chequeo de compromisos SENASA vencidos")
        threading.Event().wait(3600)  # revisa cada hora

threading.Thread(target=_chequear_acuerdos_vencidos, daemon=True).start()

# ── Informe SENASA (async) ────────────────────────────────────────────────────
@senasa_bp.route("/api/senasa/informe")
@login_required
@modulo_required("senasa")
def senasa_informe():
    """Genera el informe SENASA en background — misma arquitectura de job
    que VUA/Pad Acuático, y mismo formato visual que los informes SINTIA/
    Pad Acuático (portada compuesta, índice, encabezado/pie institucional,
    fuente unificada -- ver generar_documento.py) a pedido, 22/07/2026."""
    with get_db(HIST_DB, row_factory=True) as con:
        datos = {
            "cronologia": [dict(r) for r in con.execute("SELECT * FROM senasa_cronologia ORDER BY orden").fetchall()],
            "ejes":       [dict(r) for r in con.execute("SELECT * FROM senasa_ejes ORDER BY orden").fetchall()],
            "acuerdos":   [dict(r) for r in con.execute("SELECT * FROM senasa_acuerdos ORDER BY estado, orden").fetchall()],
        }
    job_id = str(uuid.uuid4())[:8]
    job_create(job_id, "Generando informe SENASA...", username=session.get("username", "?"))

    def _run(jid, datos):
        log = job_status[jid]["log"]
        try:
            from docx import Document as DocxDoc
            from docx.shared import Cm, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from generar_documento import (
                _unificar_fuente_documento, _agregar_encabezado, _agregar_pie_pagina,
                _ocultar_encabezado_portada, _generar_portada_compuesta,
                _agregar_imagen_portada_ajustada, _insertar_indice, _heading_indexado,
                agregar_tabla_word, kpi_box,
            )
            from generar_utils import pl

            ejes, cronologia, acuerdos = datos["ejes"], datos["cronologia"], datos["acuerdos"]
            pendientes = [a for a in acuerdos if a["estado"] != "Completado"]
            hoy = datetime.today().strftime("%d/%m/%Y")

            doc = DocxDoc()
            _unificar_fuente_documento(doc)
            for section in doc.sections:
                section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
                section.left_margin = Cm(3); section.right_margin = Cm(2.5)
            _agregar_encabezado(doc, "Dirección de Reingeniería de Procesos Aduaneros")
            _agregar_pie_pagina(doc, "Informe de Avance — Integración SENASA")
            _ocultar_encabezado_portada(doc)

            # Portada
            imagen_portada = _generar_portada_compuesta(
                titulo="INFORME DE AVANCE — INTEGRACIÓN SENASA",
                subtitulo=f"Actualizado al {hoy}",
                meta_lineas=[
                    "Dirección de Reingeniería de Procesos Aduaneros (DG ADUA)",
                    f"Última modificación: {hoy}",
                    "Elaborado por: Sección Simplificación de Procesos Operativos — DI REPA",
                ])
            if imagen_portada:
                _agregar_imagen_portada_ajustada(doc, imagen_portada)
            else:
                titulo = doc.add_heading("INFORME DE AVANCE — INTEGRACIÓN SENASA", 0)
                titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for txt, sz in [("Dirección de Reingeniería de Procesos Aduaneros (DG ADUA)", 12),
                                 (f"Actualizado al {hoy}", 11)]:
                    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(txt); run.font.size = Pt(sz); run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
                doc.add_paragraph()
                dest = doc.add_paragraph(); dest.alignment = WD_ALIGN_PARAGRAPH.CENTER
                dest.add_run("Elaborado por: ").bold = True
                dest.add_run("Sección Simplificación de Procesos Operativos — DI REPA")
            doc.add_page_break()

            # Índice -- Ejes y Cronología son fijos (1. y 2.); Compromisos
            # pendientes es condicional (mismo criterio de contador que en
            # pad_acuatico.py, no hardcodeado).
            secciones = [(1, "Resumen Ejecutivo"), (1, "1.  Ejes de trabajo"), (1, "2.  Cronología de reuniones")]
            idx_compromisos = None
            if pendientes:
                secciones.append((1, "3.  Compromisos pendientes")); idx_compromisos = 4
            _insertar_indice(doc, secciones)

            # Resumen ejecutivo
            _heading_indexado(doc, "Resumen Ejecutivo", 1, 1)
            doc.add_paragraph(
                "El presente informe resume el estado de avance de la integración PAD / SIG-Embalajes "
                f"con SENASA al {hoy}: ejes de trabajo definidos, cronología de reuniones realizadas y "
                "pendientes, y compromisos asumidos que todavía siguen abiertos.")
            en_curso = sum(1 for e in ejes if "curso" in e.get("estado", "").lower() or "análisis" in e.get("estado", "").lower() or "diseño" in e.get("estado", "").lower())
            completados = sum(1 for e in ejes if "completado" in e.get("estado", "").lower())
            pendientes_crono = sum(1 for c in cronologia if c.get("estado") == "Pendiente")
            kpi_box(doc, [
                ("EJES DE TRABAJO", str(len(ejes)), f"{completados} completados, {en_curso} en curso"),
                ("REUNIONES", str(len(cronologia)), f"{pendientes_crono} pendientes"),
                ("COMPROMISOS", str(len(acuerdos)), f"{len(pendientes)} pendientes"),
            ])
            doc.add_page_break()

            # 1. Ejes de trabajo
            _heading_indexado(doc, "1.  Ejes de trabajo", 1, 2)
            if ejes:
                doc.add_paragraph(f"Se relevaron {len(ejes)} {pl(len(ejes), 'eje de trabajo', 'ejes de trabajo')} de la integración.")
                for e in ejes:
                    p = doc.add_paragraph(); p.add_run(e["nombre"]).bold = True
                    if e.get("descripcion"): doc.add_paragraph(e["descripcion"])
                    ep = doc.add_paragraph(); ep.add_run("Estado: ").bold = True; ep.add_run(e["estado"])
            else:
                doc.add_paragraph("Todavía no hay ejes de trabajo cargados.")
            doc.add_page_break()

            # 2. Cronología de reuniones
            _heading_indexado(doc, "2.  Cronología de reuniones", 1, 3)
            if cronologia:
                doc.add_paragraph(f"Se registraron {len(cronologia)} {pl(len(cronologia), 'evento', 'eventos')} en la cronología.")
                agregar_tabla_word(doc, ["FECHA", "ACTIVIDAD", "PARTICIPANTES", "ESTADO"],
                    [[c["fecha"], c["actividad"], c.get("participantes", ""), c["estado"]] for c in cronologia],
                    col_widths=[2.2, 6.5, 4.5, 2.3])
            else:
                doc.add_paragraph("Todavía no hay entradas en la cronología.")

            # 3. Compromisos pendientes
            if pendientes:
                doc.add_page_break()
                _heading_indexado(doc, "3.  Compromisos pendientes", 1, idx_compromisos)
                doc.add_paragraph(f"{len(pendientes)} de {len(acuerdos)} {pl(len(acuerdos), 'compromiso', 'compromisos')} todavía sin cerrar.")
                agregar_tabla_word(doc, ["DESCRIPCIÓN", "RESPONSABLE", "FECHA COMPROMISO", "ESTADO"],
                    [[a["descripcion"], a.get("responsable", ""), a.get("fecha_compromiso", ""), a["estado"]] for a in pendientes],
                    col_widths=[6.5, 3.5, 3, 2.5])

            os.makedirs(OUTPUT_FOLDER, exist_ok=True)
            fname = f"Informe_SENASA_{datetime.today().strftime('%Y%m%d_%H%M')}_{jid}.docx"
            dest  = os.path.join(OUTPUT_FOLDER, fname)
            doc.save(dest)
            job_status[jid]["files"] = [dest]
            log.append(f"✓ Informe generado: {fname}")
            job_status[jid]["status"] = "done"
            _job_persist(jid)
            notificar_telegram(f"✓ Informe SENASA listo ({job_status[jid].get('username','?')})")
        except Exception as e:
            log.append(f"✗ {e}")
            job_status[jid]["status"] = "error"
            _job_persist(jid)
            notificar_telegram(f"⚠️ Informe SENASA falló ({job_status[jid].get('username','?')}): {e}")

    threading.Thread(target=_run, args=(job_id, datos)).start()
    return jsonify({"ok": True, "job_id": job_id})

@senasa_bp.route("/api/senasa/informe/download/<job_id>")
@login_required
@modulo_required("senasa")
def senasa_informe_download(job_id):
    import glob
    job = job_get(job_id)
    if job and job.get("status") == "done" and job.get("files") and os.path.exists(job["files"][0]):
        return send_file(job["files"][0], as_attachment=True,
            download_name=os.path.basename(job["files"][0]),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    if not re.match(r'^[a-zA-Z0-9]{1,32}$', job_id or ""):
        return jsonify({"ok": False, "error": "job_id inválido"}), 400
    archivos = sorted(glob.glob(os.path.join(OUTPUT_FOLDER, f"*{job_id}*.docx")), key=os.path.getmtime, reverse=True)
    if archivos:
        return send_file(archivos[0], as_attachment=True, download_name=os.path.basename(archivos[0]),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
