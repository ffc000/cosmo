"""
generar.py — Módulo de generación de informes SINTIA
Adaptado del script original generar_informe_sintia.py para uso como módulo web.
"""

import sqlite3, os, io, re
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    XLSX_OK = True
except ImportError:
    XLSX_OK = False

try:
    import matplotlib; matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    MPL_OK = True
except ImportError:
    MPL_OK = False

try:
    import anthropic, httpx
    ANT_OK = True
except ImportError:
    ANT_OK = False

# ── Constantes ─────────────────────────────────────────────────────────────────
PAISES = {"BO":"Bolivia","PY":"Paraguay","BR":"Brasil","CL":"Chile","UY":"Uruguay"}
MESES  = {"01":"Enero","02":"Febrero","03":"Marzo","04":"Abril","05":"Mayo","06":"Junio",
           "07":"Julio","08":"Agosto","09":"Septiembre","10":"Octubre","11":"Noviembre","12":"Diciembre"}
C_TRANS="#1F7DC4"; C_NO_TRANS="#C0392B"; C_TARDIO="#E67E22"
C_CARGADO="#2E86AB"; C_LASTRE="#A8DADC"
UMBRAL_VERDE=60.0; UMBRAL_AMARILLO=30.0

CAT_LIKE = {
    "ERROR ATRIBUTO/PARAMETRO INVALIDO":"%ERROR ATRIBUTO/PARAMETRO INVALIDO%",
    "EVENTO DISTINTO DE PATAI/OFTAI":"%EVENTO DISTINTO DE  PATAI O OFTAI%",
    "EVENTO YA RECIBIDO":"%EVENTO YA RECIBIDO%",
    "NRO DE MIC INEXISTENTE":"%NRO DE MIC INEXISTENTE%",
    "NRO DE MIC EXISTENTE":"%Nro de Mic existente%",
    "CONTENEDOR_VACIO":"%CONTENEDORVACIO%",
    "PAIS_DE_PASO_DUPLICADO":"%PAIS DE PASO DUPLICADO%",
    "FECHA_DEL_EVENTO":"%FECHA DEL EVENTO %",
    "CODCIUPART":"%CODCIUPART%","CODCIUENT":"%CODCIUENT%",
    "PAISESDEPASO.CODCIUSAL":"%PAISESDEPASO%CODCIUSAL%","CODCIUSAL":"%CODCIUSAL%",
    "VEHICULO":"%VEHICULO%","CONTENEDOR":"%CONTENEDOR%","CONSIGNATARIO":"%CONSIGNATARIO%",
    "CODDIVISASEG":"%CODDIVISASEG%","CARTA_PORTE_DUPLICADO":"%PORTE%DUPLICADO%",
    "CARTA PORTE":"%PORTE%","PAISESDEPASO.CODADUENT":"%PAISESDEPASO.CODADUENT%",
    "PAISESDEPASO":"%PAISESDEPASO%","CODADUEMI":"%CODADUEMI%","CODADUSAL":"%CODADUSAL%",
    "CODADUPART":"%CODADUPART%","FECHLLEGPREV":"%FECHLLEGPREV%",
    "PESOBRUTOTOTAL":"%PESOBRUTOTOTAL%","DESCRUTITINERARIOS":"%DESCRUTITINERARIOS%",
    "TIPDOCIDENT":"%TIPDOCIDENT%","PAISDEST.CODCIUDEST":"%PAISDEST.CODCIUDEST%",
    "PAISDEST.CODADUDEST":"%PAISDEST.CODADUDEST%","PAISDEST.CODADUENT":"%PAISDEST.CODADUENT%",
    "DESTINACION":"%DESTINACION%","INDNCM":"%INDNCM%",
    "CONDUCTOR.NOMBRE":"%CONDUCTOR.NOMBRE%","CRT":"%CRT%","OTROS":"%%",
}

# ── Helpers ─────────────────────────────────────────────────────────────────────
def fmt(v):
    try: return f"{int(v or 0):,}".replace(",",".")
    except: return str(v or 0)

def pct(a, total):
    if not total: return "0,0%"
    return f"{100*float(a)/float(total):.1f}%".replace(".",",")

def pct_f(a, total):
    if not total: return 0.0
    return round(100*float(a)/float(total), 1)

def n(v): return int(v or 0)

def periodo_texto(anio, mes_d, mes_h):
    if mes_d == mes_h: return f"{MESES[mes_d]} {anio}"
    return f"{MESES[mes_d]} \u2013 {MESES[mes_h]} {anio}"

def mes_label(periodo):
    yy, mm = periodo.split("-")
    return f"{MESES.get(mm,mm)[:3]} {yy[2:]}"

def mes_label_largo(periodo):
    yy, mm = periodo.split("-")
    return f"{MESES.get(mm,mm)} {yy}"

def ultimo_mes_completo():
    hoy = datetime.today()
    if hoy.month == 1: return "12", str(hoy.year-1)
    return str(hoy.month-1).zfill(2), str(hoy.year)

def color_semaforo(pct_val):
    v = float(pct_val)
    if v >= UMBRAL_VERDE: return "C6EFCE"
    if v >= UMBRAL_AMARILLO: return "FFEB9C"
    return "FFC7CE"

def fig_to_bytes(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig); buf.seek(0)
    return buf

# ── Gráficos ────────────────────────────────────────────────────────────────────
def grafico_torta(gT, gN, gTd):
    fig, ax = plt.subplots(figsize=(5.5,4), facecolor='white')
    valores=[gT,gN,gTd]; total=gT+gN+gTd
    labels=[f'Transmitidos\n{fmt(gT)}',f'No transmitidos\n{fmt(gN)}',f'Tard\u00edos\n{fmt(gTd)}']
    colors=[C_TRANS,C_NO_TRANS,C_TARDIO]
    wedges,texts,autotexts = ax.pie(valores,labels=labels,colors=colors,
        autopct=lambda v: f"{v:.1f}%".replace(".",",") if total>0 else "",
        startangle=90,pctdistance=0.75,wedgeprops=dict(edgecolor='white',linewidth=2))
    for t in texts: t.set_fontsize(9)
    for t in autotexts: t.set_fontsize(9); t.set_fontweight('bold'); t.set_color('white')
    ax.set_title("Distribuci\u00f3n MICs por estado de transmisi\u00f3n",fontsize=11,fontweight='bold',pad=12)
    fig.tight_layout()
    return fig_to_bytes(fig)

def grafico_barras_apiladas(ev_total):
    meses=[mes_label(r["MES"]) for r in ev_total]
    carg=[n(r.get("CARGADO",0)) for r in ev_total]
    lastre=[n(r.get("LASTRE",0)) for r in ev_total]
    x=range(len(meses))
    fig,ax=plt.subplots(figsize=(7,4),facecolor='white')
    ax.bar(x,carg,color=C_CARGADO,label='Cargado',edgecolor='white',linewidth=0.5)
    ax.bar(x,lastre,bottom=carg,color=C_LASTRE,label='Lastre',edgecolor='white',linewidth=0.5)
    ax.set_xticks(list(x)); ax.set_xticklabels(meses,fontsize=9)
    ax.set_ylabel("Operaciones",fontsize=9)
    ax.set_title("Evoluci\u00f3n mensual de ingreso de camiones",fontsize=11,fontweight='bold')
    ax.legend(fontsize=9); ax.yaxis.grid(True,alpha=0.3); ax.set_axisbelow(True)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    for i,(c,l) in enumerate(zip(carg,lastre)):
        t=c+l; ax.text(i,t+t*0.01,fmt(t),ha='center',va='bottom',fontsize=8,fontweight='bold')
    fig.tight_layout(); return fig_to_bytes(fig)

def grafico_lineas_pct(ev_total,ev_trans,ev_tardio,ev_no_trans):
    mt={r["MES"]:n(r.get("CARGADO",0))+n(r.get("LASTRE",0)) for r in ev_total}
    mtr={r["MES"]:n(r.get("CARGADOS",0))+n(r.get("LASTRE",0)) for r in ev_trans}
    mtd={r["MES"]:n(r.get("CARGADOS",0))+n(r.get("LASTRE",0)) for r in ev_tardio}
    mnt={r["MES"]:n(r.get("CARGADOS",0))+n(r.get("LASTRE",0)) for r in ev_no_trans}
    periodos=sorted(mt.keys()); labels=[mes_label(p) for p in periodos]
    pt=[pct_f(mtr.get(p,0),mt.get(p,0)) for p in periodos]
    ptd=[pct_f(mtd.get(p,0),mt.get(p,0)) for p in periodos]
    pnt=[pct_f(mnt.get(p,0),mt.get(p,0)) for p in periodos]
    fig,ax=plt.subplots(figsize=(7,4),facecolor='white')
    x=range(len(labels))
    ax.plot(list(x),pt, marker='o',color=C_TRANS,   linewidth=2,label='Transmitidos',markersize=5)
    ax.plot(list(x),pnt,marker='s',color=C_NO_TRANS,linewidth=2,label='No transmitidos',markersize=5)
    ax.plot(list(x),ptd,marker='^',color=C_TARDIO,  linewidth=2,label='Tard\u00edos',markersize=5)
    ax.set_xticks(list(x)); ax.set_xticklabels(labels,fontsize=9)
    ax.set_ylabel("%",fontsize=9); ax.set_ylim(0,110)
    ax.set_title("Evoluci\u00f3n mensual del estado de transmisi\u00f3n (%)",fontsize=11,fontweight='bold')
    ax.legend(fontsize=9); ax.yaxis.grid(True,alpha=0.3); ax.set_axisbelow(True)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    for i,t in enumerate(pt):
        ax.annotate(f"{t:.0f}%",(i,t),textcoords="offset points",xytext=(0,7),ha='center',fontsize=7)
    fig.tight_layout(); return fig_to_bytes(fig)

def grafico_rechazos_cat(rechazos_cat):
    cats=[r for r in rechazos_cat if r["Categoria"]!="TOTAL"][:10]
    if not cats: return None
    labels=[r["Categoria"] for r in reversed(cats)]
    valores=[n(r.get("Rechazos",0)) for r in reversed(cats)]
    fig,ax=plt.subplots(figsize=(7,max(3,len(cats)*0.5+1)),facecolor='white')
    colors=[C_TRANS if i%2==0 else C_CARGADO for i in range(len(labels))]
    bars=ax.barh(labels,valores,color=colors,edgecolor='white',linewidth=0.5)
    for bar,val in zip(bars,valores):
        ax.text(bar.get_width()+bar.get_width()*0.02,bar.get_y()+bar.get_height()/2,
                fmt(val),va='center',fontsize=9,fontweight='bold')
    ax.set_xlabel("Cantidad",fontsize=9)
    ax.set_title("Rechazos por categor\u00eda",fontsize=11,fontweight='bold')
    ax.xaxis.grid(True,alpha=0.3); ax.set_axisbelow(True)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    ax.tick_params(axis='y',labelsize=8); fig.tight_layout()
    return fig_to_bytes(fig)

def grafico_rechazos_mes(rechazos_mes):
    if not rechazos_mes: return None
    labels=[mes_label(r["periodo"]) for r in rechazos_mes]
    valores=[n(r.get("MIC_RECHAZOS",0)) for r in rechazos_mes]
    fig,ax=plt.subplots(figsize=(6,3.5),facecolor='white')
    ax.bar(range(len(labels)),valores,color=C_NO_TRANS,edgecolor='white',linewidth=0.5)
    ax.set_xticks(range(len(labels))); ax.set_xticklabels(labels,fontsize=9)
    ax.set_ylabel("MICs rechazados",fontsize=9)
    ax.set_title("Rechazos por mes",fontsize=11,fontweight='bold')
    for i,v in enumerate(valores):
        ax.text(i,v+v*0.02,fmt(v),ha='center',va='bottom',fontsize=9,fontweight='bold')
    ax.yaxis.grid(True,alpha=0.3); ax.set_axisbelow(True)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    fig.tight_layout(); return fig_to_bytes(fig)

def grafico_comparativo_meses(datos_ult, datos_ant, per_ult, per_ant):
    if not datos_ult or not datos_ant: return None
    categorias=['Transmitidos','No transmitidos','Tard\u00edos']
    ult=[n(datos_ult.get("trans",0)),n(datos_ult.get("no_trans",0)),n(datos_ult.get("tardio",0))]
    ant=[n(datos_ant.get("trans",0)),n(datos_ant.get("no_trans",0)),n(datos_ant.get("tardio",0))]
    x=range(len(categorias)); w=0.35
    fig,ax=plt.subplots(figsize=(7,4),facecolor='white')
    b1=ax.bar([i-w/2 for i in x],ant,w,label=mes_label_largo(per_ant),color='#ADB5BD',edgecolor='white')
    b2=ax.bar([i+w/2 for i in x],ult,w,label=mes_label_largo(per_ult),color=C_TRANS,edgecolor='white')
    ax.set_xticks(list(x)); ax.set_xticklabels(categorias,fontsize=10)
    ax.set_title(f"Comparativo: {mes_label_largo(per_ant)} vs {mes_label_largo(per_ult)}",fontsize=11,fontweight='bold')
    ax.legend(fontsize=9); ax.yaxis.grid(True,alpha=0.3); ax.set_axisbelow(True)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    for bar in list(b1)+list(b2):
        v=int(bar.get_height())
        if v>0: ax.text(bar.get_x()+bar.get_width()/2,v+v*0.02,fmt(v),ha='center',va='bottom',fontsize=8)
    fig.tight_layout(); return fig_to_bytes(fig)

# ── Queries ──────────────────────────────────────────────────────────────────────
def correr_queries(ruta_db, pais, anio, mes_d, mes_h, log_fn):
    tabla = f"DAT_{anio}"
    desde = f"{anio}-{mes_d}"; hasta = f"{anio}-{mes_h}"; like = f"%{pais}%"
    # per_ult = último mes del período seleccionado (mes_h), no el último mes del sistema
    mes_ult = mes_h  # el período termina en mes_h
    mes_ant = str(int(mes_ult)-1).zfill(2) if int(mes_ult) > 1 else "12"
    anio_per_ant = anio if int(mes_ult) > 1 else str(int(anio)-1)
    per_ult = f"{anio}-{mes_ult}"; per_ant = f"{anio_per_ant}-{mes_ant}"
    log_fn("Conectando a la BD...")
    con = sqlite3.connect(ruta_db); con.row_factory = sqlite3.Row
    cur = con.cursor()
    def q(sql, params=()):
        cur.execute(sql, params)
        return [dict(r) for r in cur.fetchall()]
    anio_ant = str(int(anio)-1); tabla_ant = f"DAT_{anio_ant}"
    tiene_anio_ant = bool(q("SELECT name FROM sqlite_master WHERE type='table' AND name=?", (tabla_ant,)))
    log_fn("Corriendo queries...")
    totales = q(f"""
        SELECT
            SUM(CASE WHEN CARGADO='SI' AND FECHA_TRANS_ISO IS NOT NULL AND FECHA_TRANS_ISO<=FECHA_INGRESO_ISO THEN 1 ELSE 0 END) AS CARGADO_TRANS,
            SUM(CASE WHEN CARGADO='NO' AND FECHA_TRANS_ISO IS NOT NULL AND FECHA_TRANS_ISO<=FECHA_INGRESO_ISO THEN 1 ELSE 0 END) AS LASTRE_TRANS,
            SUM(CASE WHEN CARGADO='SI' AND FECHA_TRANS_ISO IS NULL THEN 1 ELSE 0 END) AS CARGADO_NO_TRANS,
            SUM(CASE WHEN CARGADO='NO' AND FECHA_TRANS_ISO IS NULL THEN 1 ELSE 0 END) AS LASTRE_NO_TRANS,
            SUM(CASE WHEN CARGADO='SI' AND FECHA_TRANS_ISO IS NOT NULL AND FECHA_TRANS_ISO>FECHA_INGRESO_ISO THEN 1 ELSE 0 END) AS CARGADO_TARDIO,
            SUM(CASE WHEN CARGADO='NO' AND FECHA_TRANS_ISO IS NOT NULL AND FECHA_TRANS_ISO>FECHA_INGRESO_ISO THEN 1 ELSE 0 END) AS LASTRE_TARDIO
        FROM {tabla} WHERE MIC LIKE ? AND strftime('%Y-%m',FECHA_INGRESO_ISO) BETWEEN ? AND ?
    """, (like, desde, hasta))
    ev_total = q(f"""SELECT substr(FECHA_INGRESO_ISO,1,7) AS MES, SUM(CASE WHEN CARGADO='SI' THEN 1 ELSE 0 END) AS CARGADO, SUM(CASE WHEN CARGADO='NO' THEN 1 ELSE 0 END) AS LASTRE FROM {tabla} WHERE MIC LIKE ? AND strftime('%Y-%m',FECHA_INGRESO_ISO) BETWEEN ? AND ? GROUP BY MES ORDER BY MES""", (like, desde, hasta))
    ev_trans = q(f"""SELECT substr(FECHA_INGRESO_ISO,1,7) AS MES, SUM(CASE WHEN CARGADO='SI' THEN 1 ELSE 0 END) AS CARGADOS, SUM(CASE WHEN CARGADO='NO' THEN 1 ELSE 0 END) AS LASTRE FROM {tabla} WHERE MIC LIKE ? AND FECHA_TRANS_ISO IS NOT NULL AND FECHA_TRANS_ISO<=FECHA_INGRESO_ISO AND strftime('%Y-%m',FECHA_INGRESO_ISO) BETWEEN ? AND ? GROUP BY MES ORDER BY MES""", (like, desde, hasta))
    ev_tardio = q(f"""SELECT substr(FECHA_INGRESO_ISO,1,7) AS MES, SUM(CASE WHEN CARGADO='SI' THEN 1 ELSE 0 END) AS CARGADOS, SUM(CASE WHEN CARGADO='NO' THEN 1 ELSE 0 END) AS LASTRE FROM {tabla} WHERE MIC LIKE ? AND FECHA_TRANS_ISO IS NOT NULL AND FECHA_TRANS_ISO>FECHA_INGRESO_ISO AND strftime('%Y-%m',FECHA_INGRESO_ISO) BETWEEN ? AND ? GROUP BY MES ORDER BY MES""", (like, desde, hasta))
    ev_no_trans = q(f"""SELECT substr(FECHA_INGRESO_ISO,1,7) AS MES, SUM(CASE WHEN CARGADO='SI' THEN 1 ELSE 0 END) AS CARGADOS, SUM(CASE WHEN CARGADO='NO' THEN 1 ELSE 0 END) AS LASTRE FROM {tabla} WHERE MIC LIKE ? AND FECHA_TRANS_ISO IS NULL AND strftime('%Y-%m',FECHA_INGRESO_ISO) BETWEEN ? AND ? GROUP BY MES ORDER BY MES""", (like, desde, hasta))
    rechazos_mes = q("""WITH m AS (SELECT NroMic, MAX(strftime('%Y-%m',Fecha_ISO)) AS periodo FROM RECHAZOS WHERE PaisEmisor=? AND Anio=? AND Metodo='OficializarMicDta' AND strftime('%Y-%m',Fecha_ISO) BETWEEN ? AND ? GROUP BY NroMic) SELECT periodo, COUNT(*) AS MIC_RECHAZOS FROM m GROUP BY periodo ORDER BY periodo""", (pais, anio, desde, hasta))
    rechazos_cat = q("""
        WITH ranked AS (SELECT mensaje, ROW_NUMBER() OVER (PARTITION BY NroMic ORDER BY Fecha_ISO DESC) AS rn FROM RECHAZOS WHERE PaisEmisor=? AND Anio=? AND Metodo='OficializarMicDta' AND strftime('%Y-%m',Fecha_ISO) BETWEEN ? AND ?),
        res AS (SELECT CASE
            WHEN mensaje LIKE '%ERROR ATRIBUTO/PARAMETRO INVALIDO%' THEN 'ERROR ATRIBUTO/PARAMETRO INVALIDO'
            WHEN mensaje LIKE '%EVENTO DISTINTO DE  PATAI O OFTAI%' THEN 'EVENTO DISTINTO DE PATAI/OFTAI'
            WHEN mensaje LIKE '%EVENTO YA RECIBIDO%' THEN 'EVENTO YA RECIBIDO'
            WHEN mensaje LIKE '%NRO DE MIC INEXISTENTE%' THEN 'NRO DE MIC INEXISTENTE'
            WHEN mensaje LIKE '%Nro de Mic existente%' THEN 'NRO DE MIC EXISTENTE'
            WHEN mensaje LIKE '%CONTENEDORESVACIOS%' OR mensaje LIKE '%CONTENEDORVACIO%' THEN 'CONTENEDOR_VACIO'
            WHEN mensaje LIKE '%PAIS DE PASO DUPLICADO%' THEN 'PAIS_DE_PASO_DUPLICADO'
            WHEN mensaje LIKE '%FECHA DEL EVENTO %' THEN 'FECHA_DEL_EVENTO'
            WHEN mensaje LIKE '%CODCIUPART%' THEN 'CODCIUPART'
            WHEN mensaje LIKE '%CODCIUENT%' THEN 'CODCIUENT'
            WHEN mensaje LIKE '%PAISESDEPASO%CODCIUSAL%' THEN 'PAISESDEPASO.CODCIUSAL'
            WHEN mensaje LIKE '%CODCIUSAL%' THEN 'CODCIUSAL'
            WHEN mensaje LIKE '%VEHICULO%' THEN 'VEHICULO'
            WHEN mensaje LIKE '%CONTENEDOR%' THEN 'CONTENEDOR'
            WHEN mensaje LIKE '%CONSIGNATARIO%' THEN 'CONSIGNATARIO'
            WHEN mensaje LIKE '%CODDIVISASEG%' THEN 'CODDIVISASEG'
            WHEN mensaje LIKE '%PORTE%DUPLICADO%' THEN 'CARTA_PORTE_DUPLICADO'
            WHEN mensaje LIKE '%PORTE%' AND mensaje NOT LIKE '%DUPLICADO%' THEN 'CARTA PORTE'
            WHEN mensaje LIKE '%PAISESDEPASO.CODADUENT%' THEN 'PAISESDEPASO.CODADUENT'
            WHEN mensaje LIKE '%PAISESDEPASO%' THEN 'PAISESDEPASO'
            WHEN mensaje LIKE '%CODADUEMI%' THEN 'CODADUEMI'
            WHEN mensaje LIKE '%CODADUSAL%' THEN 'CODADUSAL'
            WHEN mensaje LIKE '%CODADUPART%' THEN 'CODADUPART'
            WHEN mensaje LIKE '%FECHLLEGPREV%' THEN 'FECHLLEGPREV'
            WHEN mensaje LIKE '%PESOBRUTOTOTAL%' THEN 'PESOBRUTOTOTAL'
            WHEN mensaje LIKE '%DESCRUTITINERARIOS%' THEN 'DESCRUTITINERARIOS'
            WHEN mensaje LIKE '%TIPDOCIDENT%' THEN 'TIPDOCIDENT'
            WHEN mensaje LIKE '%PAISDEST.CODCIUDEST%' THEN 'PAISDEST.CODCIUDEST'
            WHEN mensaje LIKE '%PAISDEST.CODADUDEST%' THEN 'PAISDEST.CODADUDEST'
            WHEN mensaje LIKE '%PAISDEST.CODADUENT%' THEN 'PAISDEST.CODADUENT'
            WHEN mensaje LIKE '%DESTINACION%' AND mensaje NOT LIKE '%INDNCM%' THEN 'DESTINACION'
            WHEN mensaje LIKE '%INDNCM%' THEN 'INDNCM'
            WHEN mensaje LIKE '%CONDUCTOR.NOMBRE%' THEN 'CONDUCTOR.NOMBRE'
            WHEN mensaje LIKE '%CRT%' THEN 'CRT'
            ELSE 'OTROS' END AS Categoria, COUNT(*) AS Rechazos
        FROM ranked WHERE rn=1 GROUP BY Categoria),
        unioned AS (SELECT 0 AS orden, Categoria, Rechazos FROM res UNION ALL SELECT 1,'TOTAL',SUM(Rechazos) FROM res)
        SELECT Categoria, Rechazos FROM unioned ORDER BY orden, Rechazos DESC, Categoria
    """, (pais, anio, desde, hasta))
    top_cats = [r["Categoria"] for r in rechazos_cat if r["Categoria"]!="TOTAL"][:3]
    rechazos_ej = []
    nromic_vistos = set()
    for cat in top_cats:
        like_cat = CAT_LIKE.get(cat, f"%{cat}%")
        filas = q("""SELECT Fecha_ISO, NroMic, Mensaje FROM (SELECT Fecha_ISO, NroMic, Mensaje, ROW_NUMBER() OVER (PARTITION BY NroMic ORDER BY Fecha_ISO DESC) AS rn FROM RECHAZOS WHERE PaisEmisor=? AND Anio=? AND Metodo='OficializarMicDta' AND strftime('%Y-%m',Fecha_ISO) BETWEEN ? AND ?) WHERE rn=1 AND Mensaje LIKE ? ORDER BY Fecha_ISO DESC LIMIT 5""", (pais, anio, desde, hasta, like_cat))
        for fila in filas:
            if fila.get("NroMic") not in nromic_vistos:
                nromic_vistos.add(fila.get("NroMic"))
                rechazos_ej.append(fila)
    def totales_mes(periodo):
        anio_p = periodo[:4]; tabla_p = f"DAT_{anio_p}"
        rows = q(f"""SELECT SUM(CASE WHEN FECHA_TRANS_ISO IS NOT NULL AND FECHA_TRANS_ISO<=FECHA_INGRESO_ISO THEN 1 ELSE 0 END) AS trans, SUM(CASE WHEN FECHA_TRANS_ISO IS NULL THEN 1 ELSE 0 END) AS no_trans, SUM(CASE WHEN FECHA_TRANS_ISO IS NOT NULL AND FECHA_TRANS_ISO>FECHA_INGRESO_ISO THEN 1 ELSE 0 END) AS tardio, COUNT(*) AS total FROM {tabla_p} WHERE MIC LIKE ? AND strftime('%Y-%m',FECHA_INGRESO_ISO)=?""", (like, periodo))
        return rows[0] if rows else {}
    datos_ult = totales_mes(per_ult)
    datos_ant = totales_mes(per_ant)
    impoexpo_ult = q(f"""SELECT TIPO_REGISTRO, SUM(CASE WHEN FECHA_TRANS_ISO IS NOT NULL AND FECHA_TRANS_ISO<=FECHA_INGRESO_ISO THEN 1 ELSE 0 END) AS trans, SUM(CASE WHEN FECHA_TRANS_ISO IS NULL THEN 1 ELSE 0 END) AS no_trans, SUM(CASE WHEN FECHA_TRANS_ISO IS NOT NULL AND FECHA_TRANS_ISO>FECHA_INGRESO_ISO THEN 1 ELSE 0 END) AS tardio, SUM(CASE WHEN CARGADO='SI' THEN 1 ELSE 0 END) AS cargado, SUM(CASE WHEN CARGADO='NO' THEN 1 ELSE 0 END) AS lastre, COUNT(*) AS total FROM {tabla} WHERE MIC LIKE ? AND strftime('%Y-%m',FECHA_INGRESO_ISO)=? GROUP BY TIPO_REGISTRO""", (like, per_ult))
    rechazos_ult_cat = q("""
        WITH ranked AS (SELECT mensaje, ROW_NUMBER() OVER (PARTITION BY NroMic ORDER BY Fecha_ISO DESC) AS rn FROM RECHAZOS WHERE PaisEmisor=? AND Anio=? AND Metodo='OficializarMicDta' AND strftime('%Y-%m',Fecha_ISO)=?),
        res AS (SELECT CASE
            WHEN mensaje LIKE '%Nro de Mic existente%' THEN 'NRO DE MIC EXISTENTE'
            WHEN mensaje LIKE '%NRO DE MIC INEXISTENTE%' THEN 'NRO DE MIC INEXISTENTE'
            WHEN mensaje LIKE '%VEHICULO%' THEN 'VEHICULO'
            WHEN mensaje LIKE '%CONSIGNATARIO%' THEN 'CONSIGNATARIO'
            WHEN mensaje LIKE '%PAISDEST.CODADUDEST%' THEN 'PAISDEST.CODADUDEST'
            WHEN mensaje LIKE '%PAISDEST.CODADUENT%' THEN 'PAISDEST.CODADUENT'
            WHEN mensaje LIKE '%PAISDEST.CODCIUDEST%' THEN 'PAISDEST.CODCIUDEST'
            WHEN mensaje LIKE '%PORTE%DUPLICADO%' THEN 'CARTA_PORTE_DUPLICADO'
            WHEN mensaje LIKE '%PORTE%' AND mensaje NOT LIKE '%DUPLICADO%' THEN 'CARTA PORTE'
            WHEN mensaje LIKE '%CRT%' THEN 'CRT'
            WHEN mensaje LIKE '%DESTINACION%' AND mensaje NOT LIKE '%INDNCM%' THEN 'DESTINACION'
            WHEN mensaje LIKE '%INDNCM%' THEN 'INDNCM'
            WHEN mensaje LIKE '%CODADUEMI%' THEN 'CODADUEMI'
            WHEN mensaje LIKE '%CODADUSAL%' THEN 'CODADUSAL'
            WHEN mensaje LIKE '%PAISESDEPASO%' THEN 'PAISESDEPASO'
            WHEN mensaje LIKE '%FECHA DEL EVENTO %' THEN 'FECHA_DEL_EVENTO'
            WHEN mensaje LIKE '%CONTENEDOR%' THEN 'CONTENEDOR'
            ELSE 'OTROS' END AS Categoria, COUNT(*) AS Rechazos
        FROM ranked WHERE rn=1 GROUP BY Categoria)
        SELECT Categoria, Rechazos FROM res ORDER BY Rechazos DESC
    """, (pais, anio, per_ult))
    rech_ant = q("""WITH p AS (SELECT NroMic, MAX(strftime('%Y-%m',Fecha_ISO)) AS ult_mes FROM RECHAZOS WHERE PaisEmisor=? AND Anio=? AND Metodo='OficializarMicDta' AND strftime('%Y-%m',Fecha_ISO)=? GROUP BY NroMic) SELECT COUNT(*) AS total FROM p""", (pais, anio_per_ant, per_ant))
    total_rech_ant = n(rech_ant[0].get("total",0)) if rech_ant else 0
    datos_interanual = None
    if tiene_anio_ant:
        desde_ant = f"{anio_ant}-{mes_d}"; hasta_ant = f"{anio_ant}-{mes_h}"
        rows_ia = q(f"""SELECT SUM(CASE WHEN FECHA_TRANS_ISO IS NOT NULL AND FECHA_TRANS_ISO<=FECHA_INGRESO_ISO THEN 1 ELSE 0 END) AS trans, SUM(CASE WHEN FECHA_TRANS_ISO IS NULL THEN 1 ELSE 0 END) AS no_trans, SUM(CASE WHEN FECHA_TRANS_ISO IS NOT NULL AND FECHA_TRANS_ISO>FECHA_INGRESO_ISO THEN 1 ELSE 0 END) AS tardio, COUNT(*) AS total FROM {tabla_ant} WHERE MIC LIKE ? AND strftime('%Y-%m',FECHA_INGRESO_ISO) BETWEEN ? AND ?""", (like, desde_ant, hasta_ant))
        if rows_ia: datos_interanual = rows_ia[0]
    con.close()
    log_fn("✓ Queries completadas")
    return (totales, ev_total, ev_trans, ev_tardio, ev_no_trans,
            rechazos_mes, rechazos_cat, rechazos_ej,
            datos_ult, datos_ant, datos_interanual,
            per_ult, per_ant, anio_ant,
            impoexpo_ult, rechazos_ult_cat, total_rech_ant)

def calcular_totales(totales):
    r = totales[0] if totales else {}
    cT=n(r.get("CARGADO_TRANS",0)); cN=n(r.get("CARGADO_NO_TRANS",0)); cTd=n(r.get("CARGADO_TARDIO",0))
    lT=n(r.get("LASTRE_TRANS",0));  lN=n(r.get("LASTRE_NO_TRANS",0));  lTd=n(r.get("LASTRE_TARDIO",0))
    cTot=cT+cN+cTd; lTot=lT+lN+lTd
    return (cT,cN,cTd,cTot),(lT,lN,lTd,lTot),(cT+lT,cN+lN,cTd+lTd,cTot+lTot)

# ── Word helpers ─────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); shd=OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hex_color); tcPr.append(shd)

def agregar_tabla_word(doc, headers, rows, col_widths=None, semaforo_col=None, semaforo_total_col=None):
    table=doc.add_table(rows=1,cols=len(headers)); table.style='Table Grid'; table.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr=table.rows[0]
    for i,h in enumerate(headers):
        cell=hdr.cells[i]; cell.text=h
        cell.paragraphs[0].runs[0].bold=True
        cell.paragraphs[0].runs[0].font.size=Pt(9)
        cell.paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell,"1F3864")
    for ri,row in enumerate(rows):
        tr=table.add_row(); base_fill="F2F2F2" if ri%2==0 else "FFFFFF"
        for ci,val in enumerate(row):
            cell=tr.cells[ci]; cell.text=str(val or "")
            cell.paragraphs[0].runs[0].font.size=Pt(9)
            fill=base_fill
            if semaforo_col is not None and semaforo_total_col is not None and ci==semaforo_col:
                try:
                    v_trans=n(row[semaforo_col]); v_total=n(row[semaforo_total_col])
                    if v_total>0: fill=color_semaforo(pct_f(v_trans,v_total))
                except: pass
            set_cell_bg(cell,fill)
    if col_widths:
        for i,w in enumerate(col_widths):
            for row in table.rows: row.cells[i].width=Cm(w)
    doc.add_paragraph(); return table

def insertar_grafico(doc, img_bytes, width_cm=14):
    if not img_bytes: return
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img_bytes, width=Cm(width_cm)); doc.add_paragraph()

def kpi_box(doc, kpis):
    table=doc.add_table(rows=1,cols=len(kpis)); table.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,(label,valor,sub) in enumerate(kpis):
        cell=table.rows[0].cells[i]; set_cell_bg(cell,"EBF2FA")
        p1=cell.paragraphs[0]; p1.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p1.add_run(f"{label}\n").font.size=Pt(8)
        r2=p1.add_run(f"{valor}\n"); r2.bold=True; r2.font.size=Pt(16); r2.font.color.rgb=RGBColor(0x1F,0x3D,0x64)
        r3=p1.add_run(sub); r3.font.size=Pt(8); r3.font.color.rgb=RGBColor(0x60,0x60,0x60)
    doc.add_paragraph()

# ── Narrativa IA ─────────────────────────────────────────────────────────────────
def generar_narrativa_ia(datos, api_key):
    try:
        client = anthropic.Anthropic(api_key=api_key, http_client=httpx.Client(follow_redirects=True))
        prompt = f"""ADVERTENCIA: los datos numéricos de este prompt son los ÚNICOS válidos. Si tu memoria de entrenamiento tiene números distintos para este informe, ignoralos completamente.\nSos un analista de comercio exterior de ARCA (Aduana Argentina).
GLOSARIO OBLIGATORIO — usá EXACTAMENTE estas siglas y denominaciones, sin inventar expansiones alternativas:
- SINTIA: escribilo siempre como "SINTIA" sin expandir
- MIC-DTA: "Manifiesto Internacional de Cargas - Declaración de Tránsito Aduanero"
- MIC: "Manifiesto Internacional de Cargas" (NO "Manifiesto de Importación de Carga")
- DTA: "Declaración de Tránsito Aduanero" (NO "Declaración Transitoria Aduanera")
- PAD: la primera vez que lo uses escribí "Portal Aduanero (PAD)", luego solo "PAD"
- ARCA: "Agencia de Recaudación y Control Aduanero"
- CRT: "Carta de Porte Internacional por Carretera"
- Terminología Argentina OBLIGATORIA: usar "despachantes de aduana" (NO "agentes aduanales", "agentes aduaneros" ni otras variantes foráneas); "declarante" o "transportista" según el contexto (documento de transporte internacional)
- INDNCM: "indicador de Nomenclatura Común del Mercosur"
- PATAI: "Presentación Anticipada de Transportes de Ingreso" (evento aduanero)
- OFTAI: "Oficialización de Transportes de Ingreso" (evento aduanero)
Redactá los párrafos narrativos para un informe formal sobre el estado de situación del sistema SINTIA.
El informe describe la transmisión anticipada del MIC-DTA entre {datos['pais_nombre']} y Argentina durante {datos['periodo']}.
"""
        try:
            _tot_r = int(datos['total_rechazos'].replace('.',''))
            _tot_g = int(datos['g_total'].replace('.',''))
            pct_str_rechazos = f"{round(100*_tot_r/_tot_g,1):.1f}".replace(".",",") if _tot_g>0 else "0,0"
        except Exception:
            pct_str_rechazos = "N/D"
        prompt += f"""
Datos del período:
- Total ingresos: {datos['g_total']}
- Transmitidos correctamente: {datos['g_trans']} ({datos['pct_trans']})
- No transmitidos: {datos['g_no_trans']} ({datos['pct_no_trans']})
- Tard\u00edos: {datos['g_tardio']} ({datos['pct_tardio']})
- Rechazos: {datos['total_rechazos']} sobre el total de {datos['g_total']} ingresos ({pct_str_rechazos}% del total de ingresos). IMPORTANTE: los rechazos NO son un subconjunto de los transmitidos; son eventos independientes que afectan al total de ingresos. NUNCA calcules rechazos como porcentaje de los transmitidos correctamente.
- Cargados: {datos['c_trans']} trans ({datos['pct_c_trans']}), {datos['c_no_trans']} no trans ({datos['pct_c_no_trans']}), {datos['c_tardio']} tardío ({datos['pct_c_tardio']}) de {datos['c_total']} total
- Lastre: {datos['l_trans']} trans ({datos['pct_l_trans']}), {datos['l_no_trans']} no trans ({datos['pct_l_no_trans']}), {datos['l_tardio']} tardío ({datos['pct_l_tardio']}) de {datos['l_total']} total
- Evoluci\u00f3n mensual: {datos['ev_mensual_texto']}

Gener\u00e1 exactamente 3 bloques separados por "---BLOQUE---":
1. Introducci\u00f3n: exactamente 5 p\u00e1rrafos con esta estructura:
   - P\u00e1rrafo 1: finalidad del informe, mencionar SINTIA, circuito {datos['pais_nombre']}-Argentina, cruzamiento con PAD.
   - P\u00e1rrafo 2: per\u00edodo analizado, total de ingresos, continuidad respecto al a\u00f1o anterior.
   - P\u00e1rrafo 3: rol del PAD y de SINTIA en la transmisi\u00f3n anticipada del MIC-DTA, importancia de la sincronizaci\u00f3n.
   - P\u00e1rrafo 4: resumen de desv\u00edos detectados (% trans, % tard\u00edo, % no trans, cantidad de rechazos).
   - P\u00e1rrafo 5: frase de cierre indicando que a continuaci\u00f3n se desarrollan los indicadores.
2. Estado de situaci\u00f3n (3-4 oraciones, an\u00e1lisis n\u00fameros, comparaci\u00f3n cargado vs lastre, tendencia)
3. Rechazos (1-2 oraciones)

Estilo: formal, t\u00e9cnico, espa\u00f1ol rioplatense. Sin markdown (no uses #, *, **, _, etc.). S\u00ed us\u00e1 '---BLOQUE---' como separador entre los 3 bloques. No repetir el nombre de la secci\u00f3n al inicio del texto. Evitar anglicismos (usar "desempe\u00f1o" en lugar de "performance", etc.).REGLAS ESTRICTAS DE FORMATO:
- El período analizado es EXACTAMENTE {datos['periodo']}. NO menciones meses fuera de ese período. Si los datos llegan hasta junio, NO digas "julio" ni "datos de julio no disponibles".
- SIEMPRE comenzar la introducción con "El presente informe..." — NUNCA omitir el artículo.
- Mantener tercera persona en todo el documento — PROHIBIDO "dispongamos" o cualquier primera persona del plural.
- PROHIBIDO el imperativo de voseo en recomendaciones: NO uses "Coordiná", "Recomendá", "Implementá". Usá infinitivo ("Coordinar", "Implementar") o forma impersonal ("Se recomienda").
- NO incluyas títulos ni etiquetas sueltas como "INTRODUCCIÓN", "ESTADO DE SITUACIÓN", "RECHAZOS", "Bloque 1", "1.", etc. Empezá directo con el texto de cada bloque.
- Sin markdown ni bullets.
- Usá EXACTAMENTE los porcentajes y números que te di. No recalcules ni redondees diferente.
- Cuando menciones un porcentaje, siempre aclará el denominador: "X% de los cargados" o "X% del total de operaciones" — nunca digas solo "X% del total" si el denominador real es un subconjunto (como cargados o lastre).
- Revisá la ortografía antes de responder. No uses gerundios mal formados ni palabras inventadas."""
        msg = client.messages.create(model="claude-haiku-4-5-20251001", max_tokens=1800,
            messages=[{"role":"user","content":prompt}])
        raw_text = msg.content[0].text
        # Garantizar exactamente 3 bloques — si hay más, colapsar los extra al primero
        partes = [p.strip() for p in raw_text.split("---BLOQUE---") if p.strip()]
        partes = [p.replace("---BLOQUE---","").strip() for p in partes]
        partes = [p for p in partes if p]
        if len(partes) > 3:
            # Colapsar bloques extra al primero
            partes = ["\n\n".join(partes[:-2]), partes[-2], partes[-1]]
        elif len(partes) < 3:
            # Rellenar bloques faltantes con cadena vacía
            while len(partes) < 3: partes.append("")
        # Limpiar encabezados sobrantes
        encabezados = ["INTRODUCCIÓN", "INTRODUCCION", "ESTADO DE SITUACIÓN", "ESTADO DE SITUACION",
                       "RECHAZOS", "BLOQUE 1", "BLOQUE 2", "BLOQUE 3", "Bloque 1", "Bloque 2", "Bloque 3"]
        def limpiar_encabezado(texto):
            lineas = texto.split("\n")
            while lineas and lineas[0].strip().rstrip(":").strip() in encabezados:
                lineas.pop(0)
            return "\n".join(lineas).strip()
        result = [limpiar_encabezado(p) for p in partes]
        result = [p for p in result if p]
        return [limpiar_salida_ia(b) if b else b for b in result]
    except Exception:
        return None


# ── Frases pre-calculadas (Paso 1) ──────────────────────────────────────────────
def _dir(diff, positivo="mejoró", negativo="retrocedió", neutro="se mantuvo"):
    """Elige el verbo correcto según la dirección del cambio."""
    if diff > 0.05: return positivo
    if diff < -0.05: return negativo
    return neutro

def _pp(diff):
    """Formatea una diferencia en pp con signo y coma decimal."""
    return f"{abs(diff):.1f}".replace(".", ",")

def calcular_frases(datos):
    """
    Recibe el mismo dict que va a la IA y devuelve frases comparativas
    pre-calculadas con dirección y números correctos garantizados.
    """
    def n_(v):
        try: return int(str(v).replace(".", "").replace(",", ""))
        except: return 0
    def f_(v): return str(v)

    frases = {}

    # ── Volumen ─────────────────────────────────────────────────────
    ult_t  = n_(datos.get("ult_total", 0))
    ant_t  = n_(datos.get("ant_total", 0))
    if ant_t > 0:
        var_vol = round((ult_t - ant_t) / ant_t * 100, 1)
        dir_vol = _dir(var_vol, "incremento", "caída", "nivel similar")
        frases["frase_volumen_var"] = (
            f"una {dir_vol} de {_pp(var_vol)}% respecto a {datos.get('mes_ant_nombre','el mes anterior')} "
            f"({f_(datos.get('ant_total',''))} operaciones)"
        )
    else:
        frases["frase_volumen_var"] = "sin datos del mes anterior para comparar"

    # ── Transmisión total ────────────────────────────────────────────
    pct_ult = datos.get("ult_pct_trans", "0,0%")
    pct_ant = datos.get("ant_pct_trans", "0,0%")
    def to_f(s):
        try: return float(str(s).replace("%","").replace(",","."))
        except: return 0.0
    diff_trans = round(to_f(pct_ult) - to_f(pct_ant), 1)
    dir_trans  = _dir(diff_trans, "mejorando", "retrocediendo", "sin variación")
    frases["frase_trans_var"] = (
        f"{dir_trans} {_pp(diff_trans)} pp respecto a {datos.get('mes_ant_nombre','')} "
        f"(de {pct_ant} a {pct_ult})"
    )

    # ── Cargados ─────────────────────────────────────────────────────
    carg_tot = n_(datos.get("ult_carg_tot_n", 0))
    if carg_tot > 0:
        frases["frase_carg_detalle"] = (
            f"{datos.get('ult_pct_trans_carg','')} transmitidos ({f_(datos.get('ult_carg_trans_n',''))} de {f_(datos.get('ult_carg_tot_n',''))}), "
            f"{datos.get('ult_pct_no_trans_carg','')} no transmitidos ({f_(datos.get('ult_carg_no_trans_n',''))}), "
            f"{datos.get('ult_pct_tardio_carg','')} tardíos ({f_(datos.get('ult_carg_tardio_n',''))})"
        )
    else:
        frases["frase_carg_detalle"] = "sin datos de cargados"

    # ── Lastre — el más crítico ──────────────────────────────────────
    last_tr  = n_(datos.get("ult_last_trans_n", 0))
    last_tot = n_(datos.get("ult_last_tot_n", 0))

    # Tasa de lastre no transmitido: mes actual vs anterior
    pct_last_trans_ult = to_f(datos.get("ult_pct_trans_last", "0,0%"))
    pct_last_trans_ant = to_f(datos.get("ant_last_pct_trans", "0,0%"))
    diff_last_trans = round(pct_last_trans_ult - pct_last_trans_ant, 1)

    pct_last_nt_ult = to_f(datos.get("ult_pct_no_trans_last", "0,0%"))
    # Calculamos % no trans del mes anterior directamente
    ant_last_tr  = n_(datos.get("ant_last_trans_n", 0))
    ant_last_tot = n_(datos.get("ant_last_tot_n", 0))
    ant_last_td  = n_(datos.get("ant_last_tardio_n", 0)) if "ant_last_tardio_n" in datos else 0
    # Usar ant_last_notrans_n directo si está disponible, evita errores por tardíos no informados
    ant_last_nt_direct = n_(datos.get("ant_last_notrans_n", "0"))
    if ant_last_nt_direct > 0 and ant_last_tot > 0:
        pct_last_nt_ant = round(100 * ant_last_nt_direct / ant_last_tot, 1)
    elif ant_last_tot > 0:
        pct_last_nt_ant = round(100 * (ant_last_tot - ant_last_tr - ant_last_td) / ant_last_tot, 1)
    else:
        pct_last_nt_ant = 0.0
    diff_last_nt = round(pct_last_nt_ult - pct_last_nt_ant, 1)

    if last_tot > 0:
        # Frase de detalle de lastre
        if last_tr == 0:
            frases["frase_lastre_detalle"] = (
                f"ninguna operación lastre fue transmitida (0 de {f_(datos.get('ult_last_tot_n',''))}), "
                f"{datos.get('ult_pct_no_trans_last','')} no transmitidas ({f_(datos.get('ult_last_notrans_n',''))})"
            )
        else:
            frases["frase_lastre_detalle"] = (
                f"{datos.get('ult_pct_trans_last','')} transmitido ({f_(datos.get('ult_last_trans_n',''))} de {f_(datos.get('ult_last_tot_n',''))}), "
                f"{datos.get('ult_pct_no_trans_last','')} no transmitido ({f_(datos.get('ult_last_notrans_n',''))}), "
                f"{datos.get('ult_pct_tardio_last','')} tardío ({f_(datos.get('ult_last_tardio_n',''))})"
            )

        # Frase de variación de lastre TRANSMITIDO
        dir_last = _dir(diff_last_trans, "mejoró", "retrocedió", "se mantuvo")
        ant_last_pct = datos.get("ant_last_pct_trans", "N/D")
        frases["frase_lastre_trans_var"] = (
            f"{dir_last} {_pp(abs(diff_last_trans))} pp respecto a {datos.get('mes_ant_nombre','')} "
            f"(de {ant_last_pct} a {datos.get('ult_pct_trans_last','')})"
        )

        # Frase de variación de lastre NO TRANSMITIDO
        if pct_last_nt_ant > 0:
            dir_last_nt = _dir(diff_last_nt, "subió", "bajó", "se mantuvo")
            frases["frase_lastre_notrans_var"] = (
                f"los no transmitidos en lastre {dir_last_nt} {_pp(abs(diff_last_nt))} pp "
                f"(de {str(pct_last_nt_ant).replace('.',',')}% a "
                f"{str(pct_last_nt_ult).replace('.',',')}%)"
            )
        else:
            frases["frase_lastre_notrans_var"] = (
                f"los no transmitidos en lastre representaron {datos.get('ult_pct_no_trans_last','')} "
                f"({f_(datos.get('ult_last_notrans_n',''))} operaciones)"
            )
    else:
        frases["frase_lastre_detalle"]    = "sin operaciones lastre en el período"
        frases["frase_lastre_trans_var"]  = "N/D"
        frases["frase_lastre_notrans_var"] = "N/D"

    # ── Rechazos ─────────────────────────────────────────────────────
    rech_ult = n_(datos.get("ult_rechazos", 0))
    rech_ant = n_(datos.get("ant_rechazos", 0))
    if rech_ant > 0:
        var_rech = rech_ult - rech_ant
        dir_rech = _dir(-var_rech, "disminución", "aumento", "nivel estable")  # invertido: menos rechazos = mejora
        pct_var_rech = round(abs(var_rech) / rech_ant * 100, 0)
        frases["frase_rech_var"] = (
            f"{dir_rech} de {fmt(abs(var_rech))} rechazos respecto a "
            f"{datos.get('mes_ant_nombre','')} ({f_(datos.get('ant_rechazos',''))}), "
            f"equivalente al {int(pct_var_rech)}%"
        )
    else:
        frases["frase_rech_var"] = f"{rech_ult} rechazos en el período"

    return frases


# ── Post-procesamiento de salida IA (Pasos 2, 3) ────────────────────────────────
# Correcciones fijas: typos conocidos y problemas recurrentes
_CORRECCIONES_FIJAS = [
    # Typos de siglas
    (r"\bPATIA\b",                    "PATAI"),
    (r"\bOFTAIS\b",                   "OFTAI"),
    # Guión faltante en subtítulo BR
    (r"sin transmisión situación crítica", "sin transmisión \u2014 situación crítica"),
    # Headings markdown que la IA introduce en lugar de solo negrita
    (r"^#{1,6}\s+\*\*",             "**",     re.MULTILINE),
    (r"^#{1,6}\s+",                    "",       re.MULTILINE),
    # Anglicismos recurrentes
    (r"\bperformance\b",              "desempeño"),
    (r"\bfeedback\b",                 "retroalimentación"),
    # Errores ortográficos conocidos
    (r"\bingressad",                   "ingresad"),
    (r"\bsosteniéndose\b",            "sosteniéndose"),
]

def limpiar_salida_ia(texto):
    """
    Paso 2: elimina formato markdown espurio.
    Paso 3: aplica correcciones fijas de typos y siglas.
    """
    if not texto:
        return texto
    for item in _CORRECCIONES_FIJAS:
        patron, reemplazo = item[0], item[1]
        flags = item[2] if len(item) > 2 else 0
        texto = re.sub(patron, reemplazo, texto, flags=flags)
    # Limpiar líneas vacías triples o más
    texto = re.sub(r"\n{3,}", "\n\n", texto)
    return texto.strip()

def generar_conclusion_ia(datos, api_key):
    try:
        client = anthropic.Anthropic(api_key=api_key, http_client=httpx.Client(follow_redirects=True))
        interanual = ""
        if datos.get("tiene_interanual"):
            interanual = f"\nComparativo interanual \u2014 mismo per\u00edodo {datos['anio_ant']}:\n- Total: {datos['g_total_ant']} | Trans: {datos['pct_trans_ant']} | No trans: {datos['pct_no_trans_ant']} | Tard\u00edo: {datos['pct_tardio_ant']}"
        prompt = f"""Sos un analista de comercio exterior de ARCA (Aduana Argentina).
GLOSARIO OBLIGATORIO:
- SINTIA: escribilo como "SINTIA" sin expandir
- MIC-DTA: "Manifiesto Internacional de Cargas - Declaración de Tránsito Aduanero"
- PAD: "Portal Aduanero" (NUNCA otra expansión)
- ARCA: "Agencia de Recaudación y Control Aduanero"
- CRT: "Carta de Porte Internacional por Carretera"
- Terminología Argentina OBLIGATORIA: usar "despachantes de aduana" (NO "agentes aduanales", "agentes aduaneros" ni otras variantes foráneas); "declarante" o "transportista" según el contexto
- INDNCM: "indicador de Nomenclatura Común del Mercosur"
- PATAI: "Presentación Anticipada de Transportes de Ingreso"
- OFTAI: "Oficialización de Transportes de Ingreso"
Redactá la secci\u00f3n "Conclusiones y estado actual" de un informe SINTIA para el circuito {datos['pais_nombre']}-Argentina.

\u2550\u2550\u2550 DATOS DEL \u00daCTIMO MES: {datos['mes_ult_nombre']} \u2550\u2550\u2550
Volumen: {datos['ult_total']} operaciones ({datos['ult_impo']} importaciones, {datos['ult_expo']} exportaciones)
Transmisi\u00f3n: {datos['ult_trans']} transmitidos ({datos['ult_pct_trans']}), {datos['ult_no_trans']} no transmitidos ({datos['ult_pct_no_trans']}), {datos['ult_tardio']} tard\u00edos ({datos['ult_pct_tardio']})
  \u2014 Cargados (usar EXACTAMENTE): {datos['frase_carg_detalle']}
  \u2014 Lastre (usar EXACTAMENTE estos n\u00fameros y direcci\u00f3n): {datos['frase_lastre_detalle']}
  \u2014 Variaci\u00f3n lastre transmitido (usar EXACTAMENTE): {datos['frase_lastre_trans_var']}
  \u2014 Variaci\u00f3n lastre no transmitido (usar EXACTAMENTE): {datos['frase_lastre_notrans_var']}
REGLA ABSOLUTA: las cuatro frases anteriores fueron calculadas por el sistema con los valores y direcci\u00f3n correctos. Incorporalas en la redacci\u00f3n sin modificar n\u00fameros, porcentajes ni el sentido de la variaci\u00f3n. NUNCA recalcules direcci\u00f3n de cambio de lastre por tu cuenta.
Rechazos del mes: total={datos['ult_rechazos']}, duplicados={datos['ult_rech_duplicados']}, operativos reales={datos['ult_rech_operativos']}. Variación (usar EXACTAMENTE): {datos['frase_rech_var']}. NO recalcules ni reformules la variación de rechazos.
Rechazos operativos (sin duplicados): {datos['ult_rech_operativos']}
Top categor\u00edas (SOLO estas, no inventes otras ni cambies los n\u00fameros): {datos['ult_rech_top_cats']}
Total rechazos verificado: {datos['ult_rech_total_check']} (us\u00e1 EXACTAMENTE este n\u00famero)

\u2550\u2550\u2550 MES ANTERIOR: {datos['mes_ant_nombre']} \u2550\u2550\u2550
Volumen: {datos['ant_total']} operaciones
Transmisi\u00f3n: {datos['ant_trans']} ({datos['ant_pct_trans']}), no trans {datos['ant_no_trans']} ({datos['ant_pct_no_trans']}), tard\u00edos {datos['ant_tardio']} ({datos['ant_pct_tardio']})
Rechazos: {datos['ant_rechazos']} — ESTE ES EL ÚNICO NÚMERO VÁLIDO PARA {datos['mes_ant_nombre']}. Si recordás otro número de versiones anteriores del informe, ignoralo. El mes anterior tuvo exactamente {datos['ant_rechazos']} rechazos totales

\u2550\u2550\u2550 ACUMULADO DEL PER\u00cdODO ({datos['periodo']}, {datos['cant_meses']} meses, promedio {datos['promedio_mensual']} ops/mes) \u2550\u2550\u2550
Total: {datos['g_total']} | Trans: {datos['g_trans']} ({datos['pct_trans']}) | No trans: {datos['g_no_trans']} ({datos['pct_no_trans']}) | Tard\u00edo: {datos['g_tardio']} ({datos['pct_tardio']})
{interanual}

REGLAS PARA LAS CALIFICACIONES EN SUBT\u00cdTULOS:\n- Transmisi\u00f3n anticipada: bas\u00e1 la calificaci\u00f3n EXCLUSIVAMENTE en la variaci\u00f3n del % de transmitidos vs mes anterior.\n  {datos['var_trans_pp']} pp = la diferencia. Regla:\n  * m\u00e1s de +5pp: \"mejora significativa\"\n  * +1 a +5pp: \"leve mejora\"\n  * -1 a +1pp: \"nivel estable\" (si ambos meses tienen 0,0% de transmisión, usá \"sin transmisión — situación crítica sostenida\" en lugar de cualquier calificación de mejora/retroceso)\n  * -1 a -5pp: \"leve retroceso\"\n  * menos de -5pp: \"retroceso significativo\"\n  CASO ESPECIAL: si ult_pct_trans es \"0,0%\" NO uses jamás frases como \"mejora\" o \"mejorando gradualmente\" — 0,0% de transmisión es crítico independientemente de la variación. Tampoco uses \"cayeron al 100%\" ni \"subieron al 0%\" — decí directamente los valores absolutos.\n- Rechazos: si bajaron vs mes anterior \u2192 mejora; si subieron \u2192 retroceso; siempre mencionar los duplicados.\n- REGLA CR\u00cdTICA: si un indicador BAJ\u00d3 respecto al mes anterior, NO uses \"repunte\", \"aumento\" ni similares. Si SUBI\u00d3, NO uses \"baja\" ni \"ca\u00edda\". La direcci\u00f3n del texto debe coincidir con la direcci\u00f3n de los n\u00fameros.\n- Para rechazos: si subieron de un mes a otro es un RETROCESO aunque el n\u00famero absoluto sea bajo. No uses \"mejora\" ni \"incremento\" sin calificarlo como negativo si los rechazos aumentaron. Reformulá como \"retroceso\" o \"aumento negativo\".\n- Porcentajes mes-mes: si el % actual es MAYOR al del mes anterior, es una mejora; si es MENOR, es un retroceso. Verificá la direcci\u00f3n antes de escribir.\n- Rechazos operativos reales: si bajaron de un mes al siguiente es una MEJORA — no uses \"retroceso\" ni \"aumento relativo\" si los operativos reales disminuyeron.\n\nRedact\u00e1 la secci\u00f3n con esta estructura EXACTA (us\u00e1 estos t\u00edtulos en negrita):\n\n**Impacto del mes de {datos['mes_ult_nombre']} \u2014 {datos['pais_nombre']}/Argentina (SINTIA)**\n\n**Volumen operativo**\n[P\u00e1rrafo: total operaciones, desglose impo/expo, comparaci\u00f3n con mes anterior. Si junio tiene volumen notoriamente menor, aclar\u00e1 que es mes parcial.]\n\n**Transmisi\u00f3n anticipada \u2014 [calificaci\u00f3n seg\u00fan regla arriba, en min\u00fasculas]**\n[P\u00e1rrafo: % transmitidos \u00faltimo mes vs anterior y vs acumulado. Desglose cargados (% trans, % no trans, % tard\u00edo) y lastre (% trans con n\u00famero absoluto, % no trans). Mencionar tendencia.]\n\n**Rechazos \u2014 [calificaci\u00f3n breve en min\u00fasculas]**\n[P\u00e1rrafo: total, duplicados vs operativos reales, top categor\u00edas]\n\n**Conclusi\u00f3n**\n[2-3 alertas priorizadas con recomendaciones en forma impersonal o infinitivo. PROHIBIDO el imperativo voseante: NO uses "Coordiná", "Recomendá", "Implementá" ni similares. Usá infinitivo ("Coordinar con...", "Se recomienda...", "Es necesario implementar...") o forma impersonal.]\n\nEstilo: formal, t\u00e9cnico, espa\u00f1ol rioplatense. Sin markdown extra, solo t\u00edtulos en negrita. Evitar anglicismos. Ser directo. Revisá la ortografía y la gramática antes de responder. Palabras comunes mal escritas a evitar: "ingressadas" (correcto: "ingresadas"), "campña" (correcto: "campaña"), "sosteniéndose" (no "sosteniene"). Verificá la concordancia de número: sujetos plurales requieren verbos plurales (ej: "los despachantes dispongan", no "disponga"). Los meses se escriben en minúscula en español (enero, febrero... no Enero, Febrero). El período analizado es {datos['periodo']} — NO menciones meses fuera de ese rango. "trimestre final" no tiene sentido en un informe semestral — no lo uses. PROHIBIDO hacer proyecciones temporales con años futuros ("junio 2027", "primer trimestre 2027", etc.) — si mencionás seguimiento futuro, usá frases como "en los próximos meses" o "durante el segundo semestre de {datos['anio']}". PROHIBIDO usar frases como "un aumento de -N" o "una disminución de +N" — si el valor bajó usá "disminución de N" (positivo), si subió usá "aumento de N" (positivo). DISTINGUIR interanual de intermensual: "interanual" = comparación con el mismo período del año anterior (solo si tiene_interanual=True); "intermensual" o "respecto al mes anterior" = comparación con el mes inmediatamente anterior. PROHIBIDO usar "interanual" para comparaciones mes-a-mes."""
        msg = client.messages.create(model="claude-haiku-4-5-20251001", max_tokens=1600,
            messages=[{"role":"user","content":prompt}])
        return msg.content[0].text.strip()
    except Exception:
        return None


# ── Verificación numérica (Paso 4) ──────────────────────────────────────────────
def verificar_conclusion(texto, datos_esperados, log_fn):
    """
    Verifica números y direcciones en el texto generado por la IA.
    Loguea discrepancias sin bloquear. Retorna (texto, hay_errores).
    """
    if not texto:
        return texto, False
    import re as _re

    def _n(v):
        try: return int(str(v).replace(".", "").replace(",", ""))
        except: return None

    errores = []

    # 1. Verificar absoluto lastre no transmitido
    esp_nt = _n(datos_esperados.get("ult_last_notrans_n"))
    if esp_nt and esp_nt > 0:
        # Buscar número seguido de "no transmit" cerca de "lastre" en el mismo párrafo
        for m in _re.finditer(r"(\d[\d\.]{1,9})\s*(?:operaciones?\s*)?no\s+transmitidas?\s+(?:en\s+)?lastre|lastre[^.]{0,80}(\d[\d\.]{1,9})\s*no\s+transmitid", texto, _re.IGNORECASE):
            encontrado = _n(m.group(1) or m.group(2))
            if encontrado and abs(encontrado - esp_nt) > 1:
                errores.append(f"lastre_no_trans: texto tiene {encontrado}, esperado {esp_nt}")

    # 2. Verificar dirección de variación de lastre
    frase_var = datos_esperados.get("frase_lastre_trans_var", "")
    if "mejoró" in frase_var:
        if _re.search(r"retrocedi[oó]|deterioro|ca[ií]da.{0,40}lastre|empeor.{0,40}lastre", texto, _re.IGNORECASE):
            errores.append("dir_lastre: texto dice retroceso pero la variación fue mejora")
    elif "retrocedió" in frase_var:
        if _re.search(r"mejor[oó].{0,40}lastre|increment.{0,40}lastre", texto, _re.IGNORECASE):
            errores.append("dir_lastre: texto dice mejora pero la variación fue retroceso")

    # 3. Verificar dirección de rechazos
    rech_u = _n(datos_esperados.get("ult_rechazos"))
    rech_a = _n(datos_esperados.get("ant_rechazos"))
    if rech_u is not None and rech_a is not None and rech_a > 0:
        if rech_u < rech_a:
            if _re.search(r"retroceso.{0,30}rechazo|aument.{0,30}rechazo", texto, _re.IGNORECASE):
                errores.append("dir_rechazos: texto dice retroceso/aumento pero bajaron")
        else:
            if _re.search(r"mejor[oó].{0,30}rechazo|reducci[oó]n.{0,30}rechazo", texto, _re.IGNORECASE):
                errores.append("dir_rechazos: texto dice mejora/reducción pero subieron")

    if errores:
        for e in errores:
            log_fn(f"  \u26a0 Verificaci\u00f3n: {e}")
        return texto, True
    return texto, False

# ── Generar Word ─────────────────────────────────────────────────────────────────
def _generar_word(pais, anio, mes_d, mes_h, version,
                  totales_raw, ev_total, ev_trans, ev_tardio, ev_no_trans,
                  rechazos_mes, rechazos_cat, rechazos_ej,
                  datos_ult, datos_ant, datos_interanual, per_ult, per_ant, anio_ant,
                  impoexpo_ult, rechazos_ult_cat, total_rech_ant,
                  narrativa_ia, conclusion_ia, carpeta, log_fn):

    pais_nombre=PAISES.get(pais,pais)
    # Usar el último mes con datos reales, no el mes_h declarado
    mes_h_real = per_ult[-2:] if per_ult else mes_h
    periodo=periodo_texto(anio,mes_d,mes_h_real)
    (cT,cN,cTd,cTot),(lT,lN,lTd,lTot),(gT,gN,gTd,gTot)=calcular_totales(totales_raw)
    total_rechazos=sum(n(r.get("MIC_RECHAZOS",0)) for r in rechazos_mes)

    graficos={}
    if MPL_OK:
        log_fn("Generando gr\u00e1ficos...")
        for nombre,fn in [
            ("torta",  lambda: grafico_torta(gT,gN,gTd)),
            ("barras", lambda: grafico_barras_apiladas(ev_total)),
            ("lineas", lambda: grafico_lineas_pct(ev_total,ev_trans,ev_tardio,ev_no_trans)),
            ("rech_mes", lambda: grafico_rechazos_mes(rechazos_mes)),
            ("rech_cat", lambda: grafico_rechazos_cat(rechazos_cat)),
            ("comparativo", lambda: grafico_comparativo_meses(datos_ult,datos_ant,per_ult,per_ant)),
        ]:
            try: graficos[nombre]=fn()
            except Exception as e: log_fn(f"  Gr\u00e1fico {nombre}: {e}")
        log_fn("✓ Gr\u00e1ficos generados")

    doc=Document()
    for section in doc.sections:
        section.top_margin=Cm(2.5); section.bottom_margin=Cm(2.5)
        section.left_margin=Cm(3); section.right_margin=Cm(2.5)

    # Portada
    titulo=doc.add_heading(f"ESTADO DE SITUACI\u00d3N SINTIA {anio} {pais}-AR",0)
    titulo.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for txt,sz in [("Direcci\u00f3n de Reingeniería de Procesos Aduaneros (DG ADUA)",12),(f"Per\u00edodo: {periodo}",11)]:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=p.add_run(txt); run.font.size=Pt(sz); run.font.color.rgb=RGBColor(0x40,0x40,0x40)
    doc.add_paragraph()
    meta=doc.add_paragraph()
    meta.add_run("Versi\u00f3n: ").bold=True; meta.add_run(f"{version}   ")
    meta.add_run("\u00daltima modificaci\u00f3n: ").bold=True; meta.add_run(datetime.today().strftime("%d/%m/%Y"))
    doc.add_page_break()

    # 1. Introducción
    doc.add_heading("1.  Introducci\u00f3n",level=1)
    if narrativa_ia and len(narrativa_ia)>=1:
        doc.add_paragraph(narrativa_ia[0])
    else:
        anio_ant_intro = str(int(anio)-1)
        doc.add_paragraph(f"El presente informe tiene por finalidad exponer el estado de situaci\u00f3n del sistema SINTIA en el circuito operativo entre {pais_nombre} y Argentina, a partir del an\u00e1lisis del grado de transmisi\u00f3n, rechazos y principales inconsistencias detectadas mediante el cruzamiento de informaci\u00f3n con el sistema interno Portal Aduanero (PAD).")
        doc.add_paragraph(f"Para su elaboraci\u00f3n, se consideraron las operaciones terrestres registradas durante el per\u00edodo comprendido entre {periodo}, alcanzando un total de {fmt(gTot)} ingresos al territorio nacional. Este an\u00e1lisis permite dar continuidad a los resultados observados durante el a\u00f1o {anio_ant_intro}, evidenciando tendencias operativas que se mantienen en el tiempo.")
        doc.add_paragraph("En este contexto, el PAD constituye el sistema central de registro de las operaciones de ingreso y egreso terrestre del pa\u00eds, mientras que SINTIA cumple un rol clave en la transmisi\u00f3n anticipada del MIC-DTA, elemento fundamental para la correcta gesti\u00f3n operativa. La adecuada sincronizaci\u00f3n entre ambos sistemas resulta determinante para la optimizaci\u00f3n de los circuitos operativos y la mejora en los tiempos de registro, control y despacho.")
        doc.add_paragraph(f"Del an\u00e1lisis realizado se desprende que, durante el per\u00edodo bajo estudio, persisten desv\u00edos en el proceso de transmisi\u00f3n, con una proporci\u00f3n limitada de MICs transmitidos en forma anticipada ({pct(gT,gTot)}), presencia de transmisiones tard\u00edas ({pct(gTd,gTot)}) y un volumen significativo de operaciones no transmitidas ({pct(gN,gTot)}). Asimismo, se identificaron {fmt(total_rechazos)} MICs que por inconsistencias son rechazados.")
        doc.add_paragraph("A continuaci\u00f3n, se desarrollan en detalle los principales indicadores, su evoluci\u00f3n y los resultados estad\u00edsticos obtenidos, a fin de facilitar el an\u00e1lisis de la operatoria relevada.")

    kpi_box(doc, [
        ("TOTAL INGRESOS",  fmt(gTot),          periodo),
        ("TRANSMITIDOS",    fmt(gT),             pct(gT,gTot)),
        ("NO TRANSMITIDOS", fmt(gN),             pct(gN,gTot)),
        ("TARD\u00cdOS",         fmt(gTd),            pct(gTd,gTot)),
        ("RECHAZOS",        fmt(total_rechazos), pct(total_rechazos,gTot)),
    ])

    # 2. Estado de situación
    doc.add_heading("2.  Estado de Situaci\u00f3n",level=1)
    if narrativa_ia and len(narrativa_ia)>=2:
        doc.add_paragraph(narrativa_ia[1])
    else:
        doc.add_paragraph(f"Durante el per\u00edodo ingresaron {fmt(gTot)} operaciones: {fmt(gT)} ({pct(gT,gTot)}) transmitidas correctamente, {fmt(gN)} ({pct(gN,gTot)}) no transmitidas y {fmt(gTd)} ({pct(gTd,gTot)}) tard\u00edas.")
    doc.add_paragraph()

    agregar_tabla_word(doc,
        ["","TRANS","%","NO TRANS","%","TARD\u00cdO","%","TOTAL"],
        [["CARGADO",fmt(cT),pct(cT,cTot),fmt(cN),pct(cN,cTot),fmt(cTd),pct(cTd,cTot),fmt(cTot)],
         ["LASTRE", fmt(lT),pct(lT,lTot),fmt(lN),pct(lN,lTot),fmt(lTd),pct(lTd,lTot),fmt(lTot)],
         ["TOTAL",  fmt(gT),pct(gT,gTot),fmt(gN),pct(gN,gTot),fmt(gTd),pct(gTd,gTot),fmt(gTot)]],
        col_widths=[2.5,1.8,1.2,1.8,1.2,1.8,1.2,1.8], semaforo_col=1, semaforo_total_col=7)
    if "torta" in graficos: insertar_grafico(doc,graficos["torta"],width_cm=11)

    # 2.1 Evolución mensual
    doc.add_heading("2.1.  Evoluci\u00f3n de ingreso de camiones por mes",level=2)
    agregar_tabla_word(doc,["MES","CARGADO","LASTRE","TOTAL"],
        [[mes_label_largo(r["MES"]),fmt(n(r.get("CARGADO",0))),fmt(n(r.get("LASTRE",0))),fmt(n(r.get("CARGADO",0))+n(r.get("LASTRE",0)))] for r in ev_total],
        col_widths=[3.5,2.5,2.5,2.5])
    if "barras" in graficos: insertar_grafico(doc,graficos["barras"])

    # 2.3 Transmitidos
    doc.add_page_break()
    doc.add_heading("2.2.  Evolución de transmisión anticipada por mes",level=2)
    doc.add_paragraph()
    rows_ev_trans = []
    for r in ev_total:
        mes = r["MES"]
        tot = n(r.get("CARGADO",0)) + n(r.get("LASTRE",0))
        trs = next((n(x.get("CARGADOS",0))+n(x.get("LASTRE",0)) for x in ev_trans if x.get("MES")==mes), 0)
        ntrs = next((n(x.get("CARGADOS",0))+n(x.get("LASTRE",0)) for x in ev_no_trans if x.get("MES")==mes), 0)
        tds = next((n(x.get("CARGADOS",0))+n(x.get("LASTRE",0)) for x in ev_tardio if x.get("MES")==mes), 0)
        rows_ev_trans.append([mes_label(mes), pct(trs,tot), pct(ntrs,tot), pct(tds,tot), fmt(tot)])
    agregar_tabla_word(doc,["MES","% TRANS","% NO TRANS","% TARDÍO","TOTAL"],rows_ev_trans,col_widths=[2.8,2.5,2.5,2.5,2.2])
    doc.add_paragraph()
    doc.add_heading("2.3.  Evaluaci\u00f3n del tiempo de transmisi\u00f3n \u2013 MICs Transmitidos",level=2)
    doc.add_paragraph(f"Del total ({fmt(gTot)}), {fmt(gT)} ({pct(gT,gTot)}) fueron transmitidos correctamente.")
    agregar_tabla_word(doc,["MES","CARGADOS","LASTRE","TOTAL"],
        [[mes_label_largo(r["MES"]),fmt(n(r.get("CARGADOS",0))),fmt(n(r.get("LASTRE",0))),fmt(n(r.get("CARGADOS",0))+n(r.get("LASTRE",0)))] for r in ev_trans],
        col_widths=[3.5,2.5,2.5,2.5])

    # 2.4 Tardíos
    doc.add_heading("2.4.  Evaluaci\u00f3n del tiempo de transmisi\u00f3n \u2013 MICs tard\u00edos",level=2)
    doc.add_paragraph(f"Del total ({fmt(gTot)}), {fmt(gTd)} ({pct(gTd,gTot)}) fueron transmitidos tard\u00edamente.")
    tardio_por_mes = {r["MES"]: r for r in ev_tardio}
    rows_tardio = []
    for r in ev_total:
        mes = r["MES"]
        td = tardio_por_mes.get(mes, {})
        carg = n(td.get("CARGADOS", 0))
        last = n(td.get("LASTRE", 0))
        rows_tardio.append([mes_label_largo(mes), fmt(carg), fmt(last), fmt(carg + last)])
    agregar_tabla_word(doc,["MES","CARGADOS","LASTRE","TOTAL"],rows_tardio,
        col_widths=[3.5,2.5,2.5,2.5])

    # 2.5 No transmitidos
    doc.add_heading("2.5.  Evaluaci\u00f3n del tiempo de transmisi\u00f3n \u2013 MICs no transmitidos",level=2)
    doc.add_paragraph(f"Del total ({fmt(gTot)}), {fmt(gN)} ({pct(gN,gTot)}) no fueron transmitidos.")
    agregar_tabla_word(doc,["MES","CARGADOS","LASTRE","TOTAL"],
        [[mes_label_largo(r["MES"]),fmt(n(r.get("CARGADOS",0))),fmt(n(r.get("LASTRE",0))),fmt(n(r.get("CARGADOS",0))+n(r.get("LASTRE",0)))] for r in ev_no_trans],
        col_widths=[3.5,2.5,2.5,2.5])
    if "lineas" in graficos: insertar_grafico(doc,graficos["lineas"])

    # 2.6 Rechazos
    doc.add_heading("2.6.  An\u00e1lisis de MICs Rechazados",level=2)
    if narrativa_ia and len(narrativa_ia)>=3:
        texto_26 = narrativa_ia[2] if narrativa_ia and len(narrativa_ia) > 2 and narrativa_ia[2] and len(narrativa_ia[2]) > 20 else None
        if texto_26:
            doc.add_paragraph(texto_26)
        else:
            doc.add_paragraph(f"Se registraron {fmt(total_rechazos)} rechazos en el período ({pct(total_rechazos,gTot)} de las operaciones totales).")
    else:
        doc.add_paragraph(f"Se registraron {fmt(total_rechazos)} rechazos ({pct(total_rechazos,gTot)} de las operaciones).")
    if rechazos_mes:
        doc.add_paragraph("Rechazos por mes:").runs[0].bold=True
        rech_por_mes = {r["periodo"]: n(r.get("MIC_RECHAZOS",0)) for r in rechazos_mes}
        periodos_todos = sorted(set([x["MES"] for x in ev_total] + list(rech_por_mes.keys())))
        agregar_tabla_word(doc,["MES","MIC RECHAZOS"],
            [[mes_label_largo(p), fmt(rech_por_mes.get(p,0))] for p in periodos_todos],col_widths=[4,3])
    if "rech_mes" in graficos: insertar_grafico(doc,graficos["rech_mes"],width_cm=12)
    if rechazos_cat:
        cats_sin_total=[r for r in rechazos_cat if r["Categoria"]!="TOTAL"]
        total_cat=sum(n(r.get("Rechazos",0)) for r in cats_sin_total)
        agregar_tabla_word(doc,["CATEGOR\u00cdA","RECHAZOS","%"],
            [[r.get("Categoria",""),fmt(r.get("Rechazos",0)),pct(r.get("Rechazos",0),total_cat) if r["Categoria"]!="TOTAL" else ""] for r in rechazos_cat],
            col_widths=[7,2.5,2])
    if "rech_cat" in graficos: insertar_grafico(doc,graficos["rech_cat"])
    if rechazos_ej:
        doc.add_paragraph("Ejemplos de rechazos:").runs[0].bold=True
        agregar_tabla_word(doc,["FECHA","NRO MIC","MENSAJE"],
            [[r.get("Fecha_ISO","")[:10] if r.get("Fecha_ISO") else "",r.get("NroMic",""),r.get("Mensaje","")] for r in rechazos_ej],
            col_widths=[2.5,3.0,9.0])

    # 3. Conclusiones
    doc.add_page_break()
    doc.add_heading("3.  Conclusiones y estado actual",level=1)

    ult_t=n(datos_ult.get("total",0)) if datos_ult else 0
    ult_tr=n(datos_ult.get("trans",0)) if datos_ult else 0
    ult_nt=n(datos_ult.get("no_trans",0)) if datos_ult else 0
    ult_td=n(datos_ult.get("tardio",0)) if datos_ult else 0
    ant_t=n(datos_ant.get("total",0)) if datos_ant else 0
    ant_tr=n(datos_ant.get("trans",0)) if datos_ant else 0
    ant_nt=n(datos_ant.get("no_trans",0)) if datos_ant else 0
    ant_td=n(datos_ant.get("tardio",0)) if datos_ant else 0

    if conclusion_ia:
        # Limpiar encabezados sueltos que la IA puede generar al inicio
        conc_limpia = conclusion_ia.strip()
        # Limpiar encabezados/artefactos al inicio del texto
        enc_limpiar = [
            "# 3. Conclusiones y estado actual", "# CONCLUSIONES Y ESTADO ACTUAL",
            "# Conclusiones y estado actual", "3. Conclusiones y estado actual",
            "CONCLUSIONES Y ESTADO ACTUAL", "Conclusiones y estado actual",
            "**Conclusiones y estado actual**", "**CONCLUSIONES Y ESTADO ACTUAL**",
        ]
        for enc in enc_limpiar:
            if conc_limpia.startswith(enc):
                conc_limpia = conc_limpia[len(enc):].strip()
        # Limpiar también líneas sueltas con "— Circuito..." o subtítulos redundantes
        lineas_conc = conc_limpia.split("\n")
        lineas_filtradas = []
        for lin in lineas_conc:
            ls = lin.strip()
            # Eliminar líneas que son solo subtítulos de sección o artefactos
            ls_norm = ls.strip("*#-— .:").lower()
            es_artefacto = (
                ls_norm in [
                    "conclusiones y estado actual",
                    "3 conclusiones y estado actual",
                    "3. conclusiones y estado actual",
                    "3.  conclusiones y estado actual",
                    "conclusiones",
                ] or
                ls_norm.startswith("circuito ") or
                ls_norm.startswith("conclusiones y") or
                ls_norm.startswith("3. conclusiones") or
                ls_norm.startswith("3  conclusiones") or
                (ls.startswith("**") and "conclusiones" in ls.lower() and "estado actual" in ls.lower()) or
                (ls.startswith("#") and "conclusiones" in ls.lower())
            )
            if not es_artefacto:
                lineas_filtradas.append(lin)
        conc_limpia = "\n".join(lineas_filtradas).strip()
        for bloque in conc_limpia.split("\n\n"):
            bloque=bloque.strip()
            if not bloque: continue
            # Limpiar markdown de heading (# al inicio de línea)
            lineas_bloque = bloque.split("\n")
            lineas_limpias = []
            for lin in lineas_bloque:
                lin_strip = lin.lstrip()
                if lin_strip.startswith("# "):
                    lin = lin_strip[2:].strip()
                elif lin_strip.startswith("## "):
                    lin = lin_strip[3:].strip()
                if lin: lineas_limpias.append(lin)
            bloque = "\n".join(lineas_limpias).strip()
            if not bloque: continue
            p=doc.add_paragraph()
            parts=re.split('([*][*][^*]+[*][*])',bloque)
            for part in parts:
                if part.startswith("**") and part.endswith("**"): run=p.add_run(part[2:-2]); run.bold=True
                else: p.add_run(part)
    else:
        def bold_par(doc, titulo, texto):
            p=doc.add_paragraph(); r=p.add_run(titulo); r.bold=True; p.add_run("\n"+texto)

        rult_total=sum(n(r.get("Rechazos",0)) for r in rechazos_ult_cat)
        rult_dup=next((n(r.get("Rechazos",0)) for r in rechazos_ult_cat if r.get("Categoria")=="NRO DE MIC EXISTENTE"),0)
        rult_op=rult_total-rult_dup
        top_rech=[r for r in rechazos_ult_cat if r.get("Categoria")!="NRO DE MIC EXISTENTE"][:5]
        top_txt=", ".join([f"{r['Categoria']}: {r['Rechazos']}" for r in top_rech])
        impo_r=next((r for r in impoexpo_ult if str(r.get("TIPO_REGISTRO","")).upper()=="I"),{})
        expo_r=next((r for r in impoexpo_ult if str(r.get("TIPO_REGISTRO","")).upper()=="E"),{})
        ult_impo_n=n(impo_r.get("total",0)); ult_expo_n=n(expo_r.get("total",0))
        carg_tot_n=sum(n(r.get("cargado",0)) for r in impoexpo_ult)
        last_tot_n=sum(n(r.get("lastre",0)) for r in impoexpo_ult)
        ev_trans_ult=next((r for r in ev_trans if r.get("MES")==per_ult),{})
        carg_trans_n=n(ev_trans_ult.get("CARGADOS",0)); last_trans_n=n(ev_trans_ult.get("LASTRE",0))
        ev_tardio_ult_fb=next((r for r in ev_tardio if r.get("MES")==per_ult),{})
        ev_trans_ant_fb=next((r for r in ev_trans if r.get("MES")==per_ant),{})
        ev_total_ant_fb=next((r for r in ev_total if r.get("MES")==per_ant),{})
        last_tardio_n=n(ev_tardio_ult_fb.get("LASTRE",0))
        last_notrans_n=last_tot_n-last_trans_n-last_tardio_n
        ant_last_trans_n_fb=n(ev_trans_ant_fb.get("LASTRE",0))
        ant_last_tot_n_fb=n(ev_total_ant_fb.get("LASTRE",0))
        diff_trans=pct_f(ult_tr,ult_t)-pct_f(ant_tr,ant_t)
        diff_last_trans=pct_f(last_trans_n,last_tot_n)-pct_f(ant_last_trans_n_fb,ant_last_tot_n_fb)
        pct_ult_trans=pct_f(ult_tr,ult_t)
        # Caso especial: 0,0% de transmisión no puede calificarse como mejora
        if pct_ult_trans==0.0 and pct_f(ant_tr,ant_t)==0.0: cal_trans="sin transmisi\u00f3n \u2014 situaci\u00f3n cr\u00edtica sostenida"
        elif pct_ult_trans==0.0: cal_trans="sin transmisi\u00f3n"
        elif diff_trans>=20: cal_trans="mejora significativa"
        elif diff_trans>=5: cal_trans="mejora moderada"
        elif diff_trans>=1: cal_trans="leve mejora"
        elif diff_trans>=-1: cal_trans="nivel estable"
        elif diff_trans>=-5: cal_trans="leve deterioro"
        else: cal_trans="deterioro"
        # Usar rechazos_mes como fuente primaria (misma tabla que el informe muestra)
        rech_ant_n_mes=next((n(r.get("MIC_RECHAZOS",0)) for r in rechazos_mes if r.get("periodo")==per_ant),0)
        rech_ant_n=rech_ant_n_mes if rech_ant_n_mes>0 else total_rech_ant
        if rult_total>rech_ant_n*2: cal_rech="pico preocupante"
        elif rult_total>rech_ant_n: cal_rech="incremento"
        elif rult_total<rech_ant_n: cal_rech="reducci\u00f3n"
        else: cal_rech="nivel estable"
        tots_mes_fb={r["MES"]:n(r.get("CARGADO",0))+n(r.get("LASTRE",0)) for r in ev_total}
        trs_mes_fb={r["MES"]:n(r.get("CARGADOS",0))+n(r.get("LASTRE",0)) for r in ev_trans}
        pct_por_mes={mes:pct_f(trs_mes_fb.get(mes,0),tot) for mes,tot in tots_mes_fb.items() if tot>0}
        mejor_mes=max(pct_por_mes,key=pct_por_mes.get) if pct_por_mes else per_ult
        es_mejor=mejor_mes==per_ult

        p_titulo=doc.add_paragraph()
        r_t=p_titulo.add_run(f"Impacto del mes de {mes_label_largo(per_ult)} \u2014 {PAISES.get(pais,pais)}/Argentina (SINTIA)")
        r_t.bold=True; r_t.font.size=Pt(12)
        doc.add_paragraph()

        bold_par(doc,"Volumen operativo",
            f"{mes_label_largo(per_ult)} registr\u00f3 {fmt(ult_t)} operaciones ({fmt(ult_impo_n)} importaciones y {fmt(ult_expo_n)} exportaciones), "
            f"{'una leve ca\u00edda' if ult_t<ant_t else 'un incremento'} respecto de {mes_label_largo(per_ant)} ({fmt(ant_t)}). "
            f"La participaci\u00f3n de exportaciones fue del {pct(ult_expo_n,ult_t)}.")

        mejor_txt="el mes de mayor cumplimiento del per\u00edodo analizado" if es_mejor else f"un nivel de cumplimiento {'superior' if pct_f(ult_tr,ult_t)>pct_f(gT,gTot) else 'inferior'} al promedio del per\u00edodo ({pct(gT,gTot)})"
        if pct_ult_trans==0.0:
            bold_par(doc,f"Transmisi\u00f3n anticipada \u2014 {cal_trans}",
                f"La tasa de transmisi\u00f3n anticipada se mantuvo en {pct(ult_tr,ult_t)} en {mes_label_largo(per_ult)} "
                f"({fmt(ult_tr)} operaci\u00f3n{'es' if ult_tr!=1 else ''} transmitida{'s' if ult_tr!=1 else ''} de {fmt(ult_t)} totales), "
                f"sin variaci\u00f3n respecto a {mes_label_largo(per_ant)} ({pct(ant_tr,ant_t)}). "
                f"El {pct(ult_nt,ult_t)} de las operaciones permaneci\u00f3 sin transmitir, tanto en cargados como en lastre. "
                f"No se registraron transmisiones tard\u00edas." if ult_td==0 else
                f"El {pct(ult_nt,ult_t)} de las operaciones permaneci\u00f3 sin transmitir. "
                f"Las transmisiones tard\u00edas representaron el {pct(ult_td,ult_t)} ({fmt(ult_td)} operaciones).")
        if pct_ult_trans!=0.0:
            bold_par(doc,f"Transmisi\u00f3n anticipada \u2014 {cal_trans}",
                f"{mes_label_largo(per_ult)} es {mejor_txt}: el {pct(ult_tr,ult_t)} de los MICs fue transmitido previo al arribo "
                f"({pct(carg_trans_n,carg_tot_n)} en cargados, {pct(last_trans_n,last_tot_n)} en lastre \u2014 {fmt(last_trans_n)} de {fmt(last_tot_n)}), "
                f"frente al {pct(ant_tr,ant_t)} de {mes_label_largo(per_ant)}. "
                f"Los no transmitidos {'cayeron' if ult_nt<ant_nt else 'subieron'} al {pct(ult_nt,ult_t)}. "
                f"En lastre, {fmt(last_notrans_n)} operaciones no fueron transmitidas ({pct(last_notrans_n,last_tot_n)}), "
                f"con una variaci\u00f3n de {diff_last_trans:+.1f}\u00a0pp respecto al lastre de {mes_label_largo(per_ant)} ({pct(ant_last_trans_n_fb,ant_last_tot_n_fb)}). "
                f"Las transmisiones tard\u00edas representaron el {pct(ult_td,ult_t)} ({fmt(ult_td)} operaciones vs. {fmt(ant_td)} en {mes_label_largo(per_ant)}).")

        rech_texto=(
            f"Contra la tendencia positiva en transmisi\u00f3n, {mes_label_largo(per_ult)} concentra el mayor volumen de rechazos del per\u00edodo con {fmt(rult_total)} registros"
            if cal_rech=="pico preocupante" else
            f"{mes_label_largo(per_ult)} registr\u00f3 {fmt(rult_total)} rechazos ({'m\u00e1s' if rult_total>rech_ant_n else 'menos'} que los {fmt(rech_ant_n)} de {mes_label_largo(per_ant)})")
        if rult_dup>0:
            rech_texto+=f", de los cuales {fmt(rult_dup)} corresponden a la categor\u00eda 'NRO DE MIC EXISTENTE' (MICs duplicados). Excluyendo esa categor\u00eda, los rechazos operativos suman {fmt(rult_op)}"
        if top_txt: rech_texto+=f", distribuidos principalmente en: {top_txt}"
        rech_texto+="."
        bold_par(doc,f"Rechazos \u2014 {cal_rech}",rech_texto)

        if es_mejor and diff_trans>=15:
            concl=f"{mes_label_largo(per_ult)} marca un punto de inflexi\u00f3n positivo en la transmisi\u00f3n anticipada, acerc\u00e1ndose por primera vez a niveles de cumplimiento aceptables ({pct(ult_tr,ult_t)}). "
        elif pct_ult_trans==0.0:
            concl="La transmisi\u00f3n anticipada permanece en 0,0% durante el per\u00edodo analizado, configurando un fallo sist\u00e9mico que requiere intervenci\u00f3n inmediata. "
        elif diff_trans<0:
            concl=f"{mes_label_largo(per_ult)} muestra un retroceso en la transmisi\u00f3n anticipada respecto al mes anterior, lo que requiere atenci\u00f3n para identificar las causas del deterioro. "
        else:
            concl=f"La tendencia en transmisi\u00f3n anticipada contin\u00faa mejorando gradualmente, con {mes_label_largo(per_ult)} mostrando un {pct(ult_tr,ult_t)} de cumplimiento. "
        if rult_op>0 and top_rech:
            cat_nueva=top_rech[0]["Categoria"]
            concl+=f"Sin embargo, la persistencia de rechazos en {cat_nueva} indica inconsistencias en la calidad de los datos transmitidos, lo que requiere atención focalizada."
        bold_par(doc,"Conclusi\u00f3n",concl)

    # Tabla evolución mensual sintética
    doc.add_paragraph()
    rows_ev_sint=[]
    totales_x_mes={r["MES"]:n(r.get("CARGADO",0))+n(r.get("LASTRE",0)) for r in ev_total}
    trans_x_mes={r["MES"]:n(r.get("CARGADOS",0))+n(r.get("LASTRE",0)) for r in ev_trans}
    tardio_x_mes={r["MES"]:n(r.get("CARGADOS",0))+n(r.get("LASTRE",0)) for r in ev_tardio}
    notrans_x_mes={r["MES"]:n(r.get("CARGADOS",0))+n(r.get("LASTRE",0)) for r in ev_no_trans}
    for r in ev_total:
        mes=r["MES"]; tot=totales_x_mes.get(mes,0)
        rows_ev_sint.append([mes_label(mes),pct(trans_x_mes.get(mes,0),tot),pct(notrans_x_mes.get(mes,0),tot),pct(tardio_x_mes.get(mes,0),tot),fmt(tot)])
    agregar_tabla_word(doc,["MES","% TRANS","% NO TRANS","% TARD\u00cdO","TOTAL"],rows_ev_sint,col_widths=[2.8,2.5,2.5,2.5,2.2])

    if "comparativo" in graficos: insertar_grafico(doc,graficos["comparativo"])

    if datos_ult and datos_ant:
        agregar_tabla_word(doc,
            ["PER\u00cdODO","TOTAL","TRANS","%","NO TRANS","%","TARD\u00cdO","%"],
            [[mes_label_largo(per_ant),fmt(ant_t),fmt(ant_tr),pct(ant_tr,ant_t),fmt(ant_nt),pct(ant_nt,ant_t),fmt(ant_td),pct(ant_td,ant_t)],
             [mes_label_largo(per_ult),fmt(ult_t),fmt(ult_tr),pct(ult_tr,ult_t),fmt(ult_nt),pct(ult_nt,ult_t),fmt(ult_td),pct(ult_td,ult_t)],
             [f"Acumulado {anio}",fmt(gTot),fmt(gT),pct(gT,gTot),fmt(gN),pct(gN,gTot),fmt(gTd),pct(gTd,gTot)]],
            col_widths=[3.0,1.8,1.6,1.2,1.8,1.2,1.6,1.2], semaforo_col=2, semaforo_total_col=1)

    if datos_interanual:
        ia_t=n(datos_interanual.get("total",0)); ia_tr=n(datos_interanual.get("trans",0))
        ia_nt=n(datos_interanual.get("no_trans",0)); ia_td=n(datos_interanual.get("tardio",0))
        doc.add_paragraph(f"Comparativo interanual ({periodo_texto(anio_ant,mes_d,mes_h)}):").runs[0].bold=True
        agregar_tabla_word(doc,
            ["PER\u00cdODO","TOTAL","TRANS","%","NO TRANS","%","TARD\u00cdO","%"],
            [[f"{anio_ant}",fmt(ia_t),fmt(ia_tr),pct(ia_tr,ia_t),fmt(ia_nt),pct(ia_nt,ia_t),fmt(ia_td),pct(ia_td,ia_t)],
             [f"{anio}",fmt(gTot),fmt(gT),pct(gT,gTot),fmt(gN),pct(gN,gTot),fmt(gTd),pct(gTd,gTot)]],
            col_widths=[3.0,1.8,1.6,1.2,1.8,1.2,1.6,1.2], semaforo_col=2, semaforo_total_col=1)

    nombre=f"Informe_SINTIA_{pais}_{anio}_{mes_d}-{mes_h}_v{version}.docx"
    ruta=os.path.join(carpeta,nombre); doc.save(ruta); log_fn("✓ Informe Word generado")
    return ruta

# ── Generar Excel ─────────────────────────────────────────────────────────────────
def _generar_excel(pais, anio, mes_d, mes_h, version,
                   totales_raw, ev_total, ev_trans, ev_tardio, ev_no_trans,
                   rechazos_mes, rechazos_cat, rechazos_ej,
                   datos_ult, datos_ant, datos_interanual, per_ult, per_ant, anio_ant, carpeta, log_fn):
    wb=openpyxl.Workbook(); wb.remove(wb.active)
    HDR_FILL=PatternFill("solid",fgColor="1F3864"); HDR_FONT=Font(bold=True,color="FFFFFF",size=10)
    ALT_FILL=PatternFill("solid",fgColor="EEF2F7"); NORM_FONT=Font(size=10)
    CENTER=Alignment(horizontal="center",vertical="center"); LEFT=Alignment(horizontal="left",vertical="center")
    bs=Side(style="thin",color="CCCCCC"); BORDER=Border(left=bs,right=bs,top=bs,bottom=bs)
    def add_sheet(name, headers, rows, semaforo_col=None, semaforo_total_col=None):
        ws=wb.create_sheet(name); ws.append(headers)
        for ci,h in enumerate(headers,1):
            cell=ws.cell(1,ci); cell.fill=HDR_FILL; cell.font=HDR_FONT; cell.alignment=CENTER; cell.border=BORDER
        for ri,row in enumerate(rows,2):
            for ci,val in enumerate(row,1):
                cell=ws.cell(ri,ci,val); cell.font=NORM_FONT; cell.border=BORDER
                fill=ALT_FILL if ri%2==0 else PatternFill()
                if semaforo_col is not None and semaforo_total_col is not None and ci==semaforo_col+1:
                    try:
                        v_trans=n(row[semaforo_col]); v_total=n(row[semaforo_total_col])
                        if v_total>0: fill=PatternFill("solid",fgColor=color_semaforo(pct_f(v_trans,v_total)))
                    except: pass
                cell.fill=fill; cell.alignment=LEFT
        for ci in range(1,len(headers)+1):
            col=get_column_letter(ci)
            ws.column_dimensions[col].width=max(len(str(headers[ci-1])),max((len(str(r[ci-1] if ci-1<len(r) else "")) for r in rows),default=0))+3
        return ws
    (cT,cN,cTd,cTot),(lT,lN,lTd,lTot),(gT,gN,gTd,gTot)=calcular_totales(totales_raw)
    add_sheet("Resumen Global",["Tipo","Trans","No Trans","Tardio","Total","% Trans","% No Trans","% Tardio"],
        [["Cargado",cT,cN,cTd,cTot,pct(cT,cTot),pct(cN,cTot),pct(cTd,cTot)],
         ["Lastre", lT,lN,lTd,lTot,pct(lT,lTot),pct(lN,lTot),pct(lTd,lTot)],
         ["Total",  gT,gN,gTd,gTot,pct(gT,gTot),pct(gN,gTot),pct(gTd,gTot)]])
    add_sheet("Ev Total",["Mes","Cargado","Lastre","Total"],
        [[mes_label_largo(r["MES"]),n(r.get("CARGADO",0)),n(r.get("LASTRE",0)),n(r.get("CARGADO",0))+n(r.get("LASTRE",0))] for r in ev_total])
    add_sheet("Ev Transmitidos",["Mes","Cargados","Lastre","Total"],
        [[mes_label_largo(r["MES"]),n(r.get("CARGADOS",0)),n(r.get("LASTRE",0)),n(r.get("CARGADOS",0))+n(r.get("LASTRE",0))] for r in ev_trans])
    add_sheet("Ev Tardios",["Mes","Cargados","Lastre","Total"],
        [[mes_label_largo(r["MES"]),n(r.get("CARGADOS",0)),n(r.get("LASTRE",0)),n(r.get("CARGADOS",0))+n(r.get("LASTRE",0))] for r in ev_tardio])
    add_sheet("Ev No Transmitidos",["Mes","Cargados","Lastre","Total"],
        [[mes_label_largo(r["MES"]),n(r.get("CARGADOS",0)),n(r.get("LASTRE",0)),n(r.get("CARGADOS",0))+n(r.get("LASTRE",0))] for r in ev_no_trans])
    add_sheet("Rechazos x Mes",["Periodo","MIC Rechazos"],
        [[r.get("periodo",""),n(r.get("MIC_RECHAZOS",0))] for r in rechazos_mes])
    add_sheet("Rechazos x Categoria",["Categoria","Rechazos"],
        [[r.get("Categoria",""),n(r.get("Rechazos",0))] for r in rechazos_cat])
    add_sheet("Rechazos Ejemplos",["Fecha","NroMic","Mensaje"],
        [[r.get("Fecha_ISO","")[:10] if r.get("Fecha_ISO") else "",r.get("NroMic",""),r.get("Mensaje","")] for r in rechazos_ej])
    rows_comp=[]
    if datos_ant:
        at=n(datos_ant.get("total",0)); atr=n(datos_ant.get("trans",0)); ant_=n(datos_ant.get("no_trans",0)); atd=n(datos_ant.get("tardio",0))
        rows_comp.append([mes_label_largo(per_ant),at,atr,pct(atr,at),ant_,pct(ant_,at),atd,pct(atd,at)])
    if datos_ult:
        ut=n(datos_ult.get("total",0)); utr=n(datos_ult.get("trans",0)); unt=n(datos_ult.get("no_trans",0)); utd=n(datos_ult.get("tardio",0))
        rows_comp.append([mes_label_largo(per_ult),ut,utr,pct(utr,ut),unt,pct(unt,ut),utd,pct(utd,ut)])
    rows_comp.append([f"Acumulado {anio}",gTot,gT,pct(gT,gTot),gN,pct(gN,gTot),gTd,pct(gTd,gTot)])
    if datos_interanual:
        ia_t=n(datos_interanual.get("total",0)); ia_tr=n(datos_interanual.get("trans",0))
        ia_nt=n(datos_interanual.get("no_trans",0)); ia_td=n(datos_interanual.get("tardio",0))
        rows_comp.append([f"Mismo per\u00edodo {anio_ant}",ia_t,ia_tr,pct(ia_tr,ia_t),ia_nt,pct(ia_nt,ia_t),ia_td,pct(ia_td,ia_t)])
    add_sheet("Comparativo",["Per\u00edodo","Total","Trans","%","No Trans","%","Tard\u00edo","%"],rows_comp)
    nombre=f"Informe_SINTIA_{pais}_{anio}_{mes_d}-{mes_h}_v{version}.xlsx"
    ruta=os.path.join(carpeta,nombre); wb.save(ruta); log_fn("✓ Planilla Excel generada")
    return ruta

# ── Punto de entrada principal ─────────────────────────────────────────────────
def generar_informe(ruta_db, pais, anio, mes_d, mes_h, usar_ia, api_key, carpeta, log_fn):
    nombre_base=f"Informe_SINTIA_{pais}_{anio}_{mes_d}-{mes_h}"
    version=1
    while os.path.exists(os.path.join(carpeta,f"{nombre_base}_v{version}.docx")): version+=1

    (totales,ev_total,ev_trans,ev_tardio,ev_no_trans,
     rechazos_mes,rechazos_cat,rechazos_ej,
     datos_ult,datos_ant,datos_interanual,
     per_ult,per_ant,anio_ant,
     impoexpo_ult,rechazos_ult_cat,total_rech_ant) = correr_queries(ruta_db,pais,anio,mes_d,mes_h,log_fn)

    narrativa_ia=None; conclusion_ia=None

    if usar_ia and api_key and ANT_OK:
        (cT,cN,cTd,cTot),(lT,lN,lTd,lTot),(gT,gN,gTd,gTot)=calcular_totales(totales)
        total_rechazos=sum(n(r.get("MIC_RECHAZOS",0)) for r in rechazos_mes)
        ev_txt="; ".join([f"{mes_label_largo(r['MES'])}: {n(r.get('CARGADO',0))+n(r.get('LASTRE',0))} ops" for r in ev_total])
        log_fn("Generando narrativa con IA...")
        narrativa_ia=generar_narrativa_ia({
            "pais_nombre":PAISES.get(pais,pais),"periodo":periodo_texto(anio,mes_d,per_ult[-2:] if per_ult else mes_h),
            "g_total":fmt(gTot),"g_trans":fmt(gT),"g_no_trans":fmt(gN),"g_tardio":fmt(gTd),
            "pct_trans":pct(gT,gTot),"pct_no_trans":pct(gN,gTot),"pct_tardio":pct(gTd,gTot),
            "cant_meses":str(int(mes_h)-int(mes_d)+1),
            "promedio_mensual":fmt(round(gTot/(int(mes_h)-int(mes_d)+1))) if int(mes_h)>=int(mes_d) else "N/D",
            "c_trans":fmt(cT),"c_no_trans":fmt(cN),"c_tardio":fmt(cTd),"c_total":fmt(cTot),
            "pct_c_trans":pct(cT,cTot),"pct_c_no_trans":pct(cN,cTot),"pct_c_tardio":pct(cTd,cTot),
            "l_trans":fmt(lT),"l_no_trans":fmt(lN),"l_tardio":fmt(lTd),"l_total":fmt(lTot),
            "pct_l_trans":pct(lT,lTot),"pct_l_no_trans":pct(lN,lTot),"pct_l_tardio":pct(lTd,lTot),
            "total_rechazos":fmt(total_rechazos),"ev_mensual_texto":ev_txt,
        }, api_key)
        if narrativa_ia: log_fn("✓ Narrativa generada")
        else: log_fn("  Narrativa no disponible, usando texto est\u00e1ndar")

        log_fn("Generando conclusi\u00f3n con IA...")
        ult_t=n(datos_ult.get("total",0)) if datos_ult else 0
        ult_tr=n(datos_ult.get("trans",0)) if datos_ult else 0
        ult_nt=n(datos_ult.get("no_trans",0)) if datos_ult else 0
        ult_td=n(datos_ult.get("tardio",0)) if datos_ult else 0
        ant_t=n(datos_ant.get("total",0)) if datos_ant else 0
        ant_tr=n(datos_ant.get("trans",0)) if datos_ant else 0
        ant_nt=n(datos_ant.get("no_trans",0)) if datos_ant else 0
        ant_td=n(datos_ant.get("tardio",0)) if datos_ant else 0
        impo_ult=next((r for r in impoexpo_ult if r.get("TIPO_REGISTRO","").upper()=="I"),{})
        expo_ult=next((r for r in impoexpo_ult if r.get("TIPO_REGISTRO","").upper()=="E"),{})
        ult_carg_tot=n(impo_ult.get("cargado",0))+n(expo_ult.get("cargado",0))
        ult_last_tot=n(impo_ult.get("lastre",0))+n(expo_ult.get("lastre",0))
        ev_trans_ult_ia=next((r for r in ev_trans if r.get("MES")==per_ult),{})
        ev_tardio_ult_ia=next((r for r in ev_tardio if r.get("MES")==per_ult),{})
        carg_trans_n_ia=n(ev_trans_ult_ia.get("CARGADOS",0))
        last_trans_n_ia=n(ev_trans_ult_ia.get("LASTRE",0))
        carg_tr=carg_trans_n_ia
        # Usar rechazos_mes para el total del mes ult (misma fuente que la tabla)
        ult_rech_total_mes=next((n(r.get("MIC_RECHAZOS",0)) for r in rechazos_mes if r.get("periodo")==per_ult),0)
        ult_rech_total_cat=sum(n(r.get("Rechazos",0)) for r in rechazos_ult_cat)
        ult_rech_total=ult_rech_total_mes if ult_rech_total_mes>0 else ult_rech_total_cat
        ult_rech_dup=next((n(r.get("Rechazos",0)) for r in rechazos_ult_cat if r.get("Categoria")=="NRO DE MIC EXISTENTE"),0)
        # Recalcular op con el total de la tabla para consistencia
        ult_rech_op_cat=ult_rech_total_cat-ult_rech_dup
        ult_rech_op=ult_rech_total_mes-ult_rech_dup if ult_rech_total_mes>0 and ult_rech_total_mes>=ult_rech_dup else ult_rech_op_cat
        ult_rech_top=", ".join([f"{r['Categoria']}: {r['Rechazos']}" for r in rechazos_ult_cat if r.get("Categoria")!="NRO DE MIC EXISTENTE"][:7])
        tots_mes={r["MES"]:n(r.get("CARGADO",0))+n(r.get("LASTRE",0)) for r in ev_total}
        trs_mes={r["MES"]:n(r.get("CARGADOS",0))+n(r.get("LASTRE",0)) for r in ev_trans}
        tds_mes={r["MES"]:n(r.get("CARGADOS",0))+n(r.get("LASTRE",0)) for r in ev_tardio}
        nts_mes={r["MES"]:n(r.get("CARGADOS",0))+n(r.get("LASTRE",0)) for r in ev_no_trans}
        ev_tabla_lineas=[f"{mes_label(r['MES'])}: total={fmt(tots_mes.get(r['MES'],0))}, trans={pct(trs_mes.get(r['MES'],0),tots_mes.get(r['MES'],0))}, no_trans={pct(nts_mes.get(r['MES'],0),tots_mes.get(r['MES'],0))}, tardio={pct(tds_mes.get(r['MES'],0),tots_mes.get(r['MES'],0))}" for r in ev_total]
        ia_t=n(datos_interanual.get("total",0)) if datos_interanual else 0
        ia_tr=n(datos_interanual.get("trans",0)) if datos_interanual else 0
        ia_nt=n(datos_interanual.get("no_trans",0)) if datos_interanual else 0
        ia_td=n(datos_interanual.get("tardio",0)) if datos_interanual else 0
        # Calcular ant_rechazos desde rechazos_mes (misma fuente que la tabla)
        ant_rech_mes_val=next((n(r.get("MIC_RECHAZOS",0)) for r in rechazos_mes if r.get("periodo")==per_ant),0)
        ant_rech_final=ant_rech_mes_val if ant_rech_mes_val>0 else total_rech_ant
        # Datos de lastre del mes anterior (para variación correcta lastre vs lastre)
        ev_trans_ant_ia=next((r for r in ev_trans if r.get("MES")==per_ant),{})
        ev_total_ant_ia=next((r for r in ev_total if r.get("MES")==per_ant),{})
        ant_last_trans_n=n(ev_trans_ant_ia.get("LASTRE",0))
        ant_last_tot_n=n(ev_total_ant_ia.get("LASTRE",0))
        ev_tardio_ant_ia2=next((r for r in ev_tardio if r.get("MES")==per_ant),{})
        ant_last_tardio_n2=n(ev_tardio_ant_ia2.get("LASTRE",0))
        ant_last_notrans_n=ant_last_tot_n-ant_last_trans_n-ant_last_tardio_n2
        # Valor absoluto correcto de lastre no transmitido del mes actual (tot - trans - tardio)
        ult_last_tardio_n_ia=n(ev_tardio_ult_ia.get("LASTRE",0))
        ult_last_notrans_n=ult_last_tot-last_trans_n_ia-ult_last_tardio_n_ia
        # Paso 1: pre-calcular frases comparativas con dirección garantizada
        _frases_conc = calcular_frases({
            "ult_total":fmt(ult_t), "ant_total":fmt(ant_t),
            "mes_ult_nombre":mes_label_largo(per_ult), "mes_ant_nombre":mes_label_largo(per_ant),
            "ult_pct_trans":pct(ult_tr,ult_t), "ant_pct_trans":pct(ant_tr,ant_t),
            "ult_carg_trans_n":fmt(carg_tr), "ult_carg_tot_n":fmt(ult_carg_tot),
            "ult_carg_no_trans_n":fmt(ult_carg_tot-carg_tr-n(ev_tardio_ult_ia.get("CARGADOS",0))) if ult_carg_tot else "0",
            "ult_carg_tardio_n":fmt(n(ev_tardio_ult_ia.get("CARGADOS",0))),
            "ult_pct_trans_carg":pct(carg_tr,ult_carg_tot) if ult_carg_tot else "N/D",
            "ult_pct_no_trans_carg":pct(ult_carg_tot-carg_tr-n(ev_tardio_ult_ia.get("CARGADOS",0)),ult_carg_tot) if ult_carg_tot else "N/D",
            "ult_pct_tardio_carg":pct(n(ev_tardio_ult_ia.get("CARGADOS",0)),ult_carg_tot) if ult_carg_tot else "N/D",
            "ult_last_trans_n":fmt(last_trans_n_ia), "ult_last_tot_n":fmt(ult_last_tot),
            "ult_last_notrans_n":fmt(ult_last_notrans_n), "ult_last_tardio_n":fmt(ult_last_tardio_n_ia),
            "ult_pct_trans_last":pct(last_trans_n_ia,ult_last_tot) if ult_last_tot else "N/D",
            "ult_pct_no_trans_last":pct(ult_last_notrans_n,ult_last_tot) if ult_last_tot else "N/D",
            "ult_pct_tardio_last":pct(ult_last_tardio_n_ia,ult_last_tot) if ult_last_tot else "N/D",
            "ant_last_trans_n":fmt(ant_last_trans_n), "ant_last_tot_n":fmt(ant_last_tot_n),
            "ant_last_pct_trans":pct(ant_last_trans_n,ant_last_tot_n) if ant_last_tot_n else "N/D",
            "ant_last_notrans_n":fmt(ant_last_notrans_n) if ant_last_tot_n else "0",
            "ult_rechazos":fmt(ult_rech_total), "ant_rechazos":fmt(ant_rech_final),
            "ult_rech_duplicados":fmt(ult_rech_dup), "ult_rech_operativos":fmt(ult_rech_op),
        })
        conclusion_ia=generar_conclusion_ia({
            "pais_nombre":PAISES.get(pais,pais),"periodo":periodo_texto(anio,mes_d,per_ult[-2:] if per_ult else mes_h),
            "mes_ult_nombre":mes_label_largo(per_ult),"mes_ant_nombre":mes_label_largo(per_ant),
            "ult_total":fmt(ult_t),"ult_impo":fmt(n(impo_ult.get("total",0))),"ult_pct_impo":pct(n(impo_ult.get("total",0)),ult_t) if ult_t else "N/D","ult_expo":fmt(n(expo_ult.get("total",0))),"ult_pct_expo":pct(n(expo_ult.get("total",0)),ult_t) if ult_t else "N/D",
            "ult_trans":fmt(ult_tr),"ult_pct_trans":pct(ult_tr,ult_t),
            "ult_no_trans":fmt(ult_nt),"ult_pct_no_trans":pct(ult_nt,ult_t),
            "ult_tardio":fmt(ult_td),"ult_pct_tardio":pct(ult_td,ult_t),
            "ult_pct_trans_carg":pct(carg_tr,ult_carg_tot) if ult_carg_tot else "N/D",
            "ult_pct_no_trans_carg":pct(ult_carg_tot-carg_tr-n(ev_tardio_ult_ia.get("CARGADOS",0)),ult_carg_tot) if ult_carg_tot else "N/D",
            "ult_pct_tardio_carg":pct(n(ev_tardio_ult_ia.get("CARGADOS",0)),ult_carg_tot) if ult_carg_tot else "N/D",
            "ult_carg_trans_n":fmt(carg_tr),
            "ult_carg_no_trans_n":fmt(ult_carg_tot-carg_tr-n(ev_tardio_ult_ia.get("CARGADOS",0))) if ult_carg_tot else "N/D",
            "ult_carg_tardio_n":fmt(n(ev_tardio_ult_ia.get("CARGADOS",0))),
            "ult_carg_lastre_tardio_n":fmt(n(ev_tardio_ult_ia.get("LASTRE",0))),
            "ult_carg_tot_n":fmt(ult_carg_tot),
            "ult_pct_trans_last":pct(last_trans_n_ia,ult_last_tot) if ult_last_tot else "N/D",
            "ult_last_trans_n":fmt(last_trans_n_ia),"ult_last_tot_n":fmt(ult_last_tot),
            "var_trans_pp":f"{round(pct_f(ult_tr,ult_t)-pct_f(ant_tr,ant_t),1):+.1f}".replace(".",","),
            "var_carg_trans_pp":f"{round(pct_f(carg_tr,ult_carg_tot)-pct_f(ant_tr,ant_t),1):+.1f}".replace(".",",") if ult_carg_tot and ant_t else "N/D",
            "var_last_trans_pp":f"{round(pct_f(last_trans_n_ia,ult_last_tot)-pct_f(ant_last_trans_n,ant_last_tot_n),1):+.1f}".replace(".",",") if ult_last_tot and ant_last_tot_n else "N/D","ant_last_pct_trans":pct(ant_last_trans_n,ant_last_tot_n) if ant_last_tot_n else "N/D","ant_last_trans_n":fmt(ant_last_trans_n),"ant_last_tot_n":fmt(ant_last_tot_n),"ult_last_notrans_n":fmt(ult_last_notrans_n),"ult_last_tardio_n":fmt(ult_last_tardio_n_ia),
            "ult_pct_no_trans_last":pct(ult_last_tot-last_trans_n_ia-n(ev_tardio_ult_ia.get("LASTRE",0)),ult_last_tot) if ult_last_tot else "N/D",
            "ult_pct_tardio_last":pct(n(ev_tardio_ult_ia.get("LASTRE",0)),ult_last_tot) if ult_last_tot else "N/D",
            "ult_rechazos":fmt(ult_rech_total),"ult_rech_duplicados":fmt(ult_rech_dup),
            "ult_rech_operativos":fmt(ult_rech_op),"ult_rech_top_cats":ult_rech_top,
            "ult_rech_total_check":fmt(ult_rech_total_mes if ult_rech_total_mes>0 else ult_rech_total_cat),
            "ant_total":fmt(ant_t),"ant_trans":fmt(ant_tr),"ant_pct_trans":pct(ant_tr,ant_t),
            "ant_no_trans":fmt(ant_nt),"ant_pct_no_trans":pct(ant_nt,ant_t),
            "ant_tardio":fmt(ant_td),"ant_pct_tardio":pct(ant_td,ant_t),
            "ant_rechazos":fmt(ant_rech_final),
            "var_rech_abs":fmt(abs(ult_rech_total-ant_rech_final)),
            "var_rech_abs_num":abs(ult_rech_total-ant_rech_final),
            "var_rech_abs_texto":f"{fmt(abs(ult_rech_total-ant_rech_final))} rechazos",
            "var_rech_dir":"aumento" if ult_rech_total>ant_rech_final else "disminución",
            "var_rech_pct":(f"{round(abs(ult_rech_total-ant_rech_final)/ant_rech_final*100,0):.0f}%" if ant_rech_final>0 else "N/D"),"ev_tabla_texto":" | ".join(ev_tabla_lineas),
            "g_total":fmt(gTot),"g_trans":fmt(gT),"g_no_trans":fmt(gN),"g_tardio":fmt(gTd),
            "pct_trans":pct(gT,gTot),"pct_no_trans":pct(gN,gTot),"pct_tardio":pct(gTd,gTot),
            "cant_meses":str(int(mes_h)-int(mes_d)+1),
            "promedio_mensual":fmt(round(gTot/(int(mes_h)-int(mes_d)+1))) if int(mes_h)>=int(mes_d) else "N/D",
            "tiene_interanual":datos_interanual is not None,"anio_ant":anio_ant,"anio":anio,
            "g_total_ant":fmt(ia_t),"pct_trans_ant":pct(ia_tr,ia_t),
            "pct_no_trans_ant":pct(ia_nt,ia_t),"pct_tardio_ant":pct(ia_td,ia_t),
            # Paso 1: frases pre-calculadas con dirección garantizada
            "frase_volumen_var":_frases_conc.get("frase_volumen_var",""),
            "frase_trans_var":_frases_conc.get("frase_trans_var",""),
            "frase_carg_detalle":_frases_conc.get("frase_carg_detalle",""),
            "frase_lastre_detalle":_frases_conc.get("frase_lastre_detalle",""),
            "frase_lastre_trans_var":_frases_conc.get("frase_lastre_trans_var",""),
            "frase_lastre_notrans_var":_frases_conc.get("frase_lastre_notrans_var",""),
            "frase_rech_var":_frases_conc.get("frase_rech_var",""),
        }, api_key)
        if conclusion_ia:
            conclusion_ia, _hay_err = verificar_conclusion(conclusion_ia, {
                "ult_last_notrans_n": fmt(ult_last_notrans_n),
                "frase_lastre_trans_var": _frases_conc.get("frase_lastre_trans_var",""),
                "ult_rechazos": fmt(ult_rech_total),
                "ant_rechazos": fmt(ant_rech_final),
            }, log_fn)
            log_fn("\u2713 Conclusi\u00f3n generada" + (" (\u26a0 con advertencias)" if _hay_err else ""))
        else: log_fn("  Conclusi\u00f3n no disponible, usando texto est\u00e1ndar")

    log_fn("Generando archivos...")
    archivos=[]
    if DOCX_OK:
        ruta=_generar_word(pais,anio,mes_d,mes_h,version,totales,ev_total,ev_trans,ev_tardio,ev_no_trans,
            rechazos_mes,rechazos_cat,rechazos_ej,datos_ult,datos_ant,datos_interanual,
            per_ult,per_ant,anio_ant,impoexpo_ult,rechazos_ult_cat,total_rech_ant,
            narrativa_ia,conclusion_ia,carpeta,log_fn)
        archivos.append(ruta)
    if XLSX_OK:
        ruta=_generar_excel(pais,anio,mes_d,mes_h,version,totales,ev_total,ev_trans,ev_tardio,ev_no_trans,
            rechazos_mes,rechazos_cat,rechazos_ej,datos_ult,datos_ant,datos_interanual,
            per_ult,per_ant,anio_ant,carpeta,log_fn)
        archivos.append(ruta)
    log_fn(f"✓ Proceso completado \u2014 {len(archivos)} archivo(s) listos para descargar")
    return archivos
