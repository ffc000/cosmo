"""
CosmoTools — app.py v2
Plataforma de herramientas DI REPA / ARCA
"""

import os, sqlite3, io, uuid, threading, bcrypt, logging, subprocess, json, tempfile, re, secrets, time, shutil
import xml.etree.ElementTree as ET
import urllib.request
import urllib.parse
from datetime import datetime, timedelta, date
from functools import wraps
from flask import render_template, request, jsonify, send_file, session, redirect, url_for
from werkzeug.utils import secure_filename

# app, csrf y limiter viven en core.py — ahí también los necesitan los
# blueprints, y así evitamos que "limiter" no exista todavía cuando se
# importa un blueprint que lo usa en un decorador de ruta.
from core import app, csrf, limiter, notificar_telegram

# ── Blueprints ─────────────────────────────────────────────────────────────────
# Extraídos de app.py en la Fase 2 de profesionalización (separar los ~200
# endpoints en módulos por área). Siguen: vua, senasa, finanzas — de a uno,
# corriendo los tests entre cada extracción.
from blueprints.stock import stock_bp
app.register_blueprint(stock_bp)
from blueprints.training import training_bp
app.register_blueprint(training_bp)
from blueprints.vua import vua_bp
app.register_blueprint(vua_bp)
from blueprints.senasa import senasa_bp
app.register_blueprint(senasa_bp)
from blueprints.finanzas import finanzas_bp
app.register_blueprint(finanzas_bp)

# ── Config ─────────────────────────────────────────────────────────────────────
from core import (
    HIST_DB, DB_PATH, OUTPUT_FOLDER, STOCK_REPORTS_DIR, get_db, TELEGRAM_TOKEN, _exportar_xlsx,
    login_required, admin_required, modulo_required, finanzas_owner_required,
    tiene_permiso_admin, registrar_sesion, actualizar_sesion, token_revocado,
)

# get_api_key ahora se importa de core.py (más abajo, junto con contexto_repositorio)
APP_USER      = os.environ.get("APP_USER",    "cosmo")
APP_PASS      = os.environ.get("APP_PASS",    "")
APP_USER2     = os.environ.get("APP_USER2",   "")
APP_PASS2     = os.environ.get("APP_PASS2",   "")
os.makedirs(OUTPUT_FOLDER, exist_ok=True)
os.makedirs("/data/minutas", exist_ok=True)
os.makedirs("/tmp/sintia_uploads", exist_ok=True)
os.makedirs(STOCK_REPORTS_DIR, exist_ok=True)

# ── Auth: login_required/admin_required/modulo_required/finanzas_owner_required
# ahora viven en core.py (los necesitan también los blueprints) ────────────────
def check_password(plain, hashed):
    try: return bcrypt.checkpw(plain.encode(), hashed.encode())
    except Exception: return False  # nunca comparar en texto plano

# ── Logging ────────────────────────────────────────────────────────────────────
logging.basicConfig(filename="/data/accesos.log", level=logging.INFO,
    format="%(asctime)s %(message)s", datefmt="%Y-%m-%d %H:%M:%S")

# ── Context processor ──────────────────────────────────────────────────────────
@app.context_processor
def inject_session_vars():
    return {
        "modulos": session.get("modulos", []),
        "user_role": session.get("role", ""),
    }

# ── Job queue: job_status, job_create, job_get, _job_persist, _JobLog ahora
# viven en core.py (los usan también los blueprints, ej. vua) ──────────────
import threading as _threading

# ── Recordatorio de servicios recurrentes (finanzas) ────────────────────────────
# _mes_anterior_str / _buscar_match_servicio / _estado_servicios_mes vivían acá,
# pero blueprints/finanzas.py también las necesita (para /api/finanzas/servicios)
# y no las podía importar de app.py sin generar un import circular (app.py ya
# importa finanzas_bp DESDE blueprints.finanzas). Bug encontrado en revisión:
# blueprints/finanzas.py las llamaba sin importarlas de ningún lado — tiraban
# NameError apenas se pedía GET /api/finanzas/servicios. Se movieron las tres
# a blueprints/finanzas.py (que es además donde pertenecen conceptualmente:
# son lógica de negocio de finanzas, no infraestructura genérica) y acá se
# importan en la dirección que ya existía (app.py -> blueprints.finanzas).
from blueprints.finanzas import _estado_servicios_mes

def _chequear_recordatorio_servicios():
    while True:
        try:
            hoy = datetime.now()
            if hoy.day == 2:
                mes = hoy.strftime("%Y-%m")
                with get_db(HIST_DB) as con:
                    row = con.execute("SELECT valor FROM fin_servicios_estado WHERE clave='ultimo_aviso_mes'").fetchone()
                    if not (row and row[0] == mes):
                        estado = _estado_servicios_mes(con, mes)
                        pendientes = [s["nombre"] for s in estado if not s["pagado"]]
                        if pendientes:
                            lista = "\n".join(f"• {n}" for n in pendientes)
                            notificar_telegram(f"💸 Sin pago detectado este mes ({mes}):\n\n{lista}")
                        con.execute(
                            "INSERT INTO fin_servicios_estado (clave, valor) VALUES ('ultimo_aviso_mes', ?) "
                            "ON CONFLICT(clave) DO UPDATE SET valor=excluded.valor", (mes,))
        except Exception:
            logging.exception("Error en chequeo de recordatorio de servicios")
        _threading.Event().wait(3600)  # revisa cada hora

_threading.Thread(target=_chequear_recordatorio_servicios, daemon=True).start()

# ── Backup semanal ───────────────────────────────────────────────────────────
BACKUP_DIR = "/data/backups"
os.makedirs(BACKUP_DIR, exist_ok=True)
BACKUP_FUENTES = {"pad": lambda: DB_PATH, "historial": lambda: HIST_DB}

def _ejecutar_backup(origen="manual"):
    import shutil, json as _json
    resultados = {}
    for clave, get_path in BACKUP_FUENTES.items():
        src = get_path()
        dest = os.path.join(BACKUP_DIR, f"{clave}_backup.db")
        try:
            if not os.path.exists(src):
                raise FileNotFoundError("el archivo origen no existe")
            shutil.copy2(src, dest)
            resultados[clave] = {"ok": True, "size_mb": round(os.path.getsize(dest) / (1024 * 1024), 2)}
        except Exception as e:
            resultados[clave] = {"ok": False, "error": str(e)}
    with get_db(HIST_DB) as con:
        con.execute("INSERT INTO backups (origen, resultados) VALUES (?,?)", (origen, _json.dumps(resultados)))
    todo_ok = all(r["ok"] for r in resultados.values())
    logging.info(f"BACKUP {origen} | ok={todo_ok} | {resultados}")
    if todo_ok:
        notificar_telegram(f"💾 Backup {origen} OK — " + ", ".join(f"{k}: {v['size_mb']} MB" for k, v in resultados.items()))
    else:
        errores = ", ".join(f"{k}: {v.get('error')}" for k, v in resultados.items() if not v["ok"])
        notificar_telegram(f"⚠️ Backup {origen} con errores — {errores}")
    return resultados, todo_ok

def _proximo_backup_domingo():
    hoy = datetime.now()
    dias_hasta_domingo = (6 - hoy.weekday()) % 7
    candidato = (hoy + timedelta(days=dias_hasta_domingo)).replace(hour=2, minute=0, second=0, microsecond=0)
    if candidato <= hoy:
        candidato += timedelta(days=7)
    return candidato.strftime("%Y-%m-%d %H:%M")

def _chequear_backup_automatico():
    while True:
        try:
            hoy = datetime.now()
            if hoy.weekday() == 6 and hoy.hour == 2:  # domingo 02:xx
                semana = hoy.strftime("%Y-W%W")
                with get_db(HIST_DB) as con:
                    row = con.execute(
                        "SELECT fecha FROM backups WHERE origen='auto' ORDER BY fecha DESC LIMIT 1").fetchone()
                ya_corrio_esta_semana = row and datetime.strptime(row[0][:19], "%Y-%m-%d %H:%M:%S").strftime("%Y-W%W") == semana
                if not ya_corrio_esta_semana:
                    _ejecutar_backup(origen="auto")
        except Exception:
            logging.exception("Error en backup automático semanal")
        _threading.Event().wait(1800)  # revisa cada 30 min

_threading.Thread(target=_chequear_backup_automatico, daemon=True).start()

# ── Repositorio de documentos por módulo (contexto para la IA) ─────────────────
# get_api_key, contexto_repositorio, _extraer_texto_docx,
# _extraer_cronologia_de_texto, MODULOS_REPOSITORIO, MODULOS_CON_CRONOLOGIA,
# _normalizar_fecha_a_ddmmaaaa y _validar_fecha_ddmmaaaa ahora viven en
# core.py (los usan también los blueprints, ej. vua).
from core import (
    get_api_key, contexto_repositorio, _extraer_texto_docx,
    _extraer_cronologia_de_texto, MODULOS_REPOSITORIO, MODULOS_CON_CRONOLOGIA,
    _normalizar_fecha_a_ddmmaaaa, _validar_fecha_ddmmaaaa,
    job_status, job_create, job_get, _job_persist, _JobLog,
)

def _agregar_entradas_cronologia(con, tabla, entradas, fuente):
    max_orden = con.execute(f"SELECT MAX(orden) FROM {tabla}").fetchone()[0] or 0
    agregadas = 0
    for e in entradas:
        actividad = (e.get("actividad") or "").strip()
        if not actividad:
            continue
        max_orden += 1
        fecha = _normalizar_fecha_a_ddmmaaaa(e.get("fecha", ""))
        con.execute(
            f"INSERT INTO {tabla} (fecha, actividad, participantes, estado, orden, creado, modificado) "
            f"VALUES (?,?,?,?,?,datetime('now'),datetime('now'))",
            (fecha, f"{actividad} (fuente: {fuente})",
             e.get("participantes", ""), "Pendiente", max_orden))
        agregadas += 1
    return agregadas


@app.route("/api/repositorio/<modulo>", methods=["GET"])
@login_required
def repositorio_list(modulo):
    if modulo not in MODULOS_REPOSITORIO or modulo not in session.get("modulos", []):
        return jsonify({"ok": False, "error": "Módulo no habilitado"}), 403
    with get_db(HIST_DB, row_factory=True) as con:
        rows = [dict(r) for r in con.execute(
            "SELECT id, nombre_archivo, subido_por, creado, length(contenido) as tamano "
            "FROM doc_repositorio WHERE modulo=? ORDER BY creado DESC", (modulo,)).fetchall()]
    return jsonify({"ok": True, "rows": rows})

@app.route("/api/repositorio/<modulo>", methods=["POST"])
@login_required
def repositorio_upload(modulo):
    if modulo not in MODULOS_REPOSITORIO or modulo not in session.get("modulos", []):
        return jsonify({"ok": False, "error": "Módulo no habilitado"}), 403
    if "archivo" not in request.files:
        return jsonify({"ok": False, "error": "No se recibió archivo"})
    f = request.files["archivo"]
    if not f.filename.lower().endswith(".docx"):
        return jsonify({"ok": False, "error": "Por ahora solo se aceptan archivos Word (.docx)"})
    data_bytes = f.read()
    try:
        texto = _extraer_texto_docx(io.BytesIO(data_bytes))
    except Exception as e:
        return jsonify({"ok": False, "error": f"No se pudo leer el documento: {e}"})
    if not texto.strip():
        return jsonify({"ok": False, "error": "El documento no tiene texto extraíble"})
    repo_dir = os.path.join("/data/repositorio", modulo)
    os.makedirs(repo_dir, exist_ok=True)
    # secure_filename() saca separadores de ruta, "..", y caracteres raros del
    # nombre original (antes se usaba f.filename crudo en el path — con un
    # nombre tipo "x/../../../etc/algo" se podía escribir fuera de repo_dir).
    # El nombre ORIGINAL (sin sanear) se sigue guardando en la BD para mostrarlo
    # al usuario tal cual lo subió; ruta_archivo es lo único que toca el disco.
    nombre_seguro = secure_filename(f.filename) or "documento.docx"
    ruta_archivo = os.path.join(repo_dir, f"{uuid.uuid4().hex[:12]}_{nombre_seguro}")
    with open(ruta_archivo, "wb") as out:
        out.write(data_bytes)
    with get_db(HIST_DB) as con:
        con.execute("INSERT INTO doc_repositorio (modulo, nombre_archivo, contenido, ruta_archivo, subido_por) VALUES (?,?,?,?,?)",
            (modulo, f.filename, texto, ruta_archivo, session.get("username", "?")))
        agregadas_cronologia = 0
        error_cronologia = None
        if modulo in MODULOS_CON_CRONOLOGIA:
            try:
                entradas = _extraer_cronologia_de_texto(texto, modulo)
                agregadas_cronologia = _agregar_entradas_cronologia(
                    con, MODULOS_CON_CRONOLOGIA[modulo], entradas, f.filename)
            except Exception as e:
                logging.exception(f"Error extrayendo cronología de '{f.filename}' ({modulo})")
                error_cronologia = str(e)
    logging.info(f"REPOSITORIO UPLOAD | modulo={modulo} | user={session.get('username')} | archivo={f.filename} | cronologia+={agregadas_cronologia}")
    msg = f"📁 Documento '{f.filename}' agregado al repositorio de {modulo} por {session.get('username')}"
    if agregadas_cronologia:
        msg += f"\n🗓️ Se sumaron {agregadas_cronologia} hito(s) a la cronología."
    notificar_telegram(msg)
    return jsonify({"ok": True, "agregadas_cronologia": agregadas_cronologia, "error_cronologia": error_cronologia})

@app.route("/api/repositorio/<modulo>/<int:doc_id>", methods=["DELETE"])
@login_required
def repositorio_delete(modulo, doc_id):
    if modulo not in MODULOS_REPOSITORIO or modulo not in session.get("modulos", []):
        return jsonify({"ok": False, "error": "Módulo no habilitado"}), 403
    with get_db(HIST_DB) as con:
        row = con.execute("SELECT ruta_archivo FROM doc_repositorio WHERE id=? AND modulo=?", (doc_id, modulo)).fetchone()
        con.execute("DELETE FROM doc_repositorio WHERE id=? AND modulo=?", (doc_id, modulo))
    if row and row[0] and os.path.exists(row[0]):
        try: os.remove(row[0])
        except Exception: pass
    return jsonify({"ok": True})

@app.route("/api/repositorio/<modulo>/<int:doc_id>/download")
@login_required
def repositorio_download(modulo, doc_id):
    if modulo not in MODULOS_REPOSITORIO or modulo not in session.get("modulos", []):
        return jsonify({"ok": False, "error": "Módulo no habilitado"}), 403
    with get_db(HIST_DB) as con:
        row = con.execute("SELECT nombre_archivo, ruta_archivo FROM doc_repositorio WHERE id=? AND modulo=?",
                           (doc_id, modulo)).fetchone()
    if not row or not row[1] or not os.path.exists(row[1]):
        return jsonify({"ok": False, "error": "El archivo original ya no está disponible"}), 404
    return send_file(row[1], as_attachment=True, download_name=row[0])

def _limpiar_jobs_viejos():
    """Elimina jobs de más de 2 horas para evitar memory leak (en memoria y en SQLite)."""
    while True:
        time.sleep(3600)  # cada hora
        ahora = time.time()
        viejos = [k for k, v in list(job_status.items())
                  if v.get('_ts', ahora) < ahora - 7200]
        for k in viejos:
            job_status.pop(k, None)
        try:
            with get_db(HIST_DB) as con:
                con.execute("DELETE FROM job_status_db WHERE ts < ?", (ahora - 7200,))
        except Exception:
            logging.exception("No se pudo limpiar job_status_db")
_threading.Thread(target=_limpiar_jobs_viejos, daemon=True).start()

# ── Limpieza de archivos huérfanos en disco ─────────────────────────────────
# Informes (Word/Excel), minutas y reportes de stock se guardan en disco con
# la ruta completa registrada en una columna de la base (ver ruta_archivo,
# archivo_word/excel, archivo, file_path más abajo). Si esa fila se borra
# (admin borra un item del historial, un usuario borra una minuta, un reporte
# de stock se regenera y pisa el registro con INSERT OR REPLACE) el archivo
# viejo se queda en disco para siempre — nada lo borraba hasta ahora.
_LIMPIEZA_ARCHIVOS = [
    # (directorio, [(tabla, columna), ...], recorrer subcarpetas)
    ("/data/repositorio",  [("doc_repositorio", "ruta_archivo")],                          True),
    (OUTPUT_FOLDER,        [("historial", "archivo_word"), ("historial", "archivo_excel")], False),
    (STOCK_REPORTS_DIR,    [("stock_reportes", "file_path")],                              False),
    ("/data/minutas",      [("vua_minutas", "archivo")],                                   False),
    ("/data/minutas_senasa", [("senasa_minutas", "archivo")],                              False),
]

def _limpiar_archivos_huerfanos():
    """Corre una vez por día. Por seguridad:
    - Solo borra archivos con más de 48hs de antigüedad (mtime) — evita pisar
      un archivo recién escrito cuyo INSERT a la base todavía no commiteó
      (hay una ventana entre "escribir el archivo" y "guardar la fila").
    - Nunca borra ni toca filas de la base: si una fila referencia un archivo
      que ya no está en disco, solo se loguea (puede ser intencional — por
      ejemplo alguien lo borró a mano en el servidor)."""
    while True:
        time.sleep(86400)  # una vez por día
        limite_mtime = time.time() - 48 * 3600
        for directorio, tabla_col, recursivo in _LIMPIEZA_ARCHIVOS:
            if not os.path.isdir(directorio):
                continue
            try:
                referenciados = set()
                with get_db(HIST_DB) as con:
                    for tabla, columna in tabla_col:
                        q = f"SELECT {columna} FROM {tabla} WHERE {columna} IS NOT NULL AND {columna} != ''"
                        for (ruta,) in con.execute(q):
                            referenciados.add(os.path.normpath(ruta))

                if recursivo:
                    encontrados = [os.path.join(r, n) for r, _d, fs in os.walk(directorio) for n in fs]
                else:
                    encontrados = [os.path.join(directorio, n) for n in os.listdir(directorio)]

                borrados = 0
                for ruta_disco in encontrados:
                    ruta_disco = os.path.normpath(ruta_disco)
                    if not os.path.isfile(ruta_disco) or ruta_disco in referenciados:
                        continue
                    try:
                        if os.path.getmtime(ruta_disco) < limite_mtime:
                            os.remove(ruta_disco)
                            borrados += 1
                    except FileNotFoundError:
                        pass
                if borrados:
                    logging.info(f"LIMPIEZA ARCHIVOS | {directorio} | {borrados} archivo(s) huérfano(s) eliminados")
            except Exception:
                logging.exception(f"Error limpiando archivos huérfanos en {directorio}")

_threading.Thread(target=_limpiar_archivos_huerfanos, daemon=True).start()

# ── Historial DB ───────────────────────────────────────────────────────────────
# ── Seeds separados ────────────────────────────────────────────────────────────
def _seed_ref_dira(con):
    """Carga inicial de ref_dira a partir de las 8 direcciones regionales vigentes."""
    datos = [
        ('1', 'HIDROVIA', 1),
        ('2', 'NORESTE', 2),
        ('3', 'NOROESTE', 3),
        ('4', 'CENTRAL', 4),
        ('5', 'RIO COLORADO', 5),
        ('6', 'AUSTRAL', 6),
        ('7', 'CUYO', 7),
        ('8', 'No aplica', 8),
    ]
    con.executemany("INSERT OR REPLACE INTO ref_dira (indice, nombre, orden) VALUES (?,?,?)", datos)
    con.commit()

def _seed_ref_aduanas(con):
    """Carga inicial de ref_aduanas a partir del CSV de referencia oficial (Norte/Noroeste corregidos)."""
    datos = [
        ('093','RAFAELA','1'), ('062','SANTA FE','1'), ('057','SAN LORENZO','1'),
        ('069','VILLA CONSTITUCION','1'), ('052','ROSARIO','1'), ('016','CONCORDIA','1'),
        ('041','PARANA','1'), ('020','DIAMANTE','1'), ('013','COLON','1'),
        ('015','CONCEPCION DEL URUGUAY','1'), ('026','GUALEGUAYCHU','1'), ('059','SAN NICOLAS','1'),
        ('060','SAN PEDRO','1'), ('094','VENADO TUERTO','1'),
        ('031','JUJUY','3'), ('034','LA QUIACA','3'), ('045','POCITOS','3'),
        ('053','SALTA','3'), ('066','TINOGASTA','3'), ('074','TUCUMAN','3'), ('076','ORAN','3'),
        ('010','BARRANQUERAS','2'), ('012','CLORINDA','2'), ('018','CORRIENTES','2'),
        ('024','FORMOSA','2'), ('025','GOYA','2'), ('029','IGUAZU','2'),
        ('042','PASO DE LOS LIBRES','2'), ('046','POSADAS','2'), ('054','SAN JAVIER','2'),
        ('082','BERNARDO DE IRIGOYEN','2'), ('084','SANTO TOME','2'), ('086','OBERA','2'),
        ('079','LA RIOJA','4'), ('088','GENERAL DEHEZA','4'), ('017','CORDOBA','4'),
        ('089','SANTIAGO DEL ESTERO','4'), ('090','GENERAL PICO','4'),
        ('008','CAMPANA','8'), ('073','EZEIZA','8'), ('033','LA PLATA','8'), ('001','BUENOS AIRES','8'),
        ('003','BAHIA BLANCA','5'), ('004','SAN CARLOS DE BARILOCHE','5'), ('037','MAR DEL PLATA','5'),
        ('040','NECOCHEA','5'), ('058','SAN MARTIN DE LOS ANDES','5'), ('075','NEUQUEN','5'),
        ('080','SAN ANTONIO OESTE','5'), ('085','VILLA REGINA','5'),
        ('014','COMODORO RIVADAVIA','6'), ('019','PUERTO DESEADO','6'), ('023','ESQUEL','6'),
        ('047','PUERTO MADRYN','6'), ('048','RIO GALLEGOS','6'), ('049','RIO GRANDE','6'),
        ('061','SANTA CRUZ','6'), ('067','USHUAIA','6'), ('087','CALETA OLIVIA','6'),
        ('055','SAN JUAN','7'), ('038','MENDOZA','7'), ('083','SAN LUIS','7'), ('078','SAN RAFAEL','7'),
    ]
    con.executemany("INSERT OR REPLACE INTO ref_aduanas (cod, nombre, indice_dira) VALUES (?,?,?)", datos)
    con.commit()


    """Datos iniciales de cronología VUA. Separado de init_historial para mayor claridad."""
    cronologia = [
        ("08/10/2025","Designación del referente de la DGA ante VUCEA para el proyecto VUA","DI REPA","Completado",1),
        ("14/10/2025","Primera reunión: VUCEA expone el nuevo alcance del proyecto VUA Carga","DI REPA, VUCEA","Completado",2),
        ("04/12/2025","Reunión de relevamiento del proceso de carga aérea","DI REPA, VUCEA","Completado",3),
        ("18/12/2025","Reunión de análisis de la RG 5797/2025 y su implicancia en VUA","DI REPA, VUCEA","Completado",4),
        ("11/02/2026","Demo sistema VUA: primera versión del formulario Plan de Vuelo","DI REPA, VUCEA, SENASA, ORSNA, DI ADEZ, Aerolíneas Argentinas, PSA","Completado",5),
        ("25/02/2026","Reunión presencial en Ezeiza: relevamiento operativo de circuitos de carga","DI REPA, DI ADEZ, VUCEA","Completado",6),
        ("17/03/2026","Reunión presencial con TCA: relevamiento del circuito de desconsolidación","DI REPA, DI ADEZ, VUCEA, TCA","Completado",7),
        ("25/03/2026","Reunión sobre Manifiesto Desconsolidado de Importación","DI REPA, DI ADEZ, VUCEA, TCA","Completado",8),
        ("11/05/2026","Reunión ARCA-SENASA: integración de organismos de control en PAD","DI REPA, SENASA","Completado",9),
        ("14/05/2026","Mesa de trabajo: IA, tablero de vuelos, MANE y marco normativo","DI REPA, DI SADU, DI ADEZ, VUCEA","Completado",10),
        ("21/05/2026","Reunión ampliada con aerolíneas: IA, circuito importación, desconsolidado y DAI","DGA, DI REPA, DI SADU, DI ADEZ, VUCEA, Aerolíneas Argentinas, JURCA, IATA","Completado",11),
        ("A definir","Reunión específica análisis del MANE","DI REPA, DI ADEZ, VUCEA","Pendiente",12),
        ("A definir","Análisis normativo rol de VUCEA como intermediario","DI REPA, áreas legales DGA y VUCEA","Pendiente",13),
        ("A definir","Corrección de formularios Guía Madre por parte de VUCEA","VUCEA","Pendiente",14),
    ]
    con.executemany(
        "INSERT INTO vua_cronologia (fecha,actividad,participantes,estado,orden,creado,modificado) "
        "VALUES (?,?,?,?,?,datetime('now'),datetime('now'))",
        cronologia)

def _seed_feriados(con):
    """Carga inicial de feriados migrada desde el dict DIAS_NH hardcodeado en
    stock_depositos.py (2021-2026). Sin descripción — se puede completar
    después desde el panel de administración."""
    datos = [
        ("2021-08-16", ""),
        ("2021-10-08", ""),
        ("2021-10-11", ""),
        ("2021-11-20", ""),
        ("2021-11-22", ""),
        ("2021-12-08", ""),
        ("2021-12-25", ""),
        ("2022-01-01", ""),
        ("2022-02-28", ""),
        ("2022-03-01", ""),
        ("2022-03-24", ""),
        ("2022-04-02", ""),
        ("2022-04-14", ""),
        ("2022-04-15", ""),
        ("2022-05-01", ""),
        ("2022-05-18", ""),
        ("2022-05-25", ""),
        ("2022-06-17", ""),
        ("2022-06-20", ""),
        ("2022-08-15", ""),
        ("2022-09-02", ""),
        ("2022-10-07", ""),
        ("2022-10-10", ""),
        ("2022-11-21", ""),
        ("2022-12-08", ""),
        ("2022-12-09", ""),
        ("2022-12-20", ""),
        ("2023-02-20", ""),
        ("2023-02-21", ""),
        ("2023-03-24", ""),
        ("2023-04-06", ""),
        ("2023-04-07", ""),
        ("2023-05-01", ""),
        ("2023-05-25", ""),
        ("2023-05-26", ""),
        ("2023-06-19", ""),
        ("2023-06-20", ""),
        ("2023-08-21", ""),
        ("2023-10-13", ""),
        ("2023-10-16", ""),
        ("2023-11-20", ""),
        ("2023-12-08", ""),
        ("2023-12-25", ""),
        ("2024-01-01", ""),
        ("2024-02-12", ""),
        ("2024-02-13", ""),
        ("2024-03-28", ""),
        ("2024-03-29", ""),
        ("2024-04-01", ""),
        ("2024-04-02", ""),
        ("2024-06-17", ""),
        ("2024-06-20", ""),
        ("2024-06-21", ""),
        ("2024-07-09", ""),
        ("2024-08-17", ""),
        ("2024-10-11", ""),
        ("2024-11-18", ""),
        ("2024-12-25", ""),
        ("2025-01-01", ""),
        ("2025-03-03", ""),
        ("2025-03-04", ""),
        ("2025-03-24", ""),
        ("2025-04-02", ""),
        ("2025-04-17", ""),
        ("2025-04-18", ""),
        ("2025-05-01", ""),
        ("2025-05-02", ""),
        ("2025-06-16", ""),
        ("2025-06-20", ""),
        ("2025-07-09", ""),
        ("2025-08-15", ""),
        ("2025-10-10", ""),
        ("2025-12-08", ""),
        ("2025-12-25", ""),
        ("2026-01-01", ""),
        ("2026-02-16", ""),
        ("2026-02-17", ""),
        ("2026-03-24", ""),
        ("2026-04-02", ""),
        ("2026-04-03", ""),
        ("2026-05-01", ""),
        ("2026-05-25", ""),
        ("2026-06-15", ""),
        ("2026-06-20", ""),
        ("2026-07-09", ""),
        ("2026-08-17", ""),
        ("2026-10-12", ""),
        ("2026-11-23", ""),
        ("2026-12-08", ""),
        ("2026-12-25", ""),
    ]
    con.executemany("INSERT OR IGNORE INTO feriados (fecha, descripcion) VALUES (?,?)", datos)

def init_historial():
    with get_db(HIST_DB) as con:
        con.execute("""CREATE TABLE IF NOT EXISTS historial (
            id TEXT PRIMARY KEY, fecha TEXT, usuario TEXT, pais TEXT,
            anio TEXT, mes_d TEXT, mes_h TEXT, uso_ia INTEGER,
            archivo_word TEXT, archivo_excel TEXT, revisado INTEGER DEFAULT 0,
            tipo TEXT DEFAULT 'sintia', descripcion TEXT DEFAULT ''
        )""")
        con.execute("""CREATE TABLE IF NOT EXISTS vua_cronologia (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            fecha TEXT, actividad TEXT, participantes TEXT,
            estado TEXT DEFAULT 'Pendiente', orden INTEGER DEFAULT 0,
            creado TEXT, modificado TEXT
        )""")
        con.execute("""CREATE TABLE IF NOT EXISTS vua_minutas (
            id TEXT PRIMARY KEY, fecha TEXT, asunto TEXT, lugar TEXT,
            participantes TEXT, temas TEXT, acuerdos TEXT, proximos TEXT,
            archivo TEXT, creado_por TEXT, creado TEXT
        )""")
        con.execute("""CREATE TABLE IF NOT EXISTS vua_ejes (
            id TEXT PRIMARY KEY, nombre TEXT, estado TEXT, orden INTEGER DEFAULT 0,
            descripcion TEXT DEFAULT '', propuesta_vucea TEXT DEFAULT '',
            postura_aduana TEXT DEFAULT '', recomendacion TEXT DEFAULT ''
        )""")
        # Mismo caso que 'sesiones' más arriba: estas 4 columnas ya las usaba
        # el código (vua_ejes_create/update, vua.html) pero el CREATE TABLE
        # nunca las tuvo — en una instalación nueva, crear un eje tiraba
        # "no such column: descripcion". ALTER TABLE de resguardo por si la
        # base ya existía sin ellas.
        for _col in ("descripcion", "propuesta_vucea", "postura_aduana", "recomendacion"):
            try:
                con.execute(f"ALTER TABLE vua_ejes ADD COLUMN {_col} TEXT DEFAULT ''")
            except Exception:
                pass  # ya existe
        con.execute("""CREATE TABLE IF NOT EXISTS ref_aduanas (
            cod TEXT PRIMARY KEY, nombre TEXT NOT NULL, indice_dira TEXT NOT NULL
        )""")
        con.execute("""CREATE TABLE IF NOT EXISTS ref_aduanas_log (
            id INTEGER PRIMARY KEY AUTOINCREMENT, fecha TEXT, usuario TEXT,
            accion TEXT, detalle TEXT
        )""")
        con.execute("""CREATE TABLE IF NOT EXISTS ref_dira (
            indice TEXT PRIMARY KEY, nombre TEXT NOT NULL, orden INTEGER DEFAULT 0
        )""")
        con.execute("""CREATE TABLE IF NOT EXISTS feriados (
            fecha TEXT PRIMARY KEY, descripcion TEXT DEFAULT ''
        )""")
        cur_fer = con.cursor()
        cur_fer.execute("SELECT COUNT(*) FROM feriados")
        if cur_fer.fetchone()[0] == 0:
            _seed_feriados(con)
        cur_rd = con.cursor()
        cur_rd.execute("SELECT COUNT(*) FROM ref_dira")
        if cur_rd.fetchone()[0] == 0:
            _seed_ref_dira(con)
        cur_ra = con.cursor()
        cur_ra.execute("SELECT COUNT(*) FROM ref_aduanas")
        if cur_ra.fetchone()[0] == 0:
            _seed_ref_aduanas(con)
        cur = con.cursor()
        cur.execute("SELECT COUNT(*) FROM vua_cronologia")
        if cur.fetchone()[0] == 0:
            # Bug encontrado en revisión: acá se llamaba a _seed_cronologia_vua(con),
            # una función que nunca llegó a existir en el código (a diferencia de
            # sus hermanas _seed_feriados/_seed_ref_dira/_seed_ref_aduanas, que sí
            # están definidas más arriba). En una instalación EXISTENTE nunca se
            # notó porque vua_cronologia ya tiene filas y este bloque se salteaba;
            # pero en una instalación NUEVA (base vacía) esto tiraba NameError acá
            # mismo, dentro de init_historial(), rompiendo el arranque de toda la
            # app. No se inventan acá datos históricos falsos de la cronología VUA
            # real (son hechos fechados específicos del proyecto, no algo que se
            # pueda adivinar) — se deja vacía y se carga desde el módulo VUA.
            logging.info("vua_cronologia está vacía — se deja sin sembrar (cargar desde el módulo VUA).")
        cur.execute("SELECT COUNT(*) FROM vua_ejes")
        if cur.fetchone()[0] == 0:
            ejes = [
                ("4.1","Transmisión de información anticipada — XML, sujetos obligados y marco sancionatorio","En análisis — requiere definición normativa",1),
                ("4.2","Tablero de programación de vuelos","En análisis técnico interno",2),
                ("4.3","Manifiesto de Exportación (MANE)","Pendiente — sin normativa vigente para IA de exportación",3),
                ("4.4","Manifiestos desconsolidados de importación","Pendiente — sin normativa vigente",4),
                ("4.5","Estándar de transmisión XML — Guía Madre (XFWB)","Postura definida — observaciones comunicadas a VUCEA",5),
            ]
            con.executemany("INSERT INTO vua_ejes (id, nombre, estado, orden) VALUES (?,?,?,?)", ejes)
        # ── Tablas de autenticación y sistema ────────────────────────────────────────
        con.execute("""CREATE TABLE IF NOT EXISTS usuarios (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            rol TEXT DEFAULT 'viewer',
            modulos TEXT DEFAULT 'sintia,vua,senasa',
            activo INTEGER DEFAULT 1,
            ultimo_acceso TEXT
        )""")
        con.execute("""CREATE TABLE IF NOT EXISTS sesiones (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            token TEXT UNIQUE NOT NULL,
            username TEXT NOT NULL,
            ip TEXT DEFAULT '',
            user_agent TEXT DEFAULT '',
            creado TEXT DEFAULT (datetime('now')),
            ultimo_acceso TEXT DEFAULT (datetime('now')),
            activo INTEGER DEFAULT 1
        )""")
        # Columnas agregadas después de que esta tabla ya existía en producción
        # (por eso CREATE TABLE IF NOT EXISTS no alcanza para instalaciones viejas) —
        # bug encontrado en revisión: core.py.registrar_sesion() ya insertaba en
        # ip/user_agent, y admin_sesiones_list() ya las leía, pero en una
        # instalación NUEVA (base vacía) la tabla se creaba sin esas columnas.
        # registrar_sesion() tiene un try/except a lo Pokémon, así que la falla
        # quedaba silenciada — el login "funcionaba" pero ninguna sesión se
        # registraba realmente, y /api/admin/sesiones directamente crasheaba.
        for _col, _tipo in [("ip", "TEXT DEFAULT ''"), ("user_agent", "TEXT DEFAULT ''")]:
            try:
                con.execute(f"ALTER TABLE sesiones ADD COLUMN {_col} {_tipo}")
            except Exception:
                pass  # ya existe
        con.execute("""CREATE TABLE IF NOT EXISTS tokens_revocados (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            token TEXT UNIQUE NOT NULL,
            revocado TEXT DEFAULT (datetime('now'))
        )""")
        con.execute("""CREATE TABLE IF NOT EXISTS prompts (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nombre TEXT NOT NULL,
            descripcion TEXT DEFAULT '',
            modulo TEXT DEFAULT 'general',
            contenido TEXT DEFAULT '',
            modificado TEXT DEFAULT (datetime('now'))
        )""")
        # Seed: usuario admin por defecto si tabla vacía
        cur_u = con.cursor()
        cur_u.execute("SELECT COUNT(*) FROM usuarios")
        if cur_u.fetchone()[0] == 0:
            _temp_pass = secrets.token_urlsafe(12)
            default_hash = bcrypt.hashpw(_temp_pass.encode(), bcrypt.gensalt()).decode()
            con.execute("INSERT INTO usuarios (username, password_hash, rol, modulos) VALUES (?,?,?,?)",
                ("admin", default_hash, "admin", "sintia,vua,senasa"))
            # print() a stdout, no logging.warning(): un password (aunque sea
            # temporal y de un solo uso) no debería terminar en un archivo de
            # log ni en un sistema de logging centralizado si en algún momento
            # se agrega uno — stdout del arranque es más efímero y solo lo ve
            # quien está mirando la consola en ese momento. El log sí queda,
            # pero sin el valor del password.
            print(f"\n{'='*70}\nSEED ADMIN — usuario 'admin' creado con password temporal:\n"
                  f"  {_temp_pass}\n"
                  f"Cambiala inmediatamente después de loguear (Perfil > Cambiar contraseña).\n{'='*70}\n")
            logging.warning("SEED ADMIN | tabla 'usuarios' vacía — se creó 'admin' con password temporal "
                             "(ver stdout del arranque; no se guarda en el log).")

        # ── Caché persistente de opciones de filtro de Consulta DAT ─────────────────
        # EST_MIC/ULT_ESTADO/VAR_CONTROL: se recalculan una sola vez, justo
        # cuando termina de importarse un CSV nuevo de DAT_<año> (ver
        # _recalcular_opciones_dat, enganchado al final de _procesar_csv) --
        # no en cada apertura del panel (SELECT DISTINCT en vivo era lento) ni
        # en un cron desconectado del evento real de actualización de datos.
        con.execute("""CREATE TABLE IF NOT EXISTS sintia_dat_opciones (
            tabla TEXT NOT NULL,
            campo TEXT NOT NULL,
            valores TEXT NOT NULL,
            actualizado TEXT DEFAULT (datetime('now')),
            PRIMARY KEY (tabla, campo)
        )""")

        # ── Tabla compartida de integrantes (VUA, SENASA, SINTIA) ──────────────────
        con.execute("""CREATE TABLE IF NOT EXISTS integrantes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nombre TEXT NOT NULL,
            cargo TEXT DEFAULT '',
            organismo TEXT DEFAULT '',
            email TEXT DEFAULT '',
            activo INTEGER DEFAULT 1,
            orden INTEGER DEFAULT 0,
            creado TEXT DEFAULT (datetime('now'))
        )""")
        # Migrar ROLES_PREDEFINIDOS a la tabla si está vacía
        cur2 = con.cursor()
        cur2.execute("SELECT COUNT(*) FROM integrantes")
        if cur2.fetchone()[0] == 0:
            roles_seed = [
                ("Diego Bugallo",      "Jefe Dpto. Facilitación y Simplificación de Comercio", "DI REPA",  "", 1),
                ("Martín Macías",      "Jefe Div. Modernización de Procesos Aduaneros",        "DI REPA",  "", 2),
                ("Federico Cáceres",   "Sec. Simplificación de Procesos Operativos",           "DI REPA",  "", 3),
                ("Hernán Cascón",      "Supervisor de Informática Aduanera",                   "DI SADU",  "", 4),
                ("Maximiliano Luengo", "Consejero técnico",                                    "DI ADEZ",  "", 5),
                ("Pablo Gómez Valdez", "Consejero técnico",                                    "DI ADEZ",  "", 6),
                ("Fabiola Cochello",   "Directora",                                            "VUCEA",    "", 7),
                ("Vanesa Franco",      "Jefa de Procesos",                                     "VUCEA",    "", 8),
            ]
            con.executemany(
                "INSERT INTO integrantes (nombre, cargo, organismo, email, orden) VALUES (?,?,?,?,?)",
                roles_seed)

        # Tablas VUA adicionales (pueden no existir en instalaciones anteriores)
        con.execute("""CREATE TABLE IF NOT EXISTS vua_config (
            clave TEXT PRIMARY KEY, titulo TEXT, contenido TEXT DEFAULT '',
            modificado TEXT DEFAULT (datetime('now'))
        )""")
        con.execute("""CREATE TABLE IF NOT EXISTS vua_equipo (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nombre TEXT, cargo TEXT, organismo TEXT, email TEXT,
            activo INTEGER DEFAULT 1, orden INTEGER DEFAULT 0
        )""")
        con.execute("""CREATE TABLE IF NOT EXISTS vua_glosario (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            termino TEXT, definicion TEXT, categoria TEXT DEFAULT 'general',
            orden INTEGER DEFAULT 0
        )""")
        con.execute("""CREATE TABLE IF NOT EXISTS vua_riesgos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            codigo TEXT, titulo TEXT, descripcion TEXT, mitigacion TEXT,
            probabilidad TEXT DEFAULT 'Media', impacto TEXT DEFAULT 'Alto',
            activo INTEGER DEFAULT 1, orden INTEGER DEFAULT 0
        )""")
        con.execute("""CREATE TABLE IF NOT EXISTS vua_info (
            clave TEXT PRIMARY KEY, contenido TEXT DEFAULT '', modificado TEXT
        )""")
        con.execute("""CREATE TABLE IF NOT EXISTS vua_correos_rapidos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            etiqueta TEXT, instruccion TEXT, activo INTEGER DEFAULT 1, orden INTEGER DEFAULT 0
        )""")
        con.execute("""CREATE TABLE IF NOT EXISTS vua_consultas_frecuentes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            pregunta TEXT, respuesta TEXT, activo INTEGER DEFAULT 1, orden INTEGER DEFAULT 0
        )""")
        con.execute("""CREATE TABLE IF NOT EXISTS doc_repositorio (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            modulo TEXT NOT NULL, nombre_archivo TEXT, contenido TEXT, ruta_archivo TEXT,
            subido_por TEXT, creado TEXT DEFAULT (datetime('now'))
        )""")
        try:
            con.execute("ALTER TABLE doc_repositorio ADD COLUMN ruta_archivo TEXT")
        except sqlite3.OperationalError:
            pass  # ya existe (instalación previa a este cambio)
        con.execute("""CREATE TABLE IF NOT EXISTS fin_servicios (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nombre TEXT NOT NULL, patron TEXT, activo INTEGER DEFAULT 1, orden INTEGER DEFAULT 0
        )""")
        try:
            con.execute("ALTER TABLE fin_servicios ADD COLUMN patron TEXT")
        except sqlite3.OperationalError:
            pass  # ya existe (instalación previa a este cambio)
        con.execute("""CREATE TABLE IF NOT EXISTS fin_servicios_pagos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            servicio_id INTEGER NOT NULL, mes TEXT NOT NULL,
            pagado INTEGER DEFAULT 0, fecha_pago TEXT,
            UNIQUE(servicio_id, mes)
        )""")
        con.execute("""CREATE TABLE IF NOT EXISTS fin_servicios_estado (
            clave TEXT PRIMARY KEY, valor TEXT
        )""")
        con.execute("""CREATE TABLE IF NOT EXISTS backups (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            fecha TEXT DEFAULT (datetime('now')), origen TEXT, resultados TEXT
        )""")
        con.execute("""CREATE TABLE IF NOT EXISTS fin_tarjetas_montos (
            tarjeta_id TEXT NOT NULL, mes TEXT NOT NULL, monto_a_pagar REAL DEFAULT 0,
            UNIQUE(tarjeta_id, mes)
        )""")
        cur.execute("SELECT COUNT(*) FROM fin_servicios")
        if cur.fetchone()[0] == 0:
            for i, (nombre, patron) in enumerate([
                ("Gas","gas"), ("Luz","edenor,edesur,luz"), ("Agua","aysa,agua"),
                ("Internet","fibertel,telecentro,movistar,internet"), ("Expensas","expensa"),
                ("Crédito","credito,préstamo,prestamo"), ("ABL","abl,rentas")
            ]):
                con.execute("INSERT INTO fin_servicios (nombre, patron, orden) VALUES (?,?,?)", (nombre, patron, i))
        # Seed vua_config con claves mínimas si está vacía
        cur.execute("SELECT COUNT(*) FROM vua_config")
        if cur.fetchone()[0] == 0:
            config_seed = [
                ("resumen_ejecutivo", "Resumen Ejecutivo", ""),
                ("antecedentes",      "Antecedentes",      ""),
                ("objetivo",          "Objetivo del Proyecto", ""),
                ("rol_dga",           "Rol de la DGA",     ""),
                ("alcance_operativo", "Alcance Operativo", ""),
            ]
            con.executemany(
                "INSERT OR IGNORE INTO vua_config (clave, titulo, contenido) VALUES (?,?,?)",
                config_seed)
        # ── Tablas SENASA ────────────────────────────────────────────────────────────
        con.execute("""CREATE TABLE IF NOT EXISTS senasa_cronologia (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            fecha TEXT, actividad TEXT, participantes TEXT,
            estado TEXT DEFAULT 'Pendiente', orden INTEGER DEFAULT 0,
            creado TEXT DEFAULT (datetime('now')), modificado TEXT DEFAULT (datetime('now'))
        )""")
        con.execute("""CREATE TABLE IF NOT EXISTS senasa_ejes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nombre TEXT, descripcion TEXT, estado TEXT DEFAULT 'Pendiente',
            orden INTEGER DEFAULT 0
        )""")
        con.execute("""CREATE TABLE IF NOT EXISTS senasa_minutas (
            id TEXT PRIMARY KEY, fecha TEXT, asunto TEXT, lugar TEXT,
            participantes TEXT, temas TEXT, conclusiones TEXT,
            compromisos TEXT, proximos TEXT, archivo TEXT,
            creado_por TEXT, creado TEXT DEFAULT (datetime('now'))
        )""")
        con.execute("""CREATE TABLE IF NOT EXISTS senasa_acuerdos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            descripcion TEXT, responsable TEXT, fecha_compromiso TEXT,
            estado TEXT DEFAULT 'Pendiente', orden INTEGER DEFAULT 0,
            creado TEXT DEFAULT (datetime('now'))
        )""")
        # Seed ejes SENASA si está vacío
        cur3 = con.cursor()
        cur3.execute("SELECT COUNT(*) FROM senasa_ejes")
        if cur3.fetchone()[0] == 0:
            ejes_seed = [
                ("Integración PAD", "Integración del Sistema de Información de Gestión (SIG-SENASA) con el Portal Aduanero Digital (PAD)", "En análisis", 1),
                ("Embalajes de Madera (NIMF-15)", "Implementación del control de embalajes de madera en el circuito aduanero digital", "Pendiente", 2),
                ("Intercambio de información", "Definición de protocolo de intercambio de datos entre SENASA y ARCA", "Pendiente", 3),
                ("Normativa y procedimientos", "Revisión y actualización de normativa conjunta SENASA-ARCA para comercio exterior", "Pendiente", 4),
            ]
            con.executemany("INSERT INTO senasa_ejes (nombre, descripcion, estado, orden) VALUES (?,?,?,?)", ejes_seed)

def _migrar_fechas_cronologia():
    """Normaliza fechas ya guardadas en otros formatos (aaaa-mm-dd, aaaa/mm/dd, etc.)
    al estándar único dd/mm/aaaa."""
    with get_db(HIST_DB) as con:
        for tabla in ("vua_cronologia", "senasa_cronologia"):
            rows = con.execute(f"SELECT id, fecha FROM {tabla}").fetchall()
            for rid, fecha in rows:
                if not _validar_fecha_ddmmaaaa(fecha):
                    nueva = _normalizar_fecha_a_ddmmaaaa(fecha)
                    if nueva != fecha:
                        con.execute(f"UPDATE {tabla} SET fecha=? WHERE id=?", (nueva, rid))

init_historial()
# _init_job_status_db() ya se ejecuta al importar core.py
_migrar_fechas_cronologia()

# WAL mejora mucho la concurrencia lectura/escritura de SQLite: con dos hilos de
# background (recordatorio de servicios, backup automático) más las requests
# normales escribiendo sobre las mismas bases, el modo "rollback journal" (default)
# genera bloqueos ("database is locked") más seguido de lo necesario.
# Se activa una sola vez: WAL queda grabado en el archivo de la base.
for _db in (HIST_DB, DB_PATH):
    try:
        if os.path.exists(_db):
            with get_db(_db) as _con_wal:
                _con_wal.execute("PRAGMA journal_mode=WAL")
    except Exception:
        logging.exception(f"No se pudo activar WAL en {_db}")

# ── Migración APP_PASS/APP_USER legacy → BD ────────────────────────────────────
def _migrar_usuarios_legacy():
    """
    Si APP_USER/APP_PASS están definidos en el entorno y el usuario no existe en BD,
    lo crea automáticamente con hash bcrypt y elimina la necesidad del fallback en login.
    Loggea la migración para auditoría.
    """
    for _user, _pass, _rol, _mods in [
        (APP_USER,  APP_PASS,  "admin",    "sintia,vua,senasa,stock,garmin,training"),
        (APP_USER2, APP_PASS2, "readonly", "sintia"),
    ]:
        if not _user or not _pass:
            continue
        with get_db(HIST_DB) as con:
            existing = con.execute("SELECT username FROM usuarios WHERE username=?", (_user,)).fetchone()
            if not existing:
                # Detectar si ya es hash bcrypt o texto plano
                if _pass.startswith("$2b$") or _pass.startswith("$2a$"):
                    _hash = _pass
                else:
                    _hash = bcrypt.hashpw(_pass.encode(), bcrypt.gensalt()).decode()
                    logging.warning(f"LEGACY MIGRATION | usuario '{_user}' migrado desde APP_PASS (texto plano → bcrypt). "
                                    f"Se recomienda eliminar APP_PASS del entorno y gestionar usuarios desde la UI.")
                con.execute("INSERT INTO usuarios (username, password_hash, rol, modulos) VALUES (?,?,?,?)",
                            (_user, _hash, _rol, _mods))
                logging.info(f"LEGACY MIGRATION | usuario '{_user}' creado en BD desde variables de entorno.")

_migrar_usuarios_legacy()

# get_db() (importado de core/db_utils arriba) cumple esta misma función —
# antes había acá una copia local idéntica llamada db_conn(), definida pero
# nunca usada en ningún lado de este archivo (código muerto).


@app.route("/login", methods=["GET", "POST"])
@limiter.limit("20 per 15 minutes", error_message="Demasiados intentos.")
def login():
    error = None
    if request.method == "POST":
        u = request.form.get("user","").strip()
        p = request.form.get("pass","")
        ip = request.remote_addr
        ua = request.headers.get("User-Agent","")[:200]
        user = get_user(u)
        autenticado = False; rol = "readonly"
        if user and check_password(p, user["password_hash"]):
            autenticado = True; rol = user["rol"]
        if autenticado:
            import secrets
            token = secrets.token_hex(32)
            session.update({"logged_in":True,"role":rol,"username":u,"token":token,
                "last_active":datetime.now().timestamp(),
                "modulos":user["modulos"].split(",") if user else ["sintia","vua","admin"]})
            session.permanent = True
            registrar_sesion(u, token, ip, ua)
            try:
                with get_db(HIST_DB) as con:
                    con.execute("UPDATE usuarios SET ultimo_acceso=datetime('now') WHERE username=?", (u,))
            except: pass
            logging.info("LOGIN OK | user=" + u + " | ip=" + ip)
            notificar_telegram(f"🔓 Login: {u} ({ip})")
            return redirect(url_for("index"))
        else:
            logging.warning("LOGIN FAIL | user=" + u + " | ip=" + ip)
            error = "Usuario o contraseña incorrectos"
    return render_template("login.html", error=error)


@app.route("/api/perfil/cambiar_password", methods=["POST"])
@login_required
@limiter.limit("10 per hour", error_message="Demasiados intentos de cambio de contraseña.")
def cambiar_password_propia():
    """Cada usuario puede cambiar su propia contraseña (antes solo un admin
    podía hacerlo, vía /api/admin/usuarios/<id>). Pide la contraseña actual
    para confirmar identidad, y revoca las demás sesiones activas de este
    usuario por seguridad (cualquier otro dispositivo logueado tiene que
    volver a entrar con la contraseña nueva)."""
    data = request.json or {}
    actual = data.get("actual", "")
    nueva = data.get("nueva", "")
    username = session.get("username")

    user = get_user(username)
    if not user or not check_password(actual, user["password_hash"]):
        return jsonify({"ok": False, "error": "La contraseña actual no es correcta"}), 403

    if len(nueva) < 8:
        return jsonify({"ok": False, "error": "La contraseña nueva debe tener al menos 8 caracteres"})
    if nueva == actual:
        return jsonify({"ok": False, "error": "La contraseña nueva tiene que ser distinta de la actual"})

    hashed = bcrypt.hashpw(nueva.encode(), bcrypt.gensalt()).decode()
    with get_db(HIST_DB) as con:
        con.execute("UPDATE usuarios SET password_hash=? WHERE username=?", (hashed, username))
        # Revocar todas las sesiones de este usuario (incluida la actual) — todos
        # los dispositivos tienen que volver a loguearse con la contraseña nueva.
        tokens = [r[0] for r in con.execute("SELECT token FROM sesiones WHERE username=?", (username,))]
        for t in tokens:
            con.execute("INSERT OR IGNORE INTO tokens_revocados (token) VALUES (?)", (t,))

    session.clear()
    logging.info(f"PASSWORD PROPIO CAMBIADO | user={username}")
    notificar_telegram(f"🔑 {username} cambió su contraseña")
    return jsonify({"ok": True})


@app.route("/health")
def health_check():
    """Chequeo mínimo, sin autenticación, pensado para monitoreo automático
    (uptime checks, balanceador de carga). No expone detalles internos."""
    try:
        with get_db(HIST_DB, timeout=5) as con:
            con.execute("SELECT 1")
        return jsonify({"status": "ok", "ts": datetime.now().isoformat()})
    except Exception:
        return jsonify({"status": "error", "ts": datetime.now().isoformat()}), 503

@app.route("/api/admin/health")
@login_required
@admin_required("sistema")
def health_check_detalle():
    """Chequeo detallado para admins: BD, disco, último backup, último
    informe generado, configuración de Telegram y del rate limiter."""
    detalle = {}

    # Bases de datos
    for nombre, ruta in [("historial_db", HIST_DB), ("pad_db", DB_PATH)]:
        try:
            if os.path.exists(ruta):
                with get_db(ruta, timeout=5) as con:
                    con.execute("SELECT 1")
                detalle[nombre] = {"ok": True, "size_mb": round(os.path.getsize(ruta)/1024/1024, 1)}
            else:
                detalle[nombre] = {"ok": False, "error": "archivo no existe"}
        except Exception as e:
            detalle[nombre] = {"ok": False, "error": str(e)}

    # Espacio en disco
    try:
        uso = shutil.disk_usage(os.path.dirname(HIST_DB) or "/")
        detalle["disco"] = {
            "libre_gb": round(uso.free/1024/1024/1024, 1),
            "total_gb": round(uso.total/1024/1024/1024, 1),
            "pct_libre": round(100*uso.free/uso.total, 1),
        }
    except Exception as e:
        detalle["disco"] = {"error": str(e)}

    # Último backup
    try:
        with get_db(HIST_DB, timeout=5) as con:
            row = con.execute("SELECT fecha, origen FROM backups ORDER BY id DESC LIMIT 1").fetchone()
        detalle["ultimo_backup"] = {"fecha": row[0], "origen": row[1]} if row else None
    except Exception as e:
        detalle["ultimo_backup"] = {"error": str(e)}

    # Último informe SINTIA generado
    try:
        with get_db(HIST_DB, timeout=5) as con:
            row = con.execute("SELECT fecha, usuario, descripcion FROM historial WHERE tipo='sintia' ORDER BY fecha DESC LIMIT 1").fetchone()
        detalle["ultimo_informe"] = {"fecha": row[0], "usuario": row[1], "descripcion": row[2]} if row else None
    except Exception as e:
        detalle["ultimo_informe"] = {"error": str(e)}

    # Jobs corriendo ahora mismo (en este worker)
    detalle["jobs_activos_este_worker"] = sum(1 for v in job_status.values() if v.get("status") == "running")

    # Configuración
    detalle["telegram_configurado"] = bool(TELEGRAM_TOKEN)
    detalle["rate_limiter_storage"] = os.environ.get("RATELIMIT_STORAGE_URI", "memory://")
    if detalle["rate_limiter_storage"] == "memory://":
        detalle["aviso_rate_limiter"] = ("En memoria: si corren más de 1 worker, el límite de intentos "
                                          "no se comparte entre procesos. Ver RATELIMIT_STORAGE_URI.")

    ok_general = all([
        detalle.get("historial_db", {}).get("ok"),
        detalle.get("pad_db", {}).get("ok", True),  # puede no estar cargada aún, no es un fallo
        detalle.get("disco", {}).get("pct_libre", 100) > 5,
    ])
    return jsonify({"status": "ok" if ok_general else "degraded", "detalle": detalle,
                    "ts": datetime.now().isoformat()})

@app.route("/logout")
def logout():
    logging.info(f"LOGOUT | user={session.get('username','?')} | ip={request.remote_addr}")
    session.clear()
    return redirect(url_for("login"))

# ── Index ──────────────────────────────────────────────────────────────────────
@app.route("/")
@login_required
def index():
    db_exists = os.path.exists(DB_PATH)
    db_size   = round(os.path.getsize(DB_PATH)/(1024**3),2) if db_exists else 0
    hoy = datetime.today()
    mes_ult = str(hoy.month-1).zfill(2) if hoy.month > 1 else "12"
    meses = {"01":"Enero","02":"Febrero","03":"Marzo","04":"Abril","05":"Mayo","06":"Junio",
             "07":"Julio","08":"Agosto","09":"Septiembre","10":"Octubre","11":"Noviembre","12":"Diciembre"}
    pendientes = []
    try:
        with get_db(HIST_DB, row_factory=True) as con:
            limite = (datetime.now()-timedelta(days=10)).strftime("%Y-%m-%d")
            pendientes = [dict(r) for r in con.execute(
                "SELECT * FROM historial WHERE revisado=0 AND fecha < ? ORDER BY fecha ASC",(limite,)).fetchall()]
    except: pass
    return render_template("dashboard.html",
        db_exists=db_exists, db_size=db_size, now=hoy, mes_ult=mes_ult, meses=meses,
        api_key=bool(get_api_key()), role=session.get("role","admin"),
        username=session.get("username",""), pendientes=pendientes)

# ── DB Status ──────────────────────────────────────────────────────────────────
@app.route("/api/db-status")
@login_required
def db_status():
    if not os.path.exists(DB_PATH): return jsonify({"exists":False})
    try:
        with get_db(DB_PATH) as con:
            cur = con.cursor()
            cur.execute("SELECT name FROM sqlite_master WHERE type='table'")
            tables = [r[0] for r in cur.fetchall()]
            info = {}
            for t in tables:
                try: cur.execute('SELECT COUNT(*) FROM "' + t.replace('"', '""') + '"'); info[t] = cur.fetchone()[0]
                except: pass
        return jsonify({"exists":True,"tables":info,"size_gb":round(os.path.getsize(DB_PATH)/(1024**3),2)})
    except Exception as e:
        return jsonify({"exists":True,"error":str(e)})

# ── Upload BD ──────────────────────────────────────────────────────────────────
@app.route("/api/upload-db", methods=["POST"])
@login_required
@admin_required("bd")
@limiter.limit("5 per hour", error_message="Demasiados intentos de subir la base completa.")
def upload_db():
    if "file" not in request.files: return jsonify({"ok":False,"error":"No se recibió archivo"})
    f = request.files["file"]
    if not f.filename.endswith(".db"): return jsonify({"ok":False,"error":"El archivo debe ser .db"})
    os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)
    f.save(DB_PATH)
    size_gb = round(os.path.getsize(DB_PATH)/(1024**3),2)
    logging.info(f"BD UPLOAD | user={session.get('username')} | size={os.path.getsize(DB_PATH)}")
    notificar_telegram(f"📦 BD reemplazada por {session.get('username')} ({size_gb} GB)")
    return jsonify({"ok":True,"size_gb":size_gb})

# ── Import SQL ─────────────────────────────────────────────────────────────────
@app.route("/api/import-sql", methods=["POST"])
@login_required
@admin_required("bd")
@limiter.limit("5 per hour", error_message="Demasiados intentos de reimportar la base completa.")
def import_sql():
    if "file" not in request.files: return jsonify({"ok":False,"error":"No se recibió archivo"})
    f = request.files["file"]
    if not f.filename.endswith(".sql"): return jsonify({"ok":False,"error":"El archivo debe ser .sql"})
    os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)
    tmp_sql = "/tmp/import_cosmo.sql"
    f.save(tmp_sql)
    try:
        if os.path.exists(DB_PATH): os.remove(DB_PATH)
        result = subprocess.run(["sqlite3", DB_PATH],
            stdin=open(tmp_sql,"r",encoding="utf-8",errors="replace"),
            capture_output=True, text=True, timeout=1800)
        os.remove(tmp_sql)
        if result.returncode != 0 and result.stderr:
            return jsonify({"ok":False,"error":result.stderr[:500]})
        size = round(os.path.getsize(DB_PATH)/(1024**3),2)
        logging.info(f"SQL IMPORT | user={session.get('username')} | size={size}GB")
        notificar_telegram(f"📦 BD reimportada desde .sql por {session.get('username')} ({size} GB)")
        return jsonify({"ok":True,"size_gb":size})
    except subprocess.TimeoutExpired:
        return jsonify({"ok":False,"error":"Timeout — archivo muy grande."})
    except Exception as e:
        return jsonify({"ok":False,"error":str(e)})

@app.route("/api/import-sql-agregar", methods=["POST"])
@login_required
@admin_required("bd")
@limiter.limit("5 per hour", error_message="Demasiados intentos de importar SQL.")
def import_sql_agregar():
    """Como import_sql(), pero SIN borrar la base existente primero -- corre
    el .sql tal cual contra la base actual, para poder sumar una tabla
    nueva (ej. DAT_2025) sin perder lo que ya hay (ej. DAT_2026). A pedido
    (13/07/2026): 'Importar SQL' reemplaza TODA la base, y hacía falta una
    forma de agregar sin destruir lo existente.

    Protección: si el .sql intenta crear o borrar una tabla que YA existe
    en la base, se bloquea antes de correr nada -- para no pisar datos por
    accidente sin que el usuario se entere. Si de verdad se quiere
    reemplazar una tabla puntual, primero hay que borrarla a mano (Limpiar
    tabla) o usar 'Importar SQL' si se quiere reemplazar todo."""
    if "file" not in request.files: return jsonify({"ok":False,"error":"No se recibió archivo"})
    f = request.files["file"]
    if not f.filename.endswith(".sql"): return jsonify({"ok":False,"error":"El archivo debe ser .sql"})

    contenido = f.read().decode("utf-8", errors="replace")

    tablas_en_sql = set(re.findall(r'CREATE\s+TABLE\s+(?:IF\s+NOT\s+EXISTS\s+)?["\']?(\w+)["\']?', contenido, re.I))
    tablas_en_sql |= set(re.findall(r'DROP\s+TABLE\s+(?:IF\s+EXISTS\s+)?["\']?(\w+)["\']?', contenido, re.I))
    if not tablas_en_sql:
        return jsonify({"ok": False, "error": "No se encontró ningún CREATE TABLE en el archivo."})

    tablas_existentes = set()
    if os.path.exists(DB_PATH):
        with get_db(DB_PATH) as con:
            tablas_existentes = {r[0] for r in con.execute(
                "SELECT name FROM sqlite_master WHERE type='table'").fetchall()}

    colision = tablas_en_sql & tablas_existentes
    if colision:
        return jsonify({"ok": False, "error":
            f"El .sql crea o borra la(s) tabla(s) {', '.join(sorted(colision))}, que ya existen en la base. "
            f"Para no pisar datos por accidente, este modo no lo permite. Si realmente querés reemplazar "
            f"esa tabla puntual, borrala primero desde 'Limpiar tabla'; si querés reemplazar TODA la base, "
            f"usá 'Importar SQL'."})

    os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)
    tmp_sql = "/tmp/import_cosmo_agregar.sql"
    with open(tmp_sql, "w", encoding="utf-8") as fh:
        fh.write(contenido)
    try:
        result = subprocess.run(["sqlite3", DB_PATH],
            stdin=open(tmp_sql, "r", encoding="utf-8", errors="replace"),
            capture_output=True, text=True, timeout=1800)
        os.remove(tmp_sql)
        if result.returncode != 0 and result.stderr:
            return jsonify({"ok": False, "error": result.stderr[:500]})
        size = round(os.path.getsize(DB_PATH)/(1024**3), 2)
        logging.info(f"SQL IMPORT (agregar) | user={session.get('username')} | "
                     f"tablas={sorted(tablas_en_sql)} | size={size}GB")
        notificar_telegram(f"📦 Tabla(s) {', '.join(sorted(tablas_en_sql))} agregada(s) a la BD "
                           f"por {session.get('username')} ({size} GB)")
        return jsonify({"ok": True, "size_gb": size, "tablas": sorted(tablas_en_sql)})
    except subprocess.TimeoutExpired:
        return jsonify({"ok": False, "error": "Timeout — archivo muy grande."})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)})


@app.route("/api/limpiar-tabla", methods=["POST"])
@login_required
@admin_required("bd")
@limiter.limit("10 per hour", error_message="Demasiados intentos de vaciar una tabla.")
def limpiar_tabla():
    """Borra todos los registros de una tabla DAT_YYYY o RECHAZOS, sin
    borrar la tabla en sí (estructura e índices quedan intactos)."""
    data = request.json or {}
    tabla = (data.get("tabla") or "").strip()
    if tabla not in ("REC", "RECHAZOS") and not re.match(r'^DAT_\d{4}$', tabla):
        return jsonify({"ok": False, "error": f"Tabla no permitida: '{tabla}'"})
    if not os.path.exists(DB_PATH):
        return jsonify({"ok": False, "error": "La BD no está cargada"})
    with get_db(DB_PATH) as con:
        cur = con.cursor()
        cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name=?", (tabla,))
        if not cur.fetchone():
            return jsonify({"ok": False, "error": f"La tabla {tabla} no existe"})
        cur.execute(f"SELECT COUNT(*) FROM {tabla}")
        total_antes = cur.fetchone()[0]
        cur.execute(f"DELETE FROM {tabla}")
    logging.info(f"LIMPIAR TABLA | tabla={tabla} | user={session.get('username')} | borrados={total_antes}")
    notificar_telegram(f"🗑️ Se vació la tabla {tabla} ({total_antes:,} registros borrados) por {session.get('username')}")
    return jsonify({"ok": True, "borrados": total_antes})

@app.route("/api/admin/backup/estado")
@login_required
@admin_required("bd")
def api_backup_estado():
    import json as _json
    with get_db(HIST_DB, row_factory=True) as con:
        row = con.execute("SELECT fecha, resultados FROM backups ORDER BY fecha DESC LIMIT 1").fetchone()
    ultimo = None
    if row:
        ultimo = {"fecha": row["fecha"], "resultados": _json.loads(row["resultados"])}
    archivos = {}
    for clave in BACKUP_FUENTES:
        path = os.path.join(BACKUP_DIR, f"{clave}_backup.db")
        if os.path.exists(path):
            archivos[clave] = {
                "size_mb": round(os.path.getsize(path) / (1024 * 1024), 2),
                "fecha": datetime.fromtimestamp(os.path.getmtime(path)).strftime("%Y-%m-%d %H:%M"),
            }
        else:
            archivos[clave] = None
    return jsonify({"ok": True, "ultimo": ultimo, "archivos": archivos, "proximo": _proximo_backup_domingo()})

@app.route("/api/admin/backup/ejecutar", methods=["POST"])
@login_required
@admin_required("bd")
@limiter.limit("5 per hour", error_message="Demasiados intentos de backup manual.")
def api_backup_ejecutar():
    resultados, todo_ok = _ejecutar_backup(origen="manual")
    return jsonify({"ok": todo_ok, "resultados": resultados})

@app.route("/api/admin/backup/download/<clave>")
@login_required
@admin_required("bd")
def api_backup_download(clave):
    if clave not in BACKUP_FUENTES:
        return jsonify({"ok": False, "error": "Backup no válido"}), 404
    path = os.path.join(BACKUP_DIR, f"{clave}_backup.db")
    if not os.path.exists(path):
        return jsonify({"ok": False, "error": "No hay backup disponible"}), 404
    return send_file(path, as_attachment=True, download_name=f"{clave}_backup.db")

@app.route("/api/admin/backup/restaurar", methods=["POST"])
@login_required
@admin_required("bd")
@limiter.limit("5 per hour", error_message="Demasiados intentos de restaurar backup.")
def api_backup_restaurar():
    import shutil
    data = request.json or {}
    clave = (data.get("clave") or "").strip()
    if clave not in BACKUP_FUENTES:
        return jsonify({"ok": False, "error": "Backup no válido"})
    backup_path = os.path.join(BACKUP_DIR, f"{clave}_backup.db")
    if not os.path.exists(backup_path):
        return jsonify({"ok": False, "error": "No hay backup disponible para restaurar"})
    destino = BACKUP_FUENTES[clave]()
    try:
        if os.path.exists(destino):
            shutil.copy2(destino, destino + ".pre_restore")  # por si hay que deshacer
        shutil.copy2(backup_path, destino)
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)})
    logging.info(f"BACKUP RESTAURAR | clave={clave} | user={session.get('username')}")
    notificar_telegram(f"♻️ Se restauró el backup de '{clave}' sobre la BD en uso — hecho por {session.get('username')}")
    return jsonify({"ok": True})

@app.route("/api/admin/backup/<clave>", methods=["DELETE"])
@login_required
@admin_required("bd")
def api_backup_delete(clave):
    if clave not in BACKUP_FUENTES:
        return jsonify({"ok": False, "error": "Backup no válido"}), 404
    path = os.path.join(BACKUP_DIR, f"{clave}_backup.db")
    if not os.path.exists(path):
        return jsonify({"ok": False, "error": "No hay backup para borrar"}), 404
    os.remove(path)
    logging.info(f"BACKUP DELETE | clave={clave} | user={session.get('username')}")
    notificar_telegram(f"🗑️ Se eliminó el backup de '{clave}' — hecho por {session.get('username')}")
    return jsonify({"ok": True})

# ── Update CSV (async) ─────────────────────────────────────────────────────────
@app.route("/api/update-csv", methods=["POST"])
@login_required
@admin_required("bd")
def update_csv():
    tabla = request.form.get("tabla","").strip()
    anio  = request.form.get("anio", str(datetime.today().year))
    modo  = request.form.get("modo", "reemplazar")
    if modo not in ("reemplazar", "agregar"): modo = "reemplazar"
    if tabla == "DAT": tabla = f"DAT_{anio}"
    # Whitelist: solo tablas conocidas o DAT_YYYY
    if tabla not in ("REC", "RECHAZOS") and not re.match(r'^DAT_\d{4}$', tabla):
        return jsonify({"ok": False, "error": f"Tabla no permitida: '{tabla}'"})
    if "file" not in request.files:
        return jsonify({"ok":False,"error":"No se recibió archivo"})
    if not os.path.exists(DB_PATH):
        return jsonify({"ok":False,"error":"La BD no está cargada"})
    f = request.files["file"]
    tmp_path = f"/tmp/upload_{uuid.uuid4().hex[:8]}.txt"
    f.save(tmp_path)
    size_kb = round(os.path.getsize(tmp_path)/1024,1)
    job_id = str(uuid.uuid4())[:8]
    job_create(job_id, f"Archivo recibido: {size_kb} KB",
               username=session.get("username", "?"))
    logging.info(f"CSV UPLOAD | tabla={tabla} | size={size_kb}KB | modo={modo}")
    t = threading.Thread(target=_run_csv_job, args=(job_id, tmp_path, tabla, modo))
    t.start()
    return jsonify({"ok":True,"job_id":job_id})

def _run_csv_job(job_id, tmp_path, tabla, modo="reemplazar"):
    log = job_status[job_id]["log"]
    username = job_status[job_id].get("username", "?")
    try:
        _procesar_csv(tmp_path, tabla, log, modo)
        job_status[job_id]["status"] = "done"
        _job_persist(job_id)
        ultimo = log[-1] if log else ""
        notificar_telegram(f"📄 Tabla {tabla} actualizada por {username}\n{ultimo}")
    except Exception as e:
        log.append(f"✗ Error: {e}")
        job_status[job_id]["status"] = "error"
        _job_persist(job_id)
        notificar_telegram(f"⚠️ Error actualizando tabla {tabla} ({username}): {e}")
    finally:
        try: os.remove(tmp_path)
        except: pass

def _procesar_csv(tmp_path, tabla, log, modo="reemplazar"):
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    try:
        if tabla.startswith("DAT_"):
            # Streaming línea por línea — evita cargar 650MB en RAM
            log.append(f"Procesando archivo DAT (streaming, modo: {modo})...")
            cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name=?", (tabla,))
            tabla_existe = cur.fetchone() is not None
            total_antes = 0
            if modo == "agregar" and tabla_existe:
                total_antes = cur.execute(f"SELECT COUNT(*) FROM {tabla}").fetchone()[0]
            headers = None
            placeholders = None
            batch = []
            batch_size = 2000
            inserted = 0
            with open(tmp_path, encoding="utf-8", errors="replace") as fh:
                for raw_line in fh:
                    line = raw_line.strip()
                    if not line or ";" not in line or line.startswith("---"):
                        continue
                    if line.upper().startswith("REGISTRO") and ";" not in line:
                        continue
                    if headers is None:
                        headers = [h.strip() for h in line.split(";")]
                        col_normalize = {"TIENE NOVEDAD?": "tiene_novedad", "TIENE NOVEDAD": "tiene_novedad"}
                        headers = [col_normalize.get(h, h) for h in headers]
                        insert_sql = None
                        if modo == "agregar" and tabla_existe:
                            # OPERACION_PAD_EXT (+ MIC + TIPO_REGISTRO) es el
                            # identificador único real de una operación (dato
                            # confirmado en conversación 13/07/2026) -- el
                            # archivo semanal trae el semestre ENTERO
                            # acumulado, no solo lo nuevo, así que reinsertar
                            # una operación ya vista es normal y esperado
                            # (puede venir con el estado actualizado, ej.
                            # pasó de pendiente a SAL). OR REPLACE hace que
                            # eso actualice la fila existente en vez de
                            # duplicarla o de romper por violar el índice
                            # único. El índice se crea ACÁ, antes de
                            # insertar nada -- antes se creaba recién al
                            # final, con lo cual nunca llegaba a hacer nada
                            # durante el insert (y si había datos repetidos,
                            # la creación del índice fallaba en silencio,
                            # dejando la tabla sin el índice Y con los
                            # duplicados adentro).
                            log.append("Modo agregar: se actualizan operaciones ya existentes "
                                       "(por OPERACION_PAD_EXT) y se suman las nuevas.")
                            try:
                                cur.execute(f"CREATE UNIQUE INDEX IF NOT EXISTS idx_{tabla}_key "
                                           f"ON {tabla}(OPERACION_PAD_EXT, MIC, TIPO_REGISTRO)")
                            except Exception as e:
                                log.append(f"⚠ No se pudo asegurar el índice único antes de insertar: {e}")
                            cols_quoted = ", ".join(f'"{h}"' for h in headers)
                            placeholders = ", ".join(["?" for _ in headers])
                            insert_sql = f"INSERT OR REPLACE INTO {tabla} ({cols_quoted}) VALUES ({placeholders})"
                        else:
                            cur.execute(f"DROP TABLE IF EXISTS {tabla}")
                            cols_def = ", ".join([f'"{h}" TEXT' for h in headers])
                            cur.execute(f"CREATE TABLE {tabla} ({cols_def})")
                            try:
                                cur.execute(f"CREATE UNIQUE INDEX IF NOT EXISTS idx_{tabla}_key "
                                           f"ON {tabla}(OPERACION_PAD_EXT, MIC, TIPO_REGISTRO)")
                            except Exception as e:
                                log.append(f"⚠ No se pudo crear el índice único: {e}")
                            log.append(f"Columnas: {len(headers)} — tabla creada, insertando...")
                            placeholders = ", ".join(["?" for _ in headers])
                            insert_sql = f"INSERT OR REPLACE INTO {tabla} VALUES ({placeholders})"
                        con.commit()
                        continue
                    vals = [v.strip() for v in line.split(";")]
                    while len(vals) < len(headers): vals.append(None)
                    batch.append(vals[:len(headers)])
                    if len(batch) >= batch_size:
                        cur.executemany(insert_sql, batch)
                        inserted += len(batch)
                        batch = []
                        if inserted % 50000 == 0:
                            con.commit()
                            log.append(f"  {inserted:,} filas procesadas...")
            if batch:
                cur.executemany(insert_sql, batch)
                inserted += len(batch)
            con.commit()
            if inserted == 0:
                log.append("✗ No se encontraron datos válidos"); con.close(); return
            total_despues = cur.execute(f"SELECT COUNT(*) FROM {tabla}").fetchone()[0]
            if modo == "agregar" and tabla_existe:
                nuevas = total_despues - total_antes
                actualizadas = inserted - nuevas
                log.append(f"  {inserted:,} filas del archivo procesadas: {nuevas:,} operaciones nuevas, "
                          f"{actualizadas:,} actualizadas (ya existían, se refrescó su estado). "
                          f"Total en la tabla ahora: {total_despues:,}.")
            else:
                log.append(f"  {inserted:,} filas insertadas en total. Calculando fechas ISO...")
            for col in ["FECHA_INGRESO_ISO","FECHA_TRANS_ISO"]:
                try: cur.execute(f"ALTER TABLE {tabla} ADD COLUMN {col} TEXT")
                except: pass
            # Antes esto actualizaba por rango de rowid (rowid_inicio..rowid_fin
            # de las filas recién insertadas). Con INSERT OR REPLACE eso ya no
            # sirve: reemplazar una fila existente la borra y la vuelve a
            # insertar con un rowid NUEVO, no necesariamente contiguo con las
            # demás filas de este import -- el rango calculado de antemano
            # podía quedar desalineado. Ahora se apunta directamente a "toda
            # fila que todavía no tiene la fecha ISO calculada", que es
            # exactamente el conjunto correcto sin importar cómo haya
            # quedado el rowid.
            rango = cur.execute(f"SELECT MIN(rowid), MAX(rowid) FROM {tabla} WHERE FECHA_INGRESO_ISO IS NULL").fetchone()
            if rango and rango[0] is not None:
                rowid_min, rowid_max = rango
                iso_sql = f"""UPDATE {tabla} SET
                    FECHA_INGRESO_ISO = CASE
                        WHEN FECHA_INGRESO IS NOT NULL AND length(FECHA_INGRESO)>=10 THEN
                            substr(FECHA_INGRESO,7,4)||'-'||substr(FECHA_INGRESO,4,2)||'-'||substr(FECHA_INGRESO,1,2)||
                            CASE WHEN instr(FECHA_INGRESO,' ')>0 THEN ' '||substr(FECHA_INGRESO,instr(FECHA_INGRESO,' ')+1,5) ELSE ' 00:00' END
                        ELSE NULL END,
                    FECHA_TRANS_ISO = CASE
                        WHEN FECHA_TRANS IS NOT NULL AND FECHA_TRANS NOT IN ('-','') AND length(FECHA_TRANS)>=10 THEN
                            substr(FECHA_TRANS,7,4)||'-'||substr(FECHA_TRANS,4,2)||'-'||substr(FECHA_TRANS,1,2)||
                            CASE WHEN instr(FECHA_TRANS,' ')>0 THEN ' '||substr(FECHA_TRANS,instr(FECHA_TRANS,' ')+1,5) ELSE ' 00:00' END
                        ELSE NULL END
                    WHERE rowid BETWEEN ? AND ? AND FECHA_INGRESO_ISO IS NULL"""
                batch_iso = 50000
                total_iso = rowid_max - rowid_min + 1
                for start in range(rowid_min, rowid_max + 1, batch_iso):
                    fin_batch = min(start + batch_iso - 1, rowid_max)
                    cur.execute(iso_sql, (start, fin_batch))
                    con.commit()
                    log.append(f"  Fechas ISO: hasta {min(fin_batch - rowid_min + 1, total_iso):,} / {total_iso:,}...")
            log.append("Creando índices...")
            try: cur.execute(f"CREATE INDEX IF NOT EXISTS idx_{tabla}_fecha ON {tabla}(FECHA_INGRESO_ISO)")
            except: pass
            try: cur.execute(f"CREATE INDEX IF NOT EXISTS idx_{tabla}_estado ON {tabla}(EST_MIC)")
            except: pass
            try: cur.execute(f"CREATE INDEX IF NOT EXISTS idx_{tabla}_aduana ON {tabla}(ADUANA)")
            except: pass
            con.commit(); con.close()
            log.append(f"✓ {tabla}: fechas ISO calculadas, índices creados")

            # Opciones de filtro (EST_MIC/ULT_ESTADO/VAR_CONTROL) para el
            # combo de Consulta DAT -- se recalculan acá, una sola vez por
            # import, en vez de escanear la tabla en cada apertura del panel.
            try:
                _recalcular_opciones_dat(tabla)
                log.append("✓ Opciones de filtro (Est. MIC / Ult. Estado / Var. Control) actualizadas")
            except Exception as e:
                log.append(f"⚠ No se pudieron recalcular las opciones de filtro: {e}")

        elif tabla == "RECHAZOS":
            log.append("Procesando archivo RECHAZOS (streaming)...")
            cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='RECHAZOS'")
            if not cur.fetchone():
                cur.execute("CREATE TABLE RECHAZOS (PaisEmisor TEXT, Metodo TEXT, NroMic TEXT, Fecha TEXT, Mensaje TEXT, Fecha_ISO TEXT, Mes TEXT, Anio TEXT)")
            else:
                for col in ["Fecha_ISO","Mes","Anio"]:
                    try: cur.execute(f"ALTER TABLE RECHAZOS ADD COLUMN {col} TEXT")
                    except: pass
            con.commit()

            col_map = {"Pais Emisor":"PaisEmisor","Metodo":"Metodo","Nro. MIC/DTA":"NroMic","Fecha":"Fecha","Mensaje":"Mensaje"}
            cols_utiles = ["PaisEmisor","Metodo","NroMic","Fecha","Mensaje"]
            idx_cols = None
            cols_finales = None
            placeholders = None
            batch = []
            batch_size = 5000
            inserted = 0

            with open(tmp_path, encoding="utf-8", errors="replace") as fh:
                for raw_line in fh:
                    s = raw_line.strip()
                    if not s or "@" not in s: continue
                    if len(s) >= 3 and s[:2].isdigit() and s[2] == "@": s = s[3:]
                    parts = [p.strip() for p in s.split("@")][:6]
                    if len(parts) < 2: continue

                    if idx_cols is None:
                        # Primera línea válida = headers
                        raw_headers = [col_map.get(h.strip(), h.strip()) for h in parts]
                        idx_cols = [i for i, h in enumerate(raw_headers) if h in cols_utiles]
                        cols_finales = [raw_headers[i] for i in idx_cols]
                        placeholders = ", ".join(["?" for _ in cols_finales])
                        continue

                    row = [parts[i].strip() if i < len(parts) else None for i in idx_cols]
                    batch.append(row)

                    if len(batch) >= batch_size:
                        try: cur.executemany(f"INSERT INTO RECHAZOS ({', '.join(cols_finales)}) VALUES ({placeholders})", batch); inserted += len(batch)
                        except:
                            for r in batch:
                                try: cur.execute(f"INSERT INTO RECHAZOS ({', '.join(cols_finales)}) VALUES ({placeholders})", r); inserted += 1
                                except: pass
                        con.commit()
                        batch = []
                        log.append(f"  {inserted:,} filas procesadas...")

            # Último batch
            if batch:
                try: cur.executemany(f"INSERT INTO RECHAZOS ({', '.join(cols_finales)}) VALUES ({placeholders})", batch); inserted += len(batch)
                except:
                    for r in batch:
                        try: cur.execute(f"INSERT INTO RECHAZOS ({', '.join(cols_finales)}) VALUES ({placeholders})", r); inserted += 1
                        except: pass
                con.commit()

            log.append(f"Filas insertadas: {inserted:,}")
            log.append("Calculando fechas ISO...")
            cur.execute("""UPDATE RECHAZOS SET
                Fecha_ISO = printf('%04d-%02d-%02d',
                    CAST(SUBSTR(CAST(Fecha AS TEXT),LENGTH(CAST(Fecha AS TEXT))-3,4) AS INT),
                    CAST(SUBSTR(CAST(Fecha AS TEXT),LENGTH(CAST(Fecha AS TEXT))-5,2) AS INT),
                    CAST(SUBSTR(CAST(Fecha AS TEXT),1,LENGTH(CAST(Fecha AS TEXT))-6) AS INT)),
                Anio = CAST(SUBSTR(CAST(Fecha AS TEXT),LENGTH(CAST(Fecha AS TEXT))-3,4) AS INTEGER),
                Mes = CASE CAST(SUBSTR(CAST(Fecha AS TEXT),LENGTH(CAST(Fecha AS TEXT))-5,2) AS INTEGER)
                    WHEN 1 THEN 'ENERO' WHEN 2 THEN 'FEBRERO' WHEN 3 THEN 'MARZO' WHEN 4 THEN 'ABRIL'
                    WHEN 5 THEN 'MAYO' WHEN 6 THEN 'JUNIO' WHEN 7 THEN 'JULIO' WHEN 8 THEN 'AGOSTO'
                    WHEN 9 THEN 'SEPTIEMBRE' WHEN 10 THEN 'OCTUBRE' WHEN 11 THEN 'NOVIEMBRE' WHEN 12 THEN 'DICIEMBRE'
                    END
                WHERE Fecha IS NOT NULL AND LENGTH(CAST(Fecha AS TEXT))>=6""")
            try: cur.execute("CREATE INDEX IF NOT EXISTS idx_rechazos_fecha ON RECHAZOS(Fecha_ISO)")
            except: pass
            con.commit(); con.close()
            log.append(f"✓ RECHAZOS: {inserted:,} registros insertados, fechas calculadas")
        else:
            con.close()
            log.append(f"✗ Tabla no reconocida: {tabla}")
    except Exception as e:
        log.append(f"✗ Error: {e}")
        try: con.close()
        except: pass

# ── Generar informe (async) ────────────────────────────────────────────────────
def run_job(job_id, pais, anio, mes_d, mes_h, usar_ia, username):
    log = job_status[job_id]["log"]
    inicio = time.time()
    try:
        from generar import generar_informe
        archivos = generar_informe(
            ruta_db=DB_PATH, pais=pais, anio=anio, mes_d=mes_d, mes_h=mes_h,
            usar_ia=usar_ia, api_key=get_api_key(), carpeta=OUTPUT_FOLDER,
            log_fn=lambda msg: log.append(msg), contexto_extra=contexto_repositorio("sintia"))
        # Guardar en historial
        hist_id = str(uuid.uuid4())[:8]
        word  = next((a for a in archivos if a.endswith(".docx")),"")
        excel = next((a for a in archivos if a.endswith(".xlsx")),"")
        with get_db(HIST_DB) as con:
            con.execute("INSERT INTO historial VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)",
                (hist_id, datetime.now().strftime("%Y-%m-%d %H:%M:%S"), username,
                 pais, anio, mes_d, mes_h, int(usar_ia), word, excel, 0, 'sintia',
                 f"{pais} {mes_d}-{mes_h}/{anio}"))
        logging.info(f"INFORME OK | user={username} | pais={pais} | {mes_d}-{mes_h}/{anio}")
        job_status[job_id]["status"] = "done"
        job_status[job_id]["files"]  = archivos
        _job_persist(job_id)
        duracion = round((time.time() - inicio) / 60, 1)
        if duracion >= 2:  # avisar solo si tardó lo suficiente como para que valga la pena
            notificar_telegram(f"✓ Informe SINTIA {pais} {mes_d}-{mes_h}/{anio} listo "
                                f"({duracion} min) — {username}")
    except Exception as e:
        log.append(f"✗ Error: {e}")
        job_status[job_id]["status"] = "error"
        _job_persist(job_id)
        notificar_telegram(f"⚠️ Informe SINTIA {pais} {mes_d}-{mes_h}/{anio} falló ({username}): {e}")

@app.route("/api/generar", methods=["POST"])
@login_required
@modulo_required("sintia")
def api_generar():
    if not os.path.exists(DB_PATH):
        return jsonify({"ok":False,"error":"La BD no está cargada"})
    data = request.json or {}
    pais    = data.get("pais","").upper()
    anio    = data.get("anio", str(datetime.today().year))
    mes_d   = str(data.get("mes_d","01")).zfill(2)
    mes_h   = str(data.get("mes_h","12")).zfill(2)
    usar_ia = data.get("usar_ia",True) and bool(get_api_key())
    username = session.get("username","?")
    job_id  = str(uuid.uuid4())[:8]
    job_create(job_id, "Iniciando generación...", username=username)
    t = threading.Thread(target=run_job, args=(job_id, pais, anio, mes_d, mes_h, usar_ia, username))
    t.start()
    return jsonify({"ok":True,"job_id":job_id})

def _job_autorizado(info):
    """El dueño del job o un superadmin pueden verlo/descargarlo."""
    if not info: return False
    if session.get("role") == "admin": return True
    return info.get("username") == session.get("username")

@app.route("/api/job/<job_id>")
@login_required
def job_poll(job_id):
    info = job_get(job_id)
    if not info: return jsonify({"error":"Job no encontrado"})
    if not _job_autorizado(info):
        return jsonify({"error":"Sin permiso"}), 403
    return jsonify(info)

@app.route("/api/download/<job_id>/<int:idx>")
@login_required
def download_file(job_id, idx):
    info = job_get(job_id) or {}
    if not _job_autorizado(info):
        return "Sin permiso",403
    files = info.get("files",[])
    if idx >= len(files): return "Archivo no encontrado",404
    path = files[idx]
    if not os.path.exists(path): return "Archivo no encontrado",404
    return send_file(path, as_attachment=True, download_name=os.path.basename(path))

# ── Historial ──────────────────────────────────────────────────────────────────
@app.route("/api/historial")
@login_required
@admin_required("bd")
def api_historial():
    try:
        with get_db(HIST_DB, row_factory=True) as con:
            rows = [dict(r) for r in con.execute(
                "SELECT * FROM historial ORDER BY fecha DESC LIMIT 100").fetchall()]
        return jsonify({"ok":True,"rows":rows})
    except Exception as e:
        return jsonify({"ok":False,"error":str(e)})

@app.route("/api/historial/completo")
@login_required
@admin_required("bd")
def historial_completo():
    """Mejora 4: paginación y filtros por tipo, fecha y usuario."""
    tipo    = request.args.get("tipo", "todos")      # todos | sintia | minuta | vua | aduanas_pais
    usuario = request.args.get("usuario", "")
    desde   = request.args.get("desde", "")
    hasta   = request.args.get("hasta", "")
    limit   = min(int(request.args.get("limit", 50)), 200)
    offset  = int(request.args.get("offset", 0))

    with get_db(HIST_DB, row_factory=True) as con:
        rows = []

        if tipo in ("todos", "sintia"):
            # Bug encontrado en revisión (10/07/2026): faltaba filtrar por
            # tipo='sintia' acá -- traía TODAS las filas de historial (VUA,
            # aduanas_pais, lo que sea) etiquetadas como si fueran sintia, y
            # además duplicadas con los bloques de abajo cuando tipo="todos".
            where = ["tipo='sintia'"]
            params = []
            if usuario: where.append("usuario=?"); params.append(usuario)
            if desde:   where.append("fecha>=?"); params.append(desde)
            if hasta:   where.append("fecha<=?"); params.append(hasta)
            q = ("SELECT id, fecha, usuario, 'sintia' as tipo, "
                 "pais||' '||mes_d||'-'||mes_h||'/'||anio as descripcion, "
                 "archivo_word, archivo_excel, revisado FROM historial "
                 f"WHERE {' AND '.join(where)} ORDER BY fecha DESC LIMIT ? OFFSET ?")
            rows += [dict(r) for r in con.execute(q, params+[limit, offset]).fetchall()]

        if tipo in ("todos", "minuta"):
            where = ["1=1"]
            params = []
            if usuario: where.append("creado_por=?"); params.append(usuario)
            if desde:   where.append("creado>=?"); params.append(desde)
            if hasta:   where.append("creado<=?"); params.append(hasta)
            q = ("SELECT id, creado as fecha, creado_por as usuario, 'minuta' as tipo, "
                 "asunto as descripcion, archivo as archivo_word, '' as archivo_excel, 1 as revisado "
                 f"FROM vua_minutas WHERE {' AND '.join(where)} ORDER BY creado DESC LIMIT ? OFFSET ?")
            rows += [dict(r) for r in con.execute(q, params+[limit, offset]).fetchall()]

        if tipo in ("todos", "vua"):
            where = ["h.tipo='vua'"] if "tipo" in con.execute("PRAGMA table_info(historial)").fetchone() or [] else []
            # Informes VUA del historial general
            q = ("SELECT id, fecha, usuario, 'vua' as tipo, descripcion, "
                 "archivo_word, '' as archivo_excel, revisado FROM historial "
                 "WHERE tipo='vua' ORDER BY fecha DESC LIMIT ? OFFSET ?")
            try:
                rows += [dict(r) for r in con.execute(q, [limit, offset]).fetchall()]
            except: pass

        if tipo in ("todos", "aduanas_pais"):
            where = ["tipo='aduanas_pais'"]
            params = []
            if usuario: where.append("usuario=?"); params.append(usuario)
            if desde:   where.append("fecha>=?"); params.append(desde)
            if hasta:   where.append("fecha<=?"); params.append(hasta)
            q = ("SELECT id, fecha, usuario, 'aduanas_pais' as tipo, descripcion, "
                 "archivo_word, archivo_excel, revisado FROM historial "
                 f"WHERE {' AND '.join(where)} ORDER BY fecha DESC LIMIT ? OFFSET ?")
            rows += [dict(r) for r in con.execute(q, params+[limit, offset]).fetchall()]

    todos = sorted(rows, key=lambda x: x.get("fecha",""), reverse=True)
    return jsonify({"ok": True, "rows": todos[:limit], "total": len(todos), "offset": offset})

@app.route("/api/historial/<hist_id>/revisar", methods=["POST"])
@login_required
@admin_required("bd")
def revisar_historial(hist_id):
    accion = (request.json or {}).get("accion","conservar")
    try:
        with get_db(HIST_DB) as con:
            if accion == "eliminar":
                row = con.execute("SELECT archivo_word, archivo_excel FROM historial WHERE id=?",(hist_id,)).fetchone()
                if row:
                    for f in [row[0],row[1]]:
                        if f and os.path.exists(f):
                            try: os.remove(f)
                            except: pass
                con.execute("DELETE FROM historial WHERE id=?",(hist_id,))
            else:
                con.execute("UPDATE historial SET revisado=1 WHERE id=?",(hist_id,))
        return jsonify({"ok":True})
    except Exception as e:
        return jsonify({"ok":False,"error":str(e)})

@app.route("/api/historial/<hist_id>/download/<tipo>")
@login_required
@admin_required("bd")
def download_historial(hist_id, tipo):
    try:
        with get_db(HIST_DB) as con:
            row = con.execute("SELECT archivo_word, archivo_excel FROM historial WHERE id=?",(hist_id,)).fetchone()
        if not row: return "No encontrado",404
        path = row[0] if tipo == "word" else row[1]
        if not path or not os.path.exists(path): return "Archivo no encontrado",404
        return send_file(path, as_attachment=True, download_name=os.path.basename(path))
    except: return "Error",500

# MÓDULO VUA -> blueprints/vua.py (registrado como vua_bp)

# ══════════════════════════════════════════════════════════════════════════════
# MÓDULO SENASA -> blueprints/senasa.py (registrado como senasa_bp)

# ── Integrantes compartidos (VUA, SENASA, SINTIA) ────────────────────────────
@app.route("/api/integrantes", methods=["GET"])
@login_required
def integrantes_list():
    """Lista integrantes activos. Fusiona vua_equipo si unified=1 (default)."""
    organismo = request.args.get("organismo", "")
    unified   = request.args.get("unified", "1")
    with get_db(HIST_DB, row_factory=True) as con:
        q = "SELECT * FROM integrantes WHERE activo=1"
        params = []
        if organismo:
            q += " AND organismo=?"; params.append(organismo)
        q += " ORDER BY orden, organismo, nombre"
        rows = [dict(r) for r in con.execute(q, params).fetchall()]
        # Fusionar con vua_equipo si unified=1 y la tabla existe
        if unified == "1":
            try:
                equipo = [dict(r) for r in con.execute(
                    "SELECT nombre, cargo, organismo, email FROM vua_equipo WHERE activo=1 ORDER BY organismo, nombre"
                ).fetchall()]
                nombres_existentes = {r["nombre"].lower() for r in rows}
                for e in equipo:
                    if e["nombre"].lower() not in nombres_existentes:
                        rows.append({"id": None, "nombre": e["nombre"], "cargo": e["cargo"],
                                     "organismo": e["organismo"], "email": e.get("email",""),
                                     "activo": 1, "orden": 999, "_origen": "vua_equipo"})
            except Exception:
                pass
    return jsonify({"ok": True, "rows": rows})

@app.route("/api/integrantes/migrar-equipo", methods=["POST"])
@login_required
def integrantes_migrar_equipo():
    """Migra todos los registros de vua_equipo a integrantes (ejecutar una sola vez)."""
    try:
        with get_db(HIST_DB, row_factory=True) as con:
            equipo = [dict(r) for r in con.execute(
                "SELECT nombre, cargo, organismo, email FROM vua_equipo WHERE activo=1").fetchall()]
            migrados = 0
            for e in equipo:
                existe = con.execute("SELECT id FROM integrantes WHERE LOWER(nombre)=LOWER(?)",
                                     (e["nombre"],)).fetchone()
                if not existe:
                    cur = con.cursor()
                    cur.execute("SELECT MAX(orden) FROM integrantes")
                    max_o = cur.fetchone()[0] or 0
                    con.execute("INSERT INTO integrantes (nombre,cargo,organismo,email,activo,orden) VALUES (?,?,?,?,1,?)",
                        (e["nombre"], e["cargo"], e["organismo"], e.get("email",""), max_o+1))
                    migrados += 1
    except Exception as ex:
        return jsonify({"ok": False, "error": str(ex)})
    return jsonify({"ok": True, "migrados": migrados})

@app.route("/api/integrantes", methods=["POST"])
@login_required
def integrantes_create():
    data = request.json or {}
    with get_db(HIST_DB) as con:
        cur = con.cursor()
        cur.execute("SELECT MAX(orden) FROM integrantes")
        max_orden = cur.fetchone()[0] or 0
        cur.execute(
            "INSERT INTO integrantes (nombre, cargo, organismo, email, activo, orden) VALUES (?,?,?,?,1,?)",
            (data.get("nombre",""), data.get("cargo",""),
             data.get("organismo",""), data.get("email",""), max_orden+1))
        new_id = cur.lastrowid
    return jsonify({"ok": True, "id": new_id})

@app.route("/api/integrantes/<int:iid>", methods=["PUT"])
@login_required
def integrantes_update(iid):
    data = request.json or {}
    with get_db(HIST_DB) as con:
        con.execute(
            "UPDATE integrantes SET nombre=?, cargo=?, organismo=?, email=?, activo=? WHERE id=?",
            (data.get("nombre",""), data.get("cargo",""),
             data.get("organismo",""), data.get("email",""),
             int(data.get("activo", 1)), iid))
    return jsonify({"ok": True})

@app.route("/api/integrantes/<int:iid>", methods=["DELETE"])
@login_required
def integrantes_delete(iid):
    with get_db(HIST_DB) as con:
        con.execute("UPDATE integrantes SET activo=0 WHERE id=?", (iid,))
    return jsonify({"ok": True})

@app.route("/api/integrantes/organismos", methods=["GET"])
@login_required
def integrantes_organismos():
    """Lista los organismos únicos para el filtro del selector."""
    with get_db(HIST_DB, row_factory=True) as con:
        rows = [r[0] for r in con.execute(
            "SELECT DISTINCT organismo FROM integrantes WHERE activo=1 AND organismo!='' ORDER BY organismo"
        ).fetchall()]
    return jsonify({"ok": True, "organismos": rows})

# (bloque VUA -> blueprints/vua.py, ver más abajo)



# ── Rutas Admin ───────────────────────────────────────────────────────────────
@app.route("/sintia")
@login_required
@modulo_required("sintia")
def sintia_index():
    import generar
    meses = generar.MESES
    hoy = __import__('datetime').date.today()
    mes_ult = str(hoy.month - 1).zfill(2) if hoy.month > 1 else '12' 
    db_exists = os.path.exists(DB_PATH)
    db_size = round(os.path.getsize(DB_PATH)/1e9, 2) if db_exists else 0
    api_key = get_api_key()
    return render_template("sintia.html",
        meses=meses, mes_ult=mes_ult, now=datetime.now(),
        db_exists=db_exists, db_size=db_size,
        api_key=api_key, username=session.get("username",""))

@app.route("/admin")
@login_required
@admin_required()
def admin_index():
    db_exists = os.path.exists(DB_PATH)
    db_size = round(os.path.getsize(DB_PATH)/1e9, 2) if db_exists else 0
    with get_db(HIST_DB, row_factory=True) as con:
        pendientes = con.execute(
            "SELECT * FROM historial WHERE revisado=0 AND "
            "julianday('now') - julianday(fecha) > 10").fetchall()
    return render_template("admin.html",
        db_exists=db_exists, db_size=db_size, now=datetime.now(),
        pendientes=pendientes, username=session.get("username",""),
        role=session.get("role","readonly"),
        permiso_bd=tiene_permiso_admin("bd"), permiso_sistema=tiene_permiso_admin("sistema"))

@app.route("/api/admin/usuarios", methods=["GET"])
@login_required
@admin_required("sistema")
def admin_usuarios_list():
    with get_db(HIST_DB, row_factory=True) as con:
        rows = [dict(r) for r in con.execute(
            "SELECT id, username, rol, modulos, activo, ultimo_acceso FROM usuarios ORDER BY id").fetchall()]
    return jsonify({"ok": True, "rows": rows})

@app.route("/api/admin/usuarios", methods=["POST"])
@login_required
@admin_required("sistema")
@limiter.limit("20 per hour", error_message="Demasiados intentos de crear usuarios.")
def admin_usuarios_create():
    data = request.json or {}
    username = data.get("username","").strip()
    password = data.get("password","")
    rol = data.get("rol","readonly")
    modulos = data.get("modulos","sintia,vua")
    if not username or not password:
        return jsonify({"ok": False, "error": "Usuario y password requeridos"})
    if not re.match(r'^[a-zA-Z0-9_.-]{3,32}$', username):
        return jsonify({"ok": False, "error": "Usuario inválido: solo letras, números, '.', '_', '-' (3-32 caracteres)"})
    if len(password) < 8:
        return jsonify({"ok": False, "error": "Password minimo 8 caracteres"})
    hashed = bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()
    try:
        with get_db(HIST_DB) as con:
            con.execute("INSERT INTO usuarios (username, password_hash, rol, modulos, activo) VALUES (?,?,?,?,1)",
                (username, hashed, rol, modulos))
        logging.info(f"USUARIO CREATE | by={session.get('username')} | nuevo={username} | rol={rol} | modulos={modulos}")
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"ok": False, "error": "Usuario ya existe" if "UNIQUE" in str(e) else str(e)})

@app.route("/api/admin/usuarios/<int:uid>", methods=["PUT"])
@login_required
@admin_required("sistema")
def admin_usuarios_update(uid):
    data = request.json or {}
    with get_db(HIST_DB) as con:
        fields = []; params = []
        for f in ["rol","modulos","activo"]:
            if f in data: fields.append(f + "=?"); params.append(data[f])
        cambia_pass = "password" in data and data["password"]
        if cambia_pass:
            if len(data["password"]) < 8:
                return jsonify({"ok": False, "error": "La contraseña debe tener al menos 8 caracteres"})
            hashed = bcrypt.hashpw(data["password"].encode(), bcrypt.gensalt()).decode()
            fields.append("password_hash=?"); params.append(hashed)
        if fields:
            params.append(uid)
            con.execute("UPDATE usuarios SET " + ", ".join(fields) + " WHERE id=?", params)
    logging.info(f"USUARIO UPDATE | by={session.get('username')} | uid={uid} | campos={list(data.keys())}" +
                 (" | password cambiado" if cambia_pass else ""))
    return jsonify({"ok": True})

@app.route("/api/admin/usuarios/<int:uid>", methods=["DELETE"])
@login_required
@admin_required("sistema")
def admin_usuarios_delete(uid):
    with get_db(HIST_DB) as con:
        con.execute("DELETE FROM usuarios WHERE id=?", (uid,))
    logging.info(f"USUARIO DELETE | by={session.get('username')} | uid={uid}")
    return jsonify({"ok": True})

@app.route("/api/admin/sesiones", methods=["GET"])
@login_required
@admin_required("sistema")
def admin_sesiones_list():
    current_token = session.get("token","")
    with get_db(HIST_DB, row_factory=True) as con:
        rows = con.execute(
            "SELECT username, SUBSTR(token,1,8)||'...' as token, ip, ultimo_acceso, token as full_token "
            "FROM sesiones WHERE activo=1 ORDER BY ultimo_acceso DESC").fetchall()
    result = []
    for r in rows:
        d = dict(r); d["es_propia"] = d.pop("full_token","") == current_token; result.append(d)
    return jsonify({"ok": True, "rows": result})

@app.route("/api/admin/sesiones/<token_prefix>/revocar", methods=["POST"])
@login_required
@admin_required("sistema")
def admin_sesiones_revocar(token_prefix):
    with get_db(HIST_DB) as con:
        con.execute("UPDATE sesiones SET activo=0 WHERE token LIKE ?", (token_prefix + "%",))
        con.execute("INSERT OR IGNORE INTO tokens_revocados (token) SELECT token FROM sesiones WHERE token LIKE ?",
            (token_prefix + "%",))
    return jsonify({"ok": True})

@app.route("/api/admin/sesiones/revocar-todas", methods=["POST"])
@login_required
@admin_required("sistema")
def admin_sesiones_revocar_todas():
    current_token = session.get("token","")
    with get_db(HIST_DB) as con:
        rows = con.execute("SELECT token FROM sesiones WHERE activo=1 AND token!=?", (current_token,)).fetchall()
        for r in rows:
            con.execute("INSERT OR IGNORE INTO tokens_revocados (token) VALUES (?)", (r[0],))
        con.execute("UPDATE sesiones SET activo=0 WHERE token!=?", (current_token,))
    return jsonify({"ok": True, "revocadas": len(rows)})

@app.route("/api/admin/prompts", methods=["GET"])
@login_required
@admin_required("sistema")
def admin_prompts_list():
    with get_db(HIST_DB, row_factory=True) as con:
        rows = [dict(r) for r in con.execute(
            "SELECT id, nombre, descripcion, modulo, modificado FROM prompts ORDER BY modulo, nombre").fetchall()]
    return jsonify({"ok": True, "rows": rows})

@app.route("/api/admin/prompts/<int:pid>", methods=["GET"])
@login_required
@admin_required("sistema")
def admin_prompts_get(pid):
    with get_db(HIST_DB, row_factory=True) as con:
        row = con.execute("SELECT * FROM prompts WHERE id=?", (pid,)).fetchone()
    if not row: return jsonify({"ok": False, "error": "No encontrado"})
    return jsonify({"ok": True, "prompt": dict(row)})

@app.route("/api/admin/prompts/<int:pid>", methods=["PUT"])
@login_required
@admin_required("sistema")
def admin_prompts_update(pid):
    data = request.json or {}
    contenido = data.get("contenido","").strip()
    if not contenido: return jsonify({"ok": False, "error": "Contenido vacio"})
    with get_db(HIST_DB) as con:
        con.execute("UPDATE prompts SET contenido=?, modificado=datetime('now') WHERE id=?", (contenido, pid))
    return jsonify({"ok": True})

# ── Helpers BD usuarios/sesiones ─────────────────────────────────────────────
def get_user(username):
    try:
        with get_db(HIST_DB, row_factory=True) as con:
            row = con.execute("SELECT * FROM usuarios WHERE username=? AND activo=1", (username,)).fetchone()
        return dict(row) if row else None
    except: return None

# registrar_sesion / actualizar_sesion / token_revocado ahora viven en core.py

# ══════════════════════════════════════════════════════════════════════════════
# SUBMÓDULOS CONSULTA DAT / RECHAZOS — agregar a app.py antes del if __name__
# Requiere: openpyxl   →   pip install openpyxl --break-system-packages
# ══════════════════════════════════════════════════════════════════════════════

def _resolver_tabla_dat(p):
    """Resuelve la tabla DAT_<anio> a partir del parámetro opcional 'anio' del
    request; si no viene o es inválido, usa el año actual. Valida rango para
    evitar que un valor arbitrario llegue al f-string de la query SQL."""
    import datetime as _dt
    anio = p.get("anio")
    try:
        anio = int(anio) if anio else _dt.date.today().year
    except (TypeError, ValueError):
        anio = _dt.date.today().year
    if anio < 2000 or anio > 2100:
        anio = _dt.date.today().year
    return f"DAT_{anio}", anio

def _calcular_opciones_dat(tabla):
    """SELECT DISTINCT real contra DB_PATH para EST_MIC/ULT_ESTADO/VAR_CONTROL.
    Costoso en una tabla grande -- por eso esto se llama una sola vez (al
    terminar de importar un CSV nuevo, ver _procesar_csv) y no en cada
    apertura del panel. Devuelve el dict de opciones o levanta ValueError
    si la tabla no existe."""
    with get_db(DB_PATH, row_factory=False) as con:
        existe = con.execute(
            "SELECT name FROM sqlite_master WHERE type='table' AND name=?", (tabla,)).fetchone()
        if not existe:
            raise ValueError(f"Tabla {tabla} no encontrada.")

        def _distinct(col):
            try:
                rows = con.execute(
                    f"SELECT DISTINCT {col} FROM {tabla} "
                    f"WHERE {col} IS NOT NULL AND TRIM({col}) != '' ORDER BY {col}").fetchall()
                return [r[0] for r in rows]
            except Exception:
                # La columna puede no existir en tablas de años viejos con
                # schema distinto -- se devuelve vacío, no se rompe el resto
                # del panel por una sola columna.
                return []

        return {
            "est_mic": _distinct("EST_MIC"),
            "ult_estado": _distinct("ULT_ESTADO"),
            "var_control": _distinct("VAR_CONTROL"),
        }


def _recalcular_opciones_dat(tabla):
    """Recalcula y guarda en HIST_DB (persistente, no en memoria -- sobrevive
    a un restart) las opciones de filtro de una tabla DAT_<año>. Se llama
    automáticamente al final de un import exitoso (_procesar_csv) y también
    se puede disparar a mano desde el panel (botón "Recalcular ahora")."""
    datos = _calcular_opciones_dat(tabla)
    with get_db(HIST_DB) as con:
        for campo, valores in datos.items():
            con.execute(
                "INSERT INTO sintia_dat_opciones (tabla, campo, valores, actualizado) VALUES (?,?,?,datetime('now')) "
                "ON CONFLICT(tabla, campo) DO UPDATE SET valores=excluded.valores, actualizado=excluded.actualizado",
                (tabla, campo, json.dumps(valores)))
    return datos


@app.route("/api/sintia/dat_opciones")
@login_required
@modulo_required("sintia")
def sintia_dat_opciones():
    tabla, anio = _resolver_tabla_dat(request.args.to_dict())

    try:
        with get_db(HIST_DB, row_factory=True) as con:
            filas = con.execute(
                "SELECT campo, valores, actualizado FROM sintia_dat_opciones WHERE tabla=?", (tabla,)).fetchall()

        if filas:
            datos = {r["campo"]: json.loads(r["valores"]) for r in filas}
            actualizado = filas[0]["actualizado"]
            return jsonify({"ok": True, "anio": anio, "cached": True,
                            "actualizado": actualizado, **datos})

        # Primera vez que se pide esta tabla y todavía no se calculó nunca
        # (ej. datos importados antes de que existiera este caché) -- se
        # calcula ahora mismo y se guarda, para no volver a pagar el costo
        # la próxima vez.
        if not os.path.exists(DB_PATH):
            return jsonify({"ok": False, "error": "BD no cargada."})
        datos = _recalcular_opciones_dat(tabla)
        return jsonify({"ok": True, "anio": anio, "cached": False, "actualizado": None, **datos})
    except ValueError as e:
        return jsonify({"ok": False, "error": str(e)})
    except Exception as e:
        logging.error(f"SINTIA DAT OPCIONES ERROR | {e}")
        return jsonify({"ok": False, "error": str(e)})


@app.route("/api/sintia/dat_opciones/recalcular", methods=["POST"])
@login_required
@modulo_required("sintia")
def sintia_dat_opciones_recalcular():
    """Fuerza el recálculo ahora mismo, sin esperar al próximo import --
    útil justo después de activar esta función por primera vez, o si se
    tocaron los datos por fuera del flujo normal de import."""
    tabla, anio = _resolver_tabla_dat(request.args.to_dict())
    if not os.path.exists(DB_PATH):
        return jsonify({"ok": False, "error": "BD no cargada."})
    try:
        datos = _recalcular_opciones_dat(tabla)
        logging.info(f"SINTIA DAT OPCIONES | recalculado a mano por user={session.get('username')} | tabla={tabla}")
        return jsonify({"ok": True, "anio": anio, **datos})
    except ValueError as e:
        return jsonify({"ok": False, "error": str(e)})
    except Exception as e:
        logging.error(f"SINTIA DAT OPCIONES RECALCULAR ERROR | {e}")
        return jsonify({"ok": False, "error": str(e)})


@app.route("/api/sintia/dashboard")
@login_required
@modulo_required("sintia")
def sintia_dashboard():
    if not os.path.exists(DB_PATH):
        return jsonify({"ok": False, "error": "BD no cargada."})
    import datetime as _dt
    try:
        with get_db(DB_PATH, timeout=10) as con:
            cur = con.cursor()
            anio = _dt.date.today().year
            tabla = f"DAT_{anio}"
            cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name=?", (tabla,))
            if not cur.fetchone():
                return jsonify({"ok": False, "error": f"Tabla {tabla} no encontrada."})

            hoy = _dt.date.today()
            hoy_iso = hoy.isoformat()
            d30 = (hoy - _dt.timedelta(days=30)).isoformat()
            d60 = (hoy - _dt.timedelta(days=60)).isoformat()
            d7  = (hoy - _dt.timedelta(days=7)).isoformat()

            cur.execute(f"SELECT COUNT(*) FROM {tabla} WHERE FECHA_INGRESO_ISO BETWEEN ? AND ?", (d30, hoy_iso))
            total_mes = cur.fetchone()[0]
            cur.execute(f"SELECT COUNT(*) FROM {tabla} WHERE FECHA_INGRESO_ISO BETWEEN ? AND ?", (d60, d30))
            total_ant = cur.fetchone()[0]

            paises = {"BO": "Bolivia", "PY": "Paraguay", "BR": "Brasil", "CL": "Chile", "UY": "Uruguay"}
            por_pais = {}
            for cod, nombre in paises.items():
                cur.execute(f"SELECT COUNT(*) FROM {tabla} WHERE FECHA_INGRESO_ISO BETWEEN ? AND ? AND MIC LIKE ?", (d30, hoy_iso, f"%{cod}%"))
                cnt = cur.fetchone()[0]
                if cnt: por_pais[nombre] = cnt

            cur.execute(f"SELECT EST_MIC, COUNT(*) as cnt FROM {tabla} WHERE FECHA_INGRESO_ISO BETWEEN ? AND ? GROUP BY EST_MIC ORDER BY cnt DESC", (d30, hoy_iso))
            por_estado = {r[0]: r[1] for r in cur.fetchall() if r[0]}

            rechazos_mes = 0
            rechazos_7d  = 0
            alerta_rechazos = False
            cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='RECHAZOS'")
            if cur.fetchone():
                cur.execute("SELECT COUNT(*) FROM RECHAZOS WHERE Fecha_ISO BETWEEN ? AND ?", (d30, hoy_iso))
                rechazos_mes = cur.fetchone()[0]
                cur.execute("SELECT COUNT(*) FROM RECHAZOS WHERE Fecha_ISO BETWEEN ? AND ?", (d7, hoy_iso))
                rechazos_7d = cur.fetchone()[0]
                prom_30 = rechazos_mes / 30
                prom_7  = rechazos_7d / 7
                alerta_rechazos = prom_7 > prom_30

            cur.execute(f"SELECT ADUANA, COUNT(*) as cnt FROM {tabla} WHERE FECHA_INGRESO_ISO BETWEEN ? AND ? GROUP BY ADUANA ORDER BY cnt DESC LIMIT 6", (d30, hoy_iso))
            top_aduanas = {r[0]: r[1] for r in cur.fetchall() if r[0]}

            cur.execute(f"SELECT MAX(FECHA_INGRESO_ISO) FROM {tabla}")
            ultima_fecha_bd = cur.fetchone()[0] or ''
            dias_sin_actualizar = 0
            if ultima_fecha_bd:
                try:
                    uf = _dt.date.fromisoformat(ultima_fecha_bd[:10])
                    dias_sin_actualizar = (hoy - uf).days
                except: pass

            evolucion = []
            for i in range(5, -1, -1):
                d = (hoy.replace(day=1) - _dt.timedelta(days=1)) if i > 0 else hoy
                for _ in range(i - 1):
                    d = (d.replace(day=1) - _dt.timedelta(days=1))
                mes = d.strftime("%Y-%m")
                label = d.strftime("%b %Y")
                try:
                    cur.execute("SELECT COUNT(*) FROM RECHAZOS WHERE Fecha_ISO LIKE ?", (f"{mes}%",))
                    cnt = cur.fetchone()[0]
                except: cnt = 0
                evolucion.append({"mes": label, "total": cnt})

        return jsonify({
            "ok": True,
            "label_actual": "Últimos 30 días",
            "label_ant": "30 días anteriores",
            "total_mes": total_mes,
            "total_ant": total_ant,
            "variacion": round((total_mes - total_ant) / total_ant * 100, 1) if total_ant else 0,
            "por_pais": por_pais,
            "por_estado": por_estado,
            "rechazos_mes": rechazos_mes,
            "rechazos_7d": rechazos_7d,
            "alerta_rechazos": alerta_rechazos,
            "top_aduanas": top_aduanas,
            "dias_sin_actualizar": dias_sin_actualizar,
            "evolucion": evolucion,
        })
    except Exception as e:
        logging.error(f"SINTIA DASHBOARD ERROR | {e}")
        return jsonify({"ok": False, "error": str(e)})

@app.route("/api/sintia/dat", methods=["POST"])
@login_required
@modulo_required("sintia")
def sintia_dat_query():
    if not os.path.exists(DB_PATH):
        return jsonify({"ok": False, "error": "BD no cargada."})
    p = request.json or {}

    conditions, params = [], []

    if p.get("fec_d"):
        conditions.append("FECHA_INGRESO_ISO >= ?"); params.append(p["fec_d"])
    if p.get("fec_h"):
        conditions.append("FECHA_INGRESO_ISO <= ?"); params.append(p["fec_h"])
    if p.get("ftx_d"):
        conditions.append("FECHA_TRANS_ISO >= ?");  params.append(p["ftx_d"])
    if p.get("ftx_h"):
        conditions.append("FECHA_TRANS_ISO <= ?");  params.append(p["ftx_h"])
    if p.get("aduana"):
        conditions.append("ADUANA = ?");            params.append(p["aduana"])
    if p.get("tipo"):
        conditions.append("TIPO_REGISTRO = ?");     params.append(p["tipo"])
    if p.get("cargado"):
        conditions.append("CARGADO = ?");           params.append(p["cargado"])
    if p.get("pais"):
        conditions.append("MIC LIKE ?");            params.append(f"%{p['pais']}%")
    if p.get("mic"):
        conditions.append("MIC LIKE ?");            params.append(f"%{p['mic']}%")
    if p.get("lot"):
        conditions.append("LOT LIKE ?");            params.append(f"%{p['lot']}%")
    if p.get("empresa"):
        conditions.append("EMPRESA LIKE ?");        params.append(f"%{p['empresa']}%")
    if p.get("vehiculo"):
        conditions.append("(TRACTOR LIKE ? OR SEMI LIKE ?)");
        params += [f"%{p['vehiculo']}%", f"%{p['vehiculo']}%"]
    if p.get("est"):
        conditions.append("EST_MIC = ?");           params.append(p["est"])
    if p.get("ult_estado"):
        conditions.append("ULT_ESTADO = ?");        params.append(p["ult_estado"])
    if p.get("var_control"):
        conditions.append("VAR_CONTROL = ?");       params.append(p["var_control"])
    if p.get("novedad"):
        conditions.append("tiene_novedad = ?");     params.append(p["novedad"])

    where = ("WHERE " + " AND ".join(conditions)) if conditions else ""

    # Tabla según año elegido (por defecto, el actual)
    tabla, anio = _resolver_tabla_dat(p)

    try:
        with get_db(DB_PATH, row_factory=True) as con:
            cur = con.cursor()
            # Verificar que la tabla existe
            cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name=?", (tabla,))
            if not cur.fetchone():
                return jsonify({"ok": False, "error": f"Tabla {tabla} no encontrada en la BD."})

            sql = f"SELECT * FROM {tabla} {where} LIMIT ? OFFSET ?"
            page_size = int(p.get("page_size", 500))
            offset    = int(p.get("offset", 0))
            cur.execute(sql, params + [page_size + 1, offset])
            rows_raw = cur.fetchall()

            cols = list(rows_raw[0].keys()) if rows_raw else []
            truncated = len(rows_raw) > page_size
            rows = [list(r) for r in rows_raw[:page_size]]

            resumen = {"por_pais": {}, "por_estado": {}, "total": 0}
            try:
                cur.execute(f"SELECT COUNT(*) FROM {tabla} {where}", params)
                resumen["total"] = cur.fetchone()[0]
                paises = {"Bolivia": "BO", "Paraguay": "PY", "Brasil": "BR", "Chile": "CL", "Uruguay": "UY"}
                for nombre, cod in paises.items():
                    cur.execute(f"SELECT COUNT(*) FROM {tabla} {where} {'AND' if conditions else 'WHERE'} MIC LIKE ?", params + [f"%{cod}%"])
                    cnt = cur.fetchone()[0]
                    if cnt: resumen["por_pais"][nombre] = cnt
                cur.execute(f"SELECT EST_MIC, COUNT(*) as cnt FROM {tabla} {where} GROUP BY EST_MIC ORDER BY cnt DESC", params)
                resumen["por_estado"] = {r[0]: r[1] for r in cur.fetchall()}
            except Exception as e:
                logging.warning(f"DAT RESUMEN ERROR | {e}")

        return jsonify({"ok": True, "cols": cols, "rows": rows, "truncated": truncated,
                        "resumen": resumen, "offset": offset, "page_size": page_size})

    except Exception as e:
        logging.error(f"DAT QUERY ERROR | {e}")
        return jsonify({"ok": False, "error": str(e)})


@app.route("/api/sintia/rec", methods=["POST"])
@login_required
@modulo_required("sintia")
def sintia_rec_query():
    if not os.path.exists(DB_PATH):
        return jsonify({"ok": False, "error": "BD no cargada."})
    p = request.json or {}

    conditions, params = [], []

    if p.get("fec_d"):
        conditions.append("Fecha_ISO >= ?");   params.append(p["fec_d"])
    if p.get("fec_h"):
        conditions.append("Fecha_ISO <= ?");   params.append(p["fec_h"])
    if p.get("pais"):
        conditions.append("PaisEmisor = ?");   params.append(p["pais"])
    if p.get("anio"):
        conditions.append("Anio = ?");         params.append(str(p["anio"]))
    if p.get("nromic"):
        conditions.append("NroMic LIKE ?");    params.append(f"%{p['nromic']}%")
    if p.get("mensaje"):
        conditions.append("Mensaje LIKE ?");   params.append(f"%{p['mensaje']}%")
    if p.get("metodo"):
        conditions.append("Metodo = ?");       params.append(p["metodo"])

    where = ("WHERE " + " AND ".join(conditions)) if conditions else ""

    try:
        with get_db(DB_PATH, row_factory=True) as con:
            cur = con.cursor()
            page_size = int(p.get("page_size", 500))
            offset    = int(p.get("offset", 0))
            sql = f"SELECT PaisEmisor, Anio, NroMic, Fecha_ISO, Mes, Metodo, Mensaje FROM RECHAZOS {where} ORDER BY Fecha_ISO DESC LIMIT ? OFFSET ?"
            cur.execute(sql, params + [page_size + 1, offset])
            rows_raw = cur.fetchall()

            cols = list(rows_raw[0].keys()) if rows_raw else []
            truncated = len(rows_raw) > page_size
            rows = [list(r) for r in rows_raw[:page_size]]

            resumen = {"por_pais": {}, "por_metodo": {}, "total": 0}
            try:
                cur.execute(f"SELECT COUNT(*) FROM RECHAZOS {where}", params)
                resumen["total"] = cur.fetchone()[0]
                cur.execute(f"SELECT PaisEmisor, COUNT(*) as cnt FROM RECHAZOS {where} GROUP BY PaisEmisor ORDER BY cnt DESC", params)
                resumen["por_pais"] = {r[0]: r[1] for r in cur.fetchall()}
                cur.execute(f"SELECT Metodo, COUNT(*) as cnt FROM RECHAZOS {where} GROUP BY Metodo ORDER BY cnt DESC LIMIT 6", params)
                resumen["por_metodo"] = {r[0]: r[1] for r in cur.fetchall()}
            except Exception as e:
                logging.warning(f"REC RESUMEN ERROR | {e}")

        return jsonify({"ok": True, "cols": cols, "rows": rows, "truncated": truncated,
                        "resumen": resumen, "offset": offset, "page_size": page_size})

    except Exception as e:
        logging.error(f"REC QUERY ERROR | {e}")
        return jsonify({"ok": False, "error": str(e)})


# _exportar_xlsx ahora vive en core.py (la comparten app.py y blueprints/finanzas.py).


@app.route("/api/sintia/dat/export", methods=["POST"])
@login_required
@modulo_required("sintia")
def sintia_dat_export():
    """Re-ejecuta la query sin LIMIT para exportar todos los resultados."""
    if not os.path.exists(DB_PATH):
        return jsonify({"ok": False, "error": "BD no cargada."})
    p = request.json or {}

    conditions, params = [], []
    if p.get("fec_d"):   conditions.append("FECHA_INGRESO_ISO >= ?"); params.append(p["fec_d"])
    if p.get("fec_h"):   conditions.append("FECHA_INGRESO_ISO <= ?"); params.append(p["fec_h"])
    if p.get("ftx_d"):   conditions.append("FECHA_TRANS_ISO >= ?");   params.append(p["ftx_d"])
    if p.get("ftx_h"):   conditions.append("FECHA_TRANS_ISO <= ?");   params.append(p["ftx_h"])
    if p.get("aduana"):  conditions.append("ADUANA = ?");             params.append(p["aduana"])
    if p.get("tipo"):    conditions.append("TIPO_REGISTRO = ?");      params.append(p["tipo"])
    if p.get("cargado"): conditions.append("CARGADO = ?");            params.append(p["cargado"])
    if p.get("pais"):    conditions.append("MIC LIKE ?");             params.append(f"%{p['pais']}%")
    if p.get("mic"):     conditions.append("MIC LIKE ?");             params.append(f"%{p['mic']}%")
    if p.get("lot"):     conditions.append("LOT LIKE ?");             params.append(f"%{p['lot']}%")
    if p.get("empresa"): conditions.append("EMPRESA LIKE ?");         params.append(f"%{p['empresa']}%")
    if p.get("vehiculo"):
        conditions.append("(TRACTOR LIKE ? OR SEMI LIKE ?)");
        params += [f"%{p['vehiculo']}%", f"%{p['vehiculo']}%"]
    if p.get("est"):     conditions.append("EST_MIC = ?");            params.append(p["est"])
    if p.get("ult_estado"): conditions.append("ULT_ESTADO = ?");      params.append(p["ult_estado"])
    if p.get("var_control"): conditions.append("VAR_CONTROL = ?");    params.append(p["var_control"])
    if p.get("novedad"): conditions.append("tiene_novedad = ?");      params.append(p["novedad"])

    where = ("WHERE " + " AND ".join(conditions)) if conditions else ""
    tabla, anio = _resolver_tabla_dat(p)

    try:
        with get_db(DB_PATH, row_factory=True) as con:
            cur = con.cursor()
            cur.execute(f"SELECT * FROM {tabla} {where}", params)
            rows_raw = cur.fetchall()
        cols = list(rows_raw[0].keys()) if rows_raw else []
        rows = [list(r) for r in rows_raw]
        buf = _exportar_xlsx(cols, rows)
        return send_file(buf, as_attachment=True,
                         download_name=f"DAT_{anio}_consulta.xlsx",
                         mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    except Exception as e:
        logging.error(f"DAT EXPORT ERROR | {e}")
        return jsonify({"ok": False, "error": str(e)})


@app.route("/api/sintia/rec/export", methods=["POST"])
@login_required
@modulo_required("sintia")
def sintia_rec_export():
    """Re-ejecuta la query sin LIMIT para exportar todos los resultados."""
    if not os.path.exists(DB_PATH):
        return jsonify({"ok": False, "error": "BD no cargada."})
    p = request.json or {}

    conditions, params = [], []
    if p.get("fec_d"):   conditions.append("Fecha_ISO >= ?");  params.append(p["fec_d"])
    if p.get("fec_h"):   conditions.append("Fecha_ISO <= ?");  params.append(p["fec_h"])
    if p.get("pais"):    conditions.append("PaisEmisor = ?");  params.append(p["pais"])
    if p.get("anio"):    conditions.append("Anio = ?");        params.append(str(p["anio"]))
    if p.get("nromic"):  conditions.append("NroMic LIKE ?");   params.append(f"%{p['nromic']}%")
    if p.get("mensaje"): conditions.append("Mensaje LIKE ?");  params.append(f"%{p['mensaje']}%")
    if p.get("metodo"):  conditions.append("Metodo = ?");      params.append(p["metodo"])

    where = ("WHERE " + " AND ".join(conditions)) if conditions else ""

    try:
        with get_db(DB_PATH, row_factory=True) as con:
            cur = con.cursor()
            cur.execute(f"SELECT PaisEmisor, Anio, NroMic, Fecha_ISO, Mes, Metodo, Mensaje FROM RECHAZOS {where} ORDER BY Fecha_ISO DESC", params)
            rows_raw = cur.fetchall()
        cols = list(rows_raw[0].keys()) if rows_raw else []
        rows = [list(r) for r in rows_raw]
        buf = _exportar_xlsx(cols, rows)
        return send_file(buf, as_attachment=True,
                         download_name="RECHAZOS_consulta.xlsx",
                         mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    except Exception as e:
        logging.error(f"REC EXPORT ERROR | {e}")
        return jsonify({"ok": False, "error": str(e)})


_MESES_CORTOS = ["Ene", "Feb", "Mar", "Abr", "May", "Jun", "Jul", "Ago", "Sep", "Oct", "Nov", "Dic"]

def _mes_label_corto(clave_yyyy_mm):
    """'2026-07' -> 'Jul 2026' -- para encabezados de columna legibles."""
    try:
        y, m = clave_yyyy_mm.split("-")
        return f"{_MESES_CORTOS[int(m) - 1]} {y}"
    except Exception:
        return clave_yyyy_mm


# Se sube a mano cada vez que cambia algo en cómo se arma el informe de
# Aduanas del País (Word/Excel) -- entra en la clave de caché de 24hs
# (_buscar_informe_aduanas_cacheado) para que un deploy nuevo invalide sola
# la caché vieja, y se imprime en el propio Word para poder confirmar de un
# vistazo con qué versión del generador se armó un archivo dado (encontrado
# en la práctica: sin esto, pedir "el mismo informe" después de actualizar
# el código devolvía el archivo cacheado de la versión anterior, sin
# ninguna señal de que el código sí se había actualizado).
_INFORME_ADUANAS_VERSION = "v3-2026-07-13"


def _formatear_demora(dias):
    """Convierte una demora en días (float, con fracción) a un texto legible
    en horas/minutos/segundos -- 0.75 días no dice nada de un vistazo, pero
    '18h 00m 00s' sí. Se guarda internamente en días (preciso para promediar
    y ordenar) y se formatea recién acá, al final, para mostrar."""
    if dias is None:
        return None
    total_seg = round(dias * 86400)
    h, resto = divmod(total_seg, 3600)
    m, s = divmod(resto, 60)
    return f"{h}h {m:02d}m {s:02d}s"


# FECHA_ULT_INT viene como "DD-MM-YYYY HH:MM:SS" (formato argentino), NO
# como "YYYY-MM-DD..." -- julianday() de SQLite solo reconoce este último y
# devuelve NULL en silencio con el otro (sin error, por eso el bug de la
# demora en 0h 00m 00s pasó desapercibido). Se reconstruye a mano con
# substr() antes de pasarlo a julianday(). El CASE cubre los dos formatos
# por si en algún momento conviven filas viejas ya en ISO con nuevas en
# DD-MM-YYYY. A nivel de módulo porque la usan varias funciones (tabla de
# aduanas, evolución mensual).
_FECHA_ULT_INT_ISO = (
    "(CASE WHEN FECHA_ULT_INT LIKE '____-__-__%' THEN FECHA_ULT_INT "
    "ELSE substr(FECHA_ULT_INT,7,4) || '-' || substr(FECHA_ULT_INT,4,2) "
    "|| '-' || substr(FECHA_ULT_INT,1,2) || substr(FECHA_ULT_INT,11) END)"
)


def _aduanas_nacional_datos(anio, dira_filtro=None, umbral_alerta_dias=10):
    """Arma la tabla de aduanas + indicadores nacionales.

    DAT_<año>/RECHAZOS viven en DB_PATH (se reemplazan enteras con cada
    import de PAD — ver _procesar_csv), pero ref_aduanas/ref_dira viven en
    HIST_DB (datos de referencia curados a mano, a propósito en una base
    aparte para no arriesgarlos si algún día se reimporta DB_PATH).
    SQLite no puede hacer JOIN entre archivos distintos sin ATTACH DATABASE,
    así que esto se resuelve en Python: se agrega por aduana en DB_PATH,
    se trae el catálogo de aduanas/DIRA de HIST_DB, y se cruzan acá.

    Levanta ValueError con un mensaje prolijo si la tabla del año pedido no
    existe. Devuelve (filas, indicadores, diras).
    """
    tabla = f"DAT_{anio}"
    if not os.path.exists(DB_PATH):
        raise ValueError("BD no cargada.")

    with get_db(DB_PATH, row_factory=True) as con:
        existe = con.execute(
            "SELECT name FROM sqlite_master WHERE type='table' AND name=?", (tabla,)).fetchone()
        if not existe:
            raise ValueError(f"Tabla {tabla} no encontrada.")

        # Sin ROUND acá: 2 decimales de día son ~14 minutos de error, y ahora
        # se muestra con precisión de segundos -- el redondeo se hace recién
        # al formatear para mostrar, no antes de promediar.
        #
        # El umbral (umbral_alerta_dias) ahora se aplica simétricamente:
        # - operaciones que TODAVÍA no llegaron a SAL y ya superaron el
        #   umbral desde el ingreso -> en_alerta_bandeja (como antes).
        # - operaciones que YA llegaron a SAL pero tardaron más que el
        #   umbral -> en_alerta_demora_larga (nuevo). Antes estas SÍ
        #   entraban en el promedio y lo arrastraban para arriba -- un caso
        #   de 101 horas con umbral de 2 días (48hs) seguía promediándose
        #   junto con casos normales de unas pocas horas.
        _demora_expr = f"(julianday({_FECHA_ULT_INT_ISO}) - julianday(FECHA_INGRESO_ISO))"
        agregados = [dict(r) for r in con.execute(f"""
            SELECT
                ADUANA AS aduana_cod,
                COUNT(*) AS total_operaciones,
                SUM(CASE WHEN ULT_ESTADO = 'SAL' THEN 1 ELSE 0 END) AS total_sali,
                SUM(CASE WHEN ULT_ESTADO = 'SAL' AND {_demora_expr} <= ?
                    THEN 1 ELSE 0 END) AS sali_dentro_umbral,
                AVG(CASE WHEN ULT_ESTADO = 'SAL' AND {_demora_expr} <= ?
                    THEN {_demora_expr} END) AS demora_media_dias,
                SUM(CASE WHEN ULT_ESTADO != 'SAL'
                    AND (julianday('now') - julianday(FECHA_INGRESO_ISO)) > ?
                    THEN 1 ELSE 0 END) AS en_alerta_bandeja,
                SUM(CASE WHEN ULT_ESTADO = 'SAL' AND {_demora_expr} > ?
                    THEN 1 ELSE 0 END) AS en_alerta_demora_larga
            FROM {tabla}
            GROUP BY ADUANA
        """, [umbral_alerta_dias, umbral_alerta_dias, umbral_alerta_dias, umbral_alerta_dias]).fetchall()]

    with get_db(HIST_DB, row_factory=True) as con:
        cat_aduanas = {r["cod"]: dict(r) for r in con.execute(
            "SELECT cod, nombre, indice_dira FROM ref_aduanas").fetchall()}
        cat_diras = {r["indice"]: r["nombre"] for r in con.execute(
            "SELECT indice, nombre FROM ref_dira").fetchall()}
        diras = [dict(r) for r in con.execute(
            "SELECT indice, nombre FROM ref_dira ORDER BY nombre").fetchall()]

    filas = []
    for a in agregados:
        cod = a["aduana_cod"]
        info = cat_aduanas.get(cod)
        dira_indice = info["indice_dira"] if info else None
        dira_nombre = cat_diras.get(dira_indice, "Sin DIRA asignada") if dira_indice else "Sin DIRA asignada"
        aduana_nombre = info["nombre"] if info else f"{cod} (sin nombre en ref_aduanas)"
        filas.append({
            "aduana_cod": cod,
            "aduana_nombre": aduana_nombre,
            "dira_indice": dira_indice,
            "dira_nombre": dira_nombre,
            "total_operaciones": a["total_operaciones"],
            "total_sali": a["total_sali"],
            "sali_dentro_umbral": a["sali_dentro_umbral"],
            "demora_media_dias": a["demora_media_dias"],
            "demora_media_fmt": _formatear_demora(a["demora_media_dias"]),
            "en_alerta_bandeja": a["en_alerta_bandeja"],
            "en_alerta_demora_larga": a["en_alerta_demora_larga"],
            "en_alerta_total": a["en_alerta_bandeja"] + a["en_alerta_demora_larga"],
        })

    if dira_filtro:
        filas = [f for f in filas if f["dira_indice"] == dira_filtro]
    filas.sort(key=lambda f: (f["dira_nombre"], f["aduana_nombre"]))

    total_operaciones = sum(f["total_operaciones"] for f in filas)
    total_sali = sum(f["total_sali"] for f in filas)
    en_alerta_bandeja = sum(f["en_alerta_bandeja"] for f in filas)
    en_alerta_demora_larga = sum(f["en_alerta_demora_larga"] for f in filas)
    sali_dentro_umbral_total = sum(f["sali_dentro_umbral"] for f in filas)
    # Promedio ponderado por cantidad de SALI-dentro-del-umbral por aduana
    # (no por total_sali -- las que superaron el umbral quedan afuera del
    # promedio, así que tampoco deben pesar en él). Evita además el sesgo
    # de promediar promedios simples entre aduanas con distinto volumen.
    suma_ponderada = sum((f["demora_media_dias"] or 0) * f["sali_dentro_umbral"] for f in filas)
    demora_media_nacional = (suma_ponderada / sali_dentro_umbral_total) if sali_dentro_umbral_total else None

    indicadores = {
        "total_operaciones": total_operaciones,
        "total_sali": total_sali,
        "sali_dentro_umbral": sali_dentro_umbral_total,
        "demora_media_dias": demora_media_nacional,
        "demora_media_fmt": _formatear_demora(demora_media_nacional),
        "en_alerta_bandeja": en_alerta_bandeja,
        "en_alerta_demora_larga": en_alerta_demora_larga,
        "en_alerta_total": en_alerta_bandeja + en_alerta_demora_larga,
    }
    return filas, indicadores, diras


def _aduanas_codigos_de_dira(dira_filtro):
    """Códigos de aduana (ref_aduanas.cod) que pertenecen a una DIRA -- para
    filtrar DAT_<año> por DIRA sin poder hacer JOIN entre bases (ver nota en
    _aduanas_nacional_datos)."""
    with get_db(HIST_DB, row_factory=True) as con:
        return {r["cod"] for r in con.execute(
            "SELECT cod FROM ref_aduanas WHERE indice_dira=?", (dira_filtro,)).fetchall()}


def _evolucion_mensual_nacional(anio, dira_filtro=None, umbral_alerta_dias=10, meses=6):
    """Serie mensual (últimos N meses) de demora media agregada -- nacional,
    o de la DIRA filtrada si se pasa una. Pensada para el gráfico en
    pantalla: UNA sola serie, no una por aduana (con 50+ aduanas un gráfico
    con una línea por cada una sería ilegible). El desglose fino por aduana
    va solo al Excel -- ver _evolucion_mensual_por_aduana().
    """
    tabla = f"DAT_{anio}"
    if not os.path.exists(DB_PATH):
        raise ValueError("BD no cargada.")

    aduanas_permitidas = _aduanas_codigos_de_dira(dira_filtro) if dira_filtro else None
    if dira_filtro and not aduanas_permitidas:
        return []

    filtro_aduanas_sql = ""
    params = [umbral_alerta_dias]
    if aduanas_permitidas:
        filtro_aduanas_sql = f"AND ADUANA IN ({','.join('?' * len(aduanas_permitidas))})"
        params += list(aduanas_permitidas)

    with get_db(DB_PATH, row_factory=True) as con:
        existe = con.execute(
            "SELECT name FROM sqlite_master WHERE type='table' AND name=?", (tabla,)).fetchone()
        if not existe:
            raise ValueError(f"Tabla {tabla} no encontrada.")

        _demora_expr = f"(julianday({_FECHA_ULT_INT_ISO}) - julianday(FECHA_INGRESO_ISO))"
        rows = con.execute(f"""
            SELECT substr(FECHA_INGRESO_ISO,1,7) AS mes,
                   AVG(CASE WHEN ULT_ESTADO = 'SAL' AND {_demora_expr} <= ?
                       THEN {_demora_expr} END) AS demora_media_dias
            FROM {tabla}
            WHERE 1=1 {filtro_aduanas_sql}
            GROUP BY mes
            ORDER BY mes
        """, params).fetchall()

    por_mes = {r["mes"]: r["demora_media_dias"] for r in rows if r["mes"]}

    # Completar los últimos N meses aunque no haya datos en alguno (mejor un
    # hueco visible en el gráfico que un mes salteado sin explicación).
    hoy = date.today()
    serie = []
    for i in range(meses - 1, -1, -1):
        y, m = hoy.year, hoy.month - i
        while m <= 0:
            m += 12; y -= 1
        clave = f"{y:04d}-{m:02d}"
        dias = por_mes.get(clave)
        serie.append({
            "mes": clave,
            "demora_media_dias": dias,
            "demora_media_fmt": _formatear_demora(dias),
        })
    return serie


def _evolucion_mensual_por_aduana(anio, dira_filtro=None, umbral_alerta_dias=10, meses=6):
    """Igual que _evolucion_mensual_nacional() pero desglosado por aduana --
    se usa en el export a Excel (pivot con TODAS las aduanas) y en el Word
    (tabla + gráfico de cada aduana en alerta). Cada mes trae demora Y
    cantidad de operaciones -- el gráfico de cada aduana en el Word combina
    las dos (línea de demora + barras de operaciones en un eje aparte) para
    poder ver de un vistazo si un pico de demora coincide con un pico de
    volumen, sin tener que ir a la tabla principal."""
    tabla = f"DAT_{anio}"
    if not os.path.exists(DB_PATH):
        raise ValueError("BD no cargada.")

    aduanas_permitidas = _aduanas_codigos_de_dira(dira_filtro) if dira_filtro else None
    if dira_filtro and not aduanas_permitidas:
        return [], []

    filtro_aduanas_sql = ""
    params = [umbral_alerta_dias]
    if aduanas_permitidas:
        filtro_aduanas_sql = f"AND ADUANA IN ({','.join('?' * len(aduanas_permitidas))})"
        params += list(aduanas_permitidas)

    with get_db(DB_PATH, row_factory=True) as con:
        existe = con.execute(
            "SELECT name FROM sqlite_master WHERE type='table' AND name=?", (tabla,)).fetchone()
        if not existe:
            raise ValueError(f"Tabla {tabla} no encontrada.")

        _demora_expr = f"(julianday({_FECHA_ULT_INT_ISO}) - julianday(FECHA_INGRESO_ISO))"
        rows = con.execute(f"""
            SELECT ADUANA AS aduana_cod, substr(FECHA_INGRESO_ISO,1,7) AS mes,
                   COUNT(*) AS operaciones,
                   AVG(CASE WHEN ULT_ESTADO = 'SAL' AND {_demora_expr} <= ?
                       THEN {_demora_expr} END) AS demora_media_dias
            FROM {tabla}
            WHERE 1=1 {filtro_aduanas_sql}
            GROUP BY ADUANA, mes
        """, params).fetchall()

    with get_db(HIST_DB, row_factory=True) as con:
        cat_aduanas = {r["cod"]: r["nombre"] for r in con.execute("SELECT cod, nombre FROM ref_aduanas").fetchall()}

    hoy = date.today()
    meses_cols = []
    for i in range(meses - 1, -1, -1):
        y, m = hoy.year, hoy.month - i
        while m <= 0:
            m += 12; y -= 1
        meses_cols.append(f"{y:04d}-{m:02d}")

    por_aduana = {}
    for r in rows:
        cod = r["aduana_cod"]
        por_aduana.setdefault(cod, {})[r["mes"]] = {
            "demora": r["demora_media_dias"], "operaciones": r["operaciones"]}

    filas = []
    for cod, valores_por_mes in sorted(por_aduana.items()):
        fila = {"aduana_cod": cod, "aduana_nombre": cat_aduanas.get(cod, f"{cod} (sin nombre en ref_aduanas)")}
        for mes in meses_cols:
            fila[mes] = valores_por_mes.get(mes, {"demora": None, "operaciones": 0})
        filas.append(fila)
    filas.sort(key=lambda f: f["aduana_nombre"])
    return filas, meses_cols


@app.route("/api/sintia/aduanas_nacional")
@login_required
@modulo_required("sintia")
def sintia_aduanas_nacional():
    """Indicadores + tabla de todas las aduanas del país, con demora media de
    desaduanamiento (FECHA_ULT_INT - FECHA_INGRESO_ISO para las que llegaron
    a ULT_ESTADO='SAL'). Ver _aduanas_nacional_datos() para el detalle de
    por qué esto cruza dos bases SQLite distintas.

    Supuestos sobre el schema de PAD que hay que confirmar contra datos
    reales antes de usar esto para algo que se cite hacia afuera (ver
    conversación 09-10/07/2026):
      - DAT_<año>.ULT_ESTADO = 'SAL' marca que la operación salió.
      - DAT_<año>.FECHA_ULT_INT = fecha del último movimiento/estado.
      - DAT_<año>.ADUANA coincide en formato con ref_aduanas.cod.
        Si no matchea, la fila igual aparece pero con "(sin nombre en
        ref_aduanas)" — señal de que el código no coincide.
    """
    anio = request.args.get("anio", str(date.today().year))
    dira_filtro = request.args.get("dira", "").strip() or None
    umbral_alerta_dias = int(request.args.get("umbral_dias", 10))

    try:
        filas, indicadores, diras = _aduanas_nacional_datos(anio, dira_filtro, umbral_alerta_dias)
        try:
            evolucion = _evolucion_mensual_nacional(anio, dira_filtro, umbral_alerta_dias)
        except Exception:
            evolucion = []  # que la falla de un gráfico secundario no rompa el resto del panel
        return jsonify({
            "ok": True,
            "anio": anio,
            "umbral_dias": umbral_alerta_dias,
            "indicadores": indicadores,
            "rows": filas,
            "diras": diras,
            "evolucion_mensual": evolucion,
        })
    except ValueError as e:
        return jsonify({"ok": False, "error": str(e)})
    except Exception as e:
        logging.error(f"SINTIA ADUANAS NACIONAL ERROR | {e}")
        return jsonify({"ok": False, "error": str(e)})


def _generar_narrativa_aduanas_ia(anio, dira_nombre, umbral_alerta_dias, indicadores, filas, evolucion, api_key):
    """Arma el prompt y llama a Claude para el análisis del informe de
    Aduanas del País. Mismo cuidado que generar_ia.py con el informe SINTIA:
    los únicos números válidos son los que se pasan acá, la IA no debe
    inventar ni recalcular nada.

    Ojo con el alcance: esta métrica mide tiempo entre registro y salida en
    PAD, sin desagregar la causa de la demora. No mide específicamente
    coordinación con SENASA ni ningún otro organismo puntual -- el prompt
    le pide a la IA que no haga esa atribución causal, porque los datos no
    la sostienen."""
    try:
        import anthropic, httpx
    except ImportError:
        raise RuntimeError("El paquete 'anthropic' no está instalado en el servidor.")

    top_alerta = sorted(filas, key=lambda f: f["en_alerta_total"], reverse=True)[:8]
    top_alerta_txt = "\n".join(
        f"  - {f['aduana_nombre']} ({f['dira_nombre']}): {f['en_alerta_total']} en alerta "
        f"({f['en_alerta_bandeja']} pendiente, {f['en_alerta_demora_larga']} salió tarde), "
        f"demora media {f['demora_media_fmt'] or 'sin datos'}"
        for f in top_alerta if f["en_alerta_total"] > 0
    ) or "  (ninguna aduana con alertas en este período)"

    evol_txt = "\n".join(
        f"  - {_mes_label_corto(p['mes'])}: {p['demora_media_fmt'] or 'sin datos'}"
        for p in evolucion
    )

    # Contexto específico para Misiones: hay un proyecto de ley (PYMES,
    # presentado por una diputada de Misiones) que afirma que la falta de
    # coordinación Aduana-SENASA supone un tiempo medio de liberación de
    # 24hs. La comparación numérica se calcula acá, en Python -- no se le
    # pide a la IA que haga la cuenta, solo que la redacte, para no
    # arriesgar un error de cálculo del modelo en algo potencialmente
    # citado hacia afuera.
    bloque_misiones = ""
    if "misiones" in (dira_nombre or "").lower():
        demora_dias = indicadores.get("demora_media_dias")
        if demora_dias is not None:
            demora_horas = demora_dias * 24
            comparacion = (
                f"la demora media medida en PAD ({indicadores['demora_media_fmt']}) es "
                f"{'MENOR' if demora_horas < 24 else 'MAYOR' if demora_horas > 24 else 'IGUAL'} "
                f"a esas 24hs (diferencia: {abs(24 - demora_horas):.1f} horas)."
            )
        else:
            comparacion = "no hay datos suficientes en este período para calcular esa comparación."
        bloque_misiones = f"""

CONTEXTO ADICIONAL (solo para este alcance, Misiones): existe un proyecto de ley de beneficios a
PYMES, presentado por una diputada de Misiones, que afirma que la falta de coordinación entre
Aduana y SENASA supone un tiempo medio de liberación de 24 horas. Comparación con lo medido acá: {comparacion}
IMPORTANTE: mencioná esta comparación en el análisis, pero dejá explícito que la métrica de PAD mide
tiempo total de registro a salida (no específicamente demora por falta de coordinación con SENASA,
como afirma el proyecto de ley) — son cosas relacionadas pero no exactamente lo mismo, así que la
comparación es orientativa, no una confirmación ni un desmentido directo de esa cifra."""

    prompt = f"""ADVERTENCIA: los datos numéricos de este prompt son los ÚNICOS válidos para este informe. No inventes, no redondees distinto, no calcules nada que no esté acá.

Sos un analista de procesos aduaneros de ARCA (Aduana Argentina), sección DI REPA. Redactá el análisis de un informe formal sobre tiempos de desaduanamiento a nivel país, para uso interno.

MUY IMPORTANTE sobre el alcance de la métrica: "demora" acá es exclusivamente el tiempo entre el registro de la operación en PAD (Plataforma Aduanera Digital) y su salida (estado SAL), calculado sobre datos de PAD. NO es una medición de demora específicamente atribuible a la coordinación con SENASA ni con ningún otro organismo puntual — los datos de PAD no permiten desagregar esa causa. Si mencionás alguna hipótesis sobre las causas de la demora, dejala explícitamente como hipótesis a confirmar con otras fuentes, nunca como un hecho que estos datos prueben.
{bloque_misiones}

CONTEXTO DEL REPORTE:
- Año de datos: {anio}
- Alcance: {dira_nombre}
- Umbral de alerta: {umbral_alerta_dias} días (operaciones que superan este tiempo, ya sea pendientes o recién salidas, se excluyen del promedio y se cuentan aparte como alerta)

INDICADORES NACIONALES (o del alcance filtrado):
- Operaciones totales: {indicadores['total_operaciones']}
- Salieron (SAL): {indicadores['total_sali']}
- Salieron dentro del umbral (las que arman el promedio): {indicadores['sali_dentro_umbral']}
- Demora media: {indicadores['demora_media_fmt'] or 'sin datos suficientes'}
- En alerta - pendientes (siguen en trámite hace más del umbral): {indicadores['en_alerta_bandeja']}
- En alerta - salieron tarde (superaron el umbral antes de salir): {indicadores['en_alerta_demora_larga']}
- En alerta total: {indicadores['en_alerta_total']}

EVOLUCIÓN MENSUAL (demora media, últimos meses):
{evol_txt}

ADUANAS CON MÁS ALERTAS (top 8, si las hay):
{top_alerta_txt}
{contexto_repositorio("sintia")}

Redactá en español, tono formal e institucional. Estructura:
1. Resumen ejecutivo (2-3 líneas)
2. Lectura de los indicadores nacionales
3. Tendencia (según la evolución mensual — ¿mejora, empeora, se mantiene?)
4. Puntos de atención (aduanas con más alertas, si las hay)
Máximo 350 palabras. No repitas los números en formato de lista, integralos en prosa."""

    client = anthropic.Anthropic(api_key=api_key, http_client=httpx.Client(follow_redirects=True))
    msg = client.messages.create(
        model="claude-haiku-4-5-20251001", max_tokens=1200, temperature=0.2,
        messages=[{"role": "user", "content": prompt}])
    return msg.content[0].text.strip()


def _verificar_narrativa_aduanas(texto, indicadores, log_fn):
    """Chequeo liviano: los números con más de un dígito que aparecen en la
    narrativa deberían poder rastrearse a algún valor real que le dimos a
    la IA. No corrige nada solo (a diferencia de la verificación del informe
    SINTIA, que sí corrige denominadores) -- acá el texto es más libre y
    corregir a ciegas podría romper la prosa. Solo deja constancia en el log
    para que quien revise sepa qué mirar con más atención."""
    if not texto:
        return
    valores_validos = set()
    for v in indicadores.values():
        if isinstance(v, (int, float)) and v is not None:
            valores_validos.add(str(round(v)) if isinstance(v, float) else str(v))
    numeros_en_texto = set(re.findall(r"\b\d{2,}\b", texto))
    sospechosos = numeros_en_texto - valores_validos
    if sospechosos:
        log_fn(f"  ⚠ Revisar a mano: la narrativa menciona número(s) que no matchean "
               f"exactamente ningún indicador calculado: {', '.join(sorted(sospechosos))} "
               f"(puede ser una fecha/año, un cálculo válido tipo diferencia entre dos "
               f"valores, o un número mal citado por la IA — no se pudo distinguir automáticamente).")


def _grafico_evolucion_aduana(nombre_aduana, meses_cols, valores_dias, valores_operaciones=None):
    """Gráfico combinado para una aduana puntual: línea de demora (horas,
    eje izquierdo, igual criterio que el gráfico en pantalla) + barras de
    cantidad de operaciones (eje derecho, detrás de la línea). Combinarlos
    en un solo gráfico -- en vez de uno aparte -- deja ver de un vistazo si
    un pico de demora coincide con un pico de volumen, sin duplicar
    gráficos ni alargar el informe al doble por cada aduana en alerta.
    Reusa el estilo visual del resto de la app (ver generar_graficos.py)."""
    try:
        from generar_graficos import fig_to_bytes
        import matplotlib; matplotlib.use('Agg')
        import matplotlib.pyplot as plt
    except ImportError:
        return None

    labels = [_mes_label_corto(m) for m in meses_cols]
    horas = [round(v * 24, 2) if v is not None else None for v in valores_dias]
    x = list(range(len(labels)))

    fig, ax = plt.subplots(figsize=(7, 3.2), facecolor="white")

    if valores_operaciones:
        ax2 = ax.twinx()
        ax2.bar(x, valores_operaciones, color="#94a3b8", alpha=0.35, width=0.55, zorder=1,
                label="Operaciones")
        ax2.set_ylabel("Operaciones", fontsize=9, color="#64748b")
        ax2.tick_params(axis="y", labelsize=8, labelcolor="#64748b")
        ax2.spines["top"].set_visible(False)
        max_ops = max(valores_operaciones) if valores_operaciones else 0
        if max_ops > 0:
            ax2.set_ylim(0, max_ops * 1.3)  # deja aire arriba para que no tape la línea de demora

    # spanGaps manual: matplotlib no une puntos separados por None solo, hay
    # que filtrar los huecos para la línea pero mantener el eje X completo.
    xs_validos = [xi for xi, v in zip(x, horas) if v is not None]
    ys_validos = [v for v in horas if v is not None]
    if ys_validos:
        ax.plot(xs_validos, ys_validos, marker="o", color="#2563eb", linewidth=2, markersize=5, zorder=3,
                label="Demora media")
        ax.fill_between(xs_validos, ys_validos, color="#2563eb", alpha=0.08, zorder=2)
    ax.set_zorder(ax2.get_zorder() + 1 if valores_operaciones else 1)  # línea de demora siempre encima de las barras
    ax.patch.set_visible(False)  # fondo transparente para que se vean las barras del eje de atrás
    ax.set_xticks(x); ax.set_xticklabels(labels, fontsize=8)
    ax.set_ylabel("Horas (demora)", fontsize=9, color="#2563eb")
    ax.tick_params(axis="y", labelsize=8, labelcolor="#2563eb")
    ax.set_title(f"Evolución mensual — {nombre_aduana}", fontsize=10, fontweight="bold", pad=12)
    ax.yaxis.grid(True, alpha=0.3); ax.set_axisbelow(True)
    ax.spines["top"].set_visible(False)
    if valores_operaciones:
        ax.spines["right"].set_visible(False)
        # Leyenda combinada (demora + operaciones) DEBAJO del gráfico, no
        # arriba -- ahí arriba estaba pisando el título (reportado en la
        # práctica: "el título con la descripción de las líneas quedó
        # solapado"). bbox_to_anchor con y negativo la saca del área de
        # ejes, fig_to_bytes ya recorta con bbox_inches='tight' así que no
        # queda espacio en blanco de más.
        handles1, labels1 = ax.get_legend_handles_labels()
        handles2, labels2 = ax2.get_legend_handles_labels()
        ax.legend(handles1 + handles2, labels1 + labels2,
                  loc="upper center", bbox_to_anchor=(0.5, -0.22), fontsize=7.5, frameon=False, ncol=2)
    else:
        ax.spines["right"].set_visible(False)
    if ax.get_ylim()[1] < 1:
        ax.set_ylim(0, 1)
    fig.tight_layout()
    return fig_to_bytes(fig)


def _generar_word_informe_aduanas(anio, dira_nombre, umbral_alerta_dias, indicadores, filas, evolucion,
                                   narrativa, evolucion_por_aduana=None, meses_cols=None):
    """Arma el Word del informe de Aduanas del País: filtros, indicadores,
    análisis IA, evolución mensual, y detalle por aduana. No reusa
    actas.generar_acta_word() -- esa está pensada para minutas de reunión
    (participantes/secciones con viñetas), estructura distinta a un informe
    con tablas de indicadores."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _set_cell_color(cell, hex_color):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def _tabla_simple(doc, headers, filas_datos):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for i, h in enumerate(headers):
            c = table.rows[0].cells[i]
            c.text = h
            c.paragraphs[0].runs[0].bold = True
            c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            c.paragraphs[0].runs[0].font.size = Pt(9)
            _set_cell_color(c, "242D4F")
        for fila in filas_datos:
            row = table.add_row()
            for i, val in enumerate(fila):
                row.cells[i].text = str(val) if val is not None else "—"
                for p in row.cells[i].paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)
        return table

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.2); section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5); section.right_margin = Cm(2.2)

    titulo = doc.add_paragraph(); titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("INFORME — ADUANAS DEL PAÍS")
    run.bold = True; run.font.size = Pt(16); run.font.color.rgb = RGBColor(0x24, 0x2D, 0x4F)
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("SINTIA — Tiempos de desaduanamiento (PAD)")
    sub_run.font.size = Pt(11); sub_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    doc.add_paragraph()

    for label, valor in [
        ("Generado:", datetime.now().strftime("%d/%m/%Y %H:%M")),
        ("Año (datos PAD):", str(anio)),
        ("Alcance:", dira_nombre),
        ("Umbral de alerta:", f"{umbral_alerta_dias} días"),
        ("Versión del generador:", _INFORME_ADUANAS_VERSION),
    ]:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label} "); r1.bold = True; r1.font.size = Pt(10)
        r2 = p.add_run(valor); r2.font.size = Pt(10)

    doc.add_paragraph()
    nota = doc.add_paragraph()
    nota_run = nota.add_run(
        "Metodología: demora = tiempo entre el registro de la operación en PAD y su salida (estado SAL). "
        "El promedio solo considera operaciones que salieron dentro del umbral definido arriba; las que "
        "superaron ese tiempo (pendientes o ya salidas) se excluyen del promedio y se cuentan aparte como "
        "alerta. Esta métrica no desagrega causas de demora (no mide específicamente coordinación con "
        "SENASA ni con ningún otro organismo)."
    )
    nota_run.italic = True; nota_run.font.size = Pt(8.5); nota_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()
    doc.add_paragraph().add_run("Indicadores").bold = True
    _tabla_simple(doc, ["Indicador", "Valor"], [
        ("Operaciones totales", indicadores["total_operaciones"]),
        ("Salieron (SAL)", indicadores["total_sali"]),
        ("Salieron dentro del umbral", indicadores["sali_dentro_umbral"]),
        ("Demora media", indicadores["demora_media_fmt"] or "—"),
        ("En alerta — pendiente", indicadores["en_alerta_bandeja"]),
        ("En alerta — salió tarde", indicadores["en_alerta_demora_larga"]),
        ("En alerta total", indicadores["en_alerta_total"]),
    ])

    doc.add_paragraph()
    doc.add_paragraph().add_run("Análisis").bold = True
    for parrafo in narrativa.split("\n"):
        if parrafo.strip():
            p = doc.add_paragraph(parrafo.strip())
            for r in p.runs:
                r.font.size = Pt(10.5)

    if evolucion:
        doc.add_paragraph()
        doc.add_paragraph().add_run("Evolución mensual").bold = True
        _tabla_simple(doc, ["Mes", "Demora media"],
                      [(_mes_label_corto(p["mes"]), p["demora_media_fmt"] or "—") for p in evolucion])

    if filas:
        doc.add_paragraph()
        doc.add_paragraph().add_run("Detalle por aduana").bold = True
        _tabla_simple(doc, ["Aduana", "DIRA", "Operaciones", "Demora media", "En alerta"],
                      [(f["aduana_nombre"], f["dira_nombre"], f["total_operaciones"],
                        f["demora_media_fmt"] or "—", f["en_alerta_total"]) for f in filas])

    # ── Evolución mensual de cada aduana en alerta (tabla + gráfico) ──────
    # Al final del documento, a propósito: es el detalle más fino de todos
    # (una sección completa por aduana), así que va después de todo lo
    # demás, no compitiendo por atención con el resumen ejecutivo.
    #
    # Orden: de más alertas a menos (a pedido -- distinto del orden
    # alfabético/por DIRA que usa la tabla "Detalle por aduana" de arriba).
    aduanas_en_alerta = sorted(
        [f for f in filas if f["en_alerta_total"] > 0],
        key=lambda f: f["en_alerta_total"], reverse=True)
    if aduanas_en_alerta and evolucion_por_aduana and meses_cols:
        por_cod = {f["aduana_cod"]: f for f in evolucion_por_aduana}
        doc.add_page_break()
        titulo_sec = doc.add_paragraph()
        titulo_sec_run = titulo_sec.add_run("Evolución mensual — aduanas en alerta")
        titulo_sec_run.bold = True; titulo_sec_run.font.size = Pt(13)
        titulo_sec_run.font.color.rgb = RGBColor(0x24, 0x2D, 0x4F)
        doc.add_paragraph().add_run(
            "Demora media y cantidad de operaciones por mes, de cada aduana que tiene al menos una "
            "operación en alerta (pendiente hace más del umbral, o que salió tarde). Ordenadas de "
            "mayor a menor cantidad de alertas."
        ).italic = True

        for f in aduanas_en_alerta:
            datos_aduana = por_cod.get(f["aduana_cod"])
            if not datos_aduana:
                continue
            doc.add_paragraph()
            nombre_p = doc.add_paragraph()
            nombre_run = nombre_p.add_run(f"{f['aduana_nombre']} ")
            nombre_run.bold = True; nombre_run.font.size = Pt(12)
            alerta_run = nombre_p.add_run(f"({f['en_alerta_total']} en alerta)")
            alerta_run.bold = True; alerta_run.font.size = Pt(12); alerta_run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

            valores_mes = [datos_aduana.get(m) or {"demora": None, "operaciones": 0} for m in meses_cols]
            valores_dias = [v["demora"] for v in valores_mes]
            valores_ops = [v["operaciones"] for v in valores_mes]
            _tabla_simple(doc, ["Mes", "Demora media", "Operaciones"], [
                (_mes_label_corto(m), _formatear_demora(v["demora"]) or "—", v["operaciones"])
                for m, v in zip(meses_cols, valores_mes)
            ])

            img = _grafico_evolucion_aduana(f["aduana_nombre"], meses_cols, valores_dias, valores_ops)
            if img:
                p_img = doc.add_paragraph(); p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.add_run().add_picture(img, width=Cm(13))

    return doc


def _job_informe_aduanas_nacional(job_id, anio, dira_filtro, umbral_alerta_dias, username):
    log = job_status[job_id]["log"]
    try:
        log.append("Calculando indicadores...")
        filas, indicadores, diras = _aduanas_nacional_datos(anio, dira_filtro, umbral_alerta_dias)
        dira_nombre = next((d["nombre"] for d in diras if d["indice"] == dira_filtro), dira_filtro) \
            if dira_filtro else "Todo el país"
        evolucion = _evolucion_mensual_nacional(anio, dira_filtro, umbral_alerta_dias)

        api_key = get_api_key()
        if not api_key:
            log.append("✗ API key no configurada — no se puede generar el análisis con IA.")
            job_status[job_id]["status"] = "error"
            _job_persist(job_id)
            return

        log.append("Generando análisis con IA...")
        narrativa = _generar_narrativa_aduanas_ia(
            anio, dira_nombre, umbral_alerta_dias, indicadores, filas, evolucion, api_key)
        _verificar_narrativa_aduanas(narrativa, indicadores, log.append)

        log.append("Calculando evolución por aduana...")
        evolucion_por_aduana, meses_cols = _evolucion_mensual_por_aduana(anio, dira_filtro, umbral_alerta_dias)

        log.append("Armando Word...")
        doc = _generar_word_informe_aduanas(
            anio, dira_nombre, umbral_alerta_dias, indicadores, filas, evolucion, narrativa,
            evolucion_por_aduana, meses_cols)
        os.makedirs(OUTPUT_FOLDER, exist_ok=True)
        fname = f"Informe_Aduanas_Pais_{anio}_{job_id}.docx"
        ruta = os.path.join(OUTPUT_FOLDER, fname)
        doc.save(ruta)

        # Registrar en 'historial' -- SIN esto, _limpiar_archivos_huerfanos()
        # (que corre una vez por día y borra todo lo que haya en OUTPUT_FOLDER
        # sin fila en historial.archivo_word/archivo_excel) se comería este
        # archivo a las 48hs aunque el job siga diciendo "done". Encontrado
        # en revisión post-entrega (10/07/2026) -- el resto de los generadores
        # de Word de esta app sí se registran, este se había quedado afuera.
        hist_id = str(uuid.uuid4())[:8]
        # El campo 'pais' no aplica acá (es de la época en que 'historial'
        # solo tenía informes SINTIA por país vecino) -- se reusa para
        # guardar la combinación dira+umbral+versión, así _buscar_informe_
        # aduanas_cacheado() puede encontrar una corrida previa con los
        # mismos parámetros Y la misma versión del generador, sin tener que
        # agregar una columna nueva.
        clave_cache = f"dira={dira_filtro or ''};umbral={umbral_alerta_dias};v={_INFORME_ADUANAS_VERSION}"
        with get_db(HIST_DB) as con:
            con.execute("INSERT INTO historial VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)",
                (hist_id, datetime.now().strftime("%Y-%m-%d %H:%M:%S"), username,
                 clave_cache, anio, "", "", 1, ruta, "", 0, "aduanas_pais",
                 f"Aduanas del país — {dira_nombre} ({anio})"))

        job_status[job_id]["files"] = [ruta]
        job_status[job_id]["status"] = "done"
        _job_persist(job_id)
        log.append(f"✓ Informe generado: {fname}")
        logging.info(f"INFORME ADUANAS PAIS OK | user={username} | anio={anio} | dira={dira_filtro}")
    except Exception as e:
        log.append(f"✗ Error: {e}")
        job_status[job_id]["status"] = "error"
        _job_persist(job_id)
        logging.error(f"INFORME ADUANAS PAIS ERROR | user={username} | {e}")


def _buscar_informe_aduanas_cacheado(anio, dira_filtro, umbral_alerta_dias, ttl_horas=24):
    """Si ya se generó un informe con exactamente los mismos parámetros
    (año/DIRA/umbral) Y la misma versión del generador (_INFORME_ADUANAS_
    VERSION) dentro de las últimas ttl_horas, y el archivo todavía existe en
    disco, lo devuelve -- evita recalcular todo y volver a pagar una
    llamada a la IA por algo que ya se generó hoy. Mismo criterio de
    frescura que se usó para los combos de filtro: los datos de PAD no
    cambian más seguido que una vez por semana, 24hs de caché es
    conservador. Incluir la versión evita el problema real que apareció:
    un deploy de código nuevo servía igual el archivo viejo cacheado."""
    clave_cache = f"dira={dira_filtro or ''};umbral={umbral_alerta_dias};v={_INFORME_ADUANAS_VERSION}"
    limite = (datetime.now() - timedelta(hours=ttl_horas)).strftime("%Y-%m-%d %H:%M:%S")
    with get_db(HIST_DB, row_factory=True) as con:
        fila = con.execute(
            "SELECT id, archivo_word, fecha FROM historial "
            "WHERE tipo='aduanas_pais' AND anio=? AND pais=? AND fecha >= ? "
            "ORDER BY fecha DESC LIMIT 1",
            (anio, clave_cache, limite)).fetchone()
    if fila and fila["archivo_word"] and os.path.exists(fila["archivo_word"]):
        return dict(fila)
    return None


@app.route("/api/sintia/aduanas_nacional/informe", methods=["POST"])
@login_required
@modulo_required("sintia")
def sintia_aduanas_nacional_informe():
    data = request.json or {}
    anio = data.get("anio", str(date.today().year))
    dira_filtro = (data.get("dira") or "").strip() or None
    umbral_alerta_dias = int(data.get("umbral_dias", 10))
    forzar = bool(data.get("forzar"))
    username = session.get("username", "?")

    cacheado = None if forzar else _buscar_informe_aduanas_cacheado(anio, dira_filtro, umbral_alerta_dias)
    if cacheado:
        job_id = str(uuid.uuid4())[:8]
        job_create(job_id, "", username=username)
        job_status[job_id]["log"].append(
            f"✓ Ya existe un informe con estos mismos parámetros, generado el {cacheado['fecha']} "
            f"(dentro de las últimas 24hs) — se reutiliza en vez de volver a llamar a la IA.")
        job_status[job_id]["files"] = [cacheado["archivo_word"]]
        job_status[job_id]["status"] = "done"
        _job_persist(job_id)
        return jsonify({"ok": True, "job_id": job_id, "cached": True})

    job_id = str(uuid.uuid4())[:8]
    job_create(job_id, "Iniciando informe de Aduanas del País...", username=username)
    t = threading.Thread(target=_job_informe_aduanas_nacional,
                         args=(job_id, anio, dira_filtro, umbral_alerta_dias, username))
    t.start()
    return jsonify({"ok": True, "job_id": job_id, "cached": False})


@app.route("/api/sintia/aduanas_nacional/export")
@login_required
@modulo_required("sintia")
def sintia_aduanas_nacional_export():
    """Mismos datos que sintia_aduanas_nacional(), exportados a Excel con
    dos hojas: 'Resumen' (qué filtro se usó + indicadores nacionales, para
    que el archivo tenga sentido solo, sin depender de recordar qué se tenía
    tildado en pantalla) y 'Aduanas' (el detalle fila por fila)."""
    anio = request.args.get("anio", str(date.today().year))
    dira_filtro = request.args.get("dira", "").strip() or None
    umbral_alerta_dias = int(request.args.get("umbral_dias", 10))

    try:
        filas, indicadores, diras = _aduanas_nacional_datos(anio, dira_filtro, umbral_alerta_dias)
        dira_nombre_filtro = next((d["nombre"] for d in diras if d["indice"] == dira_filtro), dira_filtro) \
            if dira_filtro else "Todas"

        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment

        wb = Workbook()

        # ── Hoja 1: Resumen ──────────────────────────────────────────────
        ws1 = wb.active
        ws1.title = "Resumen"
        ws1.append(["Reporte de Aduanas del país — SINTIA"])
        ws1["A1"].font = Font(bold=True, size=14, color="242D4F")
        ws1.append([])

        ws1.append(["Filtros aplicados"])
        ws1[f"A{ws1.max_row}"].font = Font(bold=True)
        for etiqueta, valor in [
            ("Generado", datetime.now().strftime("%d/%m/%Y %H:%M")),
            ("Año (tabla DAT)", anio),
            ("DIRA", dira_nombre_filtro),
            ("Umbral de alerta (días)", umbral_alerta_dias),
        ]:
            ws1.append([etiqueta, valor])

        ws1.append([])
        ws1.append(["Indicadores nacionales"])
        ws1[f"A{ws1.max_row}"].font = Font(bold=True)
        for etiqueta, valor in [
            ("Operaciones totales", indicadores["total_operaciones"]),
            ("Salieron (SAL)", indicadores["total_sali"]),
            ("Salieron dentro del umbral", indicadores["sali_dentro_umbral"]),
            ("Demora media (h/m/s)", indicadores["demora_media_fmt"] or "—"),
            ("Demora media (días, decimal)",
             round(indicadores["demora_media_dias"], 4) if indicadores["demora_media_dias"] is not None else None),
            ("En alerta - pendiente (bandeja)", indicadores["en_alerta_bandeja"]),
            ("En alerta - salió tarde", indicadores["en_alerta_demora_larga"]),
            ("En alerta total", indicadores["en_alerta_total"]),
        ]:
            ws1.append([etiqueta, valor])

        ws1.column_dimensions["A"].width = 32
        ws1.column_dimensions["B"].width = 26

        # ── Hoja 2: Aduanas (detalle) ────────────────────────────────────
        ws2 = wb.create_sheet("Aduanas")
        cols = ["Aduana", "DIRA", "Operaciones", "Salieron", "Salieron dentro del umbral",
                "Demora media (días, decimal)", "Demora media (h/m/s)",
                "En alerta - pendiente (bandeja)", "En alerta - salió tarde", "En alerta total"]
        ws2.append(cols)
        for cell in ws2[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="242D4F")
            cell.alignment = Alignment(horizontal="center")
        for f in filas:
            ws2.append([
                f["aduana_nombre"], f["dira_nombre"], f["total_operaciones"], f["total_sali"],
                f["sali_dentro_umbral"],
                round(f["demora_media_dias"], 4) if f["demora_media_dias"] is not None else None,
                f["demora_media_fmt"] or "",
                f["en_alerta_bandeja"], f["en_alerta_demora_larga"], f["en_alerta_total"],
            ])
        for col in ws2.columns:
            max_len = max((len(str(c.value if c.value is not None else "")) for c in col), default=8)
            ws2.column_dimensions[col[0].column_letter].width = min(max_len + 2, 42)

        # ── Hoja 3: Evolución mensual por aduana (pivot) ──────────────────
        # Solo acá, no en pantalla -- con 50+ aduanas x 6 meses ya son ~300
        # celdas, imposible como gráfico legible pero perfecto como tabla
        # para ordenar/filtrar en Excel (ver conversación 10/07/2026).
        from openpyxl.utils import get_column_letter
        filas_evol, meses_cols = _evolucion_mensual_por_aduana(anio, dira_filtro, umbral_alerta_dias)
        ws3 = wb.create_sheet("Evolución mensual")
        encabezado_evol = ["Aduana"] + [_mes_label_corto(m) for m in meses_cols]
        ws3.append(encabezado_evol)
        for cell in ws3[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="242D4F")
            cell.alignment = Alignment(horizontal="center")
        for f in filas_evol:
            fila_valores = [f["aduana_nombre"]] + [
                round(f[m]["demora"], 4) if f.get(m) and f[m].get("demora") is not None else None
                for m in meses_cols]
            ws3.append(fila_valores)
        ws3.column_dimensions["A"].width = 32
        for i in range(len(meses_cols)):
            ws3.column_dimensions[get_column_letter(2 + i)].width = 14
        if filas_evol:
            ws3.append([])
            ws3.append(["Valores en días (decimal) — demora media de ese mes, mismo criterio que la hoja Aduanas "
                        "(solo SAL dentro del umbral). Celda vacía = sin operaciones SAL ese mes."])

        # ── Hoja 4: tablas individuales por aduana en alerta ──────────────
        # A pedido: una tabla por cada aduana con alertas (mes/demora en
        # columna, no pivot como la hoja 3) -- acá solo tablas, sin gráfico
        # (los gráficos van únicamente en el Word). Mismo orden que el Word:
        # de más alertas a menos.
        aduanas_en_alerta = sorted(
            [f for f in filas if f["en_alerta_total"] > 0],
            key=lambda f: f["en_alerta_total"], reverse=True)
        if aduanas_en_alerta:
            por_cod = {f["aduana_cod"]: f for f in filas_evol}
            ws4 = wb.create_sheet("Evolución - aduanas en alerta")
            fila_actual = 1
            for f in aduanas_en_alerta:
                datos_aduana = por_cod.get(f["aduana_cod"])
                if not datos_aduana:
                    continue
                celda_titulo = ws4.cell(row=fila_actual, column=1,
                    value=f"{f['aduana_nombre']} ({f['en_alerta_total']} en alerta)")
                celda_titulo.font = Font(bold=True, size=12, color="242D4F")
                fila_actual += 1

                ws4.cell(row=fila_actual, column=1, value="Mes")
                ws4.cell(row=fila_actual, column=2, value="Demora media (días)")
                ws4.cell(row=fila_actual, column=3, value="Demora media (h/m/s)")
                ws4.cell(row=fila_actual, column=4, value="Operaciones")
                for col in (1, 2, 3, 4):
                    c = ws4.cell(row=fila_actual, column=col)
                    c.font = Font(bold=True, color="FFFFFF")
                    c.fill = PatternFill("solid", fgColor="242D4F")
                fila_actual += 1

                for mes in meses_cols:
                    valor_mes = datos_aduana.get(mes) or {"demora": None, "operaciones": 0}
                    valor_dias = valor_mes["demora"]
                    ws4.cell(row=fila_actual, column=1, value=_mes_label_corto(mes))
                    ws4.cell(row=fila_actual, column=2, value=round(valor_dias, 4) if valor_dias is not None else None)
                    ws4.cell(row=fila_actual, column=3, value=_formatear_demora(valor_dias) or "")
                    ws4.cell(row=fila_actual, column=4, value=valor_mes["operaciones"])
                    fila_actual += 1
                fila_actual += 2  # una fila en blanco entre aduanas

            ws4.column_dimensions["A"].width = 32
            ws4.column_dimensions["B"].width = 20
            ws4.column_dimensions["C"].width = 20
            ws4.column_dimensions["D"].width = 14

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return send_file(buf, as_attachment=True,
                         download_name=f"Aduanas_nacional_{anio}.xlsx",
                         mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    except Exception as e:
        logging.error(f"SINTIA ADUANAS NACIONAL EXPORT ERROR | {e}")
        return jsonify({"ok": False, "error": str(e)})


# MÓDULOS GARMIN + TRAINING -> blueprints/training.py (registrado como training_bp)



# MÓDULO STOCK -> blueprints/stock.py (registrado como stock_bp)


# ══════════════════════════════════════════════════════════════════════════════
# MÓDULO REF. ADUANAS / DIRA
# ══════════════════════════════════════════════════════════════════════════════

def _get_dira_nombres():
    """Lee el dict {indice: nombre} de direcciones regionales desde ref_dira (BD)."""
    with get_db(HIST_DB, row_factory=True) as con:
        rows = con.execute("SELECT indice, nombre FROM ref_dira ORDER BY orden, indice").fetchall()
    return {r["indice"]: r["nombre"] for r in rows}

@app.route("/api/admin/ref-dira")
@login_required
def ref_dira_list():
    with get_db(HIST_DB, row_factory=True) as con:
        rows = [dict(r) for r in con.execute(
            "SELECT indice, nombre, orden FROM ref_dira ORDER BY orden, indice"
        ).fetchall()]
    return jsonify({"ok": True, "rows": rows})

@app.route("/api/admin/ref-dira/save", methods=["POST"])
@login_required
@admin_required("bd")
def ref_dira_save():
    """Reemplaza completamente la tabla ref_dira con el array editado en vivo.
    Body esperado: {"rows": [{"indice": "1", "nombre": "HIDROVIA", "orden": 1}, ...]}"""
    data = request.get_json(force=True) or {}
    rows = data.get("rows")
    if not isinstance(rows, list) or not rows:
        return jsonify({"ok": False, "error": "No se recibieron filas para guardar"}), 400

    # Validar que ningún índice a eliminar siga referenciado en ref_aduanas
    with get_db(HIST_DB, row_factory=True) as con:
        indices_en_uso = {r["indice_dira"] for r in con.execute("SELECT DISTINCT indice_dira FROM ref_aduanas").fetchall()}

        nuevos = []
        errores = []
        indices_vistos = set()
        for i, row in enumerate(rows, 1):
            try:
                indice = str(row.get("indice", "")).strip()
                nombre = str(row.get("nombre", "")).strip().upper()
                try:
                    orden = int(row.get("orden", i))
                except (TypeError, ValueError):
                    orden = i
                if not indice:
                    errores.append(f"Fila {i}: índice vacío")
                    continue
                if not nombre:
                    errores.append(f"Fila {i} (índice {indice}): nombre vacío")
                    continue
                if indice in indices_vistos:
                    errores.append(f"Fila {i}: índice '{indice}' duplicado — se mantiene la última ocurrencia")
                indices_vistos.add(indice)
                nuevos.append((indice, nombre, orden))
            except Exception as e:
                errores.append(f"Fila {i}: {e}")

        if not nuevos:
            return jsonify({"ok": False, "error": "Ninguna fila válida para guardar", "detalle_errores": errores}), 400

        dedup = {}
        for indice, nombre, orden in nuevos:
            dedup[indice] = (nombre, orden)
        nuevos_dedup = [(indice, n, o) for indice, (n, o) in dedup.items()]
        indices_finales = set(dedup.keys())

        # Bloquear si se eliminaría un índice DIRA que todavía usa alguna aduana
        indices_eliminados = indices_en_uso - indices_finales
        if indices_eliminados:
            return jsonify({
                "ok": False,
                "error": f"No se puede eliminar el/los índice(s) {', '.join(sorted(indices_eliminados))} porque hay aduanas que todavía los usan. Reasigná esas aduanas primero en Ref. Aduanas."
            }), 400

        con.execute("DELETE FROM ref_dira")
        con.executemany("INSERT INTO ref_dira (indice, nombre, orden) VALUES (?,?,?)", nuevos_dedup)
        con.execute("INSERT INTO ref_aduanas_log (fecha, usuario, accion, detalle) VALUES (datetime('now'),?,?,?)",
                    (session.get("username","?"), "save_ref_dira",
                     f"{len(nuevos_dedup)} direcciones regionales guardadas" + (f", {len(errores)} filas con error" if errores else "")))
    logging.info(f"REF_DIRA SAVE | user={session.get('username')} | {len(nuevos_dedup)} filas | {len(errores)} errores")
    return jsonify({"ok": True, "guardadas": len(nuevos_dedup), "errores": errores[:30]})

@app.route("/api/admin/ref-aduanas")
@login_required
def ref_aduanas_list():
    dira_nombres = _get_dira_nombres()
    with get_db(HIST_DB, row_factory=True) as con:
        rows = [dict(r) for r in con.execute(
            "SELECT cod, nombre, indice_dira FROM ref_aduanas ORDER BY indice_dira, nombre"
        ).fetchall()]
    for r in rows:
        r["dira_nombre"] = dira_nombres.get(r["indice_dira"], "N/E")
    return jsonify({"ok": True, "rows": rows, "dira_nombres": dira_nombres})

@app.route("/api/admin/ref-aduanas/download")
@login_required
def ref_aduanas_download():
    """Descarga la tabla actual en el mismo formato CSV de referencia (; como separador)."""
    with get_db(HIST_DB, row_factory=True) as con:
        rows = con.execute("SELECT cod, nombre, indice_dira FROM ref_aduanas ORDER BY indice_dira, nombre").fetchall()
    lines = ["indice_dira;nombre;cod"]
    for r in rows:
        lines.append(f"{r['indice_dira']};{r['nombre']};{int(r['cod'])}")
    csv_content = "\r\n".join(lines) + "\r\n"
    from flask import Response
    return Response(csv_content, mimetype="text/csv",
                     headers={"Content-Disposition": "attachment; filename=ref_aduanas.csv"})

@app.route("/api/admin/ref-aduanas/save", methods=["POST"])
@login_required
@admin_required("bd")
def ref_aduanas_save():
    """Reemplaza completamente la tabla ref_aduanas con el array editado en vivo desde el panel.
    Body esperado: {"rows": [{"cod": "001", "nombre": "BUENOS AIRES", "indice_dira": "8"}, ...]}"""
    dira_nombres = _get_dira_nombres()
    data = request.get_json(force=True) or {}
    rows = data.get("rows")
    if not isinstance(rows, list) or not rows:
        return jsonify({"ok": False, "error": "No se recibieron filas para guardar"}), 400

    nuevos = []
    errores = []
    cods_vistos = set()
    for i, row in enumerate(rows, 1):
        try:
            cod = str(row.get("cod", "")).strip().zfill(3)
            nombre = str(row.get("nombre", "")).strip().upper()
            indice_dira = str(row.get("indice_dira", "")).strip()
            if not cod or not cod.isdigit():
                errores.append(f"Fila {i}: código inválido '{row.get('cod')}'")
                continue
            if not nombre:
                errores.append(f"Fila {i} (cod {cod}): nombre vacío")
                continue
            if indice_dira not in dira_nombres:
                errores.append(f"Fila {i} (cod {cod}): índice DIRA '{indice_dira}' inválido")
                continue
            if cod in cods_vistos:
                errores.append(f"Fila {i}: código '{cod}' duplicado en la tabla — se mantiene la última ocurrencia")
            cods_vistos.add(cod)
            nuevos.append((cod, nombre, indice_dira))
        except Exception as e:
            errores.append(f"Fila {i}: {e}")

    if not nuevos:
        return jsonify({"ok": False, "error": "Ninguna fila válida para guardar", "detalle_errores": errores}), 400

    # Deduplicar por cod conservando la última ocurrencia (consistente con el aviso de arriba)
    dedup = {}
    for cod, nombre, indice_dira in nuevos:
        dedup[cod] = (nombre, indice_dira)
    nuevos_dedup = [(cod, n, d) for cod, (n, d) in dedup.items()]

    with get_db(HIST_DB) as con:
        con.execute("DELETE FROM ref_aduanas")
        con.executemany("INSERT INTO ref_aduanas (cod, nombre, indice_dira) VALUES (?,?,?)", nuevos_dedup)
        con.execute("INSERT INTO ref_aduanas_log (fecha, usuario, accion, detalle) VALUES (datetime('now'),?,?,?)",
                    (session.get("username","?"), "save_tabla_completa",
                     f"{len(nuevos_dedup)} aduanas guardadas" + (f", {len(errores)} filas con error" if errores else "")))
    logging.info(f"REF_ADUANAS SAVE | user={session.get('username')} | {len(nuevos_dedup)} filas | {len(errores)} errores")
    return jsonify({"ok": True, "guardadas": len(nuevos_dedup), "errores": errores[:30]})


@app.route("/api/admin/feriados", methods=["GET"])
@login_required
@admin_required("bd")
def feriados_list():
    with get_db(HIST_DB, row_factory=True) as con:
        rows = [dict(r) for r in con.execute(
            "SELECT fecha, descripcion FROM feriados ORDER BY fecha"
        ).fetchall()]
    return jsonify({"ok": True, "rows": rows})

@app.route("/api/admin/feriados", methods=["POST"])
@login_required
@admin_required("bd")
def feriados_create():
    """Body esperado: {"fecha": "2027-01-01", "descripcion": "Año Nuevo"}"""
    data = request.get_json(force=True) or {}
    fecha = str(data.get("fecha", "")).strip()
    descripcion = str(data.get("descripcion", "")).strip()
    if not re.match(r"^\d{4}-\d{2}-\d{2}$", fecha):
        return jsonify({"ok": False, "error": "Fecha inválida, formato esperado AAAA-MM-DD"}), 400
    try:
        datetime.strptime(fecha, "%Y-%m-%d")
    except ValueError:
        return jsonify({"ok": False, "error": "Fecha inválida"}), 400
    with get_db(HIST_DB) as con:
        con.execute("INSERT OR REPLACE INTO feriados (fecha, descripcion) VALUES (?,?)", (fecha, descripcion))
    logging.info(f"FERIADOS CREATE | user={session.get('username')} | {fecha} | {descripcion}")
    return jsonify({"ok": True})

@app.route("/api/admin/feriados/<fecha>", methods=["DELETE"])
@login_required
@admin_required("bd")
def feriados_delete(fecha):
    with get_db(HIST_DB) as con:
        cur = con.execute("DELETE FROM feriados WHERE fecha=?", (fecha,))
        borrado = cur.rowcount > 0
    if not borrado:
        return jsonify({"ok": False, "error": "No existe ese feriado"}), 404
    logging.info(f"FERIADOS DELETE | user={session.get('username')} | {fecha}")
    return jsonify({"ok": True})


# MÓDULO FINANZAS -> blueprints/finanzas.py (registrado como finanzas_bp)


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=False)

