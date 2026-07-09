"""
core.py — Infraestructura compartida entre app.py y los blueprints de
CosmoTools (auth, sesiones, rutas de datos).

Por qué existe este archivo separado: los blueprints (blueprints/stock.py,
y los que sigan) necesitan los decoradores de auth (login_required,
modulo_required, etc.) y las rutas de datos (HIST_DB, DB_PATH). Si esas
cosas vivieran definidas en app.py, importar un blueprint DESDE app.py
crearía un import circular (app.py -> blueprints/stock.py -> app.py).
Al vivir acá, tanto app.py como cualquier blueprint importan de core.py
sin que core.py necesite saber que existen.
"""
import os
import io
import sqlite3
import json
import time
import logging
import threading
import urllib.request
import urllib.parse
from functools import wraps
from datetime import datetime, timedelta

from flask import Flask, session, redirect, url_for, request, jsonify
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from flask_wtf import CSRFProtect
from werkzeug.middleware.proxy_fix import ProxyFix

# ── App Flask + CSRF + rate limiter ─────────────────────────────────────────
# Viven acá (no en app.py) para que los blueprints puedan hacer
# `from core import app, limiter` sin generar un import circular con app.py,
# y para garantizar que `limiter` ya exista antes de que se importe
# cualquier blueprint que use @limiter.limit(...) en sus rutas.
app = Flask(__name__)

# Confía en 1 proxy (nginx) para X-Forwarded-For: reescribe request.remote_addr
# de forma segura y evita que el cliente falsee su IP con headers propios.
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=0, x_port=0, x_prefix=0)

_secret = os.environ.get("SECRET_KEY")
if not _secret:
    raise RuntimeError("SECRET_KEY no está definida en las variables de entorno. "
                       "Exportá una clave aleatoria antes de iniciar la aplicación.")
app.secret_key = _secret
app.permanent_session_lifetime = timedelta(minutes=10)
app.config['MAX_CONTENT_LENGTH'] = None  # Sin límite en Flask — nginx maneja con client_max_body_size 2G
app.config['SESSION_COOKIE_SECURE'] = True     # solo se envía por HTTPS
app.config['SESSION_COOKIE_HTTPONLY'] = True   # no accesible desde JS
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'  # mitiga CSRF básico en navegación cruzada
app.config['WTF_CSRF_TIME_LIMIT'] = None  # el timeout de sesión ya lo maneja login_required (10min)

@app.after_request
def _security_headers(resp):
    """Headers de defensa en profundidad — no reemplazan la validación de
    entrada/salida ya hecha, pero limitan el daño si algo se escapa:
    - X-Frame-Options: evita que la app se embeba en un <iframe> ajeno (clickjacking).
    - X-Content-Type-Options: evita que el navegador "adivine" un tipo de
      contenido distinto al declarado (mitiga algunos vectores de XSS).
    - Strict-Transport-Security: fuerza HTTPS en el navegador (la app ya
      requiere cookies solo por HTTPS con SESSION_COOKIE_SECURE).
    - Referrer-Policy: no filtra la URL completa (con tokens/ids) a terceros.
    """
    resp.headers.setdefault("X-Frame-Options", "DENY")
    resp.headers.setdefault("X-Content-Type-Options", "nosniff")
    resp.headers.setdefault("Referrer-Policy", "strict-origin-when-cross-origin")
    resp.headers.setdefault("Strict-Transport-Security", "max-age=31536000; includeSubDomains")
    return resp

csrf = CSRFProtect(app)

# storage_uri: en memoria por defecto (sirve con 1 solo worker). Si se corre con
# más de un worker (gunicorn -w >1), cada worker llevaría su propia cuenta y el
# límite dejaría de cumplirse — en ese caso definir RATELIMIT_STORAGE_URI=redis://...
limiter = Limiter(key_func=get_remote_address, app=app,
    default_limits=["300 per minute"],
    storage_uri=os.environ.get("RATELIMIT_STORAGE_URI", "memory://"))

# ── Rutas de datos (configurables por variable de entorno para poder
# testear sin arriesgar la base real — ver tests/conftest.py) ─────────────────
# HIST_DB, DB_PATH y get_db() viven en db_utils.py (sin dependencia de Flask) —
# ver el docstring de ese archivo para el motivo. Acá se re-exportan para que
# `from core import get_db` siga funcionando igual que antes en app.py y los
# blueprints.
from db_utils import HIST_DB, DB_PATH, get_db
OUTPUT_FOLDER     = "/data/informes"
STOCK_REPORTS_DIR = "/data/reports/stock"


# ── Migraciones ───────────────────────────────────────────────────────────────
def run_migrations(db_path, migrations_dir=None):
    """Sistema de migraciones numeradas (migrations/NNN_*.sql).

    001_init.sql es la línea base: un volcado del esquema tal como lo dejaban
    las funciones init_*_db() de siempre (core.py, finanzas.py,
    blueprints/stock.py). En una instalación EXISTENTE (la base ya tiene esas
    tablas, creadas por esas mismas funciones) se marca como aplicada sin
    ejecutarla — ejecutar sus CREATE TABLE de nuevo fallaría porque las
    tablas ya existen. En una base nueva y vacía, si esto corre antes que las
    init_*_db(), sí las crea.

    Lo que importa de acá en adelante es que cualquier cambio de esquema
    futuro se agregue como 002_*.sql, 003_*.sql, etc. — eso sí se ejecuta
    siempre (una sola vez, registrado en la tabla schema_migrations), en vez
    de sumar un CREATE TABLE IF NOT EXISTS más disperso por el código."""
    import glob as _glob
    if migrations_dir is None:
        migrations_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "migrations")
    if not os.path.isdir(migrations_dir):
        return
    try:
        con = sqlite3.connect(db_path, timeout=10)
        con.execute("""CREATE TABLE IF NOT EXISTS schema_migrations (
            version TEXT PRIMARY KEY, aplicada TEXT DEFAULT (datetime('now')))""")
        ya_aplicadas = {r[0] for r in con.execute("SELECT version FROM schema_migrations")}

        if not ya_aplicadas:
            hay_tablas_previas = con.execute(
                "SELECT COUNT(*) FROM sqlite_master WHERE type='table' "
                "AND name NOT IN ('schema_migrations','sqlite_sequence')"
            ).fetchone()[0] > 0
            if hay_tablas_previas:
                con.execute("INSERT INTO schema_migrations (version) VALUES ('001_init.sql')")
                con.commit()
                ya_aplicadas.add("001_init.sql")

        for ruta in sorted(_glob.glob(os.path.join(migrations_dir, "*.sql"))):
            nombre = os.path.basename(ruta)
            if nombre in ya_aplicadas:
                continue
            with open(ruta, encoding="utf-8") as f:
                sql = f.read()
            try:
                con.executescript(sql)
                con.execute("INSERT INTO schema_migrations (version) VALUES (?)", (nombre,))
                con.commit()
                logging.info(f"Migración aplicada: {nombre}")
            except Exception:
                con.rollback()
                logging.exception(f"Migración FALLÓ: {nombre} — no se marca como aplicada")
                raise
        con.close()
    except Exception:
        logging.exception("run_migrations: error inesperado")

run_migrations(HIST_DB)



# ── Sesiones ─────────────────────────────────────────────────────────────────
def registrar_sesion(username, token, ip, ua):
    try:
        with get_db(HIST_DB) as con:
            con.execute("INSERT OR REPLACE INTO sesiones (username, token, ip, user_agent, activo, ultimo_acceso) "
                "VALUES (?,?,?,?,1,datetime('now'))", (username, token, ip, ua[:200]))
    except: pass

def actualizar_sesion(token):
    try:
        with get_db(HIST_DB) as con:
            con.execute("UPDATE sesiones SET ultimo_acceso=datetime('now') WHERE token=?", (token,))
    except: pass

def token_revocado(token):
    try:
        with get_db(HIST_DB) as con:
            row = con.execute("SELECT 1 FROM tokens_revocados WHERE token=?", (token,)).fetchone()
        return row is not None
    except: return False


# ── Decoradores de auth ──────────────────────────────────────────────────────
def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get("logged_in"):
            return redirect(url_for("login"))
        token = session.get("token", "")
        if token and token_revocado(token):
            session.clear(); return redirect(url_for("login"))
        last = session.get("last_active")
        if last and datetime.now().timestamp() - last > 600:
            session.clear(); return redirect(url_for("login"))
        session["last_active"] = datetime.now().timestamp()
        if token:
            try: actualizar_sesion(token)
            except: pass
        session.permanent = True
        return f(*args, **kwargs)
    return decorated


def tiene_permiso_admin(seccion=None):
    """seccion=None -> algún permiso admin (bd o sistema). seccion='bd'|'sistema' -> ese en particular.
    role=='admin' es superadmin y siempre pasa (compatibilidad con usuarios existentes)."""
    if session.get("role") == "admin":
        return True
    modulos = session.get("modulos", [])
    if seccion:
        return f"admin_{seccion}" in modulos
    return "admin_bd" in modulos or "admin_sistema" in modulos


def admin_required(seccion=None):
    def decorator(f):
        @wraps(f)
        def decorated(*args, **kwargs):
            if not tiene_permiso_admin(seccion):
                if request.is_json or request.path.startswith("/api/"):
                    return jsonify({"ok": False, "error": "Sin permiso"}), 403
                return redirect(url_for("index"))
            return f(*args, **kwargs)
        return decorated
    return decorator


def modulo_required(nombre):
    """Exige que el usuario tenga `nombre` en su lista de módulos habilitados.
    Sin esto, ocultar el link del sidebar no alcanza: cualquier usuario logueado
    puede pegar la URL directamente."""
    def decorator(f):
        @wraps(f)
        def decorated(*args, **kwargs):
            if nombre not in session.get("modulos", []):
                if request.is_json or request.path.startswith("/api/"):
                    return jsonify({"ok": False, "error": "Módulo no habilitado"}), 403
                return redirect(url_for("index"))
            return f(*args, **kwargs)
        return decorated
    return decorator


_FINANZAS_ALLOWED = {u.strip() for u in os.environ.get("FINANZAS_ALLOWED_USERS", "").split(",") if u.strip()}

def finanzas_owner_required(f):
    """Los datos de fin_ddjj* son personales (una sola declaración jurada por año,
    sin columna de propietario) y NO deben quedar accesibles a cualquier usuario
    al que se le habilite el módulo 'finanzas' desde el panel de admin.
    Requiere, además del módulo habilitado, ser admin o estar en
    FINANZAS_ALLOWED_USERS (env var, usernames separados por coma)."""
    @wraps(f)
    @modulo_required("finanzas")
    def decorated(*args, **kwargs):
        u = session.get("username", "")
        if session.get("role") != "admin" and u not in _FINANZAS_ALLOWED:
            if request.is_json or request.path.startswith("/api/"):
                return jsonify({"ok": False, "error": "Sin permiso"}), 403
            return redirect(url_for("index"))
        return f(*args, **kwargs)
    return decorated


# ── Enums de estado (compartidos entre módulos con <select> de opciones fijas
# en el frontend — VUA y SENASA). Antes el backend guardaba estos campos como
# texto libre (data.get("estado", "Pendiente")), aunque el HTML solo ofrece
# 3 opciones: nada impedía que alguien mandara cualquier string por la API
# directamente (sin pasar por el <select>). El escape en el HTML ya evita que
# eso se ejecute como XSS, pero esto cierra el problema en el origen para los
# campos que SÍ son un enum real (no confundir con eje.estado, que es texto
# libre a propósito en ambos módulos — placeholder "En análisis, Pendiente,
# Completado..." — y por eso queda fuera de esta validación).
ESTADOS_TAREA = ("Pendiente", "En curso", "Completado")
NIVELES_PROBABILIDAD = ("Alta", "Media", "Baja")
NIVELES_IMPACTO = ("Alto", "Medio", "Bajo")

def _exportar_xlsx(cols, rows):
    """Genera un Excel en memoria con los datos recibidos. Genérica (no
    específica de ningún módulo) — vive acá porque tanto app.py (export de
    SINTIA) como blueprints/finanzas.py (export de movimientos) la necesitan.
    Bug encontrado en revisión: antes vivía solo en app.py, y
    blueprints/finanzas.py la llamaba sin importarla de ningún lado — tiraba
    NameError apenas se pedía exportar movimientos a Excel."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    wb = Workbook()
    ws = wb.active
    ws.append(cols)
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1E2A3B")
        cell.alignment = Alignment(horizontal="center")
    for row in rows:
        ws.append(row)
    for col in ws.columns:
        max_len = max((len(str(c.value or "")) for c in col), default=8)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 2, 40)
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def validar_enum(valor, permitidos, nombre_campo):
    """Devuelve (ok, error). ok=True si valor está entre los permitidos
    (o no vino, en cuyo caso no hay nada que validar — el default lo pone
    el INSERT/UPDATE). Usar así en cada endpoint:
        ok, err = validar_enum(data.get("estado"), ESTADOS_TAREA, "estado")
        if not ok: return jsonify({"ok": False, "error": err}), 400
    """
    if valor is None:
        return True, None
    if valor not in permitidos:
        return False, f"Valor no permitido para '{nombre_campo}': '{valor}'. Opciones válidas: {', '.join(permitidos)}."
    return True, None


# ── Job queue (progreso de generación de informes en background) ────────────
# job_status vivía solo en memoria del proceso: si la app corre con más de un
# worker (gunicorn -w >1), el worker que crea el job y el que atiende el
# polling de progreso pueden ser procesos distintos, y el segundo nunca ve el
# job (se ve como "colgado" del lado del usuario, sin ningún error visible).
# Cada cambio de estado se espeja a SQLite (tabla job_status_db), visible
# desde cualquier worker. job_status en memoria es además la caché rápida
# para el worker que efectivamente está corriendo el job.
#
# Vive en core.py (no en app.py) porque lo usan tanto app.py (informe SINTIA)
# como los blueprints (vua, y los que sigan) para sus propios informes async.
job_status = {}

def _init_job_status_db():
    try:
        with get_db(HIST_DB, timeout=10) as con:
            con.execute("""CREATE TABLE IF NOT EXISTS job_status_db (
                job_id TEXT PRIMARY KEY, status TEXT, log TEXT, files TEXT,
                username TEXT, ts REAL, progreso INTEGER DEFAULT 0)""")
    except Exception:
        logging.exception("No se pudo crear la tabla job_status_db")

_init_job_status_db()

_MARCA_HUERFANO_LOG_TXT = "✗ Proceso interrumpido: el servidor se reinició mientras este job corría."

def _marcar_jobs_huerfanos():
    """Si el proceso se reinició (deploy, crash, systemctl restart) a mitad de
    un job en background, ese job queda para siempre en status='running' en
    job_status_db — nadie lo va a terminar, y el usuario ve la barra de
    progreso colgada sin ningún error visible.

    Se corre una vez al importar core.py (o sea, una vez por worker). No basta
    con "¿está en running?" para decidir que está huérfano: con más de un
    worker (gunicorn -w >1), cada worker importa core.py por separado, y un
    worker que arranca unos segundos después de otro vería como "huérfano" un
    job que el primero recién empezó a correr — falso positivo por carrera de
    arranque, no por reinicio real.

    Por eso se usa una ventana de inactividad: un job que sigue 'running' pero
    no actualizó su 'ts' (se refresca en cada línea de log, ver
    _JobLog.append) hace más de UMBRAL_HUERFANO_SEG no tiene ningún thread
    vivo detrás, sea cual sea el motivo. Un job recién creado por otro worker
    tiene ts fresco y no se toca."""
    UMBRAL_HUERFANO_SEG = 20 * 60  # 20 min: más que el informe más largo esperado
    try:
        with get_db(HIST_DB, timeout=10, row_factory=True) as con:
            limite = time.time() - UMBRAL_HUERFANO_SEG
            rows = con.execute(
                "SELECT job_id, log FROM job_status_db WHERE status='running' AND ts < ?",
                (limite,)).fetchall()
            for r in rows:
                try:
                    log = json.loads(r["log"] or "[]")
                except Exception:
                    log = []
                log.append(_MARCA_HUERFANO_LOG_TXT)
                con.execute("UPDATE job_status_db SET status='error', log=? WHERE job_id=?",
                            (json.dumps(log), r["job_id"]))
            if rows:
                logging.warning(f"JOBS HUÉRFANOS | {len(rows)} job(s) en 'running' sin actividad "
                                 f"hace más de {UMBRAL_HUERFANO_SEG}s — marcados como 'error' al arrancar.")
    except Exception:
        logging.exception("No se pudo revisar jobs huérfanos en job_status_db")

_marcar_jobs_huerfanos()

# Progreso (%) estimado a partir del texto de los mensajes que ya emite
# generar.py (log_fn(...)) — no hace falta tocar ese archivo. Cubre además
# los mensajes típicos de VUA/SENASA/Stock ("Generando...", "completado",
# etc.), aunque el detalle fino (0-100 con hitos nombrados) es más preciso
# para los informes de generar.py. El progreso nunca retrocede.
_HITOS_PROGRESO = [
    ("proceso completado", 100), ("informe generado", 97), ("informe word generado", 90),
    ("planilla excel generada", 95), ("generando archivos", 88),
    ("gráficos generados", 85), ("generando gráficos", 75),
    ("conclusión generada", 68), ("conclusión no disponible", 68),
    ("generando conclusión", 58),
    ("narrativa generada", 50), ("narrativa no disponible", 50),
    ("generando narrativa", 40),
    ("queries completadas", 30), ("corriendo queries", 15),
    ("conectando a la bd", 5),
]

def _pct_desde_mensaje(mensaje, pct_actual):
    m = (mensaje or "").lower()
    if m.startswith("✗") or "error" in m:
        return pct_actual  # un error no debe hacer parecer que avanzó
    for clave, pct in _HITOS_PROGRESO:
        if clave in m:
            return max(pct_actual, pct)
    # Sin hito reconocido: igual dar una sensación de avance leve y acotada,
    # para que la barra no quede clavada en 0 en módulos sin hitos mapeados.
    return min(80, pct_actual + 3)

def _job_persist(job_id):
    """Espeja el estado actual de job_status[job_id] a SQLite."""
    info = job_status.get(job_id)
    if info is None:
        return
    try:
        with get_db(HIST_DB) as con:
            con.execute("""INSERT INTO job_status_db (job_id, status, log, files, username, ts, progreso)
                VALUES (?,?,?,?,?,?,?)
                ON CONFLICT(job_id) DO UPDATE SET
                    status=excluded.status, log=excluded.log, files=excluded.files,
                    ts=excluded.ts, progreso=excluded.progreso""",
                (job_id, info.get("status", "running"), json.dumps(list(info.get("log", []))),
                 json.dumps(info.get("files", [])), info.get("username", "?"), time.time(),
                 info.get("progreso", 0)))
    except Exception:
        logging.exception(f"job_status: no se pudo persistir job_id={job_id}")

class _JobLog(list):
    """Lista de progreso de un job. Persiste a SQLite en cada .append(), para
    que el progreso sea visible en vivo desde cualquier worker, no solo el
    que está corriendo el hilo de background. También actualiza el % de
    avance estimado a partir del mensaje agregado."""
    def __init__(self, job_id, inicial=None):
        super().__init__(inicial or [])
        self._job_id = job_id
    def append(self, item):
        super().append(item)
        info = job_status.get(self._job_id)
        if info is not None:
            info["progreso"] = _pct_desde_mensaje(item, info.get("progreso", 0))
        _job_persist(self._job_id)

def job_create(job_id, primer_mensaje="", username="?", status="running"):
    """Crea un job nuevo (en memoria + espejado en SQLite) y devuelve su dict,
    para usar igual que antes: job["status"]=..., job["files"]=..., etc."""
    job_status[job_id] = {
        "status": status,
        "log": _JobLog(job_id, [primer_mensaje] if primer_mensaje else []),
        "files": [],
        "username": username,
        "progreso": 0,
        "_ts": time.time(),
    }
    _job_persist(job_id)
    return job_status[job_id]

def job_get(job_id):
    """Lee el estado de un job: memoria primero (worker que lo creó, rápido);
    si no está ahí, cae a SQLite (el polling llegó a otro worker, o el
    proceso se reinició). Devuelve None si el job no existe en ningún lado."""
    info = job_status.get(job_id)
    if info is not None:
        return info
    try:
        with get_db(HIST_DB, row_factory=True) as con:
            row = con.execute(
                "SELECT status, log, files, username, progreso FROM job_status_db WHERE job_id=?",
                (job_id,)).fetchone()
        if not row:
            return None
        status, log_json, files_json, username, progreso = row
        return {"status": status, "log": json.loads(log_json or "[]"),
                "files": json.loads(files_json or "[]"), "username": username,
                "progreso": progreso or 0}
    except Exception:
        logging.exception(f"job_status: no se pudo leer job_id={job_id} de SQLite")
        return None


# ── Repositorio de documentos (contexto extra para prompts de IA) ───────────
# Compartido por vua/senasa/sintia.
MODULOS_REPOSITORIO = ("vua", "senasa", "sintia")
MODULOS_CON_CRONOLOGIA = {"vua": "vua_cronologia", "senasa": "senasa_cronologia"}

def get_api_key():
    return os.environ.get("ANTHROPIC_API_KEY", "")

def contexto_repositorio(modulo):
    """Concatena el texto completo de todos los documentos subidos para ese
    módulo, para inyectarlo como contexto adicional en los prompts de IA."""
    if modulo not in MODULOS_REPOSITORIO:
        return ""
    try:
        with get_db(HIST_DB, row_factory=True) as con:
            rows = con.execute(
                "SELECT nombre_archivo, contenido FROM doc_repositorio WHERE modulo=? ORDER BY creado",
                (modulo,)).fetchall()
    except Exception:
        return ""
    if not rows:
        return ""
    bloques = [f"--- Documento: {r['nombre_archivo']} ---\n{r['contenido']}" for r in rows]
    return ("\n\nContexto adicional: documentos de referencia subidos por el equipo "
            "(pueden incluir informes, antecedentes u otro material no estructurado). "
            "Usalos como información de fondo si son relevantes para la consulta:\n\n"
            + "\n\n".join(bloques))

def _extraer_texto_docx(file_storage):
    """Extrae texto plano (párrafos + tablas) de un .docx subido."""
    from docx import Document
    doc = Document(file_storage)
    partes = [p.text for p in doc.paragraphs if p.text.strip()]
    for tabla in doc.tables:
        for fila in tabla.rows:
            celdas = [c.text.strip() for c in fila.cells]
            if any(celdas):
                partes.append(" | ".join(celdas))
    return "\n".join(partes)

def _extraer_cronologia_de_texto(texto, modulo):
    """Le pide a la IA que identifique hechos con fecha (reuniones, hitos,
    decisiones) dentro de un documento y devuelve una lista de entradas
    {fecha, actividad, participantes} para sumar a la cronología del módulo."""
    import anthropic, httpx
    client = anthropic.Anthropic(api_key=get_api_key(), http_client=httpx.Client(follow_redirects=True))
    prompt = (
        "Del siguiente documento, identificá únicamente los hechos que tengan una fecha concreta "
        "asociada (reuniones, hitos, decisiones, entregas, cambios de estado). Ignorá contenido sin fecha. "
        "Devolvé SOLO un JSON: una lista de objetos con las claves \"fecha\" (formato DD/MM/AAAA si se puede "
        "inferir, si no la fecha tal como aparece en el texto), \"actividad\" (resumen breve, una oración) y "
        "\"participantes\" (si se mencionan, si no cadena vacía). Si no hay ningún hecho con fecha, devolvé [].\n\n"
        f"Documento:\n{texto[:12000]}"
    )
    msg = client.messages.create(model="claude-haiku-4-5-20251001", max_tokens=2000,
        system="Respondés solo con JSON válido (una lista), sin texto adicional ni markdown.",
        messages=[{"role": "user", "content": prompt}])
    import json as _json, re as _re
    texto_resp = msg.content[0].text.strip()
    texto_resp = _re.sub(r'^```(?:json)?\s*|\s*```$', '', texto_resp, flags=_re.MULTILINE).strip()
    try:
        entradas = _json.loads(texto_resp)
        return entradas if isinstance(entradas, list) else []
    except Exception:
        logging.error(f"No se pudo parsear JSON de cronología. Respuesta cruda: {texto_resp[:500]}")
        return []


# ── Fechas ────────────────────────────────────────────────────────────────────
import re as _re_mod

def _validar_fecha_ddmmaaaa(s):
    """True si s es una fecha real en formato dd/mm/aaaa (o el sentinel 'A definir')."""
    if s == "A definir":
        return True
    m = _re_mod.match(r'^(\d{2})/(\d{2})/(\d{4})$', s or "")
    if not m:
        return False
    d, mo, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
    try:
        datetime(y, mo, d)
        return True
    except ValueError:
        return False

def _normalizar_fecha_a_ddmmaaaa(s):
    """Intenta convertir distintos formatos comunes (aaaa-mm-dd, aaaa/mm/dd, dd-mm-aaaa)
    al estándar dd/mm/aaaa. Si no puede, devuelve 'A definir'."""
    s = (s or "").strip()
    if not s:
        return "A definir"
    if _validar_fecha_ddmmaaaa(s):
        return s
    m = _re_mod.match(r'^(\d{4})[-/](\d{2})[-/](\d{2})$', s)
    if m:
        y, mo, d = m.groups()
        cand = f"{d}/{mo}/{y}"
        if _validar_fecha_ddmmaaaa(cand):
            return cand
    m = _re_mod.match(r'^(\d{2})-(\d{2})-(\d{4})$', s)
    if m:
        d, mo, y = m.groups()
        cand = f"{d}/{mo}/{y}"
        if _validar_fecha_ddmmaaaa(cand):
            return cand
    return "A definir"


# ── Telegram ──────────────────────────────────────────────────────────────────
TELEGRAM_TOKEN = os.environ.get("TELEGRAM_BOT_TOKEN", "")
TELEGRAM_CHAT_ID = os.environ.get("TELEGRAM_CHAT_ID", "")

def notificar_telegram(msg: str):
    """Envía un mensaje al bot de Telegram. No rompe el flujo si falla.
    Corre en un hilo aparte para no bloquear la respuesta al usuario mientras
    espera la llamada de red (hasta 3s si Telegram está lento/caído)."""
    if not TELEGRAM_TOKEN:
        return
    def _enviar():
        try:
            urllib.request.urlopen(
                f"https://api.telegram.org/bot{TELEGRAM_TOKEN}/sendMessage"
                f"?chat_id={TELEGRAM_CHAT_ID}&text=" + urllib.parse.quote(msg),
                timeout=3
            )
        except Exception:
            pass
    threading.Thread(target=_enviar, daemon=True).start()
