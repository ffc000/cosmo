"""
actas.py — Generación del Word de "Acta de Reunión", compartida entre VUA y
SENASA (antes cada blueprint tenía su propia implementación, con estilos
distintos aunque el contenido de fondo era el mismo: reunión con
participantes/temas/acuerdos).

No se unificó al 100% porque las secciones no son idénticas — VUA usa
Temas/Acuerdos/Próximos pasos (3), SENASA usa Temas/Conclusiones/
Compromisos/Próximos pasos (4). Eso queda como parámetro (`secciones`) en
vez de hardcodeado, así cada módulo sigue mandando sus propias secciones sin
forzar un esquema común que no existe.

Nota de diseño: el formato de tabla de participantes (con encabezado azul
oscuro) es el que ya tenía VUA — al unificar, SENASA pasa a usar el mismo
formato en vez de su lista con viñetas anterior. Es un cambio de estilo
visible en los próximos Acta de SENASA que se generen (no en las ya
guardadas), a propósito: el objetivo de unificar es quedarse con la mejor
versión de las dos, no con el promedio.

Reusa header/pie de página y el estilo de tabla de generar_documento.py (Fase
8 de profesionalización -- unificar el diseño de los 3 generadores de Word
que había: SINTIA/consolidado, informe de Aduanas del país, y este). No lleva
índice (TOC): las actas son de 1-2 páginas, un índice ahí sería ruido, no
ayuda -- a diferencia de los informes largos donde sí aporta.

Sin dependencia de Flask -- generar_documento tampoco la tiene (solo
docx/openpyxl/matplotlib), así que sigue sin arrastrar Flask acá.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from generar_documento import agregar_tabla_word, _agregar_encabezado, _agregar_pie_pagina


def importar_minuta_desde_docx(file_storage, api_key, nombre_legible, contexto_ia=""):
    """Extrae texto de un .docx subido (acta/minuta externa) y usa IA para
    estructurarlo en el esquema unificado de minutas de esta app
    (temas/conclusiones/compromisos/proximos + notas_completas). Devuelve
    un dict listo para jsonify() directamente.

    Antes esto vivía duplicado a mano solo en blueprints/vua.py
    (vua_minuta_importar) -- se centraliza acá para poder ofrecer
    "Importar Acta existente" en los 4 módulos con panel de Minutas
    (unificación 23/07/2026) sin repetir la extracción de texto ni el
    prompt en cada blueprint.

    file_storage: el FileStorage de Flask (request.files['archivo']).
    nombre_legible: para el prompt y los mensajes de error, ej. "SENASA",
        "Pad Acuático", "SINTIA", "VUA".
    contexto_ia: opcional, texto adicional para el prompt del sistema
        (típicamente contexto_repositorio(modulo), documentos de
        referencia subidos para ese módulo)."""
    if not api_key:
        return {"ok": False, "error": "API key no configurada"}
    if not file_storage or not file_storage.filename:
        return {"ok": False, "error": "No se recibió archivo"}
    if not file_storage.filename.lower().endswith(".docx"):
        return {"ok": False, "error": "Solo se aceptan archivos .docx"}

    try:
        from docx import Document as DocxDoc
        import io as _io
        doc_bytes = file_storage.read()
        docx_doc = DocxDoc(_io.BytesIO(doc_bytes))
        partes = []
        for p in docx_doc.paragraphs:
            txt = p.text.strip()
            if txt: partes.append(txt)
        for tabla in docx_doc.tables:
            for fila in tabla.rows:
                celda_txt = " | ".join(c.text.strip() for c in fila.cells if c.text.strip())
                if celda_txt: partes.append(celda_txt)
        texto_completo = "\n".join(partes)
    except Exception as e:
        return {"ok": False, "error": f"Error leyendo el Word: {e}"}

    if not texto_completo.strip():
        return {"ok": False, "error": "El documento está vacío o no se pudo extraer texto"}

    try:
        import anthropic, httpx, json as _json
        client = anthropic.Anthropic(api_key=api_key, http_client=httpx.Client(follow_redirects=True))
        prompt = (
            f"El siguiente texto es una minuta/acta de reunión del proyecto {nombre_legible} "
            "(ARCA Argentina). Extraé y estructurá la información en el JSON solicitado, "
            "corrigiendo ortografía/gramática y completando frases truncadas en oraciones "
            "claras -- no te limites a copiar literal.\n\n"
            f"TEXTO DE LA MINUTA:\n{texto_completo[:6000]}\n\n"
            "Devolvé SOLO este JSON válido (sin markdown ni texto adicional):\n"
            "{\n"
            '  "asunto": "título o asunto principal de la reunión",\n'
            '  "fecha": "fecha en formato YYYY-MM-DD si la encontrás, sino vacío",\n'
            '  "lugar": "lugar o modalidad (ej: Videoconferencia, Sala 3 Paseo Colón)",\n'
            '  "participantes": [{"nombre":"...", "cargo":"...", "organismo":"..."}],\n'
            '  "temas": ["tema 1", "tema 2"],\n'
            '  "conclusiones": ["conclusión 1"],\n'
            '  "compromisos": ["ORG — compromiso concreto"],\n'
            '  "proximos": ["paso 1"],\n'
            '  "notas_completas": "el resto del contenido de la minuta, como texto corrido, '
            'para no perder matices que no encajan prolijo en las categorías de arriba"\n'
            "}\n\n"
            "Español rioplatense, con tildes y caracteres especiales correctos, sin markdown."
        )
        msg = client.messages.create(
            model="claude-haiku-4-5-20251001", max_tokens=3000,
            system=(f"Sos un asistente experto en gestión de proyectos aduaneros para ARCA "
                    f"Argentina, del área DI REPA. Extraés información estructurada de minutas "
                    f"institucionales de {nombre_legible}. Respondés solo con JSON válido."
                    + contexto_ia),
            messages=[{"role": "user", "content": prompt}])
        texto_resp = msg.content[0].text.strip().replace("```json", "").replace("```", "").strip()
        resultado = _json.loads(texto_resp)
        return {"ok": True, **resultado}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def generar_acta_word(titulo_doc, fecha, asunto, lugar, participantes, secciones,
                       roles_predefinidos=None, nombre_archivo="", notas_completas=""):
    """
    titulo_doc: encabezado del documento, ej. "ACTA DE REUNIÓN" (VUA) o
        "ACTA DE REUNIÓN — SENASA / ARCA" (SENASA).
    fecha, asunto, lugar: strings simples.
    participantes: lista de dicts {"nombre":..., "cargo":..., "organismo":...}
        (cargo/organismo opcionales) o strings sueltos (solo nombre).
    secciones: lista de tuplas (nombre_seccion, lista_de_items). Una sección
        con items=[] se omite del documento (no aparece un título vacío).
    roles_predefinidos: dict opcional {nombre: cargo} para completar el cargo
        de un participante si no vino explícito en el dict.
    nombre_archivo: opcional, para mostrar en el pie de página -- el caller
        es quien decide el nombre final del archivo (esta función no lo
        guarda), así que si no se pasa, el pie queda sin nombre de archivo.
    notas_completas: opcional -- texto corrido (ya corregido/completado por
        IA o escrito a mano) con las notas de la reunión tal cual, para
        agregar como sección final del acta. A pedido (23/07/2026): al
        estructurar la reunión en temas/conclusiones/compromisos/próximos
        se pierde matiz -- una nota puede tener contexto que no encaja
        prolijo en ninguna de esas categorías. Esta sección final conserva
        ese texto completo como respaldo, sin descartar nada. Si no se
        pasa (o viene vacío), no se agrega ninguna sección -- mismo
        comportamiento que antes para quien no la use.

    Devuelve un objeto Document (python-docx) ya armado — el caller decide
    dónde guardarlo (doc.save(ruta)).
    """
    roles_predefinidos = roles_predefinidos or {}
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)
    _agregar_encabezado(doc, "Dirección de Reingeniería de Procesos Aduaneros")
    _agregar_pie_pagina(doc, nombre_archivo)

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run(titulo_doc)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x24, 0x2D, 0x4F)
    doc.add_paragraph()

    for label, valor in [("Asunto:", asunto), ("Fecha:", fecha), ("Lugar:", lugar)]:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label} ")
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(valor or "")
        r2.font.size = Pt(11)
    doc.add_paragraph()

    if participantes:
        doc.add_paragraph().add_run("Participantes").bold = True
        filas = []
        for p in participantes:
            if isinstance(p, dict):
                nombre = p.get("nombre", "")
                cargo = p.get("cargo") or roles_predefinidos.get(nombre, "")
                organismo = p.get("organismo", "")
            else:
                nombre = str(p)
                cargo = roles_predefinidos.get(nombre, "")
                organismo = ""
            filas.append([nombre, cargo, organismo])
        agregar_tabla_word(doc, ["Nombre", "Cargo", "Organismo"], filas, col_widths=[5, 5, 5.5])

    for nombre_seccion, items in secciones:
        if not items:
            continue
        doc.add_paragraph()
        doc.add_paragraph().add_run(nombre_seccion).bold = True
        for item in items:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(item).font.size = Pt(11)

    if notas_completas and notas_completas.strip():
        doc.add_paragraph()
        doc.add_paragraph().add_run("Notas completas de la reunión").bold = True
        p_intro = doc.add_paragraph()
        r_intro = p_intro.add_run(
            "Texto de referencia con el detalle completo de la reunión, para consulta "
            "si algo no quedó reflejado en las secciones de arriba.")
        r_intro.italic = True
        r_intro.font.size = Pt(9)
        r_intro.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        for parrafo in notas_completas.strip().split("\n"):
            if parrafo.strip():
                doc.add_paragraph(parrafo.strip()).runs[0].font.size = Pt(10.5)

    return doc
