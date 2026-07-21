"""
generar_documento.py — Armado del Word y el Excel finales del informe
SINTIA (tablas, encabezados, TOC, inserción de gráficos y de la narrativa/
conclusión generadas por IA). Extraído de generar.py (Fase 3 de
profesionalización).
"""
import os
import re
import logging
import statistics
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    XLSX_OK = True
except ImportError:
    XLSX_OK = False

try:
    from PIL import Image, ImageDraw, ImageFont
    import io
    PIL_OK = True
except ImportError:
    PIL_OK = False

from generar_utils import PAISES, PAISES_CONSOLIDADO, fmt, pct, pct_f, n, pl, periodo_texto, mes_label, mes_label_largo, color_semaforo, formatear_demora
from generar_graficos import (grafico_torta, grafico_barras_apiladas, grafico_lineas_pct,
    grafico_rechazos_cat, grafico_rechazos_mes, grafico_comparativo_meses, MPL_OK,
    grafico_consolidado_pais, grafico_consolidado_impoexpo, grafico_consolidado_cargado_lastre,
    grafico_consolidado_aduana, grafico_consolidado_var_control, grafico_comparacion_interanual,
    grafico_controles_por_tipo)
from generar_ia import calcular_frases, limpiar_salida_ia
from generar_queries import calcular_totales

def set_cell_bg(cell, hex_color):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); shd=OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hex_color); tcPr.append(shd)
def agregar_tabla_word(doc, headers, rows, col_widths=None, semaforo_col=None, semaforo_total_col=None):
    table=doc.add_table(rows=1,cols=len(headers)); table.style='Table Grid'; table.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr=table.rows[0]
    for i,h in enumerate(headers):
        cell=hdr.cells[i]; cell.text=h
        cell.paragraphs[0].runs[0].bold=True
        cell.paragraphs[0].runs[0].font.size=Pt(9)
        cell.paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell,"1F3864")
    for ri,row in enumerate(rows):
        tr=table.add_row(); base_fill="F2F2F2" if ri%2==0 else "FFFFFF"
        for ci,val in enumerate(row):
            cell=tr.cells[ci]; cell.text=str(val or "")
            cell.paragraphs[0].runs[0].font.size=Pt(9)
            fill=base_fill
            if semaforo_col is not None and semaforo_total_col is not None and ci==semaforo_col:
                try:
                    v_trans=n(row[semaforo_col]); v_total=n(row[semaforo_total_col])
                    if v_total>0: fill=color_semaforo(pct_f(v_trans,v_total))
                except: pass
            set_cell_bg(cell,fill)
    if col_widths:
        for i,w in enumerate(col_widths):
            for row in table.rows: row.cells[i].width=Cm(w)
    doc.add_paragraph(); return table
def insertar_grafico(doc, img_bytes, width_cm=14):
    if not img_bytes: return
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img_bytes, width=Cm(width_cm)); doc.add_paragraph()
def _campo_word(p, instr):
    """Inserta un campo de Word (PAGE, NUMPAGES, TOC, etc.) que se calcula solo
    al abrir/actualizar el documento — no se puede calcular desde python-docx."""
    r1 = p.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin'); r1._r.append(f1)
    r2 = p.add_run()
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr
    r2._r.append(it)
    r3 = p.add_run()
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'separate'); r3._r.append(f2)
    r4 = p.add_run()
    f3 = OxmlElement('w:fldChar'); f3.set(qn('w:fldCharType'), 'end'); r4._r.append(f3)
    return r3
def _bookmark_parrafo(paragraph, ancla):
    """Marca un párrafo con un bookmark de Word (ancla interna) -- para que
    el índice manual pueda apuntarle con un hyperlink."""
    bmk_id = str(abs(hash(ancla)) % 1000000)
    start = OxmlElement('w:bookmarkStart'); start.set(qn('w:id'), bmk_id); start.set(qn('w:name'), ancla)
    end = OxmlElement('w:bookmarkEnd'); end.set(qn('w:id'), bmk_id)
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _hipervinculo_interno(paragraph, texto, ancla, sangria=False):
    """Agrega un run de hyperlink interno (a un bookmark) dentro de un
    párrafo -- python-docx no tiene esto nativo, se arma vía OXML."""
    hyperlink = OxmlElement('w:hyperlink'); hyperlink.set(qn('w:anchor'), ancla)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color'); color.set(qn('w:val'), '1F4E79')
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single')
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '20' if sangria else '22')
    rPr.append(color); rPr.append(u); rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = texto
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _heading_indexado(doc, texto, nivel, indice):
    """doc.add_heading() + bookmark, para que el índice manual pueda
    apuntarle. `indice` tiene que coincidir con la posición de este heading
    en la lista `secciones` que se le pasa a _insertar_indice() -- son el
    mismo número, es lo que conecta el link con el destino."""
    h = doc.add_heading(texto, level=nivel)
    _bookmark_parrafo(h, f"sec{indice}")
    return h


def _insertar_indice(doc, secciones):
    """Índice manual clickeable -- reemplaza el campo TOC nativo de Word
    (\"Índice \\o \"1-2\" ...\"), que depende de que el lector actualice
    campos al abrir el archivo (Word real a veces lo hace solo, pero no
    todos los visores lo hacen -- en la práctica varios informes quedaban
    con el índice vacío). Esto funciona apenas se abre el documento, en
    cualquier lector, sin pasos extra.

    secciones: lista de (nivel, texto), en el MISMO ORDEN en que después
    se crean los headings reales con _heading_indexado() más abajo en el
    documento -- la posición (1-based) es lo que conecta cada link con su
    destino."""
    doc.add_heading("Índice", level=1)
    for i, (nivel, texto) in enumerate(secciones, start=1):
        p = doc.add_paragraph()
        if nivel == 2:
            p.paragraph_format.left_indent = Cm(0.6)
        _hipervinculo_interno(p, texto, f"sec{i}", sangria=(nivel == 2))
    doc.add_page_break()


_ASSETS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets")
_LOGO_PATH = os.path.join(_ASSETS_DIR, "logo_arca.png")
_COVER_FONDO_PATH = os.path.join(_ASSETS_DIR, "cover_fondo.png")
_COVER_BANNER_PATH = os.path.join(_ASSETS_DIR, "cover_banner.png")
_FONT_BOLD = "/usr/share/fonts/truetype/crosextra/Carlito-Bold.ttf"
_FONT_REG = "/usr/share/fonts/truetype/crosextra/Carlito-Regular.ttf"
_AZUL_HEADER = RGBColor(0x1F, 0x4E, 0x79)   # color exacto tomado del template oficial ARCA
_GRIS_FOOTER = RGBColor(0x32, 0x3E, 0x4F)   # ídem
_NAVY_TITULO = (0x24, 0x2C, 0x4F)           # ídem, en RGB plano (para Pillow, no python-docx)
_AZUL_META = (0x1F, 0x4E, 0x79)


def _wrap_texto_pil(texto, font, max_width, draw):
    """Envuelve texto a un ancho máximo en píxeles, usando el font real
    (no por cantidad de caracteres -- una tipografía bold ancha entra
    distinto que una regular, así que se mide con textlength de verdad)."""
    palabras = texto.split()
    lineas, actual = [], ""
    for palabra in palabras:
        prueba = f"{actual} {palabra}".strip()
        if draw.textlength(prueba, font=font) <= max_width:
            actual = prueba
        else:
            if actual:
                lineas.append(actual)
            actual = palabra
    if actual:
        lineas.append(actual)
    return lineas


def _ocultar_encabezado_portada(doc):
    """Activa "primera página distinta" y deja el encabezado/pie de esa
    primera página en blanco -- para que la portada salga limpia (sin la
    barra de "DIRECCIÓN DE REINGENIERÍA..." arriba ni el pie abajo), y el
    encabezado/pie normal recién arranque desde la página 2 (el índice).
    A pedido, 21/07/2026 -- antes la portada mostraba encabezado y pie
    igual que el resto del documento.

    Tiene que llamarse DESPUÉS de _agregar_encabezado/_agregar_pie_pagina
    (que configuran section.header/section.footer -- una vez activado
    este flag, esos pasan a regir desde la página 2 en adelante, y
    section.first_page_header/first_page_footer son los que se ven en la
    página 1)."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    for p in section.first_page_header.paragraphs:
        p.text = ""
    for p in section.first_page_footer.paragraphs:
        p.text = ""


def _agregar_imagen_portada_ajustada(doc, imagen_bytesio, ancho_deseado_cm=16.09, margen_seguridad_cm=1.0):
    """Inserta la imagen de portada con un ancho que GARANTIZA que entre
    en una sola página, dejando aire de sobra -- encontrado en producción
    (21/07/2026): con el ancho fijo de 16.09cm, la altura resultante
    (22.76cm) dejaba apenas 0.18cm de margen contra el alto de página
    disponible en Letter (22.94cm). Cualquier diferencia mínima de cómo
    Word/LibreOffice miden el espacio empujaba la imagen a la página
    siguiente -- la portada y la página 2 quedaban casi en blanco (el
    índice terminaba recién en la página 3 en vez de la 2).

    Calcula el alto disponible de la página REAL del documento (no asume
    A4 ni Letter -- se leen page_height/top_margin/bottom_margin del
    section actual) y, si la imagen al ancho deseado no entra con
    margen_seguridad_cm de aire, la achica proporcionalmente hasta que sí
    entre -- nunca la deja al límite exacto como antes."""
    imagen_bytesio.seek(0)
    with Image.open(imagen_bytesio) as img:
        ancho_px, alto_px = img.size
    imagen_bytesio.seek(0)
    aspect_ratio = alto_px / ancho_px

    section = doc.sections[0]
    alto_disponible_cm = Emu(section.page_height).cm - Emu(section.top_margin).cm - Emu(section.bottom_margin).cm
    alto_maximo_cm = alto_disponible_cm - margen_seguridad_cm

    alto_con_ancho_deseado_cm = ancho_deseado_cm * aspect_ratio
    ancho_final_cm = (alto_maximo_cm / aspect_ratio) if alto_con_ancho_deseado_cm > alto_maximo_cm else ancho_deseado_cm

    doc.add_picture(imagen_bytesio, width=Cm(ancho_final_cm))


def _generar_portada_compuesta(titulo, subtitulo, meta_lineas):
    """Compone la imagen de portada (Fase 8: tomada del template oficial de
    ARCA -- Instructivo de Acceso Remoto SAR): fondo diagonal + banner con
    logo (assets/cover_fondo.png y cover_banner.png, extraídos en alta
    resolución del PDF oficial) con el título/subtítulo/metadata de CADA
    informe compuestos encima vía Pillow. Se genera de nuevo en cada
    informe (no es un asset estático) porque el texto cambia -- país,
    período, versión, etc.

    titulo: string, se envuelve automáticamente si no entra en una línea
        (ej. el título del consolidado es mucho más largo que "INSTRUCTIVO"
        del documento original).
    subtitulo: una línea (ej. "Período: 01/01/2026 a 31/01/2026").
    meta_lineas: lista de líneas chicas debajo (dirección, versión, fecha,
        elaborado por).

    Devuelve un BytesIO con el PNG compuesto, o None si Pillow no está
    disponible o falta algún asset -- el caller debe tener un fallback de
    portada en texto plano para ese caso (ver PIL_OK)."""
    if not PIL_OK or not os.path.exists(_COVER_FONDO_PATH) or not os.path.exists(_COVER_BANNER_PATH):
        return None
    try:
        fondo = Image.open(_COVER_FONDO_PATH).convert("RGB")
        banner = Image.open(_COVER_BANNER_PATH).convert("RGB")
        W, H = fondo.size
        escala = W / 560  # 560pt = ancho de página del PDF original de referencia

        bx0, by0 = 224.1 * escala, 88.5 * escala
        bx1, by1 = 564.9 * escala, 147.7 * escala
        banner_r = banner.resize((int(bx1 - bx0), int(by1 - by0)))
        fondo.paste(banner_r, (int(bx0), int(by0)))

        draw = ImageDraw.Draw(fondo)
        x_texto = 900  # calibrado contra el asset real (1831px de ancho)
        max_w = W - x_texto - 60

        y = int(640 * (H / 2590))
        font_titulo = ImageFont.truetype(_FONT_BOLD, int(46 * (H / 2590)))
        for linea in _wrap_texto_pil(titulo, font_titulo, max_w, draw):
            draw.text((x_texto, y), linea, font=font_titulo, fill=_NAVY_TITULO)
            y += int(58 * (H / 2590))

        y += int(40 * (H / 2590))
        font_sub = ImageFont.truetype(_FONT_REG, int(34 * (H / 2590)))
        for linea in _wrap_texto_pil(subtitulo, font_sub, max_w, draw):
            draw.text((x_texto, y), linea, font=font_sub, fill=_NAVY_TITULO)
            y += int(45 * (H / 2590))

        y += int(45 * (H / 2590))
        font_meta = ImageFont.truetype(_FONT_REG, int(26 * (H / 2590)))
        for linea in meta_lineas:
            draw.text((x_texto, y), linea, font=font_meta, fill=_AZUL_META)
            y += int(45 * (H / 2590))

        buf = io.BytesIO()
        fondo.save(buf, format="PNG")
        buf.seek(0)
        return buf
    except Exception:
        logging.exception("PORTADA COMPUESTA | fall\u00f3, se usa fallback de texto plano")
        return None


def _borde_tabla_lado(tabla, lado, color_hex="1F4E79", size=6):
    """Agrega un borde a UN lado de una tabla completa (línea horizontal de
    ancho completo) -- a diferencia de un borde de párrafo, que solo abarca
    la celda en la que está, esto cruza las 2 columnas de header/footer.
    python-docx no tiene esto nativo, se arma vía OXML."""
    tblPr = tabla._tbl.tblPr
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
    borde = OxmlElement(f'w:{lado}')
    borde.set(qn('w:val'), 'single'); borde.set(qn('w:sz'), str(size))
    borde.set(qn('w:space'), '4'); borde.set(qn('w:color'), color_hex)
    tblBorders.append(borde)


def _agregar_pie_pagina(doc, titulo_doc="", codigo_area="SSPO#DVMPAD#DESYFC#DIREPA"):
    """Pie de página institucional (Fase 8: tomado del template oficial de
    ARCA -- Instructivo de Acceso Remoto SAR, Dirección de Seguridad de la
    Información). A diferencia del original (que pone el nombre de archivo
    a la izquierda), acá va el TÍTULO del documento -- decisión explícita:
    el nombre de archivo no le dice nada a quien lo lee, el título sí.
    Código de área + página a la derecha, línea divisoria arriba."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.text = ""

    tabla = footer.add_table(rows=1, cols=2, width=Cm(16))
    tabla.autofit = False
    celda_izq, celda_der = tabla.rows[0].cells
    celda_izq.width = Cm(10); celda_der.width = Cm(6)

    r0 = celda_izq.paragraphs[0].add_run(titulo_doc)
    r0.font.size = Pt(8.5); r0.font.name = "Calibri"; r0.font.color.rgb = _GRIS_FOOTER

    p2 = celda_der.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    prefijo = f"[{codigo_area}]    " if codigo_area else ""
    r1 = p2.add_run(f"{prefijo}Pág. ")
    r1.font.size = Pt(8.5); r1.font.name = "Calibri"; r1.font.color.rgb = _GRIS_FOOTER
    r_pagina = _campo_word(p2, "PAGE")
    r_pagina.font.size = Pt(8.5); r_pagina.font.name = "Calibri"; r_pagina.font.color.rgb = _GRIS_FOOTER

    _borde_tabla_lado(tabla, "top", "1F4E79", 6)


def _agregar_encabezado(doc, direccion):
    """Encabezado institucional (Fase 8, mismo template oficial): nombre de
    la dirección a la izquierda, logo ARCA real a la derecha (assets/
    logo_arca.png, extraído en alta resolución del PDF oficial), línea
    divisoria abajo. Sin "Información Reservada Uso Interno" -- estaba en
    el original pero se decidió no incluirlo."""
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.text = ""

    tabla = header.add_table(rows=1, cols=2, width=Cm(16))
    tabla.autofit = False
    celda_izq, celda_der = tabla.rows[0].cells
    celda_izq.width = Cm(10); celda_der.width = Cm(6)

    r0 = celda_izq.paragraphs[0].add_run(direccion.upper())
    r0.font.size = Pt(9); r0.font.color.rgb = _AZUL_HEADER

    p1 = celda_der.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if os.path.exists(_LOGO_PATH):
        p1.add_run().add_picture(_LOGO_PATH, width=Cm(4.3))
    else:
        r1 = p1.add_run("ARCA")
        r1.bold = True; r1.font.size = Pt(14); r1.font.color.rgb = _AZUL_HEADER

    _borde_tabla_lado(tabla, "bottom", "1F4E79", 6)
def kpi_box(doc, kpis, col_width_cm=None):
    """col_width_cm: ancho fijo por columna (cm). Si no se pasa, se reparte
    16cm (ancho útil aprox. de la página con los márgenes de este informe)
    entre la cantidad de KPIs.

    Antes esto no fijaba ancho de columna ni autofit=False -- Word (y
    algunos visores) autoajustaban las columnas según el contenido de TODAS
    las celdas, y un valor largo como "2.517.790" (9 caracteres) en una
    tabla de 5 columnas terminaba cortado en dos líneas ("2.517.79" / "0")
    porque no hay espacio en la palabra para partir -- encontrado en el
    informe consolidado, 17/07/2026. Con ancho fijo y suficiente (además de
    autofit=False, que si no Word puede seguir angostando columnas al
    abrir/guardar) el número entra en una sola línea.

    El tamaño de fuente del valor se calcula UNA vez para TODA la fila (en
    base al valor más largo de todos los KPIs), no por KPI individual --
    la primera versión de este fix decidía el tamaño por separado para
    cada uno, y en una fila con la mayoría de valores largos y uno corto
    (ej. "2.517.790"/"1.265.071"/"1.252.719"/"1.788.694" de 9 caracteres
    junto a "729.095" de 7), ese uno quedaba en letra visiblemente más
    grande que el resto -- se veía como un error de formato en vez de una
    fila pareja (encontrado en el mismo informe, misma fecha)."""
    table=doc.add_table(rows=1,cols=len(kpis)); table.alignment=WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    w = col_width_cm or (16 / len(kpis))
    largo_max = max(len(str(valor)) for _, valor, _ in kpis)
    tam_valor = Pt(12) if largo_max >= 9 else Pt(16)
    for i,(label,valor,sub) in enumerate(kpis):
        cell=table.rows[0].cells[i]; set_cell_bg(cell,"EBF2FA")
        cell.width = Cm(w)
        p1=cell.paragraphs[0]; p1.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p1.add_run(f"{label}\n").font.size=Pt(8)
        r2=p1.add_run(f"{valor}\n"); r2.bold=True; r2.font.size=tam_valor; r2.font.color.rgb=RGBColor(0x1F,0x3D,0x64)
        r3=p1.add_run(sub); r3.font.size=Pt(8); r3.font.color.rgb=RGBColor(0x60,0x60,0x60)
    doc.add_paragraph()

# ── Narrativa IA ─────────────────────────────────────────────────────────────────
def _generar_word(pais, anio, mes_d, mes_h, version,
                  totales_raw, ev_total, ev_trans, ev_tardio, ev_no_trans,
                  rechazos_mes, rechazos_cat, rechazos_ej,
                  datos_ult, datos_ant, datos_interanual, per_ult, per_ant, anio_ant,
                  impoexpo_ult, rechazos_ult_cat, total_rech_ant,
                  narrativa_ia, conclusion_ia, carpeta, log_fn):

    pais_nombre=PAISES.get(pais,pais)
    nombre_archivo = f"Informe_SINTIA_{pais}_{anio}_{mes_d}-{mes_h}_v{version}.docx"
    # Usar el último mes con datos reales, no el mes_h declarado
    mes_h_real = per_ult[-2:] if per_ult else mes_h
    periodo=periodo_texto(anio,mes_d,mes_h_real)
    (cT,cN,cTd,cTot),(lT,lN,lTd,lTot),(gT,gN,gTd,gTot)=calcular_totales(totales_raw)
    total_rechazos=sum(n(r.get("MIC_RECHAZOS",0)) for r in rechazos_mes)

    graficos={}
    if MPL_OK:
        log_fn("Generando gr\u00e1ficos...")
        for nombre,fn in [
            ("torta",  lambda: grafico_torta(gT,gN,gTd)),
            ("barras", lambda: grafico_barras_apiladas(ev_total)),
            ("lineas", lambda: grafico_lineas_pct(ev_total,ev_trans,ev_tardio,ev_no_trans)),
            ("rech_mes", lambda: grafico_rechazos_mes(rechazos_mes)),
            ("rech_cat", lambda: grafico_rechazos_cat(rechazos_cat)),
            ("comparativo", lambda: grafico_comparativo_meses(datos_ult,datos_ant,per_ult,per_ant)),
        ]:
            try: graficos[nombre]=fn()
            except Exception as e: log_fn(f"  Gr\u00e1fico {nombre}: {e}")
        log_fn("✓ Gr\u00e1ficos generados")

    doc=Document()
    for section in doc.sections:
        section.top_margin=Cm(2.5); section.bottom_margin=Cm(2.5)
        section.left_margin=Cm(3); section.right_margin=Cm(2.5)
    _agregar_encabezado(doc, "Dirección de Reingeniería de Procesos Aduaneros")
    _agregar_pie_pagina(doc, f"Estado de Situación SINTIA {anio} {pais}-AR")
    _ocultar_encabezado_portada(doc)

    # Portada
    _imagen_portada = _generar_portada_compuesta(
        titulo=f"ESTADO DE SITUACIÓN SINTIA {anio} {pais}-AR",
        subtitulo=f"Período: {periodo}",
        meta_lineas=[
            "Dirección de Reingeniería de Procesos Aduaneros (DG ADUA)",
            f"Versión: {version}     Última modificación: {datetime.today().strftime('%d/%m/%Y')}",
            "Elaborado por: Sección Simplificación de Procesos Operativos — DI REPA",
        ])
    if _imagen_portada:
        _agregar_imagen_portada_ajustada(doc, _imagen_portada)
    else:
        titulo=doc.add_heading(f"ESTADO DE SITUACI\u00d3N SINTIA {anio} {pais}-AR",0)
        titulo.alignment=WD_ALIGN_PARAGRAPH.CENTER
        for txt,sz in [("Direcci\u00f3n de Reingeniería de Procesos Aduaneros (DG ADUA)",12),(f"Per\u00edodo: {periodo}",11)]:
            p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            run=p.add_run(txt); run.font.size=Pt(sz); run.font.color.rgb=RGBColor(0x40,0x40,0x40)
        doc.add_paragraph()
        meta=doc.add_paragraph(); meta.alignment=WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run("Versi\u00f3n: ").bold=True; meta.add_run(f"{version}   ")
        meta.add_run("\u00daltima modificaci\u00f3n: ").bold=True; meta.add_run(datetime.today().strftime("%d/%m/%Y"))
        dest=doc.add_paragraph(); dest.alignment=WD_ALIGN_PARAGRAPH.CENTER
        dest.add_run("Elaborado por: ").bold=True; dest.add_run("Secci\u00f3n Simplificaci\u00f3n de Procesos Operativos — DI REPA")
    doc.add_page_break()

    # Índice (manual, clickeable -- ver _insertar_indice)
    _insertar_indice(doc, [
        (1, "Resumen Ejecutivo"),
        (1, "1.  Introducción"),
        (1, "2.  Estado de Situación"),
        (2, "2.1.  Evolución de ingreso de camiones por mes"),
        (2, "2.2.  Evolución de transmisión anticipada por mes"),
        (2, "2.3.  Evaluación del tiempo de transmisión – MICs Transmitidos"),
        (2, "2.4.  Evaluación del tiempo de transmisión – MICs tardíos"),
        (2, "2.5.  Evaluación del tiempo de transmisión – MICs no transmitidos"),
        (2, "2.6.  Análisis de MICs Rechazados"),
        (1, "3.  Conclusiones y estado actual"),
        (1, "Anexo — Glosario de siglas"),
    ])

    # Resumen ejecutivo — KPIs + veredicto de una línea, calculado (no generado por IA)
    # para que sea 100% determinístico y no dependa de la disponibilidad de la API.
    _heading_indexado(doc, "Resumen Ejecutivo", 1, 1)
    pct_trans_num = pct_f(gT, gTot) if gTot > 0 else 0
    if pct_trans_num >= 80:
        veredicto = "El circuito muestra un desempeño sólido en la transmisión anticipada del MIC-DTA, con un volumen acotado de rechazos."
    elif pct_trans_num >= 50:
        veredicto = "El circuito muestra un desempeño parcial en la transmisión anticipada del MIC-DTA, con margen de mejora en los indicadores de sincronización."
    else:
        veredicto = "El circuito evidencia baja transmisión anticipada del MIC-DTA en el período, un punto a seguir de cerca en los próximos meses."
    doc.add_paragraph(veredicto)
    kpi_box(doc, [
        ("TOTAL INGRESOS",  fmt(gTot),          periodo),
        ("TRANSMITIDOS",    fmt(gT),             pct(gT,gTot)),
        ("NO TRANSMITIDOS", fmt(gN),             pct(gN,gTot)),
        ("TARD\u00cdOS",         fmt(gTd),            pct(gTd,gTot)),
        ("RECHAZOS ¹",        fmt(total_rechazos), pct(total_rechazos,gTot)),
    ])
    # Nota aclaratoria: rechazos no son categoría excluyente
    nota = doc.add_paragraph()
    nota_run1 = nota.add_run("¹ ")
    nota_run1.bold = True
    nota_run1.font.size = Pt(8)
    nota_run2 = nota.add_run("Los rechazos son intentos fallidos de transmisión y pueden superponerse con las otras categorías. No representan una categoría excluyente del total de ingresos.")
    nota_run2.font.size = Pt(8)
    nota_run2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    doc.add_page_break()

    # 1. Introducción
    _heading_indexado(doc, "1.  Introducci\u00f3n", 1, 2)
    if narrativa_ia and len(narrativa_ia)>=1:
        doc.add_paragraph(narrativa_ia[0])
    else:
        anio_ant_intro = str(int(anio)-1)
        doc.add_paragraph(f"El presente informe tiene por finalidad exponer el estado de situaci\u00f3n del sistema SINTIA en el circuito operativo entre {pais_nombre} y Argentina, a partir del an\u00e1lisis del grado de transmisi\u00f3n, rechazos y principales inconsistencias detectadas mediante el cruzamiento de informaci\u00f3n con el sistema interno Portal Aduanero (PAD).")
        doc.add_paragraph(f"Para su elaboraci\u00f3n, se consideraron las operaciones terrestres registradas durante el per\u00edodo comprendido entre {periodo}, alcanzando un total de {fmt(gTot)} ingresos al territorio nacional. Este an\u00e1lisis permite dar continuidad a los resultados observados durante el a\u00f1o {anio_ant_intro}, evidenciando tendencias operativas que se mantienen en el tiempo.")
        doc.add_paragraph("En este contexto, el PAD constituye el sistema central de registro de las operaciones de ingreso y egreso terrestre del pa\u00eds, mientras que SINTIA cumple un rol clave en la transmisi\u00f3n anticipada del MIC-DTA, elemento fundamental para la correcta gesti\u00f3n operativa. La adecuada sincronizaci\u00f3n entre ambos sistemas resulta determinante para la optimizaci\u00f3n de los circuitos operativos y la mejora en los tiempos de registro, control y despacho.")
        doc.add_paragraph(f"Del an\u00e1lisis realizado se desprende que, durante el per\u00edodo bajo estudio, persisten desv\u00edos en el proceso de transmisi\u00f3n, con una proporci\u00f3n limitada de MICs transmitidos en forma anticipada ({pct(gT,gTot)}), presencia de transmisiones tard\u00edas ({pct(gTd,gTot)}) y un volumen significativo de operaciones no transmitidas ({pct(gN,gTot)}). Asimismo, se identificaron {fmt(total_rechazos)} MICs que por inconsistencias son rechazados.")
        doc.add_paragraph("A continuaci\u00f3n, se desarrollan en detalle los principales indicadores, su evoluci\u00f3n y los resultados estad\u00edsticos obtenidos, a fin de facilitar el an\u00e1lisis de la operatoria relevada.")
    doc.add_paragraph("Metodolog\u00eda: los datos surgen del cruzamiento entre SINTIA y el Portal Aduanero (PAD) para el per\u00edodo y circuito indicados en la portada; los indicadores se calculan sobre el total de ingresos terrestres registrados, sin proyecciones ni estimaciones.").runs[0].font.size=Pt(9)



    # 2. Estado de situación
    _heading_indexado(doc, "2.  Estado de Situaci\u00f3n", 1, 3)
    if narrativa_ia and len(narrativa_ia)>=2:
        doc.add_paragraph(narrativa_ia[1])
    else:
        doc.add_paragraph(f"Durante el per\u00edodo ingresaron {fmt(gTot)} operaciones: {fmt(gT)} ({pct(gT,gTot)}) transmitidas correctamente, {fmt(gN)} ({pct(gN,gTot)}) no transmitidas y {fmt(gTd)} ({pct(gTd,gTot)}) tard\u00edas.")
    doc.add_paragraph()

    agregar_tabla_word(doc,
        ["","TRANS","%","NO TRANS","%","TARD\u00cdO","%","TOTAL"],
        [["CARGADO",fmt(cT),pct(cT,cTot),fmt(cN),pct(cN,cTot),fmt(cTd),pct(cTd,cTot),fmt(cTot)],
         ["LASTRE", fmt(lT),pct(lT,lTot),fmt(lN),pct(lN,lTot),fmt(lTd),pct(lTd,lTot),fmt(lTot)],
         ["TOTAL",  fmt(gT),pct(gT,gTot),fmt(gN),pct(gN,gTot),fmt(gTd),pct(gTd,gTot),fmt(gTot)]],
        col_widths=[2.5,1.8,1.2,1.8,1.2,1.8,1.2,1.8], semaforo_col=1, semaforo_total_col=7)

    # Nota informativa para meses sin transmisiones (no implica error — puede ser comportamiento esperado del circuito)
    # Se considera "sin transmisión" cuando el % redondea a 0,0%, no solo cuando es exactamente 0
    meses_sin_trans = []
    for _r in ev_total:
        _mes = _r["MES"]
        _tot = n(_r.get("CARGADO",0)) + n(_r.get("LASTRE",0))
        _trans = next((n(x.get("CARGADOS",0))+n(x.get("LASTRE",0)) for x in ev_trans if x.get("MES")==_mes), 0)
        if _tot > 0 and pct_f(_trans, _tot) == 0.0:
            meses_sin_trans.append(mes_label_largo(_mes))
    if meses_sin_trans:
        nota_m = doc.add_paragraph()
        nota_m.add_run("Nota: ").bold = True
        meses_txt = ", ".join(meses_sin_trans)
        nota_m_run = nota_m.add_run(
            f"En {meses_txt} la transmisión anticipada se mantuvo en torno a 0,0%. "
            "Este comportamiento puede responder a la dinámica propia del circuito durante ese período."
        )
        nota_m_run.font.size = Pt(9)
        nota_m_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    if "torta" in graficos: insertar_grafico(doc,graficos["torta"],width_cm=11)

    # 2.1 Evolución mensual
    _heading_indexado(doc, "2.1.  Evoluci\u00f3n de ingreso de camiones por mes", 2, 4)
    agregar_tabla_word(doc,["MES","CARGADO","LASTRE","TOTAL"],
        [[mes_label_largo(r["MES"]),fmt(n(r.get("CARGADO",0))),fmt(n(r.get("LASTRE",0))),fmt(n(r.get("CARGADO",0))+n(r.get("LASTRE",0)))] for r in ev_total],
        col_widths=[3.5,2.5,2.5,2.5])
    if "barras" in graficos: insertar_grafico(doc,graficos["barras"])

    # 2.3 Transmitidos
    doc.add_page_break()
    _heading_indexado(doc, "2.2.  Evolución de transmisión anticipada por mes", 2, 5)
    doc.add_paragraph()
    rows_ev_trans = []
    for r in ev_total:
        mes = r["MES"]
        tot = n(r.get("CARGADO",0)) + n(r.get("LASTRE",0))
        trs = next((n(x.get("CARGADOS",0))+n(x.get("LASTRE",0)) for x in ev_trans if x.get("MES")==mes), 0)
        ntrs = next((n(x.get("CARGADOS",0))+n(x.get("LASTRE",0)) for x in ev_no_trans if x.get("MES")==mes), 0)
        tds = next((n(x.get("CARGADOS",0))+n(x.get("LASTRE",0)) for x in ev_tardio if x.get("MES")==mes), 0)
        rows_ev_trans.append([mes_label(mes), pct(trs,tot), pct(ntrs,tot), pct(tds,tot), fmt(tot)])
    agregar_tabla_word(doc,["MES","% TRANS","% NO TRANS","% TARDÍO","TOTAL"],rows_ev_trans,col_widths=[2.8,2.5,2.5,2.5,2.2])
    doc.add_paragraph()
    _heading_indexado(doc, "2.3.  Evaluaci\u00f3n del tiempo de transmisi\u00f3n \u2013 MICs Transmitidos", 2, 6)
    doc.add_paragraph(f"Del total ({fmt(gTot)}), {fmt(gT)} ({pct(gT,gTot)}) fueron transmitidos correctamente.")
    agregar_tabla_word(doc,["MES","CARGADOS","LASTRE","TOTAL"],
        [[mes_label_largo(r["MES"]),fmt(n(r.get("CARGADOS",0))),fmt(n(r.get("LASTRE",0))),fmt(n(r.get("CARGADOS",0))+n(r.get("LASTRE",0)))] for r in ev_trans],
        col_widths=[3.5,2.5,2.5,2.5])

    # 2.4 Tardíos
    _heading_indexado(doc, "2.4.  Evaluaci\u00f3n del tiempo de transmisi\u00f3n \u2013 MICs tard\u00edos", 2, 7)
    doc.add_paragraph(f"Del total ({fmt(gTot)}), {fmt(gTd)} ({pct(gTd,gTot)}) fueron transmitidos tard\u00edamente.")
    tardio_por_mes = {r["MES"]: r for r in ev_tardio}
    rows_tardio = []
    for r in ev_total:
        mes = r["MES"]
        td = tardio_por_mes.get(mes, {})
        carg = n(td.get("CARGADOS", 0))
        last = n(td.get("LASTRE", 0))
        rows_tardio.append([mes_label_largo(mes), fmt(carg), fmt(last), fmt(carg + last)])
    agregar_tabla_word(doc,["MES","CARGADOS","LASTRE","TOTAL"],rows_tardio,
        col_widths=[3.5,2.5,2.5,2.5])

    # 2.5 No transmitidos
    _heading_indexado(doc, "2.5.  Evaluaci\u00f3n del tiempo de transmisi\u00f3n \u2013 MICs no transmitidos", 2, 8)
    doc.add_paragraph(f"Del total ({fmt(gTot)}), {fmt(gN)} ({pct(gN,gTot)}) no fueron transmitidos.")
    notrans_por_mes = {r["MES"]: r for r in ev_no_trans}
    rows_notrans = []
    for r in ev_total:
        mes = r["MES"]
        nt = notrans_por_mes.get(mes, {})
        carg = n(nt.get("CARGADOS", 0))
        last = n(nt.get("LASTRE", 0))
        rows_notrans.append([mes_label_largo(mes), fmt(carg), fmt(last), fmt(carg + last)])
    agregar_tabla_word(doc,["MES","CARGADOS","LASTRE","TOTAL"],rows_notrans,
        col_widths=[3.5,2.5,2.5,2.5])
    if "lineas" in graficos: insertar_grafico(doc,graficos["lineas"])

    # 2.6 Rechazos
    _heading_indexado(doc, "2.6.  An\u00e1lisis de MICs Rechazados", 2, 9)
    if narrativa_ia and len(narrativa_ia)>=3:
        texto_26 = narrativa_ia[2] if narrativa_ia and len(narrativa_ia) > 2 and narrativa_ia[2] and len(narrativa_ia[2]) > 20 else None
        if texto_26:
            doc.add_paragraph(texto_26)
        else:
            doc.add_paragraph(f"Se registraron {fmt(total_rechazos)} rechazos en el período ({pct(total_rechazos,gTot)} de las operaciones totales).")
    else:
        doc.add_paragraph(f"Se registraron {fmt(total_rechazos)} rechazos ({pct(total_rechazos,gTot)} de las operaciones).")
    if rechazos_mes:
        doc.add_paragraph("Rechazos por mes:").runs[0].bold=True
        rech_por_mes = {r["periodo"]: n(r.get("MIC_RECHAZOS",0)) for r in rechazos_mes}
        periodos_todos = sorted(set([x["MES"] for x in ev_total] + list(rech_por_mes.keys())))
        agregar_tabla_word(doc,["MES","MIC RECHAZOS"],
            [[mes_label_largo(p), fmt(rech_por_mes.get(p,0))] for p in periodos_todos],col_widths=[4,3])
    if "rech_mes" in graficos: insertar_grafico(doc,graficos["rech_mes"],width_cm=12)
    if rechazos_cat:
        cats_sin_total=[r for r in rechazos_cat if r["Categoria"]!="TOTAL"]
        total_cat=sum(n(r.get("Rechazos",0)) for r in cats_sin_total)
        agregar_tabla_word(doc,["CATEGOR\u00cdA","RECHAZOS","%"],
            [[r.get("Categoria",""),fmt(r.get("Rechazos",0)),pct(r.get("Rechazos",0),total_cat) if r["Categoria"]!="TOTAL" else ""] for r in rechazos_cat],
            col_widths=[7,2.5,2])
    if "rech_cat" in graficos: insertar_grafico(doc,graficos["rech_cat"])
    if rechazos_ej:
        doc.add_paragraph("Ejemplos de rechazos:").runs[0].bold=True
        agregar_tabla_word(doc,["FECHA","NRO MIC","MENSAJE"],
            [[r.get("Fecha_ISO","")[:10] if r.get("Fecha_ISO") else "",r.get("NroMic",""),r.get("Mensaje","")] for r in rechazos_ej],
            col_widths=[2.5,3.0,9.0])

    # 3. Conclusiones
    doc.add_page_break()
    _heading_indexado(doc, "3.  Conclusiones y estado actual", 1, 10)

    ult_t=n(datos_ult.get("total",0)) if datos_ult else 0
    ult_tr=n(datos_ult.get("trans",0)) if datos_ult else 0
    ult_nt=n(datos_ult.get("no_trans",0)) if datos_ult else 0
    ult_td=n(datos_ult.get("tardio",0)) if datos_ult else 0
    ant_t=n(datos_ant.get("total",0)) if datos_ant else 0
    ant_tr=n(datos_ant.get("trans",0)) if datos_ant else 0
    ant_nt=n(datos_ant.get("no_trans",0)) if datos_ant else 0
    ant_td=n(datos_ant.get("tardio",0)) if datos_ant else 0

    if conclusion_ia:
        # Limpiar encabezados sueltos que la IA puede generar al inicio
        conc_limpia = conclusion_ia.strip()
        # Limpiar encabezados/artefactos al inicio del texto
        enc_limpiar = [
            "# 3. Conclusiones y estado actual", "# CONCLUSIONES Y ESTADO ACTUAL",
            "# Conclusiones y estado actual", "3. Conclusiones y estado actual",
            "CONCLUSIONES Y ESTADO ACTUAL", "Conclusiones y estado actual",
            "**Conclusiones y estado actual**", "**CONCLUSIONES Y ESTADO ACTUAL**",
        ]
        for enc in enc_limpiar:
            if conc_limpia.startswith(enc):
                conc_limpia = conc_limpia[len(enc):].strip()
        # Limpiar también líneas sueltas con "— Circuito..." o subtítulos redundantes
        lineas_conc = conc_limpia.split("\n")
        lineas_filtradas = []
        for lin in lineas_conc:
            ls = lin.strip()
            # Eliminar líneas que son solo subtítulos de sección o artefactos
            ls_norm = ls.strip("*#-— .:").lower()
            es_artefacto = (
                ls_norm in [
                    "conclusiones y estado actual",
                    "3 conclusiones y estado actual",
                    "3. conclusiones y estado actual",
                    "3.  conclusiones y estado actual",
                    "conclusiones",
                ] or
                ls_norm.startswith("circuito ") or
                ls_norm.startswith("conclusiones y") or
                ls_norm.startswith("3. conclusiones") or
                ls_norm.startswith("3  conclusiones") or
                (ls.startswith("**") and "conclusiones" in ls.lower() and "estado actual" in ls.lower()) or
                (ls.startswith("#") and "conclusiones" in ls.lower())
            )
            if not es_artefacto:
                lineas_filtradas.append(lin)
        conc_limpia = "\n".join(lineas_filtradas).strip()
        for bloque in conc_limpia.split("\n\n"):
            bloque=bloque.strip()
            if not bloque: continue
            # Limpiar markdown de heading (# al inicio de línea)
            lineas_bloque = bloque.split("\n")
            lineas_limpias = []
            for lin in lineas_bloque:
                lin_strip = lin.lstrip()
                if lin_strip.startswith("# "):
                    lin = lin_strip[2:].strip()
                elif lin_strip.startswith("## "):
                    lin = lin_strip[3:].strip()
                if lin: lineas_limpias.append(lin)
            bloque = "\n".join(lineas_limpias).strip()
            if not bloque: continue
            p=doc.add_paragraph()
            parts=re.split('([*][*][^*]+[*][*])',bloque)
            for part in parts:
                if part.startswith("**") and part.endswith("**"): run=p.add_run(part[2:-2]); run.bold=True
                else: p.add_run(part)
    else:
        def bold_par(doc, titulo, texto):
            p=doc.add_paragraph(); r=p.add_run(titulo); r.bold=True; p.add_run("\n"+texto)

        # Mejora 5a: calcular frases usando la misma fuente que el bloque IA
        _frases_fb = calcular_frases({
            "ult_total":fmt(ult_t), "ant_total":fmt(ant_t),
            "mes_ult_nombre":mes_label_largo(per_ult), "mes_ant_nombre":mes_label_largo(per_ant),
            "ult_pct_trans":pct(ult_tr,ult_t), "ant_pct_trans":pct(ant_tr,ant_t),
            "ult_carg_trans_n":fmt(sum(n(r.get("CARGADOS",0)) for r in [next((x for x in ev_trans if x.get("MES")==per_ult),{})])),
            "ult_carg_tot_n":fmt(sum(n(r.get("cargado",0)) for r in impoexpo_ult)),
            "ult_carg_no_trans_n":"0","ult_carg_tardio_n":"0",
            "ult_pct_trans_carg":"N/D","ult_pct_no_trans_carg":"N/D","ult_pct_tardio_carg":"N/D",
            "ult_last_trans_n":fmt(n(next((x for x in ev_trans if x.get("MES")==per_ult),{}).get("LASTRE",0))),
            "ult_last_tot_n":fmt(sum(n(r.get("lastre",0)) for r in impoexpo_ult)),
            "ult_last_notrans_n":"0","ult_last_tardio_n":"0",
            "ult_pct_trans_last":"N/D","ult_pct_no_trans_last":"N/D","ult_pct_tardio_last":"N/D",
            "ant_last_trans_n":fmt(n(next((x for x in ev_trans if x.get("MES")==per_ant),{}).get("LASTRE",0))),
            "ant_last_tot_n":fmt(n(next((x for x in ev_total if x.get("MES")==per_ant),{}).get("LASTRE",0))),
            "ant_last_pct_trans":"N/D","ant_last_notrans_n":"0",
            "ult_rechazos":fmt(sum(n(r.get("Rechazos",0)) for r in rechazos_ult_cat)),
            "ant_rechazos":fmt(next((n(r.get("MIC_RECHAZOS",0)) for r in rechazos_mes if r.get("periodo")==per_ant),0)),
            "ult_rech_duplicados":"0","ult_rech_operativos":"0",
        })

        rult_total=sum(n(r.get("Rechazos",0)) for r in rechazos_ult_cat)
        rult_dup=next((n(r.get("Rechazos",0)) for r in rechazos_ult_cat if r.get("Categoria")=="NRO DE MIC EXISTENTE"),0)
        rult_op=rult_total-rult_dup
        top_rech=[r for r in rechazos_ult_cat if r.get("Categoria")!="NRO DE MIC EXISTENTE"][:5]
        top_txt=", ".join([f"{r['Categoria']}: {r['Rechazos']}" for r in top_rech])
        impo_r=next((r for r in impoexpo_ult if str(r.get("TIPO_REGISTRO","")).upper()=="I"),{})
        expo_r=next((r for r in impoexpo_ult if str(r.get("TIPO_REGISTRO","")).upper()=="E"),{})
        ult_impo_n=n(impo_r.get("total",0)); ult_expo_n=n(expo_r.get("total",0))
        carg_tot_n=sum(n(r.get("cargado",0)) for r in impoexpo_ult)
        last_tot_n=sum(n(r.get("lastre",0)) for r in impoexpo_ult)
        ev_trans_ult=next((r for r in ev_trans if r.get("MES")==per_ult),{})
        carg_trans_n=n(ev_trans_ult.get("CARGADOS",0)); last_trans_n=n(ev_trans_ult.get("LASTRE",0))
        ev_tardio_ult_fb=next((r for r in ev_tardio if r.get("MES")==per_ult),{})
        last_tardio_n=n(ev_tardio_ult_fb.get("LASTRE",0))
        last_notrans_n=last_tot_n-last_trans_n-last_tardio_n
        diff_trans=pct_f(ult_tr,ult_t)-pct_f(ant_tr,ant_t)
        pct_ult_trans=pct_f(ult_tr,ult_t)
        # 0,0% de transmisión se describe en términos neutrales — puede responder a la dinámica propia del circuito
        if pct_ult_trans==0.0 and pct_f(ant_tr,ant_t)==0.0: cal_trans="sin transmisi\u00f3n en el per\u00edodo"
        elif pct_ult_trans==0.0: cal_trans="sin transmisi\u00f3n"
        elif diff_trans>=20: cal_trans="mejora significativa"
        elif diff_trans>=5: cal_trans="mejora moderada"
        elif diff_trans>=1: cal_trans="leve mejora"
        elif diff_trans>=-1: cal_trans="nivel estable"
        elif diff_trans>=-5: cal_trans="leve retroceso"
        else: cal_trans="retroceso"
        # Usar rechazos_mes como fuente primaria (misma tabla que el informe muestra)
        rech_ant_n_mes=next((n(r.get("MIC_RECHAZOS",0)) for r in rechazos_mes if r.get("periodo")==per_ant),0)
        rech_ant_n=rech_ant_n_mes if rech_ant_n_mes>0 else total_rech_ant
        if rult_total>rech_ant_n*2: cal_rech="incremento notable"
        elif rult_total>rech_ant_n: cal_rech="incremento"
        elif rult_total<rech_ant_n: cal_rech="reducci\u00f3n"
        else: cal_rech="nivel estable"
        tots_mes_fb={r["MES"]:n(r.get("CARGADO",0))+n(r.get("LASTRE",0)) for r in ev_total}
        trs_mes_fb={r["MES"]:n(r.get("CARGADOS",0))+n(r.get("LASTRE",0)) for r in ev_trans}
        pct_por_mes={mes:pct_f(trs_mes_fb.get(mes,0),tot) for mes,tot in tots_mes_fb.items() if tot>0}
        mejor_mes=max(pct_por_mes,key=pct_por_mes.get) if pct_por_mes else per_ult
        es_mejor=mejor_mes==per_ult

        p_titulo=doc.add_paragraph()
        r_t=p_titulo.add_run(f"Impacto del mes de {mes_label_largo(per_ult)} \u2014 {PAISES.get(pais,pais)}/Argentina (SINTIA)")
        r_t.bold=True; r_t.font.size=Pt(12)
        doc.add_paragraph()

        bold_par(doc,"Volumen operativo",
            f"{mes_label_largo(per_ult)} registr\u00f3 {fmt(ult_t)} operaciones ({fmt(ult_impo_n)} importaciones y {fmt(ult_expo_n)} exportaciones), "
            f"{'una leve disminuci\u00f3n' if ult_t<ant_t else 'un incremento'} respecto de {mes_label_largo(per_ant)} ({fmt(ant_t)}). "
            f"La participaci\u00f3n de exportaciones fue del {pct(ult_expo_n,ult_t)}.")

        mejor_txt="el mes de mayor cumplimiento del per\u00edodo analizado" if es_mejor else f"un nivel de cumplimiento {'superior' if pct_f(ult_tr,ult_t)>pct_f(gT,gTot) else 'inferior'} al promedio del per\u00edodo ({pct(gT,gTot)})"
        if pct_ult_trans==0.0:
            bold_par(doc,f"Transmisi\u00f3n anticipada \u2014 {cal_trans}",
                f"La tasa de transmisi\u00f3n anticipada se mantuvo en {pct(ult_tr,ult_t)} en {mes_label_largo(per_ult)} "
                f"({fmt(ult_tr)} operaci\u00f3n{'es' if ult_tr!=1 else ''} transmitida{'s' if ult_tr!=1 else ''} de {fmt(ult_t)} totales), "
                f"sin variaci\u00f3n respecto a {mes_label_largo(per_ant)} ({pct(ant_tr,ant_t)}). "
                f"El {pct(ult_nt,ult_t)} de las operaciones permaneci\u00f3 sin transmitir, tanto en cargados como en lastre. "
                f"No se registraron transmisiones tard\u00edas." if ult_td==0 else
                f"El {pct(ult_nt,ult_t)} de las operaciones permaneci\u00f3 sin transmitir. "
                f"Las transmisiones tard\u00edas representaron el {pct(ult_td,ult_t)} ({fmt(ult_td)} operaciones).")
        if pct_ult_trans!=0.0:
            bold_par(doc,f"Transmisi\u00f3n anticipada \u2014 {cal_trans}",
                f"{mes_label_largo(per_ult)} es {mejor_txt}: el {pct(ult_tr,ult_t)} de los MICs fue transmitido previo al arribo "
                f"({pct(carg_trans_n,carg_tot_n)} en cargados, {pct(last_trans_n,last_tot_n)} en lastre \u2014 {fmt(last_trans_n)} de {fmt(last_tot_n)}), "
                f"frente al {pct(ant_tr,ant_t)} de {mes_label_largo(per_ant)}. "
                f"Los no transmitidos {'disminuyeron' if ult_nt<ant_nt else 'subieron'} al {pct(ult_nt,ult_t)}. "
                f"En lastre, {fmt(last_notrans_n)} operaciones no fueron transmitidas ({pct(last_notrans_n,last_tot_n)}). "
                f"{_frases_fb.get("frase_lastre_trans_var","")}. {_frases_fb.get("frase_lastre_notrans_var","")}. "
                f"Las transmisiones tardías representaron el {pct(ult_td,ult_t)} ({fmt(ult_td)} operaciones vs. {fmt(ant_td)} en {mes_label_largo(per_ant)}).")

        rech_texto=(
            f"{mes_label_largo(per_ult)} concentra el mayor volumen de rechazos del per\u00edodo, con {fmt(rult_total)} registros"
            if cal_rech=="incremento notable" else
            f"{mes_label_largo(per_ult)} registr\u00f3 {fmt(rult_total)} rechazos ({'m\u00e1s' if rult_total>rech_ant_n else 'menos'} que los {fmt(rech_ant_n)} de {mes_label_largo(per_ant)})")
        if rult_dup>0:
            rech_texto+=f", de los cuales {fmt(rult_dup)} corresponden a la categor\u00eda 'NRO DE MIC EXISTENTE' (MICs duplicados). Excluyendo esa categor\u00eda, los rechazos operativos suman {fmt(rult_op)}"
        if top_txt: rech_texto+=f", distribuidos principalmente en: {top_txt}"
        rech_texto+="."
        bold_par(doc,f"Rechazos — {cal_rech}",limpiar_salida_ia(rech_texto))

        if es_mejor and diff_trans>=15:
            concl=f"{mes_label_largo(per_ult)} marca un avance favorable en la transmisi\u00f3n anticipada, acerc\u00e1ndose a niveles de cumplimiento m\u00e1s s\u00f3lidos ({pct(ult_tr,ult_t)}). "
        elif pct_ult_trans==0.0:
            concl=f"La transmisi\u00f3n anticipada se mantuvo en 0,0% durante {mes_label_largo(per_ult)}. Se recomienda dar seguimiento a la evoluci\u00f3n de este indicador en los pr\u00f3ximos per\u00edodos. "
        elif diff_trans<0:
            concl=f"{mes_label_largo(per_ult)} muestra un retroceso en la transmisi\u00f3n anticipada respecto al mes anterior. Conviene revisar las causas asociadas a esta variaci\u00f3n. "
        else:
            concl=f"La tendencia en transmisi\u00f3n anticipada contin\u00faa mejorando de forma gradual, con {mes_label_largo(per_ult)} mostrando un {pct(ult_tr,ult_t)} de cumplimiento. "
        if rult_op>0 and top_rech:
            cat_nueva=top_rech[0]["Categoria"]
            concl+=f"Asimismo, se observa persistencia de rechazos en {cat_nueva}, lo que sugiere revisar la calidad de los datos transmitidos en esa categor\u00eda."
        bold_par(doc,"Conclusión",limpiar_salida_ia(concl))

    # Tabla evolución mensual sintética
    doc.add_paragraph()
    rows_ev_sint=[]
    totales_x_mes={r["MES"]:n(r.get("CARGADO",0))+n(r.get("LASTRE",0)) for r in ev_total}
    trans_x_mes={r["MES"]:n(r.get("CARGADOS",0))+n(r.get("LASTRE",0)) for r in ev_trans}
    tardio_x_mes={r["MES"]:n(r.get("CARGADOS",0))+n(r.get("LASTRE",0)) for r in ev_tardio}
    notrans_x_mes={r["MES"]:n(r.get("CARGADOS",0))+n(r.get("LASTRE",0)) for r in ev_no_trans}
    for r in ev_total:
        mes=r["MES"]; tot=totales_x_mes.get(mes,0)
        rows_ev_sint.append([mes_label(mes),pct(trans_x_mes.get(mes,0),tot),pct(notrans_x_mes.get(mes,0),tot),pct(tardio_x_mes.get(mes,0),tot),fmt(tot)])
    agregar_tabla_word(doc,["MES","% TRANS","% NO TRANS","% TARD\u00cdO","TOTAL"],rows_ev_sint,col_widths=[2.8,2.5,2.5,2.5,2.2])

    if "comparativo" in graficos: insertar_grafico(doc,graficos["comparativo"])

    if datos_ult and datos_ant:
        agregar_tabla_word(doc,
            ["PER\u00cdODO","TOTAL","TRANS","%","NO TRANS","%","TARD\u00cdO","%"],
            [[mes_label_largo(per_ant),fmt(ant_t),fmt(ant_tr),pct(ant_tr,ant_t),fmt(ant_nt),pct(ant_nt,ant_t),fmt(ant_td),pct(ant_td,ant_t)],
             [mes_label_largo(per_ult),fmt(ult_t),fmt(ult_tr),pct(ult_tr,ult_t),fmt(ult_nt),pct(ult_nt,ult_t),fmt(ult_td),pct(ult_td,ult_t)],
             [f"Acumulado {anio}",fmt(gTot),fmt(gT),pct(gT,gTot),fmt(gN),pct(gN,gTot),fmt(gTd),pct(gTd,gTot)]],
            col_widths=[3.0,1.8,1.6,1.2,1.8,1.2,1.6,1.2], semaforo_col=2, semaforo_total_col=1)

    if datos_interanual:
        ia_t=n(datos_interanual.get("total",0)); ia_tr=n(datos_interanual.get("trans",0))
        ia_nt=n(datos_interanual.get("no_trans",0)); ia_td=n(datos_interanual.get("tardio",0))
        doc.add_paragraph(f"Comparativo interanual ({periodo_texto(anio_ant,mes_d,mes_h)}):").runs[0].bold=True
        agregar_tabla_word(doc,
            ["PER\u00cdODO","TOTAL","TRANS","%","NO TRANS","%","TARD\u00cdO","%"],
            [[f"{anio_ant}",fmt(ia_t),fmt(ia_tr),pct(ia_tr,ia_t),fmt(ia_nt),pct(ia_nt,ia_t),fmt(ia_td),pct(ia_td,ia_t)],
             [f"{anio}",fmt(gTot),fmt(gT),pct(gT,gTot),fmt(gN),pct(gN,gTot),fmt(gTd),pct(gTd,gTot)]],
            col_widths=[3.0,1.8,1.6,1.2,1.8,1.2,1.6,1.2], semaforo_col=2, semaforo_total_col=1)

    # Apéndice — Glosario (texto fijo, no generado por IA, para que sea auditable)
    doc.add_page_break()
    _heading_indexado(doc, "Anexo — Glosario de siglas", 1, 11)
    glosario_items = [
        ("SINTIA", "Sistema de Información Aduanera — registro y control de la operatoria del circuito."),
        ("PAD", "Portal Aduanero — sistema central de registro de ingreso/egreso terrestre."),
        ("MIC", "Manifiesto Internacional de Cargas."),
        ("MIC-DTA", "Manifiesto Internacional de Cargas - Declaración de Tránsito Aduanero."),
        ("DTA", "Declaración de Tránsito Aduanero."),
        ("ARCA", "Agencia de Recaudación y Control Aduanero."),
        ("CRT", "Carta de Porte Internacional por Carretera."),
        ("INDNCM", "Indicador de Nomenclatura Común del Mercosur."),
        ("PATAI", "Presentación Anticipada de Transportes de Ingreso."),
        ("OFTAI", "Oficialización de Transportes de Ingreso."),
        ("Cargado", "Camión con mercadería declarada en el MIC-DTA."),
        ("Lastre", "Camión sin carga (vacío) en el circuito."),
    ]
    for sigla, desc in glosario_items:
        gp = doc.add_paragraph()
        gp.add_run(f"{sigla}: ").bold = True
        gp.add_run(desc)

    nombre=f"Informe_SINTIA_{pais}_{anio}_{mes_d}-{mes_h}_v{version}.docx"
    ruta=os.path.join(carpeta,nombre); doc.save(ruta); log_fn("✓ Informe Word generado")
    return ruta

# ── Generar Excel ─────────────────────────────────────────────────────────────────
def _generar_excel(pais, anio, mes_d, mes_h, version,
                   totales_raw, ev_total, ev_trans, ev_tardio, ev_no_trans,
                   rechazos_mes, rechazos_cat, rechazos_ej,
                   datos_ult, datos_ant, datos_interanual, per_ult, per_ant, anio_ant, carpeta, log_fn):
    wb=openpyxl.Workbook(); wb.remove(wb.active)
    HDR_FILL=PatternFill("solid",fgColor="1F3864"); HDR_FONT=Font(bold=True,color="FFFFFF",size=10)
    ALT_FILL=PatternFill("solid",fgColor="EEF2F7"); NORM_FONT=Font(size=10)
    CENTER=Alignment(horizontal="center",vertical="center"); LEFT=Alignment(horizontal="left",vertical="center")
    bs=Side(style="thin",color="CCCCCC"); BORDER=Border(left=bs,right=bs,top=bs,bottom=bs)
    def add_sheet(name, headers, rows, semaforo_col=None, semaforo_total_col=None):
        ws=wb.create_sheet(name); ws.append(headers)
        for ci,h in enumerate(headers,1):
            cell=ws.cell(1,ci); cell.fill=HDR_FILL; cell.font=HDR_FONT; cell.alignment=CENTER; cell.border=BORDER
        for ri,row in enumerate(rows,2):
            for ci,val in enumerate(row,1):
                cell=ws.cell(ri,ci,val); cell.font=NORM_FONT; cell.border=BORDER
                fill=ALT_FILL if ri%2==0 else PatternFill()
                if semaforo_col is not None and semaforo_total_col is not None and ci==semaforo_col+1:
                    try:
                        v_trans=n(row[semaforo_col]); v_total=n(row[semaforo_total_col])
                        if v_total>0: fill=PatternFill("solid",fgColor=color_semaforo(pct_f(v_trans,v_total)))
                    except: pass
                cell.fill=fill; cell.alignment=LEFT
        for ci in range(1,len(headers)+1):
            col=get_column_letter(ci)
            ws.column_dimensions[col].width=max(len(str(headers[ci-1])),max((len(str(r[ci-1] if ci-1<len(r) else "")) for r in rows),default=0))+3
        return ws
    (cT,cN,cTd,cTot),(lT,lN,lTd,lTot),(gT,gN,gTd,gTot)=calcular_totales(totales_raw)
    add_sheet("Resumen Global",["Tipo","Trans","No Trans","Tardio","Total","% Trans","% No Trans","% Tardio"],
        [["Cargado",cT,cN,cTd,cTot,pct(cT,cTot),pct(cN,cTot),pct(cTd,cTot)],
         ["Lastre", lT,lN,lTd,lTot,pct(lT,lTot),pct(lN,lTot),pct(lTd,lTot)],
         ["Total",  gT,gN,gTd,gTot,pct(gT,gTot),pct(gN,gTot),pct(gTd,gTot)]])
    add_sheet("Ev Total",["Mes","Cargado","Lastre","Total"],
        [[mes_label_largo(r["MES"]),n(r.get("CARGADO",0)),n(r.get("LASTRE",0)),n(r.get("CARGADO",0))+n(r.get("LASTRE",0))] for r in ev_total])
    add_sheet("Ev Transmitidos",["Mes","Cargados","Lastre","Total"],
        [[mes_label_largo(r["MES"]),n(r.get("CARGADOS",0)),n(r.get("LASTRE",0)),n(r.get("CARGADOS",0))+n(r.get("LASTRE",0))] for r in ev_trans])
    add_sheet("Ev Tardios",["Mes","Cargados","Lastre","Total"],
        [[mes_label_largo(r["MES"]),n(r.get("CARGADOS",0)),n(r.get("LASTRE",0)),n(r.get("CARGADOS",0))+n(r.get("LASTRE",0))] for r in ev_tardio])
    add_sheet("Ev No Transmitidos",["Mes","Cargados","Lastre","Total"],
        [[mes_label_largo(r["MES"]),n(r.get("CARGADOS",0)),n(r.get("LASTRE",0)),n(r.get("CARGADOS",0))+n(r.get("LASTRE",0))] for r in ev_no_trans])
    add_sheet("Rechazos x Mes",["Periodo","MIC Rechazos"],
        [[r.get("periodo",""),n(r.get("MIC_RECHAZOS",0))] for r in rechazos_mes])
    add_sheet("Rechazos x Categoria",["Categoria","Rechazos"],
        [[r.get("Categoria",""),n(r.get("Rechazos",0))] for r in rechazos_cat])
    add_sheet("Rechazos Ejemplos",["Fecha","NroMic","Mensaje"],
        [[r.get("Fecha_ISO","")[:10] if r.get("Fecha_ISO") else "",r.get("NroMic",""),r.get("Mensaje","")] for r in rechazos_ej])
    rows_comp=[]
    if datos_ant:
        at=n(datos_ant.get("total",0)); atr=n(datos_ant.get("trans",0)); ant_=n(datos_ant.get("no_trans",0)); atd=n(datos_ant.get("tardio",0))
        rows_comp.append([mes_label_largo(per_ant),at,atr,pct(atr,at),ant_,pct(ant_,at),atd,pct(atd,at)])
    if datos_ult:
        ut=n(datos_ult.get("total",0)); utr=n(datos_ult.get("trans",0)); unt=n(datos_ult.get("no_trans",0)); utd=n(datos_ult.get("tardio",0))
        rows_comp.append([mes_label_largo(per_ult),ut,utr,pct(utr,ut),unt,pct(unt,ut),utd,pct(utd,ut)])
    rows_comp.append([f"Acumulado {anio}",gTot,gT,pct(gT,gTot),gN,pct(gN,gTot),gTd,pct(gTd,gTot)])
    if datos_interanual:
        ia_t=n(datos_interanual.get("total",0)); ia_tr=n(datos_interanual.get("trans",0))
        ia_nt=n(datos_interanual.get("no_trans",0)); ia_td=n(datos_interanual.get("tardio",0))
        rows_comp.append([f"Mismo per\u00edodo {anio_ant}",ia_t,ia_tr,pct(ia_tr,ia_t),ia_nt,pct(ia_nt,ia_t),ia_td,pct(ia_td,ia_t)])
    add_sheet("Comparativo",["Per\u00edodo","Total","Trans","%","No Trans","%","Tard\u00edo","%"],rows_comp)
    nombre=f"Informe_SINTIA_{pais}_{anio}_{mes_d}-{mes_h}_v{version}.xlsx"
    ruta=os.path.join(carpeta,nombre); wb.save(ruta); log_fn("✓ Planilla Excel generada")
    return ruta

# ── Punto de entrada principal ─────────────────────────────────────────────────

# ── Informe consolidado multi-país (Fase 7) ──────────────────────────────────
def _periodo_texto_fechas(fecha_d, fecha_h):
    d = datetime.strptime(fecha_d, "%Y-%m-%d").strftime("%d/%m/%Y")
    h = datetime.strptime(fecha_h, "%Y-%m-%d").strftime("%d/%m/%Y")
    return f"{d} a {h}" if d != h else d

def _alertas_demora_aduanas(por_aduana, umbral_pct_alerta=5.0, factor_outlier_demora=3.0):
    """Identifica aduanas con demora media muy por encima del resto del
    grupo, o con un porcentaje alto de operaciones "en alerta" (superaron
    el umbral PAD). No inventa un criterio nuevo -- reusa
    DEMORA_MEDIA_DIAS/EN_ALERTA_PAD que ya calcula
    correr_queries_consolidado() (ese promedio ya excluye los outliers más
    extremos, > umbral_alerta_dias, así que un valor alto acá es "elevado
    pero no descartado", no un caso aislado).

    factor_outlier_demora: cuántas veces la MEDIANA del grupo tiene que ser
    la demora de una aduana para marcarla (mediana en vez de promedio,
    justamente para no dejar que un par de aduanas muy lentas arrastren el
    punto de comparación).
    umbral_pct_alerta: % de operaciones en alerta a partir del cual se
    marca, independientemente de si la demora media también es alta.

    Devuelve una lista de strings (uno por aduana marcada) para insertar
    como texto en el informe. Lista vacía si ninguna aduana se destaca."""
    con_datos = [r for r in por_aduana if r.get("DEMORA_MEDIA_DIAS") is not None and n(r.get("TOTAL", 0)) > 0]
    if not con_datos:
        return []
    mediana = statistics.median(r["DEMORA_MEDIA_DIAS"] for r in con_datos)
    avisos = []
    for r in con_datos:
        nombre = r.get("ADUANA_NOMBRE", r["ADUANA"])
        total = n(r.get("TOTAL", 0))
        en_alerta = n(r.get("EN_ALERTA_PAD", 0))
        pct_alerta = pct_f(en_alerta, total)
        demora = r["DEMORA_MEDIA_DIAS"]
        es_outlier_demora = mediana > 0 and demora >= mediana * factor_outlier_demora
        es_alto_pct_alerta = pct_alerta >= umbral_pct_alerta
        if not (es_outlier_demora or es_alto_pct_alerta):
            continue
        partes = []
        if es_outlier_demora:
            partes.append(f"demora media de {formatear_demora(demora)} (mediana del resto: {formatear_demora(mediana)})")
        if es_alto_pct_alerta:
            partes.append(f"{pct(en_alerta, total)} de sus operaciones en alerta ({fmt(en_alerta)} de {fmt(total)})")
        avisos.append(f"{nombre}: " + " y ".join(partes) + ".")
    return avisos


def _agrupar_por_dira(por_aduana):
    """Subtotales por DIRA (región) a partir de las filas por_aduana, para
    mostrar antes del detalle aduana por aduana -- 26 aduanas sueltas sin
    agrupar es difícil de escanear; las ~6 DIRA dan una vista de más alto
    nivel primero. Ordenado por TOTAL descendente, igual criterio que la
    tabla de detalle. Aduanas sin DIRA asignada quedan agrupadas en
    "Sin DIRA asignada" en vez de desaparecer."""
    por_dira = {}
    for r in por_aduana:
        dira = r.get("DIRA_NOMBRE") or "Sin DIRA asignada"
        acc = por_dira.setdefault(dira, {"TOTAL": 0, "IMPO": 0, "EXPO": 0, "CARGADO": 0, "LASTRE": 0})
        for campo in ("TOTAL", "IMPO", "EXPO", "CARGADO", "LASTRE"):
            acc[campo] += n(r.get(campo, 0))
    filas = [{"DIRA": dira, **vals} for dira, vals in por_dira.items()]
    filas.sort(key=lambda f: f["TOTAL"], reverse=True)
    return filas


def _agregar_controles_por_tipo(por_tipo_control):
    """Suma por_tipo_control (que viene ABIERTO por aduana, ver
    correr_queries_consolidado) agrupado solo por tipo -- para el gráfico
    resumen nacional, que muestra el total de cada tipo de control sin
    desglosar por aduana (el desglose por aduana va en la tabla, no en
    este gráfico)."""
    agregado = {}
    for r in por_tipo_control or []:
        cod = r["CODIGO"]
        agregado[cod] = agregado.get(cod, 0) + n(r.get("TOTAL", 0))
    return [{"CODIGO": k, "TOTAL": v} for k, v in agregado.items()]


def _asimetria_impoexpo_paises(por_pais, umbral_pct=65.0, min_operaciones=1000):
    """Identifica países donde IMPO o EXPO domina fuertemente sobre el
    otro -- el % del Resumen Ejecutivo es un promedio del TOTAL general, y
    puede salir cerca de 50/50 aunque países individuales estén muy lejos
    de eso (compensándose entre sí, ej. uno mayormente exportador y otro
    mayormente importador). "OTRO/SIN DATO" queda afuera a propósito, no
    es un país real. min_operaciones filtra países con muy poco volumen
    (un 100% de 3 operaciones no dice nada). Devuelve una lista de
    strings ("País: XX,X% importación/exportación"), vacía si ningún país
    supera el umbral -- se recalcula en cada informe, no son valores fijos."""
    avisos = []
    for r in por_pais:
        cod = r.get("PAIS")
        if cod == "OTRO/SIN DATO":
            continue
        total = n(r.get("TOTAL", 0))
        if total < min_operaciones:
            continue
        impo = n(r.get("IMPO", 0)); expo = n(r.get("EXPO", 0))
        nombre = PAISES_CONSOLIDADO.get(cod, cod)
        if pct_f(impo, total) >= umbral_pct:
            avisos.append(f"{nombre} ({pct(impo, total)} importación)")
        elif pct_f(expo, total) >= umbral_pct:
            avisos.append(f"{nombre} ({pct(expo, total)} exportación)")
    return avisos


def _variable_control_dominante(por_var_control, total, umbral_pct=60.0):
    """Si una variable de control concentra la mayoría de las operaciones,
    devuelve (nombre, pct_texto) para comentarlo -- si no, None. Recalculado
    en cada informe según los datos reales, no asume que siempre sea SINVC
    (podría ser cualquier otra si la operatoria cambia)."""
    if not por_var_control or not total:
        return None
    top = max(por_var_control, key=lambda r: n(r.get("TOTAL", 0)))
    pct_top = pct_f(top.get("TOTAL", 0), total)
    if pct_top < umbral_pct:
        return None
    return top["VAR_CONTROL"], pct(top.get("TOTAL", 0), total)


def _generar_word_consolidado(fecha_d, fecha_h, version, totales, por_pais, por_aduana,
                               por_var_control, comparacion_anual, carpeta, log_fn, por_tipo_control=None):
    periodo = _periodo_texto_fechas(fecha_d, fecha_h)
    nombre_archivo = f"Informe_SINTIA_Consolidado_{fecha_d}_{fecha_h}_v{version}.docx"
    total = n(totales.get("TOTAL", 0)); impo = n(totales.get("IMPO", 0)); expo = n(totales.get("EXPO", 0))
    cargado = n(totales.get("CARGADO", 0)); lastre = n(totales.get("LASTRE", 0))

    graficos = {}
    if MPL_OK:
        log_fn("Generando gr\u00e1ficos...")
        for nombre, fn in [
            ("pais",       lambda: grafico_consolidado_pais(por_pais)),
            ("impoexpo",   lambda: grafico_consolidado_impoexpo(impo, expo)),
            ("cargalast",  lambda: grafico_consolidado_cargado_lastre(cargado, lastre)),
            ("aduana",     lambda: grafico_consolidado_aduana(por_aduana)),
            ("varcontrol", lambda: grafico_consolidado_var_control(por_var_control)),
            ("controles", lambda: grafico_controles_por_tipo(_agregar_controles_por_tipo(por_tipo_control))),
            ("interanual", lambda: grafico_comparacion_interanual(comparacion_anual)),
        ]:
            try: graficos[nombre] = fn()
            except Exception as e: log_fn(f"  Gr\u00e1fico {nombre}: {e}")
        log_fn("✓ Gr\u00e1ficos generados")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3); section.right_margin = Cm(2.5)
    _agregar_encabezado(doc, "Dirección de Reingeniería de Procesos Aduaneros")
    _agregar_pie_pagina(doc, "Informe Consolidado de Operaciones Fronterizas Terrestres")
    _ocultar_encabezado_portada(doc)

    # Portada
    _imagen_portada = _generar_portada_compuesta(
        titulo="INFORME CONSOLIDADO DE OPERACIONES FRONTERIZAS TERRESTRES",
        subtitulo=f"Período: {periodo}",
        meta_lineas=[
            "Dirección de Reingeniería de Procesos Aduaneros (DG ADUA)",
            f"Versión: {version}     Última modificación: {datetime.today().strftime('%d/%m/%Y')}",
            "Elaborado por: Sección Simplificación de Procesos Operativos — DI REPA",
        ])
    if _imagen_portada:
        _agregar_imagen_portada_ajustada(doc, _imagen_portada)
    else:
        titulo = doc.add_heading("INFORME CONSOLIDADO DE OPERACIONES FRONTERIZAS TERRESTRES", 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for txt, sz in [("Dirección de Reingeniería de Procesos Aduaneros (DG ADUA)", 12),
                         (f"Período: {periodo}", 11)]:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(txt); run.font.size = Pt(sz); run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        doc.add_paragraph()
        meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run("Versión: ").bold = True; meta.add_run(f"{version}   ")
        meta.add_run("Última modificación: ").bold = True; meta.add_run(datetime.today().strftime("%d/%m/%Y"))
        dest = doc.add_paragraph(); dest.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dest.add_run("Elaborado por: ").bold = True; dest.add_run("Sección Simplificación de Procesos Operativos — DI REPA")
    doc.add_page_break()

    _insertar_indice(doc, [
        (1, "Resumen Ejecutivo"),
        (1, "1.  Operaciones por país"),
        (1, "2.  Importación vs. Exportación"),
        (1, "3.  Cargado vs. Lastre"),
        (1, "4.  Operaciones por aduana"),
        (1, "5.  Operaciones por variable de control"),
    ] + ([(1, "6.  Comparación interanual (meses cerrados)")] if comparacion_anual else []) + [
        (1, "Anexo — Glosario de siglas"),
    ])

    # Resumen ejecutivo
    _heading_indexado(doc, "Resumen Ejecutivo", 1, 1)
    doc.add_paragraph(
        f"El presente informe consolida la totalidad de las operaciones registradas en SINTIA para "
        f"todos los países del circuito durante el período {periodo}, sin discriminar por país "
        f"emisor, a diferencia del informe SINTIA estándar (que se genera por país y período).")
    kpi_box(doc, [
        ("TOTAL OPERACIONES", fmt(total),   periodo),
        ("IMPORTACIÓN",       fmt(impo),    pct(impo, total)),
        ("EXPORTACIÓN",       fmt(expo),    pct(expo, total)),
        ("CARGADO",           fmt(cargado), pct(cargado, total)),
        ("LASTRE",            fmt(lastre),  pct(lastre, total)),
    ])
    doc.add_page_break()

    # 1. Operaciones por país
    _heading_indexado(doc, "1.  Operaciones por país", 1, 2)
    doc.add_paragraph(f"Durante el período se registraron {fmt(total)} operaciones distribuidas entre "
                       f"{len(por_pais)} {pl(len(por_pais), 'país/agrupación', 'países/agrupaciones')}.")
    agregar_tabla_word(doc, ["PAÍS", "TOTAL", "IMPO", "EXPO", "CARGADO", "LASTRE"],
        [[PAISES_CONSOLIDADO.get(r["PAIS"], r["PAIS"]), fmt(r.get("TOTAL", 0)), fmt(r.get("IMPO", 0)),
          fmt(r.get("EXPO", 0)), fmt(r.get("CARGADO", 0)), fmt(r.get("LASTRE", 0))] for r in por_pais],
        col_widths=[3.5, 2, 1.8, 1.8, 2, 1.8])
    if "pais" in graficos: insertar_grafico(doc, graficos["pais"])

    # 2. Importación vs. Exportación
    _heading_indexado(doc, "2.  Importación vs. Exportación", 1, 3)
    doc.add_paragraph(f"Del total de operaciones, {fmt(impo)} ({pct(impo, total)}) corresponden a "
                       f"importación y {fmt(expo)} ({pct(expo, total)}) a exportación.")
    asimetrias = _asimetria_impoexpo_paises(por_pais)
    if asimetrias:
        doc.add_paragraph(
            "Este porcentaje es un promedio del total general: individualmente, algunos países están "
            "lejos del 50/50 (se compensan entre sí en el agregado) — " + ", ".join(asimetrias) + ".")
    if "impoexpo" in graficos: insertar_grafico(doc, graficos["impoexpo"], width_cm=11)

    # 3. Cargado vs. Lastre
    _heading_indexado(doc, "3.  Cargado vs. Lastre", 1, 4)
    doc.add_paragraph(f"{fmt(cargado)} ({pct(cargado, total)}) operaciones fueron con mercadería "
                       f"cargada y {fmt(lastre)} ({pct(lastre, total)}) en lastre (vacío).")
    if "cargalast" in graficos: insertar_grafico(doc, graficos["cargalast"], width_cm=11)
    doc.add_page_break()

    # 4. Operaciones por aduana
    _heading_indexado(doc, "4.  Operaciones por aduana", 1, 5)
    doc.add_paragraph(
        f"Se relevaron operaciones en {len(por_aduana)} {pl(len(por_aduana), 'aduana')} durante el período. "
        f"\"Demora media\" y \"En alerta\" son la misma métrica PAD (tiempo entre ingreso y salida) "
        f"que usa el informe \"Aduanas del país\" — se muestran acá para cruzar en una sola tabla "
        f"volumen de operaciones (SINTIA) con tiempos de desaduanamiento (PAD) por aduana.")

    por_dira = _agrupar_por_dira(por_aduana)
    if len(por_dira) > 1:
        doc.add_paragraph().add_run("Subtotales por DIRA").bold = True
        agregar_tabla_word(doc, ["DIRA", "TOTAL", "IMPO", "EXPO", "CARGADO", "LASTRE"],
            [[f["DIRA"], fmt(f["TOTAL"]), fmt(f["IMPO"]), fmt(f["EXPO"]), fmt(f["CARGADO"]), fmt(f["LASTRE"])]
             for f in por_dira],
            col_widths=[4.5, 2.2, 2.0, 2.0, 2.2, 2.0])

    avisos_demora = _alertas_demora_aduanas(por_aduana)
    if avisos_demora:
        p_alerta = doc.add_paragraph()
        p_alerta.add_run("⚠ Aduanas con demora fuera de lo habitual: ").bold = True
        for aviso in avisos_demora:
            doc.add_paragraph(aviso, style="List Bullet")

    doc.add_paragraph().add_run("Detalle por aduana").bold = True
    agregar_tabla_word(doc, ["ADUANA", "DIRA", "TOTAL", "IMPO", "EXPO", "CARGADO", "LASTRE",
                              "DEMORA MEDIA", "EN ALERTA"],
        [[r.get("ADUANA_NOMBRE", r["ADUANA"]), r.get("DIRA_NOMBRE", "—"), fmt(r.get("TOTAL", 0)),
          fmt(r.get("IMPO", 0)), fmt(r.get("EXPO", 0)), fmt(r.get("CARGADO", 0)), fmt(r.get("LASTRE", 0)),
          r.get("DEMORA_MEDIA_FMT") or "—", fmt(r.get("EN_ALERTA_PAD", 0))]
         for r in por_aduana],
        col_widths=[3.0, 2.3, 1.3, 1.2, 1.2, 1.4, 1.3, 2.0, 1.3])
    if "aduana" in graficos: insertar_grafico(doc, graficos["aduana"])
    doc.add_page_break()

    # 5. Variables de control
    _heading_indexado(doc, "5.  Operaciones por variable de control", 1, 6)
    sin_dato = next((n(r.get("TOTAL", 0)) for r in por_var_control if r["VAR_CONTROL"] == "SIN VARIABLE DE CONTROL"), 0)
    doc.add_paragraph(
        f"Se identificaron {len(por_var_control)} "
        f"{pl(len(por_var_control), 'variable de control distinta', 'variables de control distintas')} en el período."
        + (f" {fmt(sin_dato)} {pl(sin_dato, 'operación', 'operaciones')} ({pct(sin_dato, total)}) no "
           f"{pl(sin_dato, 'tiene', 'tienen')} variable de control registrada." if sin_dato else ""))
    dominante = _variable_control_dominante(por_var_control, total)
    if dominante:
        var_nombre, var_pct = dominante
        if var_nombre == "SINVC":
            doc.add_paragraph(
                f"El {var_pct} de las operaciones no pasó por ningún criterio de selectividad (SINVC) — "
                f"solo una minoría fue seleccionada por algún criterio operativo, aleatorio o determinístico.")
        else:
            doc.add_paragraph(
                f"\"{var_nombre}\" concentra el {var_pct} de las operaciones, muy por encima del resto "
                f"de las variables de control del período.")
    agregar_tabla_word(doc, ["VARIABLE DE CONTROL", "TOTAL", "%"],
        [[r["VAR_CONTROL"], fmt(r.get("TOTAL", 0)), pct(r.get("TOTAL", 0), total)] for r in por_var_control],
        col_widths=[7, 2.5, 2])
    if "varcontrol" in graficos: insertar_grafico(doc, graficos["varcontrol"])

    if por_tipo_control:
        doc.add_paragraph()
        doc.add_paragraph().add_run("Tipos de control efectuados").bold = True
        doc.add_paragraph(
            "Cantidad de operaciones de cada aduana sobre las que se efectuó cada tipo de control en "
            "el período (catálogo de tipos editable en /admin — no todas las combinaciones de aduana/"
            "tipo tienen por qué haberse dado en este período puntual; solo se muestran las que sí).")
        filas_ordenadas = sorted(por_tipo_control, key=lambda r: (r.get("DIRA_NOMBRE", ""), r.get("ADUANA_NOMBRE", ""), r["CODIGO"]))
        agregar_tabla_word(doc, ["ADUANA", "TIPO DE CONTROL", "CANT. CONTROLES", "CANT. OPERACIONES"],
            [[r.get("ADUANA_NOMBRE", r.get("ADUANA", "")), r["CODIGO"], fmt(r.get("TOTAL", 0)),
              fmt(r.get("CANT_OPERACIONES", 0))] for r in filas_ordenadas],
            col_widths=[4.5, 3.5, 3, 3])
        if "controles" in graficos: insertar_grafico(doc, graficos["controles"])

    # 6. Comparación interanual (condicional -- solo si hay al menos un mes
    # cerrado del año en curso para comparar; ver comparacion_anual_meses_completos)
    indice_anexo = 7
    if comparacion_anual:
        doc.add_page_break()
        _heading_indexado(doc, "6.  Comparación interanual (meses cerrados)", 1, 7)
        anio_act = comparacion_anual[0]["anio_actual"]; anio_ant = comparacion_anual[0]["anio_anterior"]
        doc.add_paragraph(
            f"Comparación mes a mes entre {anio_act} y {anio_ant}, considerando únicamente los meses "
            f"de {anio_act} que ya terminaron (no se incluye el mes en curso, cuya cifra todavía está "
            f"a medio transcurrir y no es comparable contra un mes completo).")
        agregar_tabla_word(doc, ["MES", f"{anio_ant}", f"{anio_act}", "VAR. %"],
            [[r["mes_label"], fmt(r["total_anterior"]) if r["total_anterior"] is not None else "—",
              fmt(r["total_actual"]),
              (f"{'+' if r['variacion_pct'] >= 0 else ''}{r['variacion_pct']}%".replace(".", ",")
               if r["variacion_pct"] is not None else "—")]
             for r in comparacion_anual],
            col_widths=[4, 3, 3, 3])
        if "interanual" in graficos: insertar_grafico(doc, graficos["interanual"], width_cm=11)
        indice_anexo = 8

    # Anexo — Glosario (mismo que el informe SINTIA estándar)
    doc.add_page_break()
    _heading_indexado(doc, "Anexo — Glosario de siglas", 1, indice_anexo)
    glosario_items = [
        ("SINTIA", "Sistema de Información Aduanera — registro y control de la operatoria del circuito."),
        ("PAD", "Portal Aduanero — sistema central de registro de ingreso/egreso terrestre."),
        ("MIC", "Manifiesto Internacional de Cargas."),
        ("MIC-DTA", "Manifiesto Internacional de Cargas - Declaración de Tránsito Aduanero."),
        ("Cargado", "Camión con mercadería declarada en el MIC-DTA."),
        ("Lastre", "Camión sin carga (vacío) en el circuito."),
        ("Variable de control", "Criterio/regla de selectividad aplicado a la operación en SINTIA."),
        ("SINVC", "Sin variable de control asignada — la operación no fue seleccionada por ningún criterio de selectividad."),
        ("OPERA", "Variable de control aplicada por decisión operativa (no aleatoria ni determinística)."),
        ("ALEAT", "Variable de control aplicada por sorteo aleatorio."),
        ("DETER", "Variable de control aplicada por regla determinística (criterio fijo, no aleatorio)."),
        ("DIRA", "Dirección Regional Aduanera — agrupa varias aduanas bajo una misma jurisdicción regional."),
        ("OTRO/SIN DATO", "Operaciones cuyo MIC no permitió identificar el país emisor con el criterio de detección actual (ver Sección 1)."),
    ]
    for sigla, desc in glosario_items:
        gp = doc.add_paragraph()
        gp.add_run(f"{sigla}: ").bold = True
        gp.add_run(desc)

    nombre = f"Informe_SINTIA_Consolidado_{fecha_d}_{fecha_h}_v{version}.docx"
    ruta = os.path.join(carpeta, nombre); doc.save(ruta); log_fn("✓ Informe Word generado")
    return ruta


def _generar_excel_consolidado(fecha_d, fecha_h, version, totales, por_pais, por_aduana,
                                por_var_control, comparacion_anual, carpeta, log_fn, por_tipo_control=None):
    wb = openpyxl.Workbook(); wb.remove(wb.active)
    HDR_FILL = PatternFill("solid", fgColor="1F3864"); HDR_FONT = Font(bold=True, color="FFFFFF", size=10)
    ALT_FILL = PatternFill("solid", fgColor="EEF2F7"); NORM_FONT = Font(size=10)
    CENTER = Alignment(horizontal="center", vertical="center"); LEFT = Alignment(horizontal="left", vertical="center")
    bs = Side(style="thin", color="CCCCCC"); BORDER = Border(left=bs, right=bs, top=bs, bottom=bs)

    def add_sheet(name, headers, rows):
        ws = wb.create_sheet(name); ws.append(headers)
        for ci, h in enumerate(headers, 1):
            cell = ws.cell(1, ci); cell.fill = HDR_FILL; cell.font = HDR_FONT; cell.alignment = CENTER; cell.border = BORDER
        for ri, row in enumerate(rows, 2):
            for ci, val in enumerate(row, 1):
                cell = ws.cell(ri, ci, val); cell.font = NORM_FONT; cell.border = BORDER
                cell.fill = ALT_FILL if ri % 2 == 0 else PatternFill(); cell.alignment = LEFT
        for ci in range(1, len(headers) + 1):
            col = get_column_letter(ci)
            ws.column_dimensions[col].width = max(
                len(str(headers[ci - 1])),
                max((len(str(r[ci - 1] if ci - 1 < len(r) else "")) for r in rows), default=0)) + 3
        return ws

    total = n(totales.get("TOTAL", 0))
    add_sheet("Resumen General", ["Indicador", "Valor", "%"], [
        ["Total operaciones", total, "100,0%"],
        ["Importación", n(totales.get("IMPO", 0)), pct(totales.get("IMPO", 0), total)],
        ["Exportación", n(totales.get("EXPO", 0)), pct(totales.get("EXPO", 0), total)],
        ["Cargado", n(totales.get("CARGADO", 0)), pct(totales.get("CARGADO", 0), total)],
        ["Lastre", n(totales.get("LASTRE", 0)), pct(totales.get("LASTRE", 0), total)],
    ])
    add_sheet("Por País", ["País", "Total", "Impo", "Expo", "Cargado", "Lastre"],
        [[PAISES_CONSOLIDADO.get(r["PAIS"], r["PAIS"]), n(r.get("TOTAL", 0)), n(r.get("IMPO", 0)),
          n(r.get("EXPO", 0)), n(r.get("CARGADO", 0)), n(r.get("LASTRE", 0))] for r in por_pais])

    por_dira = _agrupar_por_dira(por_aduana)
    if len(por_dira) > 1:
        add_sheet("Por DIRA", ["DIRA", "Total", "Impo", "Expo", "Cargado", "Lastre"],
            [[f["DIRA"], f["TOTAL"], f["IMPO"], f["EXPO"], f["CARGADO"], f["LASTRE"]] for f in por_dira])

    # Misma detección que en el Word (_alertas_demora_aduanas) pero acá se
    # anota fila por fila en vez de como párrafo -- una columna "Alerta"
    # con el motivo, vacía si la aduana no se destaca.
    con_datos_demora = [r for r in por_aduana if r.get("DEMORA_MEDIA_DIAS") is not None and n(r.get("TOTAL", 0)) > 0]
    mediana_demora = statistics.median(r["DEMORA_MEDIA_DIAS"] for r in con_datos_demora) if con_datos_demora else 0

    def _motivo_alerta(r):
        if r.get("DEMORA_MEDIA_DIAS") is None or not n(r.get("TOTAL", 0)):
            return ""
        total = n(r.get("TOTAL", 0)); en_alerta = n(r.get("EN_ALERTA_PAD", 0))
        pct_alerta = pct_f(en_alerta, total)
        motivos = []
        if mediana_demora > 0 and r["DEMORA_MEDIA_DIAS"] >= mediana_demora * 3:
            motivos.append("demora muy por encima de la mediana")
        if pct_alerta >= 5.0:
            motivos.append(f"{pct(en_alerta, total)} en alerta")
        return "; ".join(motivos)

    add_sheet("Por Aduana", ["Aduana", "DIRA", "Total", "Impo", "Expo", "Cargado", "Lastre",
                              "Demora media PAD", "En alerta PAD", "Alerta"],
        [[r.get("ADUANA_NOMBRE", r["ADUANA"]), r.get("DIRA_NOMBRE", "—"), n(r.get("TOTAL", 0)),
          n(r.get("IMPO", 0)), n(r.get("EXPO", 0)), n(r.get("CARGADO", 0)), n(r.get("LASTRE", 0)),
          r.get("DEMORA_MEDIA_FMT") or "—", n(r.get("EN_ALERTA_PAD", 0)), _motivo_alerta(r)]
         for r in por_aduana])
    add_sheet("Por Variable de Control", ["Variable de Control", "Total", "%"],
        [[r["VAR_CONTROL"], n(r.get("TOTAL", 0)), pct(r.get("TOTAL", 0), total)] for r in por_var_control])

    if por_tipo_control:
        filas_ordenadas = sorted(por_tipo_control, key=lambda r: (r.get("DIRA_NOMBRE", ""), r.get("ADUANA_NOMBRE", ""), r["CODIGO"]))
        add_sheet("Tipos de Control", ["Aduana", "Tipo de Control", "Cant. Controles", "Cant. Operaciones"],
            [[r.get("ADUANA_NOMBRE", r.get("ADUANA", "")), r["CODIGO"], n(r.get("TOTAL", 0)),
              n(r.get("CANT_OPERACIONES", 0))] for r in filas_ordenadas])

    if comparacion_anual:
        anio_act = comparacion_anual[0]["anio_actual"]; anio_ant = comparacion_anual[0]["anio_anterior"]
        add_sheet("Comparación Interanual", ["Mes", f"{anio_ant}", f"{anio_act}", "Var. %"],
            [[r["mes_label"], r["total_anterior"] if r["total_anterior"] is not None else "—",
              r["total_actual"],
              (f"{r['variacion_pct']}%".replace(".", ",") if r["variacion_pct"] is not None else "—")]
             for r in comparacion_anual])

    nombre = f"Informe_SINTIA_Consolidado_{fecha_d}_{fecha_h}_v{version}.xlsx"
    ruta = os.path.join(carpeta, nombre); wb.save(ruta); log_fn("✓ Planilla Excel generada")
    return ruta
